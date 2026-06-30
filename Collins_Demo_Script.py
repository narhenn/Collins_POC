"""Generate Collins_Demo_Script.docx — a click-by-click demo storyboard."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x0C, 0x23, 0x40)  # Collins navy

NAVY = RGBColor(0x0C, 0x23, 0x40)
BLUE = RGBColor(0x4B, 0x8B, 0xF5)
TEAL = RGBColor(0x18, 0xA9, 0x99)
RED  = RGBColor(0xE2, 0x56, 0x4E)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_say(text):
    """Add a 'SAY:' block — what the presenter speaks."""
    p = doc.add_paragraph()
    run = p.add_run('SAY:  ')
    run.bold = True
    run.font.color.rgb = TEAL
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)

def add_do(text):
    """Add a 'DO:' block — what the presenter clicks."""
    p = doc.add_paragraph()
    run = p.add_run('DO:  ')
    run.bold = True
    run.font.color.rgb = BLUE
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)

def add_show(text):
    """Add a 'ON SCREEN:' block — what appears."""
    p = doc.add_paragraph()
    run = p.add_run('ON SCREEN:  ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x6D, 0xF0)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)

def add_note(text):
    """Add a presenter note."""
    p = doc.add_paragraph()
    run = p.add_run('NOTE:  ')
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(8)

def add_time(text):
    p = doc.add_paragraph()
    run = p.add_run(f'[{text}]')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run('─' * 72)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def shade_cell(cell, color_hex):
    """Shade a table cell."""
    shading = cell._tc.get_or_add_tcPr()
    sh = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:fill'): color_hex,
    })
    shading.append(sh)


# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COLLINS AEROSPACE')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MRO Facility Intelligence')
run.font.size = Pt(32)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Live Demo Script & Storyboard')
run.font.size = Pt(18)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NextXR Digital Twin  +  AUTOMIND  +  GoalCert')
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Meta table
tbl = doc.add_table(rows=5, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta = [
    ('Duration', '12-15 minutes'),
    ('Audience', 'Collins Aerospace — Engineering & MRO Leadership'),
    ('Presenter', 'GoalCert Team'),
    ('Date', datetime.date.today().strftime('%B %d, %Y')),
    ('URL', 'http://localhost:5174'),
]
for i, (k, v) in enumerate(meta):
    c0, c1 = tbl.rows[i].cells
    c0.text = k
    c1.text = v
    c0.paragraphs[0].runs[0].bold = True
    for c in (c0, c1):
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].runs[0].font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Demo Flow', level=1)
toc_items = [
    ('Pre-Demo Checklist', '30 sec'),
    ('Act 1 — Opening Hook', '1 min'),
    ('Act 2 — The Twins Library: "One Platform, Any Asset"', '2 min'),
    ('Act 3 — Wire EDM: Live Physics & AI Co-Pilot', '3 min'),
    ('Act 4 — Fault Injection & Real-Time Detection', '2 min'),
    ('Act 5 — The Agent Army: Diagnosis to Work Order', '3 min'),
    ('Act 6 — Scenario Sandbox: What-If Without Risk', '2 min'),
    ('Act 7 — Cross-Industry Proof: Same Platform, Different Twin', '1 min'),
    ('Act 8 — Closing: The 18-Step Vision', '1 min'),
    ('Recovery Playbook', '—'),
    ('Tough Questions & Answers', '—'),
]
toc = doc.add_table(rows=len(toc_items), cols=2)
for i, (title, time_) in enumerate(toc_items):
    toc.rows[i].cells[0].text = title
    toc.rows[i].cells[1].text = time_
    toc.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    toc.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    toc.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# PRE-DEMO CHECKLIST
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Pre-Demo Checklist', level=1)
p = doc.add_paragraph('Complete all items before the client enters the room.')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.size = Pt(10)

checks = [
    'Docker Desktop running (Neo4j, Postgres, Redis containers)',
    'NextXR server on port 8000 — verify http://localhost:8000/api/v1/health returns "healthy"',
    'AUTOMIND server on port 8001 — verify http://localhost:8001/docs loads',
    'GoalCert server on port 8002 — verify http://localhost:8002/api/health returns "ok"',
    'Orchestrator on port 8090 — verify http://localhost:8090/api/health shows all platforms connected',
    'Frontend on port 5174 — open http://localhost:5174 in Chrome (full screen, no bookmarks bar)',
    'Melbourne HQ seeded (cross-tenant intelligence demo)',
    'Browser zoom at 100%, dark mode OFF',
    'Close Slack, email, notifications — nothing should pop up during demo',
    'Pre-create one Wire EDM twin and one Gas Turbine twin so the Twins library shows them immediately',
    'Run 5-10 physics steps on the EDM twin so it has sensor history when you open it',
]
for c in checks:
    doc.add_paragraph(c, style='List Bullet')

startup = """cd /Users/narhen/Collins_POC
docker compose up -d neo4j postgres redis-am
cd nextxr-ontology && NEO4J_URI=bolt://localhost:7687 NEO4J_USER=neo4j NEO4J_PASSWORD=nextxr2026 NXR_REDIS_URL=redis://localhost:6379/0 python -m server.main &
cd ../automind-engine/backend && DATABASE_URL="postgresql+asyncpg://automind:automind2026@localhost:5433/automind" REDIS_URL="redis://localhost:6380/0" SECRET_KEY="collins-poc-secret-2026" DEBUG=true python -m uvicorn app.main:app --host 0.0.0.0 --port 8001 &
cd ../../goalcert-engine/backend && python -m uvicorn app.main:app --host 0.0.0.0 --port 8002 &
cd ../../collins-demo/orchestrator && python -c "import uvicorn; uvicorn.run('main:app', host='0.0.0.0', port=8090)" &
cd ../web && npx vite &
open http://localhost:5174"""

p = doc.add_paragraph()
run = p.add_run('Startup commands:')
run.bold = True
run.font.size = Pt(10)
p2 = doc.add_paragraph()
run2 = p2.add_run(startup)
run2.font.name = 'JetBrains Mono'
run2.font.size = Pt(8)
run2.font.color.rgb = GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 1 — OPENING HOOK
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 1 — Opening Hook', level=1)
add_time('1 minute')

add_say(
    '"Thank you for your time today. What we are about to show you is a live system, '
    'not slides, not a mockup. Everything you see is running right now, against real '
    'physics models, real graph databases, and real AI agents.'
)

add_say(
    '"Collins operates some of the most complex MRO facilities in the world. Turbine '
    'test cells, hydraulic rigs, avionics bays, each with hundreds of sensors generating '
    'data every second. The problem is not collecting that data. The problem is turning it '
    'into decisions before something breaks."'
)

add_say(
    '"We built a platform that does exactly that. Three products working together. '
    'NextXR builds a living digital twin of your facility. AUTOMIND orchestrates AI agents '
    'that diagnose, predict, and generate work orders. GoalCert trains your technicians on '
    'scenarios derived from real faults. Let me show you how they work."'
)

add_note('Keep the energy confident but calm. You are showing capability, not selling.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 2 — TWINS LIBRARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 2 — The Twins Library: "One Platform, Any Asset"', level=1)
add_time('2 minutes')

add_say(
    '"This is the Twins library. Every card you see is a different type of facility '
    'or machine that our platform can model. The same ontology backbone, the same physics '
    'engine, the same AI agents, adapted to each domain."'
)

add_do('Point to each card briefly: Wire EDM Machine, Helix Data Center, St. Vera Hospital, Forge Plant 7.')

add_show(
    '5 twin cards in a grid. Each shows: domain tag, signal count, "live" or "full physics" badges.'
)

add_say(
    '"For Collins, we are going to focus on two things: a precision Wire EDM machine, '
    'which is the kind of equipment your MRO shops use every day, and a gas turbine engine '
    'on a test rig. But I want you to notice that the same platform also models data centers, '
    'hospitals, manufacturing plants. The architecture is domain-agnostic. We are not rebuilding '
    'anything for each vertical."'
)

add_do('Click "Open twin" on the Wire EDM Machine card.')

add_show('Dashboard loads. Sensor cards start populating. 3D Wire EDM model renders.')

add_note(
    'If the twin was pre-created, the dashboard loads instantly with sensor history. '
    'If you need to create fresh, it takes 2-3 seconds. Have one ready.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 3 — LIVE PHYSICS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 3 — Wire EDM: Live Physics & AI Co-Pilot', level=1)
add_time('3 minutes')

add_say(
    '"This is a live digital twin of a CNC wire-cut EDM machine. Every signal you see '
    'is computed by a real physics model, not scripted values. Cutting speed, spark gap '
    'voltage, dielectric temperature, wire tension, short-circuit rate, all coupled together."'
)

add_do('Point to the sensor grid. Highlight the live indicator dots pulsing green.')

add_show(
    '8 sensor cards: Cutting Speed, Short Circuit Rate, Wire Break Risk, Dielectric Temp, '
    'Dielectric Conductivity, Wire Tension, Gap Voltage, Dielectric Pressure. All streaming.'
)

add_say(
    '"The 3D model on screen is not a static render. It is the actual Wire EDM machine '
    'modelled down to its C-frame, wire guides, dielectric tank, and CNC pendant. '
    'The discharge sparks you see at the cutting zone are driven by the physics engine."'
)

add_do('Point to the 3D Wire EDM model. Let it rotate slowly.')

add_divider()

add_say(
    '"Now look at the AI Co-Pilot section. This is not a chatbot waiting for you to type something. '
    'It is continuously watching every sensor stream and generating real-time observations, '
    'like an experienced engineer standing next to the machine."'
)

add_do('Scroll to the AI Co-Pilot card. Wait for a narration to appear.')

add_show(
    'AI Co-Pilot card with timestamps and narration lines like: '
    '"Dielectric temperature stable at 28C, cutting speed nominal at 4.2 mm\u00b2/min, '
    'wire tension within tolerance. No anomalies detected."'
)

add_say(
    '"Every observation is grounded in the actual sensor readings. Not hallucinated. '
    'The AI agent receives the live telemetry snapshot and reasons over it. '
    'And it is powered by Claude, Anthropic\'s latest model."'
)

add_note(
    'The narration refreshes every 8 seconds. If Claude API is slow, the spinner shows. '
    'Do not rush, let it arrive naturally.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 4 — FAULT INJECTION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 4 — Fault Injection & Real-Time Detection', level=1)
add_time('2 minutes')

add_say(
    '"Now here is where it gets interesting. I am going to deliberately inject a fault '
    'into the running machine and show you what happens in real time."'
)

add_do(
    'In the top bar, find the "Inject fault..." dropdown. '
    'Select "Wire breakage" or "Dielectric contamination".'
)

add_show(
    'Sensor cards start changing. Wire Break Risk climbs. Dielectric Conductivity drifts. '
    'Cards turn amber, then red. Health index drops from ~95% to below 70%.'
)

add_say(
    '"Watch the health index drop. The physics engine is propagating the fault through '
    'the coupled model. Wire break risk is climbing because the dielectric is contaminated, '
    'which changes the spark gap stability, which stresses the wire tension."'
)

add_do('Wait 5-10 seconds. Let findings appear in the Active Findings panel.')

add_show(
    'Active Findings panel populates: [critical] "Wire break risk elevated", '
    '[warning] "Dielectric conductivity above threshold". '
    'Incidents panel shows grouped incidents.'
)

add_say(
    '"Three tiers of detection just fired. Tier C caught the threshold breach. '
    'Tier B detected the statistical deviation from baseline. And Tier A, the physics tier, '
    'identified that the degradation pattern matches wire breakage, not just random noise. '
    'The findings are automatically grouped into incidents."'
)

add_do('Point to the severity badges on each finding (critical vs warning).')

add_note(
    'This is the "wow" moment. Let the findings accumulate visually before moving on. '
    'Do not click away too fast.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 5 — THE AGENT ARMY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 5 — The Agent Army: Diagnosis to Work Order', level=1)
add_time('3 minutes')

add_say(
    '"Now that we have a real fault detected, I want to show you the AI agents that '
    'take action. This is where the three products start working together."'
)

# --- Diagnosis ---
p = doc.add_paragraph()
run = p.add_run('5A. Diagnosis Agent')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY

add_do('Click "Twin Intelligence" in the sidebar. Click "Run Diagnosis".')

add_show(
    'Full per-component, per-sensor diagnostic report. Health percentages for each '
    'subsystem. Root cause identification. Recommended actions.'
)

add_say(
    '"The diagnosis agent just analyzed every component and every sensor in the twin. '
    'It identified the root cause as dielectric contamination leading to wire break risk, '
    'not just a symptom. And it cross-referenced the physics model to validate its reasoning."'
)

add_divider()

# --- Work Order ---
p = doc.add_paragraph()
run = p.add_run('5B. Maintenance Work Order Agent')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY

add_do('Go back to Dashboard. Scroll to the Work Order section. Click "Generate Work Order".')

add_show(
    'AS9100-compliant work order: WO number, ATA chapter reference, priority (AOG/Routine), '
    'fault description, root cause, numbered repair steps with pass/fail criteria and safety warnings, '
    'parts list, sign-off line.'
)

add_say(
    '"This is a full AS9100-compliant maintenance work order generated in seconds. '
    'ATA chapter reference, step-by-step repair procedure with safety warnings, parts list, '
    'and a compliance sign-off. Your technician can print this and walk to the machine."'
)

add_note('Let the client read the work order. They will recognize the format. This builds credibility.')

add_divider()

# --- Cascade ---
p = doc.add_paragraph()
run = p.add_run('5C. Cascade Analysis Agent')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY

add_do('Scroll to Cascade Analysis. Click "Run Cascade Analysis".')

add_show(
    'Multi-paragraph analysis of how the current fault propagates to other subsystems. '
    'Which components are at risk next, estimated time to secondary failures.'
)

add_say(
    '"The cascade agent is thinking about second-order effects. If the dielectric system '
    'fails, what happens to the wire transport? To the workpiece finish? To the machine '
    'downtime cascade? This is the kind of reasoning that usually requires a senior '
    'engineer with 20 years of experience."'
)

add_divider()

# --- More agents ---
p = doc.add_paragraph()
run = p.add_run('5D. Additional Agents (mention, demo if time)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY

agents_table = doc.add_table(rows=7, cols=3)
agents_table.style = 'Light Grid Accent 1'
headers = ['Agent', 'What It Does', 'Where']
for i, h in enumerate(headers):
    agents_table.rows[0].cells[i].text = h
    agents_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

agent_rows = [
    ('Predictive Alert', 'Proactive warning before a limit is breached', 'Dashboard banner'),
    ('Parts Procurement', 'Identifies replacement parts with part numbers', 'Twin Intelligence'),
    ('Incident Report', 'AS9100 incident report with regulatory refs', 'Twin Intelligence'),
    ('Troubleshooting Chatbot', 'Multi-turn diagnostic conversation', 'Twin Intelligence'),
    ('Asset Status', 'Deep-dive on a single component (click in 3D)', 'Dashboard'),
    ('Scenario Author', 'Generates runnable what-if specs from text', 'Scenario & Faults'),
]
for i, (a, w, loc) in enumerate(agent_rows):
    agents_table.rows[i+1].cells[0].text = a
    agents_table.rows[i+1].cells[1].text = w
    agents_table.rows[i+1].cells[2].text = loc

add_say(
    '"In total, the platform has 12 AI agents working across the three products. '
    'Every agent is grounded in the twin\'s live data, not generic advice. '
    'And they are orchestrated by AUTOMIND, our visual workflow engine."'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 6 — SCENARIO SANDBOX
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 6 — Scenario Sandbox: What-If Without Risk', level=1)
add_time('2 minutes')

add_say(
    '"What if you could test failure scenarios on a running machine without actually '
    'breaking anything? That is what the Scenario Sandbox does."'
)

add_do('Click "Scenario & Faults" in the sidebar.')

add_show(
    'Two tabs: Scenarios (external situations) and Faults (component failures). '
    'Preset scenario cards: "Summer heatwave", "New dielectric batch", "Aggressive roughing run".'
)

add_do('In the "Author with agent" section, type: "The operator runs an aggressive roughing cut through thick stock for 4 hours without clearing debris."')

add_do('Click "Author spec".')

add_show(
    'Claude generates a runnable what-if specification: fault type, severity, control parameters, '
    'horizon. The spec is grounded in this machine\'s physics model.'
)

add_do('Click "Simulate" on the authored spec.')

add_show(
    'Trajectory chart appears: projected sensor values over 2 hours. Health curve declining. '
    'Predicted detections listed with timestamps.'
)

add_say(
    '"The scenario engine just forked the current twin state, injected the hypothetical fault, '
    'and projected it forward on the real physics model. No risk to the actual machine. '
    'The AI analyzed the outcome and told us exactly when the wire would break and '
    'what the cascade effects would be.'
)

add_say(
    '"This is where GoalCert comes in. Once we know the failure mode, GoalCert generates '
    'a training scenario so your technicians can practice the response before it happens '
    'in the field. Detect, diagnose, train, respond. That is the closed loop."'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 7 — CROSS-INDUSTRY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 7 — Cross-Industry Proof: Same Platform, Different Twin', level=1)
add_time('1 minute')

add_say(
    '"Before I close, let me show you something that proves the architecture. '
    'I am going to switch to a completely different domain in one click."'
)

add_do('Click "Twins" in the sidebar. Click "Open twin" on "Helix Data Center".')

add_show(
    'Dashboard loads with data center signals: Rack Load, Inlet Temp, Cooling COP, '
    'UPS Charge, PUE. 3D scene shows a server room layout. AI Co-Pilot narrates data center health.'
)

add_say(
    '"Same platform. Same physics engine. Same AI agents. Different ontology, different sensors, '
    'different 3D model. I did not change a line of code. The domain knowledge is loaded from the '
    'ontology, and everything adapts. That is the power of a domain-agnostic architecture."'
)

add_do('Optionally show Hospital or Manufacturing twin briefly. Do not linger.')

add_note(
    'This act is about showing scalability and architectural elegance. '
    'Keep it fast. The point is made by the instant domain switch.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# ACT 8 — CLOSING
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Act 8 — Closing: The 18-Step Vision', level=1)
add_time('1 minute')

add_say(
    '"Let me put what you just saw into the bigger picture. Prem\'s vision for Collins '
    'is an 18-step workflow that goes from the moment a sensor fires to the moment that '
    'knowledge is permanently wired into the system for every future technician."'
)

# 18-step summary table
steps_table = doc.add_table(rows=7, cols=3)
steps_table.style = 'Light Grid Accent 1'
step_headers = ['Phase', 'Steps', 'Product']
for i, h in enumerate(step_headers):
    steps_table.rows[0].cells[i].text = h
    steps_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

step_rows = [
    ('Detection & Visualization', '1-4: Sensor alert, classification, twin visualization, scenario analysis', 'NextXR Digital Twin'),
    ('Training & Skilling', '5-8: AI training generation, LMS delivery, notification, JIT training', 'GoalCert'),
    ('Field Response', '9-11: Technician dispatch, AR assistance, expert guidance', 'GoalCert + AR'),
    ('Orchestration', '4, 12: Workflow orchestration, interaction monitoring', 'AUTOMIND'),
    ('Resolution', '13: Maintenance completion, sensor confirmation, twin resync', 'NextXR'),
    ('Knowledge Capture', '14-18: Session archive, report generation, knowledge conversion, indexing, future reuse', 'GoalCert + NextXR'),
]
for i, (phase, steps, product) in enumerate(step_rows):
    steps_table.rows[i+1].cells[0].text = phase
    steps_table.rows[i+1].cells[1].text = steps
    steps_table.rows[i+1].cells[2].text = product

doc.add_paragraph()

add_say(
    '"Today you saw Steps 1 through 4 live: sensor alert, classification, twin visualization, '
    'AI diagnosis, scenario analysis, and automated work order generation. You saw 12 AI agents '
    'working together. You saw the physics engine detecting a fault that three tiers of behavior '
    'rules caught independently. And you saw all of it adapt to a different industry in one click."'
)

add_say(
    '"The remaining steps, training generation, AR assistance, knowledge capture, '
    'are built on the same architecture. GoalCert already generates training scenarios from the '
    'fault data. The foundation is live and working. The 18-step vision is not a roadmap. '
    'It is an architecture that is already running."'
)

add_say('"We would love to discuss next steps with your team. Thank you."')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# RECOVERY PLAYBOOK
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Recovery Playbook', level=1)
p = doc.add_paragraph('If something goes wrong during the demo, use these fallbacks.')
p.runs[0].font.color.rgb = GRAY
p.runs[0].font.size = Pt(10)

recovery = [
    ('Twin takes too long to create',
     'Say "The graph is seeding the full equipment hierarchy" and switch to a pre-created twin from the library.'),
    ('AI Co-Pilot is not narrating',
     'Check the top bar for "Agent: stub" vs "Agent: Claude". If stub, the Anthropic API key is missing or expired. Say "The agent needs API access" and move on to the manual agents (Diagnosis, Work Order).'),
    ('Sensor cards are not updating',
     'Click the throttle buttons (Low/Mid/High) in the top bar to push physics steps. If no throttle buttons, you are on a simulated twin, not a live one.'),
    ('Work Order shows generic content',
     'Inject a fault first so there are active findings. The work order agent needs findings to ground its output.'),
    ('Frontend shows blank page',
     'Open browser console (Cmd+Option+I). Check for CORS errors. Restart the orchestrator: kill the process on 8090 and relaunch.'),
    ('NextXR connection refused',
     'Check the orchestrator .env file. NEXTXR_URL must be http://localhost:8000 (not 8080).'),
    ('"All platforms connected" but agents fail',
     'The Anthropic API key may have exhausted credits. Check collins-demo/orchestrator/.env for ANTHROPIC_API_KEY.'),
    ('3D model not loading',
     'Three.js requires WebGL. Check that Chrome hardware acceleration is on. If on a VM, the 3D model will not render.'),
]

for problem, fix in recovery:
    p = doc.add_paragraph()
    run = p.add_run(f'Problem: {problem}')
    run.bold = True
    run.font.size = Pt(10)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'Fix: {fix}')
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_after = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TOUGH QUESTIONS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Tough Questions & Answers', level=1)

qa = [
    ('How is this different from Azure Digital Twins or Bentley iTwin?',
     'Azure Digital Twins requires you to manually define DTDL models and wire up IoT Hub. '
     'Bentley iTwin requires Revit or MicroStation models. Neither has integrated AI agents. '
     'Our platform seeds a full twin from a domain template in seconds, runs real coupled physics, '
     'and has 12 AI agents that diagnose, predict, generate work orders, and author training scenarios, '
     'all from the same live data. No other platform does detection-to-training in one loop.'),
    ('Is the physics real or simulated?',
     'The physics is real. The turbine engine runs a Brayton-cycle two-spool model with '
     'compressor maps, combustor heat release, and turbine expansion. The Wire EDM runs a '
     'full discharge-generator model with Lazarenko pulse dynamics, dielectric thermodynamics, '
     'and wire transport mechanics. These are not scripted curves. They are differential equations '
     'solved at each timestep.'),
    ('Can this ingest real sensor data from our equipment?',
     'Yes. The platform has an ingest API that accepts any sensor payload. Today we are driving it '
     'with the physics engine as a stand-in for real IoT data. In production, your PLCs or IoT gateway '
     'would POST sensor readings to the same endpoint, and everything downstream, detection, diagnosis, '
     'work orders, stays the same.'),
    ('What about data security? Our MRO data is ITAR-controlled.',
     'The platform runs entirely on-premise. No data leaves your network. The AI agents use Claude via API, '
     'but the prompts contain only sensor readings and physics outputs, no proprietary design data. '
     'For ITAR compliance, we can run with a local LLM or air-gapped deployment.'),
    ('How long to deploy this for a real facility?',
     'The platform is production-grade today for the demo scope. For a full deployment: '
     'Phase 1 (4-6 weeks) connects real sensors and validates the physics models against your equipment. '
     'Phase 2 (4-6 weeks) adds the GoalCert training loop and AUTOMIND workflow orchestration. '
     'The architecture does not change between PoC and production.'),
    ('Can we add our own equipment types?',
     'Yes. The ontology is extensible. We define new equipment classes, sensor types, and behavior rules '
     'in SHACL/TTL files. The platform picks them up at startup. We have already built turbine, wire EDM, '
     'data center, hospital, and manufacturing domains on the same framework.'),
    ('What happens when the AI is wrong?',
     'Every AI agent output is grounded in the physics model and the twin\'s live data. '
     'The diagnosis agent cites specific sensor readings and threshold values. The work order agent '
     'references ATA chapters and compliance standards. If the physics model says a component is healthy, '
     'the AI cannot override that. The twin is the source of truth, not the AI.'),
]

for q, a in qa:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: "{q}"')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'A: {a}')
    run2.font.size = Pt(10)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# PLATFORM STATS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Platform Stats (Reference)', level=1)

stats_table = doc.add_table(rows=12, cols=2)
stats_table.style = 'Light Grid Accent 1'
stats_data = [
    ('AI Agents', '12 Claude-powered agents across 3 products'),
    ('Physics Models', 'Brayton-cycle turbine + Lazarenko wire-EDM + coupled facility models'),
    ('Twin Domains', '5 (Turbine, Wire EDM, Data Center, Hospital, Manufacturing)'),
    ('Ontology Classes', '113 classes, 30+ predicates, SHACL validation'),
    ('Detection Tiers', '3 (Tier C threshold, Tier B statistical, Tier A physics)'),
    ('Sensor Signals', '8 per machine domain, live physics-driven'),
    ('Compliance', 'AS9100, EASA Part 145, ATA chapter references'),
    ('AI Model', 'Claude Sonnet 4 (Anthropic) — configurable'),
    ('3D Engine', 'Three.js + React Three Fiber, procedural scene generation'),
    ('Codebase', '55,000+ lines of source code'),
    ('Products Integrated', 'NextXR Digital Twin + AUTOMIND + GoalCert'),
]
stats_table.rows[0].cells[0].text = 'Metric'
stats_table.rows[0].cells[1].text = 'Value'
for c in stats_table.rows[0].cells:
    c.paragraphs[0].runs[0].bold = True

for i, (metric, value) in enumerate(stats_data):
    stats_table.rows[i+1].cells[0].text = metric
    stats_table.rows[i+1].cells[1].text = value

# ── Save ────────────────────────────────────────────────────────────────
out = '/Users/narhen/Collins_POC/Collins_Demo_Script.docx'
doc.save(out)
print(f'Saved: {out}')
