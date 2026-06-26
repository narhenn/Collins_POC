"""
Generate GoalCert AutoMind BRD (Business Requirements Document)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──────────────────────────────────────────────────────
PURPLE = RGBColor(73, 2, 162)
DARK_GRAY = RGBColor(51, 51, 51)
MED_GRAY = RGBColor(102, 102, 102)
WHITE = RGBColor(255, 255, 255)
LIGHT_BG = "F5F0FF"
FONT = "Calibri"
BODY_SIZE = Pt(11)
H1_SIZE = Pt(16)
H2_SIZE = Pt(13)
H3_SIZE = Pt(11)
SAVE_PATH = "/Users/narhen/automind/docs/GoalCert_AutoMind_BRD_v1_0.docx"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" '
            f'w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def style_table_header_row(row, col_count):
    for i in range(col_count):
        cell = row.cells[i]
        set_cell_shading(cell, "4902A2")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "4902A2")

    # Body
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_BG)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading1(doc, number, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.clear()
    run = p.add_run(f"{number}. {title}")
    run.font.name = FONT
    run.font.size = H1_SIZE
    run.font.color.rgb = PURPLE
    run.bold = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_heading2(doc, number, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.clear()
    run = p.add_run(f"{number} {title}")
    run.font.name = FONT
    run.font.size = H2_SIZE
    run.font.color.rgb = PURPLE
    run.bold = True
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading3(doc, number, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    p.clear()
    run = p.add_run(f"{number} {title}")
    run.font.name = FONT
    run.font.size = H3_SIZE
    run.font.color.rgb = PURPLE
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


def add_cover_page(doc):
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # GoalCert header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GoalCert")
    run.font.name = FONT
    run.font.size = Pt(36)
    run.font.color.rgb = PURPLE
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AutoMind — Agentic AI Workforce Platform")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = MED_GRAY

    doc.add_paragraph()

    # Purple divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = PURPLE
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Business Requirements Document (BRD)")
    run.font.name = FONT
    run.font.size = Pt(22)
    run.font.color.rgb = PURPLE
    run.bold = True

    # Description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Comprehensive business requirements for the GoalCert AutoMind platform,\n"
        "an enterprise agentic AI workforce solution for automating knowledge work."
    )
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = MED_GRAY

    for _ in range(4):
        doc.add_paragraph()

    # Confidentiality notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(180, 0, 0)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This document contains proprietary information belonging to GoalCert Pte Ltd.\n"
        "Unauthorized distribution, reproduction, or use of this document is strictly prohibited."
    )
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = MED_GRAY

    doc.add_page_break()


def add_document_control(doc):
    p = doc.add_paragraph()
    run = p.add_run("Document Control")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = PURPLE
    run.bold = True

    doc.add_paragraph()

    rows = [
        ("Document Type", "Business Requirements Document (BRD)"),
        ("Product", "GoalCert AutoMind"),
        ("Version", "1.0"),
        ("Date", "23 June 2026"),
        ("Classification", "CONFIDENTIAL"),
        ("Prepared By", "Engineering"),
        ("Reviewed By", "Product Management"),
        ("Approved By", "CTO"),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(rows):
        # Label cell
        cell_l = table.rows[i].cells[0]
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell_l, "4902A2")
        cell_l.width = Cm(5)

        # Value cell
        cell_r = table.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        cell_r.width = Cm(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Version History
    p = doc.add_paragraph()
    run = p.add_run("Version History")
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = PURPLE
    run.bold = True

    doc.add_paragraph()

    add_styled_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ("0.1", "10 June 2026", "Engineering", "Initial draft — platform scope and objectives"),
            ("0.5", "16 June 2026", "Engineering", "Added business requirements BR-01 through BR-15, user personas, and revenue model"),
            ("0.9", "20 June 2026", "Engineering", "Incorporated stakeholder review feedback, finalized competitive analysis and risk matrix"),
            ("1.0", "23 June 2026", "Engineering", "Final release — approved by CTO and Product Management"),
        ],
        col_widths=[2, 3, 3, 8],
    )

    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = PURPLE
    run.bold = True

    doc.add_paragraph()

    toc_items = [
        ("1.", "Introduction"),
        ("   1.1", "Purpose"),
        ("   1.2", "Background"),
        ("   1.3", "Scope"),
        ("   1.4", "Definitions & Abbreviations"),
        ("2.", "Business Objectives & Goals"),
        ("3.", "Stakeholders"),
        ("4.", "Market Context & Competitive Landscape"),
        ("   4.1", "Market Overview"),
        ("   4.2", "Competitive Analysis"),
        ("   4.3", "Differentiation Strategy"),
        ("5.", "Platform Model — Business Requirements"),
        ("6.", "User Personas"),
        ("7.", "Revenue Model & Pricing"),
        ("   7.1", "Pricing Tiers"),
        ("   7.2", "Revenue Projections"),
        ("8.", "Success Metrics & KPIs"),
        ("9.", "Constraints & Assumptions"),
        ("   9.1", "Constraints"),
        ("   9.2", "Assumptions"),
        ("10.", "Risks & Mitigations"),
        ("11.", "Dependencies"),
        ("12.", "Glossary"),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        run.font.name = FONT
        run.font.size = Pt(11) if not num.startswith("   ") else Pt(10)
        run.font.color.rgb = DARK_GRAY if not num.startswith("   ") else MED_GRAY
        run.bold = not num.startswith("   ")
        p.paragraph_format.space_after = Pt(2)
        if num.startswith("   "):
            p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()


def build_section_1(doc):
    add_heading1(doc, "1", "Introduction")

    add_heading2(doc, "1.1", "Purpose")
    add_body(doc,
        "This Business Requirements Document (BRD) defines the business needs, objectives, and high-level "
        "requirements for GoalCert AutoMind, an enterprise-grade agentic AI workforce platform. The document "
        "serves as the primary reference for all stakeholders involved in the planning, development, and "
        "deployment of the platform. It establishes the business justification, scope boundaries, and success "
        "criteria that will guide subsequent functional and technical specifications."
    )
    add_body(doc,
        "The intended audience includes executive leadership, product management, engineering teams, sales "
        "and marketing stakeholders, and external partners who require a clear understanding of the platform's "
        "business rationale and strategic direction. This document does not prescribe implementation details; "
        "those are covered in the companion Functional Requirements Document (FRD) and Technical Architecture "
        "Document (TAD)."
    )

    add_heading2(doc, "1.2", "Background")
    add_body(doc,
        "The enterprise automation landscape has undergone a fundamental transformation with the emergence of "
        "large language models (LLMs) and agentic AI architectures. Traditional workflow automation tools such "
        "as Zapier and Make have enabled businesses to connect applications and automate simple trigger-action "
        "sequences for over a decade. However, these platforms are inherently limited to deterministic, rule-based "
        "workflows that cannot reason about ambiguous inputs, synthesize unstructured information, or adapt their "
        "behavior based on context."
    )
    add_body(doc,
        "Simultaneously, open-source AI agent frameworks like LangGraph, CrewAI, and AutoGen have demonstrated "
        "the power of autonomous AI systems that can plan, execute, and iterate on complex tasks. These frameworks, "
        "however, require significant software engineering expertise to deploy and operate, placing them beyond the "
        "reach of business users and operations teams who would benefit most from intelligent automation."
    )
    add_body(doc,
        "GoalCert AutoMind addresses this gap by combining the accessibility of visual no-code workflow builders "
        "with the intelligence of agentic AI systems. The platform enables non-technical users to design, deploy, "
        "and manage AI-powered agents that autonomously execute multi-step workflows involving web research, AI "
        "analysis, decision-making, and cross-platform integrations. AutoMind is positioned as a core product "
        "within the GoalCert suite, alongside the Simulation Engine, Digital Twin, and AR Remote Assistance "
        "offerings, creating a comprehensive enterprise intelligence platform."
    )

    add_heading2(doc, "1.3", "Scope")

    add_heading3(doc, "1.3.1", "In Scope")
    in_scope = [
        "Visual drag-and-drop workflow builder supporting seven distinct node types: Trigger, Web Search, AI Action, Decision, Integration, Escalation, and Code Execution",
        "Multi-model AI integration supporting OpenAI GPT-4, Anthropic Claude, and future model providers",
        "Autonomous scheduled execution via cron-based scheduling (RedBeat) with manual and webhook trigger options",
        "Real-time execution monitoring with server-sent events (SSE) streaming for live console output",
        "Persistent agent memory system where execution summaries are stored and injected into future workflow runs",
        "Conversational agent chat interface for interactive debriefing and ad-hoc queries to deployed agents",
        "Integration ecosystem encompassing email (SMTP), Slack, generic webhooks, and extensible connector framework",
        "Template marketplace for sharing and discovering pre-built agent workflows across the organization",
        "Multi-tenant architecture with strict data isolation between users and organizations",
        "Cost tracking and optimization dashboards showing per-execution and per-node LLM token consumption",
        "Reports and analytics module for aggregate insights across agent populations",
        "Role-based access control (RBAC) with user, admin, and super-admin roles",
        "Production deployment pipeline with Docker containerization and horizontal scaling",
        "JILLA AI concierge for guided client demonstrations and onboarding",
        "Plugin and node SDK for third-party extensibility",
    ]
    for item in in_scope:
        add_bullet(doc, item)

    add_heading3(doc, "1.3.2", "Out of Scope")
    out_scope = [
        "On-premises deployment model (cloud-only for v1.0; on-prem evaluated for v2.0)",
        "Mobile native applications (responsive web only for initial release)",
        "Custom LLM fine-tuning infrastructure (platform consumes pre-trained models via API)",
        "Real-time collaboration features such as simultaneous multi-user workflow editing",
        "Compliance certifications (SOC 2, ISO 27001) beyond foundational security controls — targeted for post-launch",
        "Integration with GoalCert Simulation Engine and Digital Twin (planned for v1.5 roadmap)",
    ]
    for item in out_scope:
        add_bullet(doc, item)

    add_heading2(doc, "1.4", "Definitions & Abbreviations")
    add_styled_table(doc,
        ["Term", "Definition"],
        [
            ("Agent", "An autonomous AI entity configured with a workflow, schedule, and memory that executes tasks without continuous human intervention"),
            ("Workflow", "A directed acyclic graph (DAG) of interconnected nodes that defines the sequence of operations an agent performs"),
            ("Node", "A discrete processing unit within a workflow that performs a specific action (e.g., web search, AI analysis, integration call)"),
            ("Trigger", "An event or schedule that initiates agent execution — may be cron-based, manual, or webhook-driven"),
            ("BFS", "Breadth-First Search — the graph traversal algorithm used by the execution engine to process workflow nodes in topological order"),
            ("SSE", "Server-Sent Events — a unidirectional streaming protocol used for real-time execution monitoring"),
            ("RAG", "Retrieval-Augmented Generation — a technique where relevant stored context is injected into LLM prompts to improve response quality"),
            ("RedBeat", "A Celery beat scheduler backed by Redis that manages cron-based agent execution schedules"),
            ("RBAC", "Role-Based Access Control — a security model where permissions are assigned to roles rather than individual users"),
            ("LLM", "Large Language Model — foundation AI models such as GPT-4 and Claude that power agent intelligence"),
            ("JILLA", "GoalCert's AI concierge system designed to guide new users and demonstrate platform capabilities"),
            ("MoSCoW", "Prioritization framework: Must have, Should have, Could have, Won't have (this time)"),
        ],
        col_widths=[3, 13],
    )


def build_section_2(doc):
    add_heading1(doc, "2", "Business Objectives & Goals")

    add_body(doc,
        "The following business objectives define the strategic outcomes that GoalCert AutoMind must achieve "
        "to justify investment and demonstrate product-market fit. Each objective is measurable and time-bound "
        "to enable clear progress tracking against the platform roadmap."
    )

    objectives = [
        (
            "OBJ-01: Democratize AI Agent Creation",
            "Enable non-technical business users to create, configure, and deploy autonomous AI agents without "
            "writing code, reducing the time from concept to deployed agent from weeks (with custom development) "
            "to under 30 minutes. The visual workflow builder must achieve a System Usability Scale (SUS) score "
            "of 75 or higher among non-technical test participants within the first six months of launch."
        ),
        (
            "OBJ-02: Reduce Enterprise Knowledge Work Costs",
            "Deliver measurable cost savings for enterprise customers by automating repetitive knowledge work tasks "
            "such as market research, competitor monitoring, lead qualification, and report generation. Target a "
            "minimum 40% reduction in time spent on automated tasks compared to manual execution, validated through "
            "customer case studies within the first year of general availability."
        ),
        (
            "OBJ-03: Achieve Product-Market Fit in the SMB and Mid-Market Segments",
            "Acquire 200 paying customers within the first 12 months of launch across the Starter ($99/mo) and "
            "Pro ($499/mo) tiers. Maintain a monthly churn rate below 5% and achieve a Net Promoter Score (NPS) "
            "of 40 or higher, indicating strong product-market fit and customer satisfaction."
        ),
        (
            "OBJ-04: Establish a Defensible Platform Ecosystem",
            "Build a template marketplace with at least 50 community-contributed agent templates within 12 months "
            "and release a plugin/node SDK enabling third-party developers to extend platform capabilities. The "
            "ecosystem must demonstrate network effects where each new template or integration increases platform "
            "value for all users."
        ),
        (
            "OBJ-05: Deliver Enterprise-Grade Reliability and Security",
            "Achieve 99.9% platform uptime (excluding scheduled maintenance), maintain P95 API response latency "
            "under 200ms for non-AI operations, and implement comprehensive multi-tenant data isolation. Complete "
            "foundational security controls sufficient to pass enterprise procurement security questionnaires for "
            "80% of prospect engagements."
        ),
        (
            "OBJ-06: Generate Sustainable Revenue Growth",
            "Reach $50,000 in monthly recurring revenue (MRR) within 12 months of launch through a combination of "
            "self-serve SaaS subscriptions and enterprise contracts. Achieve a customer lifetime value (LTV) to "
            "customer acquisition cost (CAC) ratio of at least 3:1 by month 18, validating the unit economics "
            "of the go-to-market strategy."
        ),
    ]

    for title, description in objectives:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = FONT
        run.font.size = Pt(11)
        run.font.color.rgb = PURPLE
        run.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        add_body(doc, description)


def build_section_3(doc):
    add_heading1(doc, "3", "Stakeholders")

    add_body(doc,
        "The following table identifies the key stakeholders for the GoalCert AutoMind platform, their roles "
        "within the project, and their primary areas of interest. Effective communication with all stakeholder "
        "groups is critical to aligning expectations and ensuring the platform delivers value across the organization."
    )

    add_styled_table(doc,
        ["Stakeholder", "Role", "Interest"],
        [
            ("CEO / Founder", "Executive Sponsor", "Strategic alignment with GoalCert product suite, revenue growth, market positioning against competitors, and investor narrative"),
            ("CTO", "Technical Authority", "Architecture scalability, security posture, infrastructure cost management, technical debt governance, and integration with existing GoalCert products"),
            ("VP Product", "Product Owner", "Feature prioritization, roadmap planning, user research outcomes, competitive differentiation, and pricing strategy validation"),
            ("Engineering Lead", "Technical Delivery", "Development velocity, code quality, CI/CD pipeline health, on-call burden, and cross-team technical dependencies"),
            ("Sales Director", "Revenue Generation", "Demo readiness, feature parity with competitors for deal closure, pricing flexibility for enterprise negotiations, and JILLA concierge effectiveness"),
            ("Marketing Manager", "Market Awareness", "Messaging clarity, content generation capabilities, website copy, case study development, and template marketplace promotion"),
            ("Customer Success", "Retention & Expansion", "Onboarding friction reduction, support ticket volume, feature adoption metrics, churn prediction, and upsell opportunity identification"),
            ("Enterprise Customers", "End Users (Buyer)", "ROI justification, security compliance, SSO integration, SLA guarantees, and dedicated support channels"),
            ("SMB Customers", "End Users (Self-Serve)", "Ease of use, time-to-value, pricing transparency, template availability, and community support quality"),
            ("Third-Party Developers", "Ecosystem Partners", "SDK documentation quality, API stability, marketplace monetization, and developer support responsiveness"),
        ],
        col_widths=[3.5, 3, 9.5],
    )


def build_section_4(doc):
    add_heading1(doc, "4", "Market Context & Competitive Landscape")

    add_heading2(doc, "4.1", "Market Overview")
    add_body(doc,
        "The global intelligent automation market is projected to reach $25.2 billion by 2027, growing at a "
        "compound annual growth rate (CAGR) of 14.6% from 2023. Within this broader market, the agentic AI "
        "segment represents the fastest-growing category, driven by enterprise adoption of large language models "
        "and the emergence of autonomous AI systems capable of multi-step reasoning and tool use."
    )
    add_body(doc,
        "Three distinct market segments are converging to create the opportunity that AutoMind addresses. First, "
        "the no-code/low-code automation market (led by Zapier, Make, and n8n) has established that business users "
        "will adopt visual workflow tools for process automation. Second, the AI agent framework market (led by "
        "LangGraph, CrewAI, and AutoGen) has demonstrated that LLM-powered agents can autonomously execute complex "
        "tasks. Third, the enterprise AI platform market (led by Relevance AI, Cognigy, and Kore.ai) has proven "
        "that organizations will pay premium prices for managed AI solutions with enterprise security and compliance."
    )
    add_body(doc,
        "AutoMind occupies the intersection of these three segments — a visual, no-code platform for building "
        "enterprise-grade AI agents. This positioning addresses the unmet need of business users who want the power "
        "of agentic AI without the complexity of code-first frameworks, and who require the security, reliability, "
        "and support guarantees that open-source tools cannot provide."
    )

    add_heading2(doc, "4.2", "Competitive Analysis")
    add_styled_table(doc,
        ["Competitor", "Category", "Strengths", "Weaknesses", "AutoMind Advantage"],
        [
            ("Zapier", "No-Code Automation", "Massive integration library (6,000+ apps), strong brand recognition, simple trigger-action model", "No AI reasoning capability, no agent memory, limited to linear workflows, no autonomous execution", "AI-native architecture with agent memory and autonomous scheduled execution"),
            ("Make (Integromat)", "No-Code Automation", "Visual workflow builder, complex branching logic, competitive pricing", "No native LLM integration, no agent persistence, steep learning curve for advanced scenarios", "Purpose-built for AI agents with conversational interface and memory"),
            ("n8n", "Open-Source Automation", "Self-hosted option, extensible node system, active community", "Requires technical expertise to deploy, no managed cloud offering with SLA, limited AI capabilities", "Fully managed platform with enterprise SLA and native AI integration"),
            ("LangGraph", "AI Agent Framework", "Powerful graph-based agent orchestration, state management, production-ready", "Code-only interface, requires Python expertise, no visual builder, complex deployment", "Visual no-code builder accessible to non-technical users"),
            ("CrewAI", "AI Agent Framework", "Multi-agent collaboration, role-based agents, growing ecosystem", "Code-only, limited production tooling, no execution monitoring, no persistent memory", "Production-grade monitoring, scheduling, and cost tracking"),
            ("AutoGen", "AI Agent Framework", "Microsoft backing, multi-agent conversations, research-grade capabilities", "Primarily research-focused, limited production deployment tools, complex configuration", "Enterprise-focused with billing, RBAC, and multi-tenancy"),
            ("Relevance AI", "AI Platform", "Closest competitor, visual agent builder, tool integration, managed platform", "Smaller integration ecosystem, less mature workflow builder, limited scheduling options", "Superior workflow builder UX, deeper memory system, template marketplace"),
        ],
        col_widths=[2.5, 2.5, 3.5, 3.5, 4],
    )

    add_heading2(doc, "4.3", "Differentiation Strategy")
    add_body(doc,
        "AutoMind differentiates through three strategic pillars that collectively create a defensible market "
        "position. The first pillar is accessibility: the visual workflow builder with seven specialized node types "
        "makes it possible for operations managers, sales leads, and marketing professionals to build sophisticated "
        "AI agents without engineering support. The second pillar is intelligence: the agent memory system, which "
        "stores and injects execution summaries into future runs, enables agents to learn and improve over time — "
        "a capability absent from traditional automation platforms. The third pillar is ecosystem: the template "
        "marketplace and plugin SDK create network effects that increase platform value with each new user and "
        "contributor, building a competitive moat that is difficult for new entrants to replicate."
    )


def build_section_5(doc):
    add_heading1(doc, "5", "Platform Model — Business Requirements")

    add_body(doc,
        "This section enumerates the core business requirements for the GoalCert AutoMind platform. Each requirement "
        "is assigned a unique identifier, priority level (using MoSCoW classification), and rationale. Requirements "
        "are traced to business objectives defined in Section 2 and will be decomposed into functional requirements "
        "in the companion FRD."
    )

    requirements = [
        ("BR-01", "Visual Workflow Builder", "Must",
         "The platform shall provide a visual drag-and-drop workflow builder that enables users to design agent "
         "workflows by placing and connecting seven node types on a canvas: Trigger, Web Search, AI Action, Decision, "
         "Integration, Escalation, and Code Execution. Each node shall have a dedicated configuration panel where "
         "users can set parameters, select AI models, define decision conditions, and configure integration targets. "
         "The builder must support variable substitution using {variable_name} syntax so that outputs from upstream "
         "nodes flow as inputs to downstream nodes. The workflow graph must be validated before deployment to ensure "
         "connectivity and prevent cycles. This requirement directly supports OBJ-01 (Democratize AI Agent Creation) "
         "by eliminating the need for users to write code to orchestrate AI workflows.",
         "OBJ-01"),

        ("BR-02", "AI Model Integration", "Must",
         "The platform shall integrate with multiple large language model providers, initially OpenAI (GPT-4, GPT-4o) "
         "and Anthropic (Claude 3.5 Sonnet, Claude 3 Opus), with an extensible provider architecture that allows "
         "additional models to be added without modifying the core execution engine. Users must be able to select the "
         "AI model for each AI Action node independently, compare model outputs, and switch models without rebuilding "
         "workflows. The integration layer must handle rate limiting, token counting, error retries, and cost tracking "
         "transparently. This requirement supports OBJ-02 (Reduce Knowledge Work Costs) by giving users access to the "
         "most capable and cost-effective models for each task.",
         "OBJ-02"),

        ("BR-03", "Scheduled Autonomous Execution", "Must",
         "The platform shall enable agents to execute autonomously on user-defined schedules using cron expressions "
         "managed by RedBeat (a Redis-backed Celery beat scheduler). Users must be able to configure schedules through "
         "a natural-language-friendly interface that translates inputs like 'every weekday at 9am' into cron syntax. "
         "In addition to scheduled execution, agents must support manual trigger (on-demand run from the dashboard) and "
         "webhook trigger (execution initiated by an external HTTP POST request with a payload). The scheduling system "
         "must handle timezone awareness, missed execution recovery, and concurrent execution prevention for agents "
         "that should not run in parallel. This requirement supports OBJ-01 and OBJ-02.",
         "OBJ-01, OBJ-02"),

        ("BR-04", "Real-Time Execution Monitoring", "Must",
         "The platform shall provide real-time visibility into agent execution through a live console that streams "
         "execution events using Server-Sent Events (SSE). The console must display node-by-node progress including "
         "node entry, processing status, intermediate outputs, token consumption, execution duration, and error "
         "messages. Users must be able to view execution history with full logs for past runs, compare execution "
         "outcomes across runs, and identify performance bottlenecks. The monitoring system must support concurrent "
         "viewing by multiple users without performance degradation. This requirement supports OBJ-05 (Enterprise-Grade "
         "Reliability) by enabling rapid issue identification and resolution.",
         "OBJ-05"),

        ("BR-05", "Agent Memory and Learning", "Must",
         "The platform shall implement a persistent memory system where each agent maintains a history of past "
         "execution summaries. After each execution, the system generates a structured summary using an LLM (covering "
         "key findings, decisions made, outputs produced, and anomalies detected) and stores it in the agent's memory "
         "store. On subsequent executions, relevant past summaries are retrieved and injected into the agent's context "
         "window, enabling the agent to reference historical data, avoid repeating past mistakes, and build upon "
         "previous analyses. The memory system must support configurable retention policies (number of summaries to "
         "retain, summarization depth) and provide a memory viewer in the UI for users to inspect what an agent "
         "remembers. This requirement supports OBJ-01 and OBJ-04.",
         "OBJ-01, OBJ-04"),

        ("BR-06", "Conversational Agent Interface", "Must",
         "The platform shall provide a chat-based interface where users can engage in natural-language conversations "
         "with their deployed agents. The chat interface must allow users to ask agents about their latest execution "
         "results, request re-analysis of specific data points, ask follow-up questions about findings, and instruct "
         "agents to perform ad-hoc tasks within their workflow capability. Chat responses must be streamed in real time "
         "via SSE to provide immediate feedback. The agent must incorporate its memory and most recent execution context "
         "when responding to chat queries, ensuring continuity between automated runs and interactive sessions. This "
         "requirement supports OBJ-01 by making agents accessible through a familiar conversational paradigm.",
         "OBJ-01"),

        ("BR-07", "Integration Ecosystem", "Should",
         "The platform shall provide native integrations for email delivery (via SMTP), Slack messaging (via Slack API), "
         "and generic webhook calls (HTTP POST to arbitrary endpoints). Each integration must be configurable as an "
         "Integration node within the workflow builder, supporting dynamic content through variable substitution. The "
         "integration framework must be extensible to allow additional connectors (e.g., Microsoft Teams, Jira, Salesforce, "
         "Google Sheets) to be added via the plugin SDK without modifying the core platform. Integration nodes must "
         "support retry logic, timeout configuration, and error escalation. This requirement supports OBJ-02 and OBJ-04.",
         "OBJ-02, OBJ-04"),

        ("BR-08", "Template Marketplace", "Should",
         "The platform shall include a template marketplace where users can discover, preview, and instantiate pre-built "
         "agent workflows for common use cases. Initial templates must cover competitive intelligence monitoring, lead "
         "qualification and enrichment, content generation pipelines, customer feedback analysis, regulatory change tracking, "
         "and market research briefings. Users must be able to publish their own workflows as templates (with optional "
         "anonymization of sensitive configuration), rate and review templates, and fork templates for customization. "
         "The marketplace must support categories, search, and popularity-based ranking. This requirement supports OBJ-04 "
         "(Defensible Platform Ecosystem).",
         "OBJ-04"),

        ("BR-09", "Multi-Tenant Isolation", "Must",
         "The platform shall enforce strict multi-tenant data isolation ensuring that agents, workflows, execution logs, "
         "memory stores, and configuration data belonging to one user or organization are never accessible to another. "
         "Isolation must be enforced at the database level (row-level security or schema-per-tenant), API level (request "
         "authentication and authorization), and execution level (separate Celery task queues or worker pools per tenant "
         "where required). The isolation model must be validated through automated security testing as part of the CI/CD "
         "pipeline. This requirement supports OBJ-05.",
         "OBJ-05"),

        ("BR-10", "Cost Tracking and Optimization", "Must",
         "The platform shall track and display the cost of every agent execution broken down by individual node, showing "
         "input tokens, output tokens, model used, and computed cost based on provider pricing. Users must have access to "
         "a cost dashboard showing spending trends over time, per-agent cost comparisons, and projected monthly costs based "
         "on current usage patterns. The system must support configurable cost alerts that notify users when spending "
         "exceeds defined thresholds. Cost data must be exportable for financial reporting. This requirement supports OBJ-02 "
         "and OBJ-06.",
         "OBJ-02, OBJ-06"),

        ("BR-11", "Reports and Analytics", "Should",
         "The platform shall provide a comprehensive analytics module that aggregates execution data across all agents "
         "to surface insights such as overall success rates, average execution duration, most active agents, common failure "
         "patterns, and usage trends. Reports must be viewable within the platform dashboard and exportable as PDF or CSV. "
         "The analytics engine must support custom date ranges, agent-level and organization-level views, and comparison "
         "across time periods. This requirement supports OBJ-03 and OBJ-06.",
         "OBJ-03, OBJ-06"),

        ("BR-12", "Role-Based Access Control", "Must",
         "The platform shall implement a role-based access control (RBAC) system with at minimum three roles: User (can "
         "create and manage own agents), Admin (can manage all agents within an organization and access organization-level "
         "analytics), and Super Admin (platform-wide administration including user management and system configuration). "
         "Permissions must be granular enough to control access to specific features (e.g., template publishing, cost "
         "dashboards, integration management). The RBAC system must integrate with enterprise identity providers via SAML "
         "or OIDC for single sign-on (SSO) in enterprise deployments. This requirement supports OBJ-05.",
         "OBJ-05"),

        ("BR-13", "Production Deployment and Scaling", "Must",
         "The platform shall be deployable as a containerized application using Docker with orchestration via Kubernetes "
         "or Docker Compose. The architecture must support horizontal scaling of the API layer, execution workers (Celery), "
         "and the real-time streaming layer independently. The deployment must include automated health checks, graceful "
         "shutdown handling for in-flight executions, database migration management (Alembic), and zero-downtime deployment "
         "capability. Infrastructure must be provisioned to support at least 500 concurrent agent executions with P99 "
         "latency under 5 seconds for non-AI operations. This requirement supports OBJ-05.",
         "OBJ-05"),

        ("BR-14", "JILLA AI Concierge", "Could",
         "The platform shall include JILLA, an AI-powered concierge that guides new users through the platform during "
         "onboarding and assists sales teams during client demonstrations. JILLA must be able to walk users through "
         "creating their first agent, explain platform features in natural language, answer questions about capabilities "
         "and limitations, and demonstrate pre-built template workflows in a guided interactive format. JILLA must be "
         "accessible from any page within the platform via a persistent chat widget. This requirement supports OBJ-03 "
         "by reducing onboarding friction and improving demo conversion rates.",
         "OBJ-03"),

        ("BR-15", "Extensibility via Plugin/Node SDK", "Should",
         "The platform shall provide a documented Software Development Kit (SDK) that enables third-party developers to "
         "create custom node types, integration connectors, and workflow templates. The SDK must include a development "
         "server for local testing, a packaging format for distribution, a submission process for the template marketplace, "
         "and versioning support. The SDK must enforce security sandboxing to prevent custom nodes from accessing other "
         "tenants' data or platform internals. This requirement supports OBJ-04 by enabling ecosystem growth beyond the "
         "core engineering team's capacity.",
         "OBJ-04"),
    ]

    add_styled_table(doc,
        ["ID", "Requirement", "Priority", "Description", "Traces To"],
        [(r[0], r[1], r[2], r[3], r[4]) for r in requirements],
        col_widths=[1.5, 3, 1.5, 8.5, 1.5],
    )


def build_section_6(doc):
    add_heading1(doc, "6", "User Personas")

    add_body(doc,
        "The following user personas represent the primary target users of the GoalCert AutoMind platform. Each "
        "persona captures the user's context, goals, pain points, and how AutoMind addresses their needs. These "
        "personas inform feature prioritization, UX design decisions, and go-to-market messaging."
    )

    personas = [
        ("CEO / Executive", "Sarah Chen, CEO of a 150-person B2B SaaS company",
         "Sarah needs to stay informed about market trends, competitor moves, and customer sentiment without spending "
         "hours each week reading reports. She wants AI agents that autonomously monitor her competitive landscape, "
         "summarize board-ready briefings, and alert her to critical developments. Her primary concern is trustworthiness "
         "of AI outputs — she needs to understand what the agent did and verify its sources before sharing findings with "
         "her board. She evaluates tools based on ROI impact and expects the platform to pay for itself within the first "
         "quarter of use through time savings and better decision-making.",
         "Competitive intelligence agent, executive briefing agent, board report generator"),

        ("Operations Manager", "David Park, Head of Operations at a logistics company",
         "David manages complex workflows spanning procurement, vendor management, and compliance reporting. He currently "
         "relies on a combination of spreadsheets, email chains, and manual data entry to track operational metrics. He "
         "wants AI agents that automate data collection from multiple sources, flag anomalies in operational KPIs, and "
         "generate weekly status reports for his team. His primary pain point is the time spent on data aggregation — "
         "he estimates 15 hours per week across his team is spent copying data between systems. He needs a platform "
         "that integrates with existing tools (email, Slack, internal dashboards) and requires minimal IT involvement "
         "to set up.",
         "Operations monitoring agent, vendor compliance checker, KPI dashboard agent"),

        ("Sales Lead", "Maria Rodriguez, VP of Sales at a mid-market technology company",
         "Maria's sales team struggles with lead qualification — they spend significant time researching prospects "
         "manually before discovery calls, and many qualified leads go stale because follow-up is delayed. She wants "
         "AI agents that automatically enrich incoming leads with company data, news, and social signals, then score "
         "and route leads to the appropriate sales rep with a pre-call briefing. She also needs agents that monitor "
         "existing customer accounts for expansion signals (new funding rounds, hiring surges, product launches). Her "
         "success metric is pipeline velocity — she needs to demonstrate that automated lead qualification reduces the "
         "time from lead capture to first meeting.",
         "Lead enrichment agent, prospect research agent, account monitoring agent"),

        ("Marketing Manager", "James Thompson, Marketing Manager at a D2C brand",
         "James is responsible for content creation, social media monitoring, and campaign performance analysis. He "
         "needs AI agents that generate first drafts of blog posts based on trending topics in his industry, monitor "
         "brand mentions across social platforms, and compile weekly marketing performance reports. His key challenge "
         "is maintaining a consistent content cadence with a small team — he produces far less content than competitors "
         "with larger marketing departments. He values the template marketplace because he wants to start with "
         "proven workflows rather than building from scratch.",
         "Content generation agent, social monitoring agent, campaign analytics agent"),

        ("IT Administrator", "Priya Sharma, IT Director at a financial services firm",
         "Priya's primary concern is security and compliance. She needs to ensure that any new platform meets the "
         "firm's data governance requirements, integrates with their existing identity provider (Okta), and provides "
         "audit trails for all AI agent activities. She evaluates platforms based on their security architecture, "
         "data residency options, and ability to enforce access controls. She is interested in AutoMind for automating "
         "IT operations tasks such as security alert triage, compliance document generation, and vendor risk assessment, "
         "but will not approve deployment without satisfactory answers to her security questionnaire.",
         "Security alert triage agent, compliance report agent, vendor risk agent"),

        ("Developer", "Alex Kim, Senior Software Engineer building internal tools",
         "Alex has experience with LangChain and LangGraph but spends too much time on infrastructure — deployment, "
         "monitoring, scheduling, and cost management — rather than on agent logic. He wants a platform that handles "
         "the operational complexity of running AI agents in production while giving him the flexibility to write "
         "custom code when the visual builder is insufficient. He is the most likely user of the plugin/node SDK and "
         "template marketplace as both a consumer and contributor. His evaluation criteria focus on API quality, "
         "extensibility, and whether the platform constrains him less than building from scratch.",
         "Custom code execution nodes, API-driven agent management, plugin development"),
    ]

    for title, description, narrative, use_cases in personas:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.color.rgb = PURPLE
        run.bold = True
        p.paragraph_format.space_before = Pt(12)

        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.color.rgb = MED_GRAY
        run.italic = True

        add_body(doc, narrative)

        p = doc.add_paragraph()
        run = p.add_run("Key Use Cases: ")
        run.font.name = FONT
        run.font.size = BODY_SIZE
        run.font.color.rgb = DARK_GRAY
        run.bold = True
        run2 = p.add_run(use_cases)
        run2.font.name = FONT
        run2.font.size = BODY_SIZE
        run2.font.color.rgb = DARK_GRAY


def build_section_7(doc):
    add_heading1(doc, "7", "Revenue Model & Pricing")

    add_heading2(doc, "7.1", "Pricing Tiers")
    add_body(doc,
        "AutoMind follows a tiered SaaS pricing model designed to capture value across the SMB, mid-market, and "
        "enterprise segments. Pricing is structured to encourage adoption at the Starter tier with clear upgrade "
        "incentives as usage scales. All tiers include a 14-day free trial with full feature access."
    )

    add_styled_table(doc,
        ["Feature", "Starter ($99/mo)", "Pro ($499/mo)", "Enterprise (Custom)"],
        [
            ("Agents", "5", "25", "Unlimited"),
            ("Executions / Month", "500", "5,000", "Unlimited"),
            ("AI Models", "GPT-4o, Claude 3.5 Sonnet", "All models incl. GPT-4, Opus", "All models + custom fine-tuned"),
            ("Agent Memory", "Last 10 summaries", "Last 100 summaries", "Unlimited + custom RAG"),
            ("Integrations", "Email, Slack", "All integrations", "All + custom connectors"),
            ("Templates", "Browse only", "Browse + Publish", "Private org templates"),
            ("Users", "1", "5", "Unlimited"),
            ("RBAC", "Basic", "Full", "Full + custom roles"),
            ("Support", "Community + Email", "Priority email + chat", "Dedicated CSM + SLA"),
            ("Cost Tracking", "Basic", "Detailed + alerts", "Detailed + API export"),
            ("JILLA Concierge", "Onboarding only", "Full access", "Custom-trained"),
            ("SSO / SAML", "—", "—", "Included"),
            ("Data Residency", "US (default)", "US / EU", "Custom"),
            ("Plugin SDK Access", "—", "Read-only", "Full + private registry"),
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5],
    )

    add_heading2(doc, "7.2", "Revenue Projections")
    add_body(doc,
        "Revenue projections are based on a bottoms-up model assuming organic growth supplemented by targeted "
        "outbound sales for the enterprise segment. The model assumes a 3% monthly visitor-to-trial conversion "
        "rate, 25% trial-to-paid conversion, and 5% monthly expansion rate from Starter to Pro upgrades."
    )

    add_styled_table(doc,
        ["Metric", "Month 3", "Month 6", "Month 12"],
        [
            ("Starter Customers", "35", "80", "150"),
            ("Pro Customers", "5", "15", "40"),
            ("Enterprise Customers", "0", "2", "5"),
            ("Starter MRR", "$3,465", "$7,920", "$14,850"),
            ("Pro MRR", "$2,495", "$7,485", "$19,960"),
            ("Enterprise MRR", "$0", "$6,000", "$15,000"),
            ("Total MRR", "$5,960", "$21,405", "$49,810"),
            ("Annualized ARR", "$71,520", "$256,860", "$597,720"),
        ],
        col_widths=[5, 3.5, 3.5, 3.5],
    )


def build_section_8(doc):
    add_heading1(doc, "8", "Success Metrics & KPIs")

    add_body(doc,
        "The following key performance indicators (KPIs) will be tracked to measure the success of the AutoMind "
        "platform against its business objectives. Metrics are organized by category and include target values for "
        "the first 12 months post-launch."
    )

    add_styled_table(doc,
        ["Category", "KPI", "Target (12-Month)", "Measurement Method"],
        [
            ("Adoption", "Total registered users", "1,000+", "User database count"),
            ("Adoption", "Monthly active users (MAU)", "400+", "Users with at least one execution per month"),
            ("Adoption", "Agents created", "2,500+", "Agent database count"),
            ("Revenue", "Monthly Recurring Revenue (MRR)", "$50,000", "Stripe subscription data"),
            ("Revenue", "LTV:CAC ratio", "3:1 or higher", "Cohort analysis (finance team)"),
            ("Revenue", "Monthly churn rate", "<5%", "Subscription cancellation tracking"),
            ("Engagement", "Average executions per agent per week", "5+", "Execution log aggregation"),
            ("Engagement", "Template marketplace installs", "500+", "Template install events"),
            ("Engagement", "Chat sessions per user per week", "3+", "Chat session tracking"),
            ("Quality", "Execution success rate", ">95%", "Successful vs. failed execution ratio"),
            ("Quality", "P95 API latency (non-AI)", "<200ms", "Application Performance Monitoring (APM)"),
            ("Quality", "Platform uptime", "99.9%", "Infrastructure monitoring (uptime checks)"),
            ("Satisfaction", "Net Promoter Score (NPS)", "40+", "Quarterly NPS survey"),
            ("Satisfaction", "System Usability Scale (SUS)", "75+", "Bi-annual usability testing"),
            ("Satisfaction", "Support ticket resolution time", "<24 hours", "Help desk SLA tracking"),
        ],
        col_widths=[2.5, 4, 3, 6.5],
    )


def build_section_9(doc):
    add_heading1(doc, "9", "Constraints & Assumptions")

    add_heading2(doc, "9.1", "Constraints")
    constraints = [
        "Budget: The initial development and launch must be completed within a total engineering budget of $250,000, covering a team of four full-time engineers over six months plus infrastructure costs. Post-launch infrastructure costs must not exceed $5,000/month at the 500-customer scale.",
        "Timeline: The minimum viable product (MVP) must be ready for closed beta within 4 months and general availability within 6 months from the project start date. This timeline is driven by competitive pressure from Relevance AI's recent funding announcement and upcoming feature releases.",
        "Technology Stack: The platform must be built on the existing GoalCert technology foundation (React frontend, FastAPI backend, PostgreSQL database, Redis cache, Celery task queue) to leverage team expertise and shared infrastructure. Migration to alternative technologies requires CTO approval and business justification.",
        "AI Provider Dependency: The platform is dependent on third-party LLM providers (OpenAI, Anthropic) for core AI capabilities. Provider API changes, pricing modifications, or service disruptions directly impact platform functionality and cost structure. The architecture must abstract provider-specific logic to mitigate lock-in risk.",
        "Regulatory: The platform must comply with GDPR for EU customers and PDPA for Singapore-based operations. Data processing agreements must be established with all LLM providers to ensure customer data is handled in accordance with applicable regulations.",
        "Team Size: The engineering team is limited to four developers (two backend, one frontend, one full-stack), one product manager, and one designer. Feature prioritization must account for this resource constraint.",
    ]
    for c in constraints:
        add_bullet(doc, c)

    add_heading2(doc, "9.2", "Assumptions")
    assumptions = [
        "LLM API pricing will remain stable or decrease over the projection period, maintaining the viability of the current pricing model. A 20% increase in LLM costs can be absorbed without pricing changes; larger increases require tier price adjustments.",
        "Enterprise customers will accept cloud-hosted (multi-tenant) deployment for v1.0, with dedicated single-tenant or on-premises options available in v2.0. At least 80% of enterprise prospects in the first year will not require on-premises deployment.",
        "The existing GoalCert customer base (from Simulation Engine and Digital Twin products) will provide early adopters for AutoMind, reducing customer acquisition cost for the first 50 customers.",
        "Users of the visual workflow builder will be comfortable with basic logical concepts (if-then conditions, variable references) even if they cannot write code. Onboarding materials and JILLA concierge will bridge the gap for users unfamiliar with these concepts.",
        "The template marketplace will achieve a self-sustaining contribution rate (at least 5 new templates per month from community users) within 9 months of launch, reducing the burden on the core team to maintain template freshness.",
        "Third-party integration providers (Slack, email services) will maintain stable APIs throughout the development period. Breaking changes from integration partners are estimated to require less than 20 engineering hours per incident to resolve.",
    ]
    for a in assumptions:
        add_bullet(doc, a)


def build_section_10(doc):
    add_heading1(doc, "10", "Risks & Mitigations")

    add_body(doc,
        "The following risk register identifies the key risks to the successful delivery and adoption of the "
        "AutoMind platform, along with their likelihood, impact, and planned mitigation strategies."
    )

    add_styled_table(doc,
        ["ID", "Risk", "Likelihood", "Impact", "Mitigation Strategy"],
        [
            ("R-01", "LLM provider API outage or degradation causes agent execution failures", "Medium", "High",
             "Implement multi-provider failover (if OpenAI is unavailable, fall back to Anthropic and vice versa). Cache recent successful responses for idempotent queries. Maintain a provider health dashboard with automated alerting."),
            ("R-02", "LLM API pricing increases erode platform margins", "Medium", "High",
             "Abstract provider layer to enable rapid model switching. Implement cost optimization features (prompt caching, model routing based on task complexity). Maintain 20% pricing buffer in tier pricing. Negotiate volume discounts with providers."),
            ("R-03", "Competitor (Relevance AI) achieves feature parity before AutoMind reaches GA", "Medium", "Medium",
             "Prioritize differentiated features (agent memory, template marketplace, JILLA concierge) that are difficult to replicate quickly. Accelerate timeline by reducing scope of non-critical features (move BR-14 and BR-15 to post-GA). Leverage GoalCert's existing customer relationships for early traction."),
            ("R-04", "Multi-tenant data leakage due to isolation failure", "Low", "Critical",
             "Implement row-level security at the database layer. Automated security testing in CI/CD pipeline. Quarterly third-party penetration testing. Bug bounty program post-GA. Separate execution environments for enterprise-tier customers."),
            ("R-05", "Template marketplace fails to achieve critical mass of community contributions", "Medium", "Medium",
             "Seed marketplace with 25+ high-quality templates built by the core team. Offer template creation bounties during the first 6 months. Feature top contributors in marketing materials. Simplify the template publishing workflow to reduce friction."),
            ("R-06", "Complex workflow builder UX deters non-technical users", "Medium", "High",
             "Conduct iterative usability testing with target persona representatives starting from prototype stage. Implement progressive disclosure (simple mode vs. advanced mode). Invest in JILLA concierge guided walkthroughs. Maintain a library of video tutorials for common workflows."),
            ("R-07", "Celery worker scaling issues under high concurrent execution load", "Medium", "Medium",
             "Load test execution engine to validate 500 concurrent agent target. Implement execution queue prioritization and backpressure mechanisms. Design worker autoscaling based on queue depth. Maintain performance benchmarks in CI pipeline."),
            ("R-08", "Key engineering staff departure during critical development phase", "Low", "High",
             "Maintain comprehensive documentation and code review practices. Cross-train team members across frontend and backend. Use pair programming for complex features. Establish relationships with contract developers who can backfill if needed."),
        ],
        col_widths=[1, 4, 2, 1.5, 7.5],
    )


def build_section_11(doc):
    add_heading1(doc, "11", "Dependencies")

    add_body(doc,
        "The following dependencies represent external factors and internal deliverables that the AutoMind project "
        "relies upon. Each dependency is tracked with an owner and current status."
    )

    add_styled_table(doc,
        ["ID", "Dependency", "Type", "Owner", "Status"],
        [
            ("D-01", "OpenAI API access with sufficient rate limits for production usage (Tier 4+)", "External", "Engineering", "Active — current Tier 3, upgrade request submitted"),
            ("D-02", "Anthropic API access with production-grade rate limits", "External", "Engineering", "Active — production access approved"),
            ("D-03", "GoalCert shared authentication service (SSO, user management)", "Internal", "Platform Team", "In Development — ETA July 2026"),
            ("D-04", "AWS infrastructure provisioning (ECS, RDS, ElastiCache, SES)", "External", "DevOps", "Complete — staging environment operational"),
            ("D-05", "Slack App Directory submission and review for integration publishing", "External", "Product", "Not Started — planned for Month 4"),
            ("D-06", "Legal review of Terms of Service, Privacy Policy, and Data Processing Agreement", "Internal", "Legal", "In Review — first draft circulated"),
            ("D-07", "UX design system and component library for workflow builder", "Internal", "Design", "In Development — 60% complete"),
            ("D-08", "GoalCert marketing website and product landing page", "Internal", "Marketing", "Not Started — planned for Month 5"),
            ("D-09", "Payment processing integration (Stripe) for subscription billing", "External", "Engineering", "Active — Stripe account configured, integration in development"),
            ("D-10", "Redis cluster provisioning for RedBeat scheduler and agent memory store", "External", "DevOps", "Complete — production Redis cluster operational"),
        ],
        col_widths=[1, 6, 2, 2.5, 4.5],
    )


def build_section_12(doc):
    add_heading1(doc, "12", "Glossary")

    add_body(doc,
        "This glossary supplements the definitions provided in Section 1.4 with additional terms used throughout "
        "this document."
    )

    add_styled_table(doc,
        ["Term", "Definition"],
        [
            ("Agentic AI", "An AI system architecture where autonomous agents plan, execute, and iterate on tasks with minimal human intervention, using tools and memory to accomplish complex objectives"),
            ("ARR", "Annual Recurring Revenue — the annualized value of active subscriptions, calculated as MRR multiplied by 12"),
            ("BFS Traversal", "Breadth-First Search traversal — the algorithm used to execute workflow nodes layer by layer, ensuring all inputs to a node are available before it executes"),
            ("CAC", "Customer Acquisition Cost — the total sales and marketing cost required to acquire a single paying customer"),
            ("Celery", "An asynchronous task queue framework for Python used to manage distributed execution of agent workflows across worker processes"),
            ("CRON Expression", "A scheduling syntax (e.g., '0 9 * * 1-5') that defines recurring execution times, used by RedBeat to schedule agent runs"),
            ("DAG", "Directed Acyclic Graph — a graph structure with directed edges and no cycles, used to represent workflow node dependencies"),
            ("LTV", "Customer Lifetime Value — the projected total revenue a customer generates over their entire relationship with the platform"),
            ("MRR", "Monthly Recurring Revenue — the predictable revenue generated from active subscriptions each month"),
            ("NPS", "Net Promoter Score — a customer satisfaction metric ranging from -100 to 100, measured by asking 'How likely are you to recommend this product?'"),
            ("P95/P99 Latency", "The response time at the 95th or 99th percentile — meaning 95% or 99% of requests are served within this time"),
            ("SUS", "System Usability Scale — a standardized questionnaire measuring perceived usability, scored from 0 to 100"),
            ("Variable Substitution", "A templating mechanism where {variable_name} placeholders in node configurations are replaced with actual values from upstream node outputs at execution time"),
        ],
        col_widths=[3.5, 12.5],
    )


def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT
    font.size = BODY_SIZE
    font.color.rgb = DARK_GRAY

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_cover_page(doc)
    add_document_control(doc)
    add_toc(doc)

    build_section_1(doc)
    doc.add_page_break()
    build_section_2(doc)
    doc.add_page_break()
    build_section_3(doc)
    doc.add_page_break()
    build_section_4(doc)
    doc.add_page_break()
    build_section_5(doc)
    doc.add_page_break()
    build_section_6(doc)
    doc.add_page_break()
    build_section_7(doc)
    doc.add_page_break()
    build_section_8(doc)
    doc.add_page_break()
    build_section_9(doc)
    doc.add_page_break()
    build_section_10(doc)
    doc.add_page_break()
    build_section_11(doc)
    doc.add_page_break()
    build_section_12(doc)

    doc.save(SAVE_PATH)
    print(f"BRD saved to {SAVE_PATH}")


if __name__ == "__main__":
    main()
