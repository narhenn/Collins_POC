"""
Generate GoalCert AutoMind FRD (Functional Requirements Document)
Exact Flex Coach format matching all styling rules.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Constants (Flex Coach palette) ──
NAVY = RGBColor(0x1B, 0x2A, 0x4A)       # #1B2A4A
BLUE = RGBColor(0x2E, 0x74, 0xB5)        # #2E74B5
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
META_LEFT_HEX = "E8EEF4"
META_RIGHT_HEX = "FFFFFF"
HDR_HEX = "1B2A4A"
ALT_ROW_HEX = "F0F4F8"
FONT = "Calibri"
BODY_SIZE = Pt(11)
SAVE_PATH = "/Users/narhen/automind/docs/GoalCert_AutoMind_FRD_v1_0.docx"


# ── Helpers ───────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _run(p, text, font_name=FONT, size=Pt(11), bold=False, color=None):
    r = p.add_run(text)
    r.font.name = font_name
    r.font.size = size
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def add_body(doc, text):
    p = doc.add_paragraph()
    _run(p, text, size=BODY_SIZE)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    _run(p, text, size=BODY_SIZE)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


def add_heading1(doc, text):
    """Heading 1 = 20pt bold #1B2A4A"""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.clear()
    _run(p, text, size=Pt(20), bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    return p


def add_heading2(doc, text):
    """Heading 2 = 15pt bold #2E74B5"""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.clear()
    _run(p, text, size=Pt(15), bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_styled_table(doc, headers, rows, col_widths=None):
    """Header fill #1B2A4A white bold text, alternating rows #F0F4F8."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(p, h, size=Pt(10), bold=True, color=WHITE)
        set_cell_shading(cell, HDR_HEX)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _run(p, str(val), size=Pt(10))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_HEX)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_metadata_table(doc, rows):
    """Left col fill #E8EEF4, right col white."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        cell_l = table.rows[i].cells[0]
        cell_l.text = ""
        _run(cell_l.paragraphs[0], label, size=Pt(10), bold=True)
        set_cell_shading(cell_l, META_LEFT_HEX)
        cell_l.width = Cm(5)

        cell_r = table.rows[i].cells[1]
        cell_r.text = ""
        _run(cell_r.paragraphs[0], value, size=Pt(10))
        set_cell_shading(cell_r, META_RIGHT_HEX)
        cell_r.width = Cm(11)
    return table


def add_req_table(doc, reqs):
    """Each requirement as a 7-row table: FR ID, Requirement, Input, Process, Output, Priority, Validation Rules."""
    for req in reqs:
        fr_id, requirement, inp, process, output, priority, validation = req
        rows_data = [
            ("FR ID", fr_id),
            ("Requirement", requirement),
            ("Input", inp),
            ("Process", process),
            ("Output", output),
            ("Priority", priority),
            ("Validation Rules", validation),
        ]
        table = doc.add_table(rows=len(rows_data), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(rows_data):
            cell_l = table.rows[i].cells[0]
            cell_l.text = ""
            _run(cell_l.paragraphs[0], label, size=Pt(10), bold=True, color=WHITE)
            set_cell_shading(cell_l, HDR_HEX)
            cell_l.width = Cm(3.5)

            cell_r = table.rows[i].cells[1]
            cell_r.text = ""
            _run(cell_r.paragraphs[0], value, size=Pt(10))
            cell_r.width = Cm(12.5)
        doc.add_paragraph()  # spacer


# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════

def add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "GoalCert AutoMind", size=Pt(32), bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Functional Requirements Document (FRD)", size=Pt(20), color=BLUE)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Agentic AI Workforce Platform", size=Pt(14), color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Version 1.0  |  25 June 2026", size=Pt(11), color=BLUE)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "CONFIDENTIAL", size=Pt(12), bold=True, color=RGBColor(180, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "GoalCert Pte Ltd. All rights reserved.", size=Pt(9), color=BLUE)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# METADATA + REVISION HISTORY
# ═══════════════════════════════════════════════════════════════════

def add_document_control(doc):
    add_heading1(doc, "Document Information")

    add_metadata_table(doc, [
        ("Document Title", "Functional Requirements Document (FRD)"),
        ("Product", "GoalCert AutoMind"),
        ("Version", "1.0"),
        ("Date", "25 June 2026"),
        ("Classification", "Confidential"),
        ("Prepared By", "Engineering Team"),
        ("Reviewed By", "Product Management"),
        ("Approved By", "CTO"),
    ])

    doc.add_paragraph()
    add_heading2(doc, "Revision History")

    add_styled_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ("0.1", "12 June 2026", "Engineering", "Initial draft with FR-AUTH and FR-AGT requirements"),
            ("0.5", "17 June 2026", "Engineering", "Added workflow builder, execution engine, memory, and chat modules"),
            ("0.9", "21 June 2026", "Engineering", "Incorporated review feedback, added NFRs and error handling matrix"),
            ("1.0", "25 June 2026", "Engineering", "Final release approved by CTO and Product Management"),
        ],
        col_widths=[2, 3, 3, 8],
    )

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (Manual Normal style paragraphs)
# ═══════════════════════════════════════════════════════════════════

def add_toc(doc):
    p = doc.add_paragraph()
    _run(p, "Table of Contents", size=Pt(20), bold=True, color=NAVY)
    doc.add_paragraph()

    toc_l1 = [
        "1. Introduction",
        "2. System Actors & Permissions",
        "3. Functional Requirements",
        "4. Non-Functional Requirements",
        "5. UI Requirements",
        "6. Error Handling Matrix",
        "7. Glossary",
    ]
    toc_l2 = {
        "1. Introduction": ["1.1 Purpose", "1.2 Audience", "1.3 Scope"],
        "3. Functional Requirements": [
            "3.1 FR-AUTH: Authentication & Authorization",
            "3.2 FR-AGT: Agent Management",
            "3.3 FR-WFB: Workflow Builder",
            "3.4 FR-EXE: Execution Engine",
            "3.5 FR-SCH: Scheduled Execution",
            "3.6 FR-CHT: Agent Chat",
            "3.7 FR-MEM: Agent Memory",
            "3.8 FR-RPT: Reports & Analytics",
            "3.9 FR-TPL: Templates",
            "3.10 FR-INT: Integrations",
        ],
    }

    for item in toc_l1:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        _run(p, item, size=Pt(11), bold=True, color=NAVY)
        p.paragraph_format.space_after = Pt(2)

        if item in toc_l2:
            for sub in toc_l2[item]:
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                _run(p, sub, size=Pt(11), bold=False, color=NAVY)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(1)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# SECTION 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════════

def build_section_1(doc):
    add_heading1(doc, "1. Introduction")

    add_heading2(doc, "1.1 Purpose")
    add_body(doc,
        "This Functional Requirements Document (FRD) specifies the detailed functional behavior of the GoalCert "
        "AutoMind platform. It decomposes the business requirements defined in the companion BRD v1.0 into "
        "module-level functional requirements that are implementable, testable, and traceable. The FRD serves as "
        "the primary contract between product management and engineering."
    )

    add_heading2(doc, "1.2 Audience")
    add_body(doc,
        "This document is intended for software engineers, QA engineers, UX designers, technical architects, and "
        "project managers who need precise specifications for implementation, testing, and project planning. "
        "Business stakeholders may reference this document to validate that their requirements are accurately captured."
    )

    add_heading2(doc, "1.3 Scope")
    add_body(doc,
        "GoalCert AutoMind is a web-based platform that enables users to create, configure, and deploy autonomous "
        "AI agents through a visual workflow builder. The platform consists of:"
    )
    components = [
        "React Frontend (SPA): Visual workflow builder canvas, agent management dashboard, execution console, chat interface, analytics, reports, templates, integrations, and settings pages",
        "FastAPI Backend: RESTful API server handling authentication (JWT/bcrypt), agent CRUD, workflow persistence (JSON DAGs), execution orchestration, SSE streaming, and integration management",
        "PostgreSQL Database: Primary data store for users, agents, workflows, executions, execution node logs, agent memory summaries, templates, and integrations",
        "Redis: Message broker for Celery task queue, RedBeat schedule storage, and real-time execution event pub/sub for SSE streaming",
        "Celery Workers: Distributed task execution workers processing agent workflow runs via BFS traversal, publishing progress events to Redis channels",
        "RedBeat Scheduler: Celery beat scheduler backed by Redis managing cron-based agent execution schedules with timezone support",
        "OpenAI Integration: GPT-4o-mini for AI Action node execution, agent chat responses, and memory summary generation",
    ]
    for c in components:
        add_bullet(doc, c)


# ═══════════════════════════════════════════════════════════════════
# SECTION 2: SYSTEM ACTORS & PERMISSIONS
# ═══════════════════════════════════════════════════════════════════

def build_section_2(doc):
    add_heading1(doc, "2. System Actors & Permissions")

    add_body(doc,
        "The AutoMind platform defines four actor types. Access control is enforced at the API layer via JWT-based "
        "authentication with user-scoped database queries ensuring complete data isolation."
    )

    add_styled_table(doc,
        ["Actor", "Description", "Permissions"],
        [
            ("Admin",
             "Organization administrator with full platform access",
             "All User permissions; manage organization settings; access organization-wide analytics; manage user accounts; publish templates; configure integrations globally"),
            ("User",
             "Standard platform user who creates and manages their own agents",
             "Register/login; create, edit, delete own agents; configure workflows; trigger executions; view own execution logs and analytics; chat with agents; manage own integrations; browse templates; view own cost data"),
            ("Agent",
             "Autonomous AI entity that executes workflows on behalf of a user",
             "Execute assigned workflow nodes; call external APIs (OpenAI, Resend, Slack); read/write execution variables; generate memory summaries; publish SSE events; access integration credentials"),
            ("System",
             "Platform infrastructure components (Celery workers, RedBeat scheduler, Redis)",
             "Dispatch and execute Celery tasks; manage scheduled task entries; publish/subscribe to Redis channels; update execution records; generate memory entries; enforce timeouts"),
        ],
        col_widths=[2.5, 5, 8.5],
    )


# ═══════════════════════════════════════════════════════════════════
# SECTION 3: FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════

def build_section_3(doc):
    add_heading1(doc, "3. Functional Requirements")

    # ── 3.1 FR-AUTH ──────────────────────────────────────────────────
    add_heading2(doc, "3.1 FR-AUTH: Authentication & Authorization")

    add_req_table(doc, [
        ("FR-AUTH-01",
         "User Registration",
         "Email, password, display name via POST /api/auth/register",
         "Validate email uniqueness; hash password with bcrypt; create User record; generate JWT access token with user ID and email in claims",
         "AuthResponse containing UserResponse (id, email, name, created_at) and JWT token; 201 Created on success; 409 Conflict if email already exists",
         "Must",
         "Email must be valid format; password minimum 8 characters; name is optional; duplicate email returns 409"),

        ("FR-AUTH-02",
         "User Login",
         "Email and password via POST /api/auth/login",
         "Look up user by email; verify password against stored bcrypt hash; generate JWT access token on success",
         "AuthResponse with UserResponse and JWT token; 401 Unauthorized if email not found or password mismatch",
         "Must",
         "Generic 401 error without revealing whether email exists; password verification uses bcrypt.checkpw"),

        ("FR-AUTH-03",
         "Get Current User Profile",
         "JWT Bearer token via GET /api/auth/me",
         "Decode JWT; extract user ID from sub claim; query User table; return profile data",
         "UserResponse (id, email, name, created_at); 401 if token expired or invalid",
         "Must",
         "Token must contain valid sub claim; user must exist in database"),

        ("FR-AUTH-04",
         "JWT Token Management",
         "User ID and email for token generation",
         "Create JWT with sub (user_id), email, and exp claims using HS256 algorithm; configurable expiry via JWT_EXPIRY_HOURS setting",
         "Signed JWT string; decoded payload on validation; HTTPException on expiry or invalid signature",
         "Must",
         "Token algorithm is HS256; secret key from environment; expired tokens raise 401; invalid tokens raise 401"),

        ("FR-AUTH-05",
         "Route-Level Authorization",
         "JWT Bearer token on all protected endpoints",
         "OAuth2PasswordBearer dependency extracts token; get_current_user dependency decodes and validates; all agent/workflow/execution endpoints require authenticated user",
         "Current User object injected into route handler; 401 if token missing, expired, or user not found",
         "Must",
         "Every API endpoint except /api/auth/register, /api/auth/login, and /api/templates requires authentication"),
    ])

    # ── 3.2 FR-AGT ──────────────────────────────────────────────────
    add_heading2(doc, "3.2 FR-AGT: Agent Management")

    add_req_table(doc, [
        ("FR-AGT-01",
         "Create Agent",
         "Agent name, type (sales|marketing|support|custom), description, optional template_id via POST /api/agents",
         "Create Agent record with user_id, name, type, description, status=draft; create associated Workflow with empty definition or template workflow if template_id provided; enforce unique (user_id, name) constraint",
         "AgentResponse with id, user_id, name, description, type, status, schedule fields, timestamps, total_executions=0, success_rate=null; 201 Created; 409 if duplicate name",
         "Must",
         "Name max 255 chars; type must be one of sales, marketing, support, custom; template_id is optional; duplicate name per user returns 409"),

        ("FR-AGT-02",
         "List Agents",
         "JWT token via GET /api/agents",
         "Query all agents where user_id = current_user.id ordered by created_at desc; for each agent compute total_executions (count of executions) and success_rate (avg of success=1.0, failed=0.0)",
         "Array of AgentResponse objects with computed execution stats",
         "Must",
         "Only returns agents owned by authenticated user; stats computed from executions table using aggregate queries"),

        ("FR-AGT-03",
         "Get Agent Detail",
         "Agent ID via GET /api/agents/{agent_id}",
         "Query agent by ID and user_id; compute execution stats; return 404 if not found or not owned by user",
         "Single AgentResponse with computed stats; 404 if not found",
         "Must",
         "Agent must belong to authenticated user; 404 returned for non-existent or unauthorized access"),

        ("FR-AGT-04",
         "Update Agent",
         "Agent ID and partial update fields via PATCH /api/agents/{agent_id}",
         "Load agent by ID and user_id; apply only provided fields (exclude_unset); flush and refresh",
         "Updated AgentResponse with recomputed stats; 404 if not found",
         "Must",
         "Only provided fields are updated; supports updating name, description, type, status, schedule_cron, schedule_timezone"),

        ("FR-AGT-05",
         "Delete Agent",
         "Agent ID via DELETE /api/agents/{agent_id}",
         "Load agent by ID and user_id; unschedule any RedBeat entry; cascade delete agent, workflow, executions, node logs, and memory",
         "204 No Content; 404 if not found",
         "Must",
         "Cascade delete removes all associated records; RedBeat entry removed before deletion"),

        ("FR-AGT-06",
         "Pause Agent",
         "Agent ID via POST /api/agents/{agent_id}/pause",
         "Set agent status to paused; remove RedBeat scheduled entry via unschedule_agent",
         "Updated AgentResponse with status=paused; 404 if not found",
         "Must",
         "Agent must belong to user; unschedule removes RedBeat entry silently if not found"),

        ("FR-AGT-07",
         "Resume Agent",
         "Agent ID via POST /api/agents/{agent_id}/resume",
         "Set agent status to active; if agent has schedule_cron, register schedule with RedBeat via schedule_agent",
         "Updated AgentResponse with status=active; 404 if not found",
         "Must",
         "Only re-registers schedule if schedule_cron is set; schedule_agent creates RedBeat entry with cron and timezone"),

        ("FR-AGT-08",
         "Generate Agent from Description",
         "Natural language description via POST /api/agents/generate",
         "Send description to OpenAI GPT-4o-mini with system prompt defining available node types (trigger, ai_action, integration, decision, escalation, web_search, code_exec); parse JSON response containing agent_name, agent_type, agent_description, schedule, and workflow definition; create Agent and Workflow records",
         "AgentResponse with generated name, type, description, and populated workflow; 201 Created; 422 if generation fails; 500 on unexpected error",
         "Must",
         "Description must be non-empty; agent_type validated against allowed values; workflow must contain nodes array; JSON parsing handles code fences"),
    ])

    # ── 3.3 FR-WFB ──────────────────────────────────────────────────
    add_heading2(doc, "3.3 FR-WFB: Workflow Builder")

    add_req_table(doc, [
        ("FR-WFB-01",
         "Node Type Support",
         "User selects from 7 node types in component sidebar",
         "System supports trigger, ai_action, web_search, decision, integration, escalation, and code_exec node types; each node type has a dedicated executor class with specific configuration schema",
         "Nodes rendered on React Flow canvas with type-specific icons, labels, and connection ports",
         "Must",
         "Every node must have id, type, data (containing label and config), and position; trigger node required as workflow entry point"),

        ("FR-WFB-02",
         "Drag-and-Drop Node Addition",
         "User drags node type from ComponentSidebar onto WorkflowCanvas",
         "React Flow onDrop handler creates new node with unique ID, default config, and drop position; node added to canvas state",
         "New node appears on canvas at drop position with default configuration",
         "Must",
         "Node ID must be unique within workflow; default config varies by node type; canvas auto-fits on initial load"),

        ("FR-WFB-03",
         "Edge Connection Management",
         "User connects output port of one node to input port of another",
         "React Flow onConnect handler creates edge with source and target node IDs; decision nodes have sourceHandle for true/false branching",
         "Curved Bezier edge rendered between nodes; decision edges labeled true/false",
         "Must",
         "Self-loops should be prevented; decision nodes require sourceHandle true or false; edges can be deleted"),

        ("FR-WFB-04",
         "Node Configuration Panel",
         "User selects a node on canvas",
         "NodeConfigPanel component renders type-specific configuration form; fields vary by node type (prompt template for ai_action, query for web_search, condition fields for decision, service/action/recipients for integration, code for code_exec); changes update node data via handleUpdateNodeData callback",
         "Configuration panel displays on right side with editable fields; changes mark workflow as dirty",
         "Must",
         "Config panel only shows when a node is selected; builderStore tracks selectedNodeId and dirty state; all changes auto-applied"),

        ("FR-WFB-05",
         "Workflow Canvas Operations",
         "User interacts with canvas (pan, zoom, select, move nodes)",
         "React Flow canvas supports pan, zoom, grid snapping, minimap, and node selection; viewport state (x, y, zoom) persisted in workflow definition",
         "Interactive canvas with full pan/zoom/select capabilities",
         "Must",
         "Viewport state saved with workflow definition; canvas background uses grid pattern"),

        ("FR-WFB-06",
         "Save Workflow",
         "Save button clicked or auto-save triggered",
         "PUT /api/agents/{agent_id}/workflow sends current nodes, edges, and viewport as JSON definition to backend; backend updates Workflow record",
         "WorkflowResponse with updated definition, status, timestamps; dirty flag reset",
         "Must",
         "Definition stored as JSONB in PostgreSQL; includes nodes array, edges array, and viewport object"),

        ("FR-WFB-07",
         "Deploy Workflow",
         "Deploy button clicked via POST /api/agents/{agent_id}/workflow/deploy",
         "Set workflow status to active; set deployed_at timestamp; set agent status to active; if agent has schedule_cron, register with RedBeat scheduler",
         "WorkflowResponse with status=active and deployed_at set; agent status updated to active",
         "Must",
         "Deploy activates both workflow and agent; schedule registration only occurs if schedule_cron is configured"),
    ])

    # ── 3.4 FR-EXE ──────────────────────────────────────────────────
    add_heading2(doc, "3.4 FR-EXE: Execution Engine")

    add_req_table(doc, [
        ("FR-EXE-01",
         "Manual Execution Trigger",
         "Agent ID via POST /api/agents/{agent_id}/execute",
         "Verify agent ownership; load workflow; validate workflow is active or draft and has nodes; create Execution record with status=pending, triggered_by=manual; dispatch Celery task execute_workflow with execution_id and workflow definition",
         "ExecutionResponse with execution ID, status=pending; 202 Accepted; 404 if agent/workflow not found; 400 if workflow invalid",
         "Must",
         "Workflow must have status active or draft; workflow must contain at least one node; Celery task dispatched asynchronously"),

        ("FR-EXE-02",
         "Celery Task Execution",
         "execution_id and workflow_definition from Celery task queue",
         "execute_workflow_task runs asyncio.run to drive async executor; creates fresh SQLAlchemy engine per event loop; WorkflowExecutor processes workflow; on failure, marks execution as failed in database",
         "Task result dict with status (completed/failed) and execution_id; execution record updated in database",
         "Must",
         "Each task creates its own async engine (asyncpg not reusable across event loops); uncaught exceptions trigger _mark_failed fallback"),

        ("FR-EXE-03",
         "BFS Workflow Traversal",
         "Workflow definition containing nodes and edges",
         "Parse nodes and edges from definition; identify root node (no incoming edges, prefer trigger type); BFS traversal using queue; for each node execute via type-specific executor; decision nodes determine branch (true/false) to follow; skip nodes on non-taken branches",
         "All reachable nodes executed in BFS order; variables dict accumulated across nodes; final status success or failed",
         "Must",
         "Exactly one start node identified; visited set prevents re-execution; decision branch filtering uses edge label/sourceHandle matching"),

        ("FR-EXE-04",
         "Node Execution Lifecycle",
         "Node config, workflow variables, memory context",
         "For each node: create ExecutionNodeLog with status=running; resolve variable substitutions in config using {variable} syntax; call executor.execute(); record output and timing; merge output_variables into workflow variables; track LLM cost for ai_action/web_search/code_exec nodes",
         "ExecutionNodeLog updated with status (success/failed), output_data, duration_ms, llm_usage; workflow variables updated with node outputs",
         "Must",
         "Variable interpolation supports dot-notation for nested access and array indexing; unresolvable placeholders left as-is; cost accumulated in total_cost"),

        ("FR-EXE-05",
         "Variable Interpolation Engine",
         "Template string with {variable} placeholders and variables dictionary",
         "Regex matches {placeholder} patterns; resolves via dot-separated path walking (dict keys, list indices); if entire string is single placeholder, returns raw value (preserves types); otherwise string substitution; recursively interpolates dicts and lists",
         "Resolved values with types preserved for single-placeholder strings; string interpolation for mixed templates; unresolved placeholders left unchanged",
         "Must",
         "Supports nested paths like {leads.0.name}; type preservation for int/list/dict when entire string is one placeholder; recursive interpolation for dict and list config values"),

        ("FR-EXE-06",
         "SSE Real-Time Streaming",
         "Client connects to GET /api/executions/{execution_id}/stream",
         "Create async Redis pub/sub subscription on channel execution:{execution_id}:logs; yield SSE events as messages arrive; close stream on __STREAM_END__ sentinel or client disconnect",
         "Server-Sent Events with JSON payloads containing timestamp, message, node_id, and status fields; stream closes on completion",
         "Must",
         "No authentication required (execution IDs are unguessable UUIDs); Redis pub/sub used for real-time event delivery; executor publishes events for each node start, completion, and error"),

        ("FR-EXE-07",
         "Execution Error Handling",
         "Exception during node execution",
         "Catch exception; log to ExecutionNodeLog with status=failed and error_message; publish error event to SSE; set execution status to failed with error_message; publish __STREAM_END__ sentinel",
         "Execution marked failed with error details; partial results preserved for completed nodes; SSE stream terminated cleanly",
         "Must",
         "Node failure stops entire execution (fail-fast); error message includes node label and type; _mark_failed fallback catches top-level Celery task failures"),

        ("FR-EXE-08",
         "Execution Cost Tracking",
         "LLM usage data from ai_action, web_search, code_exec node executors",
         "Each AI node executor returns llm_usage dict with cost field; WorkflowExecutor accumulates total_cost; on completion, total_cost stored as Decimal(10,6) in execution record",
         "Execution record contains total_cost; per-node llm_usage stored in ExecutionNodeLog; cost visible in execution detail and dashboard",
         "Must",
         "Cost stored as Decimal for precision; llm_usage JSON includes model, tokens, and cost; cost data aggregated in dashboard stats and analytics"),
    ])

    # ── 3.5 FR-SCH ──────────────────────────────────────────────────
    add_heading2(doc, "3.5 FR-SCH: Scheduled Execution")

    add_req_table(doc, [
        ("FR-SCH-01",
         "Cron Schedule Configuration",
         "Agent schedule_cron field (5-part cron expression) and schedule_timezone",
         "Agent model stores schedule_cron (String 100) and schedule_timezone (String 50, default UTC); cron expression validated as exactly 5 space-separated parts",
         "Agent record with cron and timezone; validation error if cron format invalid",
         "Must",
         "Cron format: minute hour day_of_month month day_of_week; timezone defaults to UTC; stored on Agent model"),

        ("FR-SCH-02",
         "RedBeat Schedule Registration",
         "Agent ID, cron expression, timezone",
         "schedule_agent function: remove any existing entry via unschedule_agent; create RedBeatSchedulerEntry with name=agent:{agent_id}, task=execute_workflow_scheduled, schedule=parsed crontab, args=[agent_id]",
         "RedBeat entry stored in Redis; Celery beat triggers task on schedule",
         "Must",
         "Existing entry removed before creating new one to avoid duplicates; entry name format is agent:{uuid}; task name is execute_workflow_scheduled"),

        ("FR-SCH-03",
         "Scheduled Execution Task",
         "agent_id from RedBeat trigger via execute_workflow_scheduled Celery task",
         "Look up agent by ID; verify status=active; look up workflow; verify workflow has valid definition with nodes; create Execution record with triggered_by=schedule; run WorkflowExecutor",
         "Execution completed; agent with status!=active or missing workflow silently skipped",
         "Must",
         "Inactive agents silently skipped without error; empty workflow silently skipped; execution creation and workflow run use fresh SQLAlchemy engine"),

        ("FR-SCH-04",
         "Timezone Support",
         "schedule_timezone field on Agent model",
         "Agent stores timezone string (e.g., UTC, America/New_York, Asia/Singapore); timezone passed to schedule_agent for RedBeat entry; Celery crontab handles timezone conversion",
         "Scheduled executions fire at correct local time regardless of server timezone",
         "Must",
         "Default timezone is UTC; timezone stored as IANA timezone string; Celery crontab supports timezone parameter"),

        ("FR-SCH-05",
         "Pause and Resume Scheduling",
         "Pause via POST /api/agents/{agent_id}/pause; Resume via POST /api/agents/{agent_id}/resume",
         "Pause: set status=paused, call unschedule_agent to remove RedBeat entry; Resume: set status=active, if schedule_cron exists call schedule_agent to re-register",
         "Agent paused stops scheduled executions; resume restores schedule; manual execution still works regardless of status",
         "Must",
         "unschedule_agent silently no-ops if entry does not exist; resume only re-registers if schedule_cron is set"),
    ])

    # ── 3.6 FR-CHT ──────────────────────────────────────────────────
    add_heading2(doc, "3.6 FR-CHT: Agent Chat")

    add_req_table(doc, [
        ("FR-CHT-01",
         "Chat with Agent (Streaming)",
         "Agent ID, user message, conversation history via POST /api/agents/{agent_id}/chat",
         "Verify agent ownership; build system prompt from agent name, type, description, status, schedule; retrieve agent memory context (last 15 memories); construct OpenAI chat messages with system prompt, last 10 history messages, and user message; stream response via GPT-4o-mini",
         "SSE stream of JSON events: {type: token, content: text_chunk} for each token, {type: done} on completion, {type: error, content: message} on failure",
         "Must",
         "History limited to last 10 messages; memory context limited to 15 most recent entries; model is gpt-4o-mini; max_tokens=1024; temperature=0.7"),

        ("FR-CHT-02",
         "Memory-Augmented Context",
         "Agent ID for memory retrieval",
         "get_agent_context queries AgentMemory table for agent, ordered by created_at desc, limited to N entries; formats as chronological list with timestamps: [YYYY-MM-DD HH:MM] summary",
         "Formatted context string injected into chat system prompt under 'Previous execution context' header",
         "Must",
         "Memory context provides execution history awareness; empty string returned if no memories exist"),

        ("FR-CHT-03",
         "Conversation History Support",
         "Array of {role: user|assistant, content: string} in request body",
         "Frontend maintains conversation state in component; sends last 10 messages as history array; backend prepends system prompt and appends current user message",
         "Contextual responses that reference prior conversation turns",
         "Must",
         "History is client-side managed; backend does not persist chat sessions; last 10 messages included for context window efficiency"),

        ("FR-CHT-04",
         "Agent Persona in Chat",
         "Agent name, type, description, status, schedule from Agent model",
         "System prompt constructed: 'You are [agent.name], an AI agent on the GoalCert AutoMind platform'; includes type, description, status, schedule; instructs agent to reference execution data and explain capabilities",
         "Chat responses reflect agent's identity, role, and execution history",
         "Must",
         "System prompt dynamically generated per chat request; includes current agent status and schedule"),

        ("FR-CHT-05",
         "Chat Fallback without API Key",
         "OpenAI API key not configured",
         "Check settings.OPENAI_API_KEY; if not set, return static JSON response indicating chat requires API key configuration",
         "JSON response: {response: 'Chat requires an OpenAI API key. Please configure it in Settings.'}",
         "Must",
         "Non-streaming fallback; no SSE; directs user to Settings page"),
    ])

    # ── 3.7 FR-MEM ──────────────────────────────────────────────────
    add_heading2(doc, "3.7 FR-MEM: Agent Memory")

    add_req_table(doc, [
        ("FR-MEM-01",
         "Execution Summary Generation",
         "Execution results: agent_name, status, duration_ms, node_count, key_outputs",
         "After execution completes, save_execution_memory generates summary; if OpenAI key available, send to GPT-4o-mini with summarization prompt requesting 2-3 concise sentences; fallback to template: 'Execution {status} in {duration_ms}ms with {node_count} nodes'; store as AgentMemory record with agent_id, execution_id, summary, key_outputs JSONB, memory_type=execution_summary",
         "AgentMemory record created; key_outputs stored as JSONB; memory available for future execution context and chat",
         "Must",
         "key_outputs excludes internal variables (triggered_by, branch, condition_result, _-prefixed); LLM summary temperature=0.3, max_tokens=200; fallback works without API key"),

        ("FR-MEM-02",
         "Memory Retrieval for Execution Context",
         "Agent ID and limit parameter",
         "get_agent_context queries AgentMemory ordered by created_at desc with configurable limit (default 10 for executions, 15 for chat); formats as timestamped list",
         "Formatted context string: 'Previous execution context (most recent first):\\n- [YYYY-MM-DD HH:MM] summary' for each memory",
         "Must",
         "Returns empty string if no memories; context injected at start of execution before first AI Action node; also used by chat system prompt"),

        ("FR-MEM-03",
         "List Agent Memories",
         "Agent ID, pagination (limit, offset) via GET /api/agents/{agent_id}/memory",
         "Verify agent ownership; count total memories; query paginated memories ordered by created_at desc",
         "MemoryListResponse containing memories array (id, agent_id, execution_id, summary, key_outputs, memory_type, created_at) and total count",
         "Must",
         "Default limit=20, offset=0; only agent owner can view memories; total count supports frontend pagination"),

        ("FR-MEM-04",
         "Clear Agent Memories",
         "Agent ID via DELETE /api/agents/{agent_id}/memory",
         "Verify agent ownership; delete all AgentMemory records where agent_id matches",
         "204 No Content; all memories permanently removed",
         "Must",
         "Irreversible operation; frontend shows confirmation dialog before calling; 404 if agent not found"),
    ])

    # ── 3.8 FR-RPT ──────────────────────────────────────────────────
    add_heading2(doc, "3.8 FR-RPT: Reports & Analytics")

    add_req_table(doc, [
        ("FR-RPT-01",
         "Dashboard Statistics",
         "JWT token via GET /api/dashboard/stats",
         "Count total and active agents for user; count successful executions; compute average duration_ms; calculate estimated_savings as tasks_completed * 0.5 hours * $50/hr",
         "DashboardStats: total_agents, active_agents, tasks_completed, estimated_savings, avg_response_time",
         "Must",
         "Stats computed in real-time from database; estimated_savings is a calculated metric; avg_response_time may be null if no executions"),

        ("FR-RPT-02",
         "Activity Feed",
         "JWT token via GET /api/dashboard/activity",
         "Query last 20 executions for user's agents joined with agent name and type; ordered by created_at desc",
         "Array of ActivityEvent: execution_id, agent_id, agent_name, agent_type, status, triggered_by, started_at, ended_at, duration_ms, total_cost, created_at",
         "Must",
         "Limited to 20 most recent; includes cost data per execution; cost converted to float from Decimal"),

        ("FR-RPT-03",
         "Analytics Page Metrics",
         "Aggregate data from agents and executions via frontend",
         "Frontend computes: total agents, active agents, total executions (sum across agents), avg success rate, tasks completed (from dashboard stats), estimated savings; displays in 6-metric grid",
         "Six metric cards with icons: Total Agents, Active Agents, Total Executions, Avg Success Rate, Tasks Completed, Estimated Savings",
         "Must",
         "All metrics derived from existing API endpoints; no dedicated analytics API; frontend aggregates client-side"),

        ("FR-RPT-04",
         "Cost by Agent Visualization",
         "Activity feed data aggregated by agent name",
         "Frontend aggregates total_cost and execution count per agent from activity feed; renders horizontal bar chart with percentage-based width; displays cost amount and execution count per agent",
         "Cost by Agent panel with bar chart showing relative cost distribution across agents",
         "Must",
         "Bar width proportional to highest-cost agent; minimum 2% bar width for visibility; cost formatted with formatCost utility"),

        ("FR-RPT-05",
         "Activity Heatmap",
         "Activity feed data grouped by day of week",
         "Frontend groups executions by day-of-week (Mon-Sun); counts per day; renders 7-cell grid with intensity-based coloring (rgba opacity mapped to count/maxCount)",
         "Weekly heatmap grid with color intensity indicating execution frequency; legend from Less to More",
         "Must",
         "Day mapping: JS getDay adjusted to Mon=0 through Sun=6; color uses rgba(73,2,162,alpha); includes count label in each cell"),

        ("FR-RPT-06",
         "Consolidated Agent Reports Page",
         "Agent list and latest execution logs per agent",
         "Reports page shows executive summary banner (total agents, total executions, total cost, avg success rate); per-agent report cards showing status, last run, execution count, success rate, cost, and latest AI output extracted from execution node logs",
         "Executive summary with 4 metrics; per-agent cards with AI output rendered as formatted markdown; agents without executions shown as empty state",
         "Must",
         "AI output extracted from ai_action node logs by finding longest output text; markdown rendering supports headers (# ## ###), bold (**text**), and line breaks"),
    ])

    # ── 3.9 FR-TPL ──────────────────────────────────────────────────
    add_heading2(doc, "3.9 FR-TPL: Templates")

    add_req_table(doc, [
        ("FR-TPL-01",
         "List Templates",
         "GET /api/templates (no authentication required)",
         "Query all AgentTemplate records ordered by name; return template details including workflow_definition, features list, icon, and color",
         "Array of TemplateResponse: id, name, description, type, workflow_definition, icon, color, features, created_at",
         "Must",
         "Templates are public (no auth required); templates seeded in database; ordered alphabetically by name"),

        ("FR-TPL-02",
         "Get Template Detail",
         "Template ID via GET /api/templates/{template_id}",
         "Query AgentTemplate by ID; return full template with workflow_definition",
         "Single TemplateResponse; 404 if not found",
         "Must",
         "No authentication required; workflow_definition contains full React Flow JSON"),

        ("FR-TPL-03",
         "Template Duplication (Create Agent from Template)",
         "Template ID passed as template_id in POST /api/agents create request",
         "During agent creation (FR-AGT-01), if template_id provided, look up template; copy template workflow_definition into new agent's workflow",
         "New agent created with template's workflow pre-loaded; agent has status=draft",
         "Must",
         "Template workflow deep-copied (not referenced); agent name and type still user-specified; template features displayed on template card in UI"),

        ("FR-TPL-04",
         "Template Display and Configuration",
         "Template data rendered in TemplatesPage frontend",
         "Templates displayed as card grid with gradient header, type icon, type badge, name, description, feature list with checkmarks, and 'Duplicate & Configure' button; clicking opens CreateAgentModal with preselectedTemplate; 'Build from Scratch' card also available",
         "Visual template marketplace with type-specific icons (Briefcase=sales, PenTool=marketing, Mail=support, Users=custom), colored badges, feature lists, and one-click duplication",
         "Must",
         "Template cards show hover lift effect; features rendered as checklist with green checkmarks; Build from Scratch card uses dashed border"),
    ])

    # ── 3.10 FR-INT ─────────────────────────────────────────────────
    add_heading2(doc, "3.10 FR-INT: Integrations")

    add_req_table(doc, [
        ("FR-INT-01",
         "Integration Node Types in Workflow",
         "Integration node configuration: service (email|slack), action, recipients, subject, body with {variable} placeholders",
         "IntegrationNodeExecutor resolves variable substitutions in all config fields; looks up user's integration credentials from database; for email (Resend): sends via Resend API with api_key; for Slack: sends via webhook_url; records delivery status",
         "Node output with delivery status (sent/failed) and service response; integration credentials masked in API responses (last 4 chars for keys, origin-only for webhook URLs)",
         "Must",
         "Allowed services: resend, slack; sensitive config keys: api_key, webhook_url, secret, token; integration looked up by user_id and service"),

        ("FR-INT-02",
         "Connect Integration",
         "Service name and config via POST /api/integrations",
         "Validate service against allowed list (resend, slack); check for existing integration (upsert); if exists, update config and set status=active; if new, create Integration record with user_id, service, config, status=active",
         "IntegrationResponse with id, user_id, service, masked config, status, timestamps; 201 Created for new; 200 for update",
         "Must",
         "Unique constraint on (user_id, service); upsert semantics - updating existing integration reactivates it; config stored as JSONB"),

        ("FR-INT-03",
         "Disconnect and List Integrations",
         "DELETE /api/integrations/{integration_id} and GET /api/integrations",
         "List: query all integrations for current user ordered by created_at desc, return with masked config; Disconnect: verify ownership, delete record",
         "List returns array of IntegrationResponse with masked sensitive values; Delete returns 204 No Content; 404 if not found",
         "Must",
         "Config masking: api_key shows ****{last4}; webhook_url shows scheme://host/****; sensitive keys: api_key, webhook_url, secret, token"),
    ])


# ═══════════════════════════════════════════════════════════════════
# SECTION 4: NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════

def build_section_4(doc):
    add_heading1(doc, "4. Non-Functional Requirements")

    add_body(doc,
        "The following non-functional requirements define the quality attributes and operational constraints "
        "for the GoalCert AutoMind platform."
    )

    add_styled_table(doc,
        ["NFR ID", "Category", "Requirement", "Target"],
        [
            ("NFR-01", "Performance",
             "API response time for CRUD endpoints (agents, workflows, integrations)",
             "P95 < 200ms under 100 concurrent users"),
            ("NFR-02", "Performance",
             "SSE event delivery latency from Celery worker to connected client",
             "< 500ms via Redis pub/sub"),
            ("NFR-03", "Security",
             "Password hashing algorithm and implementation",
             "bcrypt via python bcrypt library; plaintext never logged or stored"),
            ("NFR-04", "Security",
             "JWT token security: algorithm, expiry, and secret management",
             "HS256 algorithm; configurable expiry; secret from environment variable"),
            ("NFR-05", "Security",
             "Multi-tenant data isolation at query layer",
             "All queries scoped to authenticated user_id via SQLAlchemy ORM; no cross-tenant data access possible"),
            ("NFR-06", "Scalability",
             "Celery worker pool horizontal scaling",
             "Workers independently scalable based on queue depth; each task creates own async engine"),
            ("NFR-07", "Scalability",
             "Database capacity for expected 12-month data volume",
             "PostgreSQL with JSONB for flexible schema; UUID primary keys; indexed foreign keys"),
            ("NFR-08", "Availability",
             "Graceful degradation when external services unavailable",
             "OpenAI unavailable: memory summary falls back to template string; Redis unavailable: SSE streaming degrades; Celery unavailable: manual execution queued"),
        ],
        col_widths=[1.5, 2, 6.5, 6],
    )


# ═══════════════════════════════════════════════════════════════════
# SECTION 5: UI REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════

def build_section_5(doc):
    add_heading1(doc, "5. UI Requirements")

    add_body(doc,
        "The following table lists all pages in the GoalCert AutoMind frontend application with their key UI elements "
        "and functional components."
    )

    add_styled_table(doc,
        ["Page", "Route", "Key UI Elements"],
        [
            ("Login", "/login",
             "Email and password form; GoalCert logo; link to signup; error message display; Calibri font throughout"),
            ("Signup", "/signup",
             "Name, email, password form; GoalCert logo; link to login; validation feedback"),
            ("Dashboard", "/",
             "Welcome banner with gradient background and 'Launch an Agent' CTA; StatsCards (total agents, active agents, tasks completed, estimated savings); System Status bar (active count, error count, worker health, total agents, last execution); ActivityFeed panel; Agent Summary with status dots; Agent Grid with AgentCard components; CreateAgentModal"),
            ("Agent Detail", "/agents/:id",
             "Gradient header with agent name, status badge, action buttons (Run Now, Edit Workflow, Chat, Pause/Resume, Delete); Stats cards (success rate, total executions, total cost, avg duration); Execution Timeline (20-square colored grid); Agent Memory panel with list and clear button; LiveExecutionPanel; LiveConsole for active executions; Agent Details (created, updated, schedule, timezone); Execution History table (status, triggered_by, started, duration, cost) with expandable rows; AgentChat sliding panel"),
            ("Workflow Builder", "/agents/:id/builder",
             "Three-panel layout: ComponentSidebar (node type palette), WorkflowCanvas (React Flow with pan/zoom/grid), NodeConfigPanel (type-specific configuration forms); Save and Deploy buttons in canvas toolbar"),
            ("Analytics", "/analytics",
             "6-metric grid (Total Agents, Active Agents, Total Executions, Avg Success Rate, Tasks Completed, Estimated Savings); Cost by Agent bar chart; Executions Over Time weekly heatmap; Top Performers ranked list with success rate bars; Agent Performance Breakdown table (name, type, status, executions, success rate)"),
            ("Reports", "/reports",
             "Page header with date; Executive Summary gradient banner (total agents, executions, cost, success rate); Per-Agent Report Cards showing status, last run, execution count, cost, and latest AI output with markdown rendering"),
            ("Templates", "/templates",
             "Template card grid with gradient headers, type icons, type badges, descriptions, feature checklists, 'Duplicate & Configure' buttons; Build from Scratch dashed-border card; CreateAgentModal with preselected template"),
            ("Integrations", "/integrations",
             "Integration cards for Resend (Email) and Slack; connection status indicator (green dot = connected); config form with masked values; connect/disconnect/save buttons; disconnect confirmation dialog"),
            ("Settings", "/settings",
             "Profile section (name, email); API Configuration (OpenAI key, Resend key); Notifications toggles (execution completed, failed, agent errors, weekly report); Security (change password); Account section showing workspace and data isolation notice"),
            ("Execution Detail", "/executions/:id",
             "Execution metadata (status, trigger, timing, cost); Node log timeline with per-node status, input/output data, timing, and LLM usage; error details for failed nodes"),
        ],
        col_widths=[2.5, 3, 10.5],
    )


# ═══════════════════════════════════════════════════════════════════
# SECTION 6: ERROR HANDLING MATRIX
# ═══════════════════════════════════════════════════════════════════

def build_section_6(doc):
    add_heading1(doc, "6. Error Handling Matrix")

    add_body(doc,
        "The following matrix defines error codes, conditions, user-facing messages, and system actions for all "
        "error scenarios in the AutoMind platform."
    )

    add_styled_table(doc,
        ["Error Code", "Condition", "User Message", "System Action"],
        [
            ("401 Unauthorized",
             "Missing, expired, or invalid JWT token",
             "Session expired. Please sign in again.",
             "Return 401 response; frontend redirects to /login; clear stored token"),
            ("401 Unauthorized",
             "Invalid email or password during login",
             "Invalid email or password",
             "Return generic 401 without revealing if email exists; log failed attempt"),
            ("404 Not Found",
             "Agent, workflow, execution, or integration not found or not owned by user",
             "Resource not found",
             "Return 404 response; no data leakage about other users' resources"),
            ("409 Conflict",
             "Duplicate email during registration or duplicate agent name per user",
             "A user/agent with this name already exists",
             "Return 409 response; rollback transaction; no partial data created"),
            ("400 Bad Request",
             "Workflow has no nodes or is not in active/draft status for execution",
             "Workflow is not ready for execution",
             "Return 400 with specific detail message; no execution record created"),
            ("400 Bad Request",
             "Integration service not in allowed list (resend, slack)",
             "Service must be one of: resend, slack",
             "Return 400 response; reject invalid service type"),
            ("422 Unprocessable Entity",
             "Agent generation from description fails (invalid LLM output)",
             "Failed to generate agent from description",
             "Return 422 response; log generation error; no partial agent created"),
            ("500 Internal Server Error",
             "Unexpected error during agent generation or execution",
             "An unexpected error occurred. Please try again.",
             "Return 500 response; log full exception with traceback; execution marked as failed if applicable"),
            ("SSE Error Event",
             "Error during node execution in running workflow",
             "Error in node [label]: [error message]",
             "Publish error event to Redis channel; update ExecutionNodeLog with error; continue to mark execution as failed"),
            ("SSE Stream End",
             "Execution completed (success or failure)",
             "__STREAM_END__ sentinel",
             "Publish __STREAM_END__ to Redis channel; SSE endpoint closes stream; client disconnects cleanly"),
            ("Chat Error",
             "OpenAI API error during chat streaming",
             "Error message from OpenAI exception",
             "Yield SSE error event with {type: error, content: message}; log exception; stream terminates"),
            ("Memory Generation Failure",
             "OpenAI unavailable for memory summary generation",
             "No user-facing message (background operation)",
             "Fallback to template summary string; log warning; execution not affected"),
            ("Schedule Parse Error",
             "Invalid cron expression (not exactly 5 parts)",
             "Invalid cron expression",
             "Raise ValueError; schedule not registered; agent remains in current state"),
            ("Celery Task Failure",
             "Uncaught exception in Celery workflow execution task",
             "Workflow execution failed",
             "execute_workflow_task catches exception; calls _mark_failed to update execution status; returns failure result dict"),
        ],
        col_widths=[2.5, 4, 3.5, 6],
    )


# ═══════════════════════════════════════════════════════════════════
# SECTION 7: GLOSSARY
# ═══════════════════════════════════════════════════════════════════

def build_section_7(doc):
    add_heading1(doc, "7. Glossary")

    add_styled_table(doc,
        ["Term", "Definition"],
        [
            ("Agent", "An autonomous AI entity configured with a name, type, description, workflow, and optional schedule that executes tasks on behalf of a user"),
            ("Workflow", "A directed acyclic graph (DAG) of interconnected nodes defining the processing pipeline an agent follows during execution"),
            ("Node", "A single processing unit within a workflow; one of seven types: trigger, ai_action, web_search, decision, integration, escalation, code_exec"),
            ("Edge", "A connection between two nodes defining data flow direction; decision nodes have labeled edges (true/false) for branching"),
            ("Execution", "A single run of an agent's workflow, triggered manually or by schedule; records status, timing, cost, and per-node results"),
            ("ExecutionNodeLog", "A detailed record of a single node's execution within a workflow run, including input/output data, timing, LLM usage, and errors"),
            ("Variable Interpolation", "The system that replaces {placeholder} tokens in node configurations with values from upstream node outputs at runtime"),
            ("Agent Memory", "Persistent execution summaries stored per agent that provide historical context for future executions and chat conversations"),
            ("RedBeat", "A Celery beat scheduler backed by Redis that manages cron-based scheduled task entries for agent execution"),
            ("SSE (Server-Sent Events)", "A unidirectional streaming protocol used to push real-time execution progress and chat responses from server to client"),
            ("BFS (Breadth-First Search)", "The graph traversal algorithm used by the execution engine to determine node execution order from the trigger node"),
            ("React Flow", "The frontend graph visualization library used to render the interactive workflow builder canvas with drag-and-drop node management"),
            ("Celery", "A distributed task queue framework used to process agent workflow executions asynchronously in worker processes"),
            ("JWT (JSON Web Token)", "The authentication token format used for API authorization, containing user ID and email claims signed with HS256"),
            ("JSONB", "PostgreSQL binary JSON data type used to store flexible schema data including workflow definitions, execution variables, and integration configs"),
            ("Template", "A pre-built agent workflow definition that users can browse and duplicate to quickly create new agents"),
            ("Integration", "A configured connection to an external service (Resend email or Slack) that agents can use in workflow integration nodes"),
        ],
        col_widths=[3, 13],
    )


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT
    font.size = BODY_SIZE

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_cover_page(doc)
    add_document_control(doc)
    add_toc(doc)

    build_section_1(doc)
    doc.add_page_break()
    build_section_2(doc)
    doc.add_page_break()
    build_section_3(doc)
    doc.add_page_break()
    build_section_4(doc)
    doc.add_page_break()
    build_section_5(doc)
    doc.add_page_break()
    build_section_6(doc)
    doc.add_page_break()
    build_section_7(doc)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(SAVE_PATH), exist_ok=True)
    doc.save(SAVE_PATH)
    print(f"FRD saved to {SAVE_PATH}")


if __name__ == "__main__":
    main()
