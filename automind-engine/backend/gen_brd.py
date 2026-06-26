"""
Generate GoalCert AutoMind Business Requirements Document (BRD) v1.0
Matches Flex Coach BRD format exactly.

Output: /Users/narhen/automind/docs/GoalCert_AutoMind_BRD_v1_0.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants matching Flex Coach BRD exactly ─────────────────────
DARK_NAVY = RGBColor(0x1B, 0x2A, 0x4A)     # #1B2A4A - Heading 1, cover title
BLUE_ACCENT = RGBColor(0x2E, 0x74, 0xB5)    # #2E74B5 - Heading 2, subtitle
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)      # #333333 - Normal text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "Calibri"

# Cell fills
META_LABEL_FILL = "E8EEF4"
HEADER_ROW_FILL = "1B2A4A"
ALT_ROW_FILL = "F0F4F8"

SAVE_PATH = "/Users/narhen/automind/docs/GoalCert_AutoMind_BRD_v1_0.docx"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_styled_table(doc, headers, rows, col_widths_emu=None):
    """Create a table matching the Flex Coach BRD style exactly:
    - Header row: fill #1B2A4A, white bold text
    - Alternating body rows: even rows white, odd rows #F0F4F8
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, HEADER_ROW_FILL)

    # Body rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = BODY_COLOR
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_FILL)

    if col_widths_emu:
        for i, w in enumerate(col_widths_emu):
            for row in table.rows:
                row.cells[i].width = w

    return table


def add_body_paragraph(doc, text):
    """Add body text in Calibri 11pt."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_COLOR
    return p


def add_bullet(doc, text):
    """Add a List Bullet paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    return p


# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
def build_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    # Title: Calibri 32pt bold #1B2A4A centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GoalCert AutoMind")
    run.font.name = FONT_NAME
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = DARK_NAVY

    # Subtitle: Calibri 20pt #2E74B5 centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Business Requirements Document (BRD)")
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.color.rgb = BLUE_ACCENT

    for _ in range(4):
        doc.add_paragraph()

    # ── Metadata table (9x2) ─────────────────────────────────────
    meta_rows = [
        ("Document Title", "GoalCert AutoMind \u2013 Business Requirements Document (BRD)"),
        ("Document Type", "Business Requirements Document"),
        ("Version", "1.0"),
        ("Author", "NextXR Technologies"),
        ("Project", "GoalCert AutoMind"),
        ("Classification", "Confidential"),
        ("Status", "Draft"),
        ("Created", "June 2026"),
        ("Last Updated", "June 2026"),
    ]

    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(meta_rows):
        cell_l = table.rows[i].cells[0]
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_shading(cell_l, META_LABEL_FILL)

        cell_r = table.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_COLOR

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# REVISION HISTORY
# ══════════════════════════════════════════════════════════════════
def build_revision_history(doc):
    doc.add_heading("Document Revision History", level=1)

    make_styled_table(doc,
        ["Version", "Date", "Author", "Change Description", "Reviewed/\nApproved by", "Date Approved"],
        [
            ("1.0", "June 2026", "NextXR Technologies",
             "Initial version \u2013 Complete business requirements for AutoMind Agentic AI Platform",
             "Prem Kumar", "June 2026"),
        ],
    )

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
def build_toc(doc):
    doc.add_heading("Table of Contents", level=1)

    toc_entries = [
        (False, "1.  Executive Summary"),
        (False, "2.  Business Objectives"),
        (False, "3.  Stakeholders"),
        (False, "4.  Scope"),
        (True,  "  4.1  In Scope"),
        (True,  "  4.2  Out of Scope"),
        (False, "5.  Business Context & Current State"),
        (False, "6.  Use Case Diagram"),
        (False, "7.  Business Requirements"),
        (True,  "  7.1  Agent Management"),
        (True,  "  7.2  Workflow Design"),
        (True,  "  7.3  Execution & Monitoring"),
        (True,  "  7.4  Scheduling & Automation"),
        (True,  "  7.5  Chat & Memory"),
        (True,  "  7.6  Analytics & Reporting"),
        (True,  "  7.7  Administration"),
        (False, "8.  Business Rules"),
        (False, "9.  Assumptions & Constraints"),
        (False, "10.  Dependencies & Risks"),
        (False, "11.  Success Criteria"),
        (False, "12.  Glossary"),
    ]

    for is_sub, text in toc_entries:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.color.rgb = DARK_NAVY
        if is_sub:
            run.bold = False
        else:
            run.bold = True

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════
def build_section_1(doc):
    doc.add_heading("1. Executive Summary", level=1)

    add_body_paragraph(doc,
        "GoalCert AutoMind is an agentic AI workflow automation platform that enables organizations to "
        "design, deploy, and manage autonomous AI agents through a visual workflow builder. The platform "
        "empowers non-technical business users to create intelligent agents that autonomously execute "
        "multi-step workflows involving web research, AI-powered analysis, conditional decision-making, "
        "sandboxed code execution, and cross-platform integrations \u2013 all without writing a single "
        "line of code."
    )
    add_body_paragraph(doc,
        "Built on a modern technology stack comprising FastAPI (Python 3.12 async backend), React with "
        "TypeScript (frontend), Celery with Redis (distributed task execution), and PostgreSQL with "
        "asyncpg (persistent storage), AutoMind provides enterprise-grade capabilities including "
        "real-time execution monitoring via Server-Sent Events (SSE) over Redis pub/sub, persistent "
        "agent memory that enables cross-run learning, conversational chat interfaces for interactive "
        "agent debriefing, cron-based scheduled automation via RedBeat, per-execution cost tracking "
        "with LLM token accounting across 8 supported models, and a consolidated analytics dashboard."
    )
    add_body_paragraph(doc,
        "The platform supports seven distinct workflow node types \u2013 Trigger, AI Action, Web Search, "
        "Decision, Integration, Escalation, and Code Execution \u2013 which can be composed into "
        "sophisticated automation workflows using a drag-and-drop React Flow canvas. Agents can be "
        "created manually, instantiated from pre-built templates, or generated entirely by AI from a "
        "natural language description. The execution engine performs BFS (Breadth-First Search) graph "
        "traversal, supports conditional branching via Decision nodes, and provides variable "
        "interpolation across all node configurations using {placeholder} syntax with nested path "
        "resolution."
    )
    add_body_paragraph(doc,
        "AutoMind integrates with external services including OpenAI (GPT-4o, GPT-4o-mini, GPT-4.1, "
        "GPT-4.1-mini, GPT-4.1-nano, o3-mini), Anthropic (Claude Sonnet 4, Claude Haiku 4.5), "
        "Resend (transactional email), Slack (incoming webhooks), and DuckDuckGo (web search). "
        "The platform features JWT-based authentication with bcrypt password hashing, per-user data "
        "isolation, a template library for rapid agent creation, and a comprehensive settings "
        "interface for API key and integration management."
    )
    add_body_paragraph(doc,
        "This Business Requirements Document captures the functional and non-functional requirements, "
        "business objectives, scope boundaries, use cases, business rules, and success criteria for "
        "GoalCert AutoMind version 1.0."
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 2: BUSINESS OBJECTIVES
# ══════════════════════════════════════════════════════════════════
def build_section_2(doc):
    doc.add_heading("2. Business Objectives", level=1)

    add_body_paragraph(doc,
        "The following table outlines the primary business objectives that GoalCert AutoMind aims "
        "to achieve, along with priority levels and measurable success criteria."
    )

    make_styled_table(doc,
        ["Obj ID", "Business Objective", "Priority", "Success Measure"],
        [
            ("BO-01",
             "Reduce manual repetitive task time by automating knowledge work workflows with autonomous AI agents",
             "High",
             "40% reduction in time spent on automated tasks compared to manual execution within 6 months of adoption"),
            ("BO-02",
             "Enable non-technical users to create and deploy AI agents without coding through a visual drag-and-drop workflow builder powered by React Flow",
             "High",
             "80% of agents created by users with no software development background; average time from concept to deployed agent under 30 minutes"),
            ("BO-03",
             "Provide real-time execution monitoring and full audit trails for all agent operations via SSE streaming and per-node execution logging",
             "High",
             "Live SSE console available for 100% of executions; node-level status, duration, input/output data, and LLM cost visible within 2 seconds of completion"),
            ("BO-04",
             "Ensure cost-effective AI operations with per-execution cost tracking using model-specific pricing tables across 8 supported LLM models (OpenAI and Anthropic)",
             "High",
             "Average execution cost below $0.03 per run; cost dashboard accurate to within 1% of actual LLM provider billing; costs tracked with 6-digit decimal precision"),
            ("BO-05",
             "Support fully autonomous scheduled agent execution via cron-based scheduling using RedBeat (Redis-backed Celery Beat) with timezone support",
             "Medium",
             "95% scheduled execution reliability (on-time within 60-second window); support for minute, hourly, daily, and weekly schedules with 5-field cron syntax"),
            ("BO-06",
             "Enable rapid agent creation through a pre-built template library and AI-powered natural language agent generation using GPT-4o-mini",
             "Medium",
             "Template marketplace with pre-built workflows for sales, marketing, support, and custom agent types; 70% of AI-generated agents deployable without manual edits"),
            ("BO-07",
             "Build persistent agent memory that enables cross-run learning, providing context-aware AI responses and data-grounded agent chat interactions",
             "Medium",
             "Agent chat responses reference specific data from past executions; memory summaries auto-generated after each execution (LLM-enhanced when API key available)"),
            ("BO-08",
             "Integrate with existing business tools (email via Resend, Slack via webhooks) to embed AI agents into operational communication workflows",
             "Medium",
             "Native integration nodes for Resend email and Slack messaging; integration setup completed within 5 minutes per service; credential masking in API responses"),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 3: STAKEHOLDERS
# ══════════════════════════════════════════════════════════════════
def build_section_3(doc):
    doc.add_heading("3. Stakeholders", level=1)

    add_body_paragraph(doc,
        "The following stakeholders have been identified for the GoalCert AutoMind project."
    )

    make_styled_table(doc,
        ["Stakeholder", "Role", "Interest / Responsibility"],
        [
            ("Product Owner (Prem Kumar)",
             "Project Sponsor / Decision Maker",
             "Defines product vision and roadmap; approves feature scope and prioritization; represents business stakeholder interests; manages GoalCert product suite strategy and alignment"),
            ("Platform Engineers (Development Team)",
             "Technical Leads \u2013 Full Stack",
             "Design and implement backend architecture (FastAPI, Celery, PostgreSQL); build frontend React application with workflow builder; develop execution engine, memory system, and scheduling service; manage deployment infrastructure"),
            ("Business Users",
             "Primary Platform Users (Department Heads / Operations Managers)",
             "Primary end users who create and manage agents via the visual workflow builder; define automation requirements; evaluate platform ROI and productivity gains; provide feedback on usability and feature requests"),
            ("AI Agents",
             "Automated Actors",
             "Autonomous software entities that execute assigned workflows on schedule or on demand; consume LLM API tokens; produce reports, summaries, and outputs; interact with external integrations (email, Slack)"),
            ("System Administrators",
             "Platform Operations",
             "Manage user accounts and API key configurations; monitor system health, Redis connectivity, and Celery worker status; configure integrations at the organization level; handle deployment and scaling operations"),
            ("Integration Partners (OpenAI, Anthropic, Resend, Slack)",
             "Technology Partners",
             "Provide external API services consumed by the platform; API compatibility, rate limit compliance, and SLA adherence directly impact platform functionality; usage-based billing accuracy affects cost tracking"),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 4: SCOPE
# ══════════════════════════════════════════════════════════════════
def build_section_4(doc):
    doc.add_heading("4. Scope", level=1)

    doc.add_heading("4.1 In Scope", level=2)
    in_scope = [
        "Agent creation via three methods: manual configuration with name/type/description, template-based instantiation from pre-built workflow blueprints, and AI-generated agents from natural language descriptions via GPT-4o-mini",
        "Visual workflow builder with drag-and-drop React Flow canvas supporting 7 node types: Trigger (workflow entry point with schedule config), AI Action (LLM call with prompt templating and model selection), Web Search (DuckDuckGo HTML scraping), Decision (conditional branching with comparison operators), Integration (Resend email and Slack messaging), Escalation (human notification via email), and Code Execution (sandboxed Python with forbidden import restrictions)",
        "Distributed workflow execution engine using Celery workers with BFS graph traversal, variable interpolation using {placeholder} syntax with nested path resolution (e.g., {leads.0.name}), and conditional branching via Decision node true/false edge routing",
        "Real-time execution monitoring via Redis pub/sub and Server-Sent Events (SSE) streaming, publishing node-level progress events (node entry, status, output, completion) with a __STREAM_END__ sentinel signaling execution completion",
        "Cron-based agent scheduling using RedBeat (Redis-backed Celery Beat) with 5-field cron expression support, configurable timezone (default UTC), and dynamic schedule management (create on deploy, remove on pause/delete, re-create on resume)",
        "Persistent agent memory service that auto-generates LLM-summarized execution records (with template fallback when API unavailable), stores summaries and key_outputs as JSONB, and injects cross-run context into AI Action node prompts and agent chat conversations",
        "Agent chat interface with SSE-streamed conversational responses powered by GPT-4o-mini, incorporating agent metadata (name, type, status, schedule) and up to 15 recent memory summaries as context, with 10-message conversation history support",
        "Per-execution and per-node cost tracking with model-specific pricing tables supporting 8 LLM models across OpenAI (GPT-4o, GPT-4o-mini, GPT-4.1, GPT-4.1-mini, GPT-4.1-nano, o3-mini) and Anthropic (Claude Sonnet 4, Claude Haiku 4.5), with 6-digit decimal precision",
        "External integrations: transactional email via Resend API, Slack notifications via incoming webhooks, web search via DuckDuckGo HTML scraping, and sandboxed Python code execution with restricted imports (blocking os, sys, subprocess, shutil, socket, requests, urllib, eval, exec, open)",
    ]
    for item in in_scope:
        add_bullet(doc, item)

    doc.add_heading("4.2 Out of Scope", level=2)
    out_scope = [
        "Multi-tenant SaaS with organization-level isolation, team workspaces, and role-based access control beyond single-user ownership \u2013 planned for future release",
        "Webhook-triggered agent execution (triggered_by='webhook' defined in types but not implemented in v1.0 API endpoints)",
        "Custom LLM fine-tuning or model training infrastructure \u2013 platform consumes pre-trained models via provider APIs exclusively",
        "Mobile native application (iOS / Android) \u2013 responsive web interface only for v1.0",
        "Marketplace for third-party agent plugins, custom node types, or node type development SDK \u2013 planned for future release",
        "Container-level code execution isolation (Docker/Firecracker) \u2013 v1.0 uses pattern-based import restriction in subprocess; OS-level sandboxing deferred",
    ]
    for item in out_scope:
        add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════
# SECTION 5: BUSINESS CONTEXT & CURRENT STATE
# ══════════════════════════════════════════════════════════════════
def build_section_5(doc):
    doc.add_heading("5. Business Context & Current State", level=1)

    add_body_paragraph(doc,
        "The AutoMind platform has been developed as a full-stack application with a FastAPI backend "
        "(routers: agents, workflows, executions, chat, memory, dashboard, integrations, templates, auth) "
        "and a React/TypeScript frontend (pages: Dashboard, WorkflowBuilder, AgentDetail, ExecutionDetail, "
        "Analytics, Reports, Templates, Integrations, Settings, Login, Signup). The following table "
        "summarizes the current implementation status of each major functional area as of June 2026."
    )

    make_styled_table(doc,
        ["Area", "Current State", "Status"],
        [
            ("Agent Management",
             "Full CRUD via /api/agents endpoints. Supports manual creation (AgentCreate schema with name, type, description, template_id), template-based instantiation, and AI-generated agents via /api/agents/generate endpoint using GPT-4o-mini. Agent types: sales, marketing, support, custom. Statuses: draft, active, paused, error. Per-agent statistics (total_executions, success_rate) computed via SQL aggregation.",
             "Implemented"),
            ("Visual Workflow Builder",
             "React Flow-based drag-and-drop canvas (WorkflowBuilderPage) with ComponentSidebar listing 7 node types, WorkflowCanvas for graph editing, and NodeConfigPanel for selected node configuration. Workflow definitions stored as JSONB (nodes, edges, viewport). Save via PUT /api/agents/{id}/workflow, deploy via POST /api/agents/{id}/workflow/deploy.",
             "Implemented"),
            ("Execution Engine",
             "WorkflowExecutor class in app/engine/executor.py implements BFS graph traversal via deque. 7 node executors registered: TriggerNodeExecutor, AIActionNodeExecutor, IntegrationNodeExecutor, DecisionNodeExecutor, EscalationNodeExecutor, WebSearchNodeExecutor, CodeExecNodeExecutor. Variable interpolation via app/engine/variables.py supports {placeholder} syntax with nested path resolution. Graph utilities in app/engine/graph.py provide parse_workflow, topological_sort, and get_next_nodes functions.",
             "Implemented"),
            ("Scheduling System",
             "RedBeat scheduler (app/services/scheduler_service.py) manages cron schedules. schedule_agent creates RedBeatSchedulerEntry, unschedule_agent removes it. Scheduled executions via execute_workflow_scheduled_task Celery task which validates agent status before execution. Timezone-aware with UTC default.",
             "Implemented"),
            ("LLM Integration",
             "AIActionNodeExecutor supports multi-provider LLM calls. OpenAI models: gpt-4o ($2.50/$10.00 per M tokens), gpt-4o-mini ($0.15/$0.60), gpt-4.1 ($2.00/$8.00), gpt-4.1-mini ($0.40/$1.60), gpt-4.1-nano ($0.10/$0.40), o3-mini ($1.10/$4.40). Anthropic models: claude-sonnet-4-20250514 ($3.00/$15.00), claude-haiku-4-5-20251001 ($0.80/$4.00). Mock fallback when API keys not configured.",
             "Implemented"),
            ("Agent Memory",
             "AgentMemory model stores execution summaries in agent_memory table. save_execution_memory generates LLM-enhanced summaries via GPT-4o-mini (template fallback without API key). get_agent_context retrieves up to 10/15 recent memories for prompt injection. Memory API (/api/agents/{id}/memory) supports paginated listing and bulk deletion.",
             "Implemented"),
            ("Chat Interface",
             "SSE-streamed chat endpoint at POST /api/agents/{id}/chat. System prompt includes agent identity (name, type, description, status, schedule) and memory context (up to 15 summaries). Uses GPT-4o-mini with max_tokens=1024, temperature=0.7. Conversation history (last 10 messages) maintained per session.",
             "Implemented"),
            ("Analytics & Reports",
             "Dashboard API (/api/dashboard/stats, /api/dashboard/activity) provides aggregate metrics: total/active agents, tasks completed, estimated savings (tasks * $25), avg response time. Frontend AnalyticsPage renders cost-by-agent breakdown, top performers by success rate, and activity heatmap by day of week. ReportsPage displays per-agent AI output with markdown formatting.",
             "Implemented"),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 6: USE CASE DIAGRAM
# ══════════════════════════════════════════════════════════════════
def build_section_6(doc):
    doc.add_heading("6. Use Case Diagram", level=1)

    add_body_paragraph(doc,
        "The following Use Case Catalog enumerates the primary interactions between actors and the "
        "GoalCert AutoMind system. Each use case is assigned a unique identifier, mapped to its primary "
        "actor, and prioritized according to business impact."
    )

    doc.add_heading("Use Case Catalog", level=2)

    make_styled_table(doc,
        ["UC ID", "Use Case Name", "Primary Actor", "Priority"],
        [
            ("UC-01", "Register New User Account (POST /api/auth/register)", "Business User", "High"),
            ("UC-02", "Authenticate and Manage Session (POST /api/auth/login, GET /api/auth/me)", "Business User", "High"),
            ("UC-03", "Create Agent Manually (POST /api/agents)", "Business User", "High"),
            ("UC-04", "Create Agent from Pre-built Template", "Business User", "High"),
            ("UC-05", "Generate Agent from Natural Language Description (POST /api/agents/generate)", "Business User", "High"),
            ("UC-06", "Design Workflow via Visual Drag-and-Drop Builder (React Flow Canvas)", "Business User", "High"),
            ("UC-07", "Configure Workflow Node Parameters (Trigger, AI Action, Decision, Integration, Web Search, Code Exec, Escalation)", "Business User", "High"),
            ("UC-08", "Save and Deploy Workflow (PUT /api/agents/{id}/workflow, POST .../deploy)", "Business User", "High"),
            ("UC-09", "Trigger Manual Agent Execution (POST /api/agents/{id}/execute)", "Business User", "High"),
            ("UC-10", "Monitor Execution in Real-Time via SSE (GET /api/executions/{id}/stream)", "Business User", "High"),
            ("UC-11", "View Execution History and Per-Node Logs (GET /api/agents/{id}/executions, GET /api/executions/{id})", "Business User", "Medium"),
            ("UC-12", "Configure Scheduled Execution with Cron Expression", "Business User", "Medium"),
            ("UC-13", "Pause and Resume Agent (POST /api/agents/{id}/pause, .../resume)", "Business User", "Medium"),
            ("UC-14", "Chat with Agent Using Memory Context (POST /api/agents/{id}/chat)", "Business User", "Medium"),
            ("UC-15", "View Analytics Dashboard (GET /api/dashboard/stats, .../activity)", "Business User", "Medium"),
            ("UC-16", "Connect External Integration \u2013 Resend / Slack (POST /api/integrations)", "System Administrator", "Medium"),
            ("UC-17", "Browse and Instantiate Agent Templates (GET /api/templates)", "Business User", "Medium"),
            ("UC-18", "View and Clear Agent Memory (GET/DELETE /api/agents/{id}/memory)", "Business User", "Low"),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 7: BUSINESS REQUIREMENTS
# ══════════════════════════════════════════════════════════════════
def build_section_7(doc):
    doc.add_heading("7. Business Requirements", level=1)

    add_body_paragraph(doc,
        "This section details the business requirements organised by functional area. Each requirement "
        "includes a unique identifier, description, priority level, and acceptance criteria derived "
        "from the implemented codebase."
    )

    # ── 7.1 Agent Management ─────────────────────────────────────
    doc.add_heading("7.1 Agent Management", level=2)
    make_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Acceptance Criteria"],
        [
            ("BR-AM-01",
             "System shall allow users to create new agents by specifying a name (max 255 chars), type (sales/marketing/support/custom), and optional description. Agent names must be unique per user, enforced by database constraint uq_agent_user_name.",
             "High",
             "Agent created with draft status; duplicate name returns HTTP 409 Conflict; agent persisted with auto-generated UUID and user_id from JWT; empty workflow definition (nodes=[], edges=[], viewport) created automatically"),
            ("BR-AM-02",
             "System shall support agent creation from pre-built templates stored in the agent_templates table, where the template's workflow_definition (JSONB with nodes, edges, icon, color, features) is copied into the new agent's workflow.",
             "High",
             "Template selection via template_id in AgentCreate schema populates workflow canvas with template nodes and edges; template metadata (icon, color, features) available for UI display; agent functional immediately after template instantiation"),
            ("BR-AM-03",
             "System shall enable AI-powered agent generation via POST /api/agents/generate, accepting a natural language description and producing a complete agent with workflow definition, name, type, description, and optional cron schedule using GPT-4o-mini.",
             "High",
             "Generated JSON validated for required fields (workflow, nodes); agent_type constrained to allowed values with fallback to 'custom'; workflow includes trigger node, action nodes, correctly wired edges, and React Flow viewport; schedule extracted if provided"),
            ("BR-AM-04",
             "System shall support full agent lifecycle management: update (PATCH with partial fields), pause (sets status to 'paused' and unschedules RedBeat entry), resume (sets status to 'active' and re-creates schedule if cron configured), and delete (cascades to workflow, executions, node logs, and memories via ON DELETE CASCADE).",
             "High",
             "Pause calls unschedule_agent before status change; resume calls schedule_agent if agent.schedule_cron is set; delete calls unschedule_agent then removes agent (cascade deletes all children); all operations verify user ownership via user_id check"),
            ("BR-AM-05",
             "System shall compute and expose per-agent statistics: total_executions (COUNT of execution records) and success_rate (AVG of CASE(success=1.0, failed=0.0) * 100, rounded to 1 decimal place) on both list and detail API responses.",
             "Medium",
             "Statistics computed via SQL aggregation on each request; success_rate is None when no executions exist; values included in AgentResponse schema; list endpoint returns stats for all user agents"),
        ],
    )

    # ── 7.2 Workflow Design ──────────────────────────────────────
    doc.add_heading("7.2 Workflow Design", level=2)
    make_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Acceptance Criteria"],
        [
            ("BR-WD-01",
             "System shall provide a visual drag-and-drop workflow builder using React Flow (via @xyflow/react) with three panels: ComponentSidebar (node type palette), WorkflowCanvas (graph editor with useNodesState/useEdgesState), and NodeConfigPanel (configuration form for selected node).",
             "High",
             "All 7 node types available in sidebar; nodes draggable onto canvas; edges connectable between node ports; selected node opens config panel on right; builderStore tracks selectedNodeId and dirty state"),
            ("BR-WD-02",
             "System shall support 7 node types with dedicated configuration schemas: Trigger (frequency, cron, timezone), AI Action (prompt, model, max_tokens, temperature, output_variable, system_prompt), Web Search (query, max_results, output_variable), Decision (left_operand, operator [==, !=, >, <, >=, <=, contains], right_operand), Integration (service [email/slack], action, recipients, subject, body, message), Escalation (recipient_email, message_template, subject), and Code Execution (code, timeout [max 30s], output_variable).",
             "High",
             "Each node type renders correctly on canvas; config panel shows type-specific fields; changes persist on save; all node types execute correctly in the workflow engine with their respective NodeExecutor implementations"),
            ("BR-WD-03",
             "System shall support variable interpolation using {variable_name} syntax across all node configurations via the interpolation engine (app/engine/variables.py), including nested path resolution (e.g., {leads.0.name}), type-preserving substitution for single-placeholder strings, and recursive interpolation into dicts and lists.",
             "High",
             "Variables from upstream nodes (output_variable fields) resolved in downstream node prompts, email bodies, code templates, and search queries at execution time; unresolved placeholders left as-is; full string {var} returns raw value (preserving int/list/dict types)"),
            ("BR-WD-04",
             "System shall persist workflow definitions as JSONB (nodes, edges, viewport) via PUT /api/agents/{id}/workflow, and support deployment via POST /api/agents/{id}/workflow/deploy which sets workflow.status='active', workflow.deployed_at=now(), agent.status='active', and registers cron schedule with RedBeat if agent.schedule_cron is configured.",
             "High",
             "Workflow save updates definition JSONB; deploy endpoint transitions both workflow and agent status; deployed_at timestamp recorded; schedule_agent called with agent's cron and timezone on deploy; workflow must be active or draft to execute"),
        ],
    )

    # ── 7.3 Execution & Monitoring ───────────────────────────────
    doc.add_heading("7.3 Execution & Monitoring", level=2)
    make_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Acceptance Criteria"],
        [
            ("BR-EM-01",
             "System shall execute workflows asynchronously via Celery workers (execute_workflow task) using a WorkflowExecutor that performs BFS graph traversal via collections.deque, starting from the trigger node (first node with no incoming edges, preferring type='trigger'), and following edges respecting Decision node branch routing (true/false via edge labels/sourceHandles).",
             "High",
             "Execution creates Execution record with status='pending', triggered_by='manual'/'schedule'; Celery task processes graph in BFS order; Decision nodes route to correct branch based on condition evaluation; all reachable nodes executed; execution completes with 'success' or 'failed' status"),
            ("BR-EM-02",
             "System shall provide real-time execution streaming via GET /api/executions/{id}/stream SSE endpoint, publishing JSON events through Redis pub/sub channel 'execution:{id}:logs' with fields: timestamp (ISO 8601), message (human-readable), node_id (nullable), and status (info/running/success/error/done). A '__STREAM_END__' sentinel message signals execution completion.",
             "High",
             "SSE endpoint returns EventSourceResponse; events published by _publish_log method in WorkflowExecutor; each node start and completion triggers an event; stream closes on __STREAM_END__; no auth required on SSE endpoint (UUID unguessability)"),
            ("BR-EM-03",
             "System shall create ExecutionNodeLog records for each node execution, capturing: node_id, node_type, node_label, status (pending/running/success/failed/skipped), started_at, ended_at, duration_ms, input_data (JSONB), output_data (JSONB), error_message, and llm_usage (JSONB with model, input_tokens, output_tokens, cost). Logs retrievable via GET /api/executions/{id}/logs ordered chronologically.",
             "High",
             "Every executed node produces an ExecutionNodeLog row; node_log created with status='running' at start, updated with final status and output at completion; unknown node types logged as 'skipped'; logs include all specified fields"),
            ("BR-EM-04",
             "System shall track LLM token usage and cost per execution using a model-specific pricing table (_COST_TABLE) covering 8 models. Cost formula: (input_tokens * input_rate + output_tokens * output_rate) / 1,000,000. Total cost accumulated in WorkflowExecutor.total_cost and stored as Numeric(10,6) on the Execution record.",
             "High",
             "AI Action nodes return llm_usage dict with model, input_tokens, output_tokens, cost; total_cost is sum of all AI node costs; unknown models use default pricing ($2.50/$10.00 per M tokens); cost displayed in execution detail and analytics"),
            ("BR-EM-05",
             "System shall update Agent.last_execution_at upon execution completion, persist accumulated workflow variables on Execution.variables (JSONB), and auto-generate an AgentMemory record with LLM-summarized execution summary, key_outputs (non-internal variables), and execution_id reference.",
             "Medium",
             "Agent timestamp updated in _finish_execution; variables contain all accumulated state from node outputs; memory record created with summary text and key_outputs; internal variables (triggered_by, branch, condition_result, _prefixed) excluded from key_outputs"),
        ],
    )

    # ── 7.4 Scheduling & Automation ──────────────────────────────
    doc.add_heading("7.4 Scheduling & Automation", level=2)
    make_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Acceptance Criteria"],
        [
            ("BR-SA-01",
             "System shall support scheduled agent execution using 5-field cron expressions (minute, hour, day_of_month, month_of_year, day_of_week) parsed by scheduler_service.parse_cron into celery.schedules.crontab objects, managed by RedBeatSchedulerEntry with task name 'execute_workflow_scheduled'.",
             "High",
             "Cron expression validated for 5 parts; ValueError raised for invalid format; RedBeatSchedulerEntry created with name='agent:{agent_id}'; scheduled task fires within 60 seconds of configured time"),
            ("BR-SA-02",
             "System shall dynamically manage schedules at runtime: schedule_agent creates/replaces RedBeat entry (unschedule first to avoid duplicates), unschedule_agent removes entry by key 'redbeat:agent:{id}' (no-op if not found). Schedule lifecycle: created on workflow deploy, removed on agent pause/delete, re-created on agent resume.",
             "High",
             "Deploy with schedule_cron calls schedule_agent; pause calls unschedule_agent; resume with schedule_cron calls schedule_agent; delete calls unschedule_agent before DB deletion; no orphaned schedules after any lifecycle operation"),
            ("BR-SA-03",
             "Scheduled executions shall be processed by execute_workflow_scheduled_task Celery task which: looks up agent by ID, verifies agent.status == 'active' (skips if paused/deleted), retrieves current workflow definition, creates Execution record with triggered_by='schedule', and runs WorkflowExecutor with a dedicated async DB engine (asyncpg pool_size=5).",
             "High",
             "Scheduled task creates fresh async engine for each run; agent status checked before execution; inactive agents skipped with log message; execution record created with triggered_by='schedule'; engine disposed after execution completes or fails"),
        ],
    )

    # ── 7.5 Chat & Memory ────────────────────────────────────────
    doc.add_heading("7.5 Chat & Memory", level=2)
    make_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Acceptance Criteria"],
        [
            ("BR-CM-01",
             "System shall provide an SSE-streamed chat endpoint (POST /api/agents/{id}/chat) for each agent, using GPT-4o-mini (max_tokens=1024, temperature=0.7) with a system prompt incorporating agent identity (name, type, description, status, schedule) and memory context from get_agent_context (up to 15 recent summaries). Chat accepts message and history (list of role/content dicts), using last 10 history messages.",
             "High",
             "Chat endpoint streams token-by-token SSE events with type='token'/'done'/'error'; system prompt includes agent metadata and formatted memory context; conversation history limited to last 10 messages; requires OPENAI_API_KEY; returns fallback message without key"),
            ("BR-CM-02",
             "System shall maintain persistent agent memory via save_execution_memory (app/services/memory_service.py), generating summaries after each execution. When OPENAI_API_KEY is available and not placeholder, summary generated by GPT-4o-mini (max_tokens=200, temperature=0.3) from SUMMARIZE_PROMPT template; otherwise falls back to template string 'Execution {status} in {duration}ms with {nodes} nodes'.",
             "High",
             "AgentMemory record created with agent_id, execution_id, summary, key_outputs (JSONB), memory_type='execution_summary'; LLM summary captures what was accomplished, key findings, and data points; key_outputs exclude internal variables"),
            ("BR-CM-03",
             "System shall provide paginated memory listing (GET /api/agents/{id}/memory with limit/offset, returns MemoryListResponse with memories and total count) and bulk memory deletion (DELETE /api/agents/{id}/memory removes all AgentMemory records for the agent). Both endpoints verify agent ownership.",
             "Medium",
             "GET returns paginated list ordered by created_at descending; total count computed via SQL COUNT; DELETE removes all records and returns 204; both endpoints return 404 if agent not found or not owned by current user"),
        ],
    )

    # ── 7.6 Analytics & Reporting ────────────────────────────────
    doc.add_heading("7.6 Analytics & Reporting", level=2)
    make_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Acceptance Criteria"],
        [
            ("BR-AR-01",
             "System shall provide GET /api/dashboard/stats returning DashboardStats: total_agents, active_agents (COUNT WHERE status='active'), tasks_completed (COUNT WHERE execution.status='success'), estimated_savings (tasks_completed * 0.5 * 50), and avg_response_time (AVG of execution.duration_ms, rounded to 1 decimal). All queries scoped to current user.",
             "High",
             "Stats computed from live database queries; estimated_savings uses fixed rate of $25 per task; avg_response_time is None when no executions exist; all values refresh on each API call"),
            ("BR-AR-02",
             "System shall provide GET /api/dashboard/activity returning the 20 most recent executions as ActivityEvent objects with: execution_id, agent_id, agent_name, agent_type, status, triggered_by, started_at, ended_at, duration_ms, total_cost, and created_at. Results ordered by created_at descending, scoped to current user.",
             "Medium",
             "Activity feed joins Execution and Agent tables; total_cost cast to float from Decimal; limited to 20 entries; each event includes both execution and agent metadata"),
            ("BR-AR-03",
             "Frontend AnalyticsPage shall render: (1) metrics cards showing total agents, active agents, total executions, avg success rate, tasks completed, and estimated savings; (2) cost-by-agent horizontal bar chart aggregated from execution activity data; (3) top performers sorted by success rate; (4) activity heatmap by day of week computed from execution timestamps.",
             "Medium",
             "Analytics page fetches data via useAgents, useDashboardStats, and useDashboardActivity hooks; cost breakdown computed from activity data grouped by agent_name; heatmap maps JS getDay() to Mon-Sun labels; all charts use real execution data with fallback to agent list when no activity exists"),
            ("BR-AR-04",
             "Frontend ReportsPage shall display per-agent report cards with agent status, last execution timestamp, execution count, success rate, cumulative cost, and the rendered AI output from the latest execution's ai_action node output_data with markdown formatting support.",
             "Medium",
             "Reports page fetches latest execution with node logs for each agent; AI output extracted from ai_action/web_search node output_data; markdown rendered with headers, bold, and paragraphs; agents without executions show 'No executions yet' placeholder"),
        ],
    )

    # ── 7.7 Administration ───────────────────────────────────────
    doc.add_heading("7.7 Administration", level=2)
    make_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Acceptance Criteria"],
        [
            ("BR-AD-01",
             "System shall provide JWT-based authentication via /api/auth endpoints: register (POST with email, password, name; returns AuthResponse with JWT token and UserResponse), login (POST with email, password; validates credentials with bcrypt verify_password), and session check (GET /api/auth/me returns current user). Passwords hashed with bcrypt; tokens expire after JWT_EXPIRY_HOURS (default 24 hours) configured via environment variable.",
             "High",
             "Registration creates User with hashed password and returns JWT; duplicate email returns 409; login validates credentials and returns JWT; invalid credentials return 401; /me endpoint extracts user from JWT; expired tokens rejected with 401"),
            ("BR-AD-02",
             "System shall enforce per-user data isolation: all agent, workflow, execution, and memory queries filter by current_user.id extracted from JWT. Integration CRUD scoped by user_id. No global list endpoints expose cross-user data.",
             "High",
             "All API endpoints filter by user_id; agent creation sets user_id from JWT context; accessing another user's resources returns 404 (not 403, to prevent enumeration); integration uniqueness enforced by uq_integration_user_service constraint"),
            ("BR-AD-03",
             "System shall support external service integration management via /api/integrations: list (GET, returns masked credentials), connect/upsert (POST with service name from allowlist [resend, slack] and config JSONB), and disconnect (DELETE by integration_id). Sensitive config fields (api_key, webhook_url, secret, token) masked in responses using _mask_config: tokens show '****' + last 4 chars, webhook URLs show scheme://host/****.",
             "Medium",
             "POST creates or updates integration for user+service pair (upsert behavior); service validated against ALLOWED_SERVICES set; unsupported service returns 400; credentials stored as JSONB; GET response masks sensitive fields; DELETE removes integration record and returns 204"),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 8: BUSINESS RULES
# ══════════════════════════════════════════════════════════════════
def build_section_8(doc):
    doc.add_heading("8. Business Rules", level=1)

    add_body_paragraph(doc,
        "The following business rules govern the behaviour of the GoalCert AutoMind platform, "
        "derived from the implemented codebase constraints and design decisions."
    )

    make_styled_table(doc,
        ["Rule ID", "Category", "Business Rule"],
        [
            ("BR-01", "Agent Naming",
             "Each agent name must be unique per user, enforced by database constraint uq_agent_user_name. Agent names are limited to 255 characters (String(255)). Attempting to create a duplicate returns HTTP 409 Conflict with message 'An agent named {name} already exists'."),
            ("BR-02", "Agent-Workflow Relationship",
             "Each agent has exactly one workflow (one-to-one relationship enforced by unique constraint on Workflow.agent_id). Workflow created automatically on agent creation with empty definition. Workflow cascades on agent deletion (ON DELETE CASCADE)."),
            ("BR-03", "Workflow Validation",
             "Workflow execution requires at least one node in the definition. Empty workflows (no nodes or no 'nodes' key) are rejected with HTTP 400 'Workflow has no nodes defined'. Workflow must have status 'active' or 'draft' to be executed."),
            ("BR-04", "Code Execution Security",
             "Code Execution nodes prohibit 12 forbidden patterns: import os, import sys, import subprocess, import shutil, __import__, eval(, exec(, open(, import socket, import requests, import urllib. Code runs in subprocess with configurable timeout (default 10s, capped at 30s). Only simple-type local variables (str, int, float, bool, list, dict) captured as output."),
            ("BR-05", "Cost Tracking Precision",
             "Execution costs tracked as Numeric(10,6) with 6-digit decimal precision. Cost per AI node: (input_tokens * input_rate + output_tokens * output_rate) / 1,000,000. Unknown models fall back to default pricing (input: $2.50/M, output: $10.00/M). Only ai_action, web_search, and code_exec nodes have llm_usage tracking."),
            ("BR-06", "Integration Allowlist",
             "Integration services restricted to ALLOWED_SERVICES set: {'resend', 'slack'}. Requests for unsupported services rejected with HTTP 400. Per-user service uniqueness enforced by uq_integration_user_service database constraint (upsert behavior on duplicate)."),
            ("BR-07", "Cascade Deletion",
             "Deleting an agent cascades to: workflow (ON DELETE CASCADE), all executions (ON DELETE CASCADE), all execution node logs (ON DELETE CASCADE via execution), and agent memory entries (ON DELETE CASCADE). RedBeat schedule entry removed via unschedule_agent before database deletion."),
            ("BR-08", "Memory Retention",
             "Agent memory summaries retained indefinitely unless explicitly cleared via DELETE /api/agents/{id}/memory. Memory context injection limited to most recent entries (10 for workflow execution, 15 for chat). Execution memory deletion uses ON DELETE SET NULL for execution_id reference (memory persists if individual execution deleted)."),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 9: ASSUMPTIONS & CONSTRAINTS
# ══════════════════════════════════════════════════════════════════
def build_section_9(doc):
    doc.add_heading("9. Assumptions & Constraints", level=1)

    doc.add_heading("Assumptions", level=2)
    assumptions = [
        "Users have access to at least one LLM API key (OpenAI or Anthropic) for AI Action nodes to produce meaningful results. The platform degrades gracefully without keys: AI nodes return mock responses, agent generation fails with a clear error, memory summaries use template fallback, and chat returns a configuration reminder.",
        "The deployment environment provides PostgreSQL 14+ with gen_random_uuid() support, Redis 6+ for Celery task queue and pub/sub, and a Python 3.12+ runtime with asyncpg driver for async database access.",
        "Business users interacting with the visual workflow builder are comfortable with basic logical concepts (if-then conditions, variable references using {placeholder} syntax) even if they cannot write code.",
        "LLM API pricing from OpenAI and Anthropic will remain stable or decrease during the initial deployment period. The cost tracking model uses a hardcoded _COST_TABLE that must be updated manually when pricing changes.",
        "DuckDuckGo's HTML search endpoint (html.duckduckgo.com/html/) will maintain a stable HTML structure for web search node scraping. Changes to DuckDuckGo's response format would require updates to the WebSearchNodeExecutor regex patterns.",
        "External integration providers (Resend for email, Slack incoming webhooks) will maintain stable API contracts throughout the deployment period.",
        "Estimated cost savings in the analytics dashboard use a fixed rate of $25 per completed task (0.5 hours at $50/hour) as a baseline approximation of manual work displacement.",
    ]
    for a in assumptions:
        add_bullet(doc, a)

    doc.add_heading("Constraints", level=2)
    constraints = [
        "The platform is designed as a single-user-per-account system in v1.0 with per-user data isolation. Multi-tenant organization support, team workspaces, and shared agent management are deferred to future releases.",
        "Code Execution nodes run Python code in a subprocess on the Celery worker host with pattern-based import restriction (not OS-level sandboxing). The FORBIDDEN list blocks 12 patterns but does not prevent all possible exploits. Container-level isolation (Docker/Firecracker) is planned for future releases.",
        "Real-time SSE streaming endpoints (/api/executions/{id}/stream) do not require authentication, relying on UUID unguessability for security. This is a deliberate trade-off to avoid SSE authentication complexity.",
        "Integration credentials are stored as plaintext JSONB in PostgreSQL. At-rest encryption of sensitive fields (api_key, webhook_url, token, secret) is not implemented in v1.0. Credentials are masked in API responses but stored unencrypted.",
        "The technology stack is fixed to FastAPI (Python 3.12), React with TypeScript, Celery with Redis, and PostgreSQL with asyncpg. The stack leverages pydantic-settings for configuration management with .env file support.",
        "Agent chat is limited to text-based interaction using GPT-4o-mini with a 1024 max_tokens response limit and 10-message conversation history window per session. Memory context is capped at 15 most recent summaries.",
        "The platform currently supports two integration services (Resend and Slack) via an allowlist. Adding new integrations requires backend code changes to IntegrationNodeExecutor and the ALLOWED_SERVICES set.",
    ]
    for c in constraints:
        add_bullet(doc, c)


# ══════════════════════════════════════════════════════════════════
# SECTION 10: DEPENDENCIES & RISKS
# ══════════════════════════════════════════════════════════════════
def build_section_10(doc):
    doc.add_heading("10. Dependencies & Risks", level=1)

    doc.add_heading("10.1 Dependencies", level=2)
    make_styled_table(doc,
        ["Dep ID", "Dependency", "Type", "Impact if Unavailable"],
        [
            ("DEP-01",
             "PostgreSQL 14+ with asyncpg driver (postgresql+asyncpg://)",
             "Infrastructure",
             "Complete platform outage. All CRUD operations, data persistence, and user authentication require database connectivity. No fallback available."),
            ("DEP-02",
             "Redis 6+ (aioredis for async pub/sub, RedBeat for scheduling, Celery broker)",
             "Infrastructure",
             "No task queuing (executions queue but don't process), no cron scheduling (RedBeat inoperable), no real-time SSE streaming (pub/sub unavailable). Critical single point of failure."),
            ("DEP-03",
             "Celery Workers (execute_workflow, execute_workflow_scheduled tasks)",
             "Infrastructure",
             "Workflow executions queue in Redis but do not process. Manual and scheduled triggers create Execution records with status='pending' that never transition to 'running'. No workflow processing capacity."),
            ("DEP-04",
             "OpenAI API (GPT-4o, GPT-4o-mini, GPT-4.1 family, o3-mini)",
             "External Service",
             "AI Action nodes with OpenAI models return mock responses. Agent generation via /api/agents/generate fails. Chat returns configuration reminder. Memory summaries use template fallback. Platform functional but AI-degraded."),
            ("DEP-05",
             "Anthropic API (Claude Sonnet 4, Claude Haiku 4.5)",
             "External Service",
             "AI Action nodes with Anthropic models return mock responses. Does not affect OpenAI-model nodes, agent generation, chat, or memory summarization (which use OpenAI)."),
            ("DEP-06",
             "Resend API (transactional email)",
             "External Service",
             "Integration and Escalation nodes return mock results (emails_sent count but no actual delivery). Workflow execution continues with mock status. Non-blocking degradation."),
            ("DEP-07",
             "Slack Incoming Webhooks",
             "External Service",
             "Slack integration nodes return mock results (message_sent=True, mock=True). Workflow execution continues. Non-blocking degradation."),
            ("DEP-08",
             "DuckDuckGo HTML Search Endpoint (html.duckduckgo.com/html/)",
             "External Service",
             "Web Search nodes return empty results list. Workflows depending on web research produce incomplete or empty data for downstream nodes."),
        ],
    )

    doc.add_heading("10.2 Risks", level=2)
    make_styled_table(doc,
        ["Risk ID", "Risk Description", "Probability", "Impact", "Mitigation Strategy"],
        [
            ("RSK-01",
             "LLM API rate limiting, quota exhaustion, or provider outage causes cascading execution failures across multiple agents",
             "Medium", "High",
             "Implement retry logic with exponential backoff in AIActionNodeExecutor; support model fallback chain (e.g., GPT-4o -> GPT-4o-mini); existing mock mode provides degraded-but-functional fallback; monitor API error rates"),
            ("RSK-02",
             "DuckDuckGo changes HTML structure, breaking WebSearchNodeExecutor regex parsing and returning empty results",
             "Medium", "Medium",
             "Abstract search behind BaseNodeExecutor interface; prepare migration path to dedicated search API (SerpAPI, Brave Search); add integration tests with snapshot comparison for HTML structure changes"),
            ("RSK-03",
             "Code Execution node sandbox bypass via creative import patterns, indirect function calls, or dynamic code generation that circumvents the FORBIDDEN list",
             "Low", "High",
             "Migrate to RestrictedPython or containerized execution (Docker with resource limits); add process-level restrictions (seccomp, namespace isolation); audit FORBIDDEN list against known bypass techniques; implement execution output size limits"),
            ("RSK-04",
             "Redis becomes single point of failure: simultaneous loss of task queue, cron scheduler, and real-time streaming",
             "Low", "High",
             "Deploy Redis Sentinel or Redis Cluster for high availability; implement circuit breaker for SSE fallback to polling; add health check endpoints for Redis connectivity monitoring; persist critical state in PostgreSQL as backup"),
            ("RSK-05",
             "Uncontrolled LLM costs from misconfigured agents with tight cron schedules, verbose prompts, or expensive model selection (e.g., GPT-4o at $2.50/$10.00 per M tokens)",
             "Medium", "Medium",
             "Implement per-agent and per-user cost caps; add budget alert notifications via Escalation nodes; display cost projections in UI before deployment; default to cost-effective GPT-4o-mini ($0.15/$0.60 per M tokens); track costs with 6-digit precision"),
            ("RSK-06",
             "JWT secret compromise (default 'dev-secret-change-me' in config) enables unauthorized access and token forgery",
             "Low", "High",
             "Enforce strong JWT_SECRET via environment variable validation; remove default placeholder; implement token refresh and revocation; add JWT expiry monitoring; use separate signing keys for production and development environments"),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 11: SUCCESS CRITERIA
# ══════════════════════════════════════════════════════════════════
def build_section_11(doc):
    doc.add_heading("11. Success Criteria", level=1)

    add_body_paragraph(doc,
        "The following criteria define what constitutes a successful delivery of GoalCert AutoMind v1.0."
    )

    make_styled_table(doc,
        ["#", "Success Criterion", "Measurement Method", "Target"],
        [
            ("SC-01",
             "All 7 node types (Trigger, AI Action, Web Search, Decision, Integration, Escalation, Code Execution) execute correctly in multi-node workflow pipelines with variable passing and conditional branching",
             "End-to-end integration tests for each node type in isolation and in composite workflows with Decision branching",
             "100% pass rate across all node type combinations"),
            ("SC-02",
             "Natural-language agent generation via GPT-4o-mini produces structurally valid and deployable workflow definitions from diverse descriptions",
             "Test with 20 diverse agent descriptions spanning sales, marketing, support, and custom types",
             "70% of generated workflows deployable without manual node/edge edits"),
            ("SC-03",
             "Scheduled agents via RedBeat fire within acceptable tolerance of cron specification and successfully create and execute workflow runs",
             "Monitor scheduled vs actual trigger times over 30-day period; verify Execution records with triggered_by='schedule'",
             "95% on-time within 60-second tolerance; 99% successful execution creation"),
            ("SC-04",
             "Execution audit trails are complete: every executed node produces an ExecutionNodeLog record with all required fields populated (status, duration, input/output data, LLM usage)",
             "Verify node_logs count matches expected executed nodes for 100 randomly sampled executions; check all fields non-null for completed nodes",
             "100% completeness; no orphaned or missing node log records"),
            ("SC-05",
             "Real-time SSE streaming delivers execution events with acceptable latency from Redis publish to client receipt",
             "Measure timestamp delta between _publish_log call and SSE client event receipt under normal load",
             "P95 latency < 2 seconds; __STREAM_END__ delivered for 100% of completed executions"),
            ("SC-06",
             "LLM cost tracking accuracy: platform-computed costs via _COST_TABLE match actual API provider billing within acceptable variance",
             "Compare platform-computed total_cost values against OpenAI/Anthropic usage dashboard for a billing period",
             "Within 5% variance for all tracked models; zero untracked cost leakage"),
            ("SC-07",
             "Platform handles concurrent agent executions without degradation in execution duration, SSE delivery, or API response times",
             "Load test with 50 simultaneous workflow triggers via Celery worker pool; monitor queue depth and execution latency",
             "All executions complete within 2x normal duration; SSE latency within target; API response time < 500ms P95"),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# SECTION 12: GLOSSARY
# ══════════════════════════════════════════════════════════════════
def build_section_12(doc):
    doc.add_heading("12. Glossary", level=1)

    make_styled_table(doc,
        ["Term", "Definition"],
        [
            ("Agent", "An autonomous AI entity configured with a workflow, schedule, and memory that executes tasks without continuous human intervention. Agents have a type (sales/marketing/support/custom), status (draft/active/paused/error), and optional cron schedule."),
            ("Workflow", "A directed acyclic graph (DAG) of interconnected nodes and edges defining the sequence of operations an agent performs. Stored as a JSONB definition in React Flow format (nodes, edges, viewport)."),
            ("Node", "A discrete processing unit within a workflow that performs a specific action. AutoMind supports 7 types: Trigger, AI Action, Web Search, Decision, Integration, Escalation, and Code Execution."),
            ("Execution", "A single run of an agent's workflow, tracked with status lifecycle (pending -> running -> success/failed), timing, cost, per-node logs, and accumulated variable state."),
            ("BFS", "Breadth-First Search \u2013 the graph traversal algorithm used by WorkflowExecutor to process workflow nodes layer by layer from the trigger node, following edges and respecting decision branches."),
            ("SSE", "Server-Sent Events \u2013 a unidirectional HTTP streaming protocol used for real-time execution monitoring (/api/executions/{id}/stream) and chat response streaming (/api/agents/{id}/chat)."),
            ("Celery", "An asynchronous distributed task queue framework for Python. AutoMind uses Celery workers to process workflow executions (execute_workflow, execute_workflow_scheduled tasks) decoupled from the FastAPI web server."),
            ("RedBeat", "A Redis-backed Celery Beat scheduler that enables dynamic cron schedule management without restarting the scheduler process. Stores schedule entries as RedBeatSchedulerEntry objects in Redis."),
            ("Variable Interpolation", "The process of replacing {placeholder} tokens in node configurations with runtime values from the workflow's variable store. Supports nested path resolution (e.g., {leads.0.name}) and type-preserving substitution."),
            ("Agent Memory", "A persistent store of LLM-generated execution summaries (AgentMemory table) that provides agents with context from previous runs. Injected into AI Action node system prompts and agent chat conversations."),
            ("LLM", "Large Language Model \u2013 foundation AI models (e.g., GPT-4o-mini, Claude Sonnet 4) called by AI Action nodes to generate, analyse, or transform text and data within workflows."),
            ("React Flow", "An open-source React library (@xyflow/react) for building interactive node-based editors. Powers the AutoMind visual workflow builder with drag-and-drop canvas, node state management, and edge connections."),
        ],
    )


# ══════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════
def main():
    doc = Document()

    # ── Configure document styles to match Flex Coach BRD ────────
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.color.rgb = BODY_COLOR

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = DARK_NAVY

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = BLUE_ACCENT

    # Page margins: 1 inch all sides
    for section in doc.sections:
        section.top_margin = 914400
        section.bottom_margin = 914400
        section.left_margin = 914400
        section.right_margin = 914400

    # ── Build document ───────────────────────────────────────────
    build_cover_page(doc)
    build_revision_history(doc)
    build_toc(doc)
    build_section_1(doc)
    build_section_2(doc)
    build_section_3(doc)
    build_section_4(doc)
    build_section_5(doc)
    build_section_6(doc)
    build_section_7(doc)
    build_section_8(doc)
    build_section_9(doc)
    build_section_10(doc)
    build_section_11(doc)
    build_section_12(doc)

    # Save
    os.makedirs(os.path.dirname(SAVE_PATH), exist_ok=True)
    doc.save(SAVE_PATH)
    print(f"BRD saved to {SAVE_PATH}")


if __name__ == "__main__":
    main()
