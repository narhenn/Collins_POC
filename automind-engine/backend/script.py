"""Generate GoalCert AutoMind Data Flow Diagram (DFD) document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# -- Colour constants --
PURPLE = RGBColor(73, 2, 162)
WHITE = RGBColor(255, 255, 255)
GRAY = RGBColor(128, 128, 128)
LIGHT_GRAY = RGBColor(245, 245, 245)
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(51, 51, 51)
LIGHT_PURPLE = RGBColor(240, 230, 255)

doc = Document()

# ---------- default style ----------
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = DARK_GRAY
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.space_before = Pt(0)
paragraph_format.line_spacing = 1.15

# ---------- heading styles ----------
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hf = hs.font
    hf.name = "Calibri"
    hf.color.rgb = PURPLE
    hf.bold = True
    if level == 1:
        hf.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hf.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hf.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


# ===================================================================
#  Helper functions
# ===================================================================

def set_cell_shading(cell, hex_color):
    """Apply background shading to a cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs keys: top, bottom, left, right, insideH, insideV
    values: dict with sz, val, color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_purple_header_row(table, headers):
    """Style the first row of a table as a purple header."""
    row = table.rows[0]
    for idx, cell in enumerate(row.cells):
        set_cell_shading(cell, "4902A2")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(headers[idx])
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)


def add_table_row(table, values, bold_first=False, shade_alt=False, row_idx=0):
    """Populate a row in a table."""
    row = table.rows[row_idx] if row_idx < len(table.rows) else table.add_row()
    for idx, cell in enumerate(row.cells):
        if idx < len(values):
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(values[idx]))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY
            if bold_first and idx == 0:
                run.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    if shade_alt:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")
    return row


def create_table(doc, rows, cols, headers, data, col_widths=None):
    """Create a fully styled table with purple header and alternating rows."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_purple_header_row(table, headers)
    for i, row_data in enumerate(data):
        row = table.add_row() if i + 1 >= rows else None
        actual_row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = actual_row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        if i % 2 == 1:
            for cell in actual_row.cells:
                set_cell_shading(cell, "F5F0FF")
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width
    return table


def add_body_paragraph(doc, text, bold=False, italic=False):
    """Add a body paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(8)
    return p


def add_figure_caption(doc, text):
    """Add a figure caption in italic, centered."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)


# ===================================================================
#  COVER PAGE
# ===================================================================

# Add several blank paragraphs for vertical centering
for _ in range(6):
    doc.add_paragraph()

# GoalCert
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GoalCert")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = PURPLE
run.font.name = "Calibri"

# AutoMind subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AutoMind \u2014 Agentic AI Workforce Platform")
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

# Add spacing
doc.add_paragraph()

# Document title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Data Flow Diagram (DFD)")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = DARK_GRAY
run.font.name = "Calibri"

# Description line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A comprehensive mapping of data flows across the AutoMind platform,\nfrom user input through AI processing to output delivery.")
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

# Add spacing
for _ in range(4):
    doc.add_paragraph()

# Confidentiality notice
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = PURPLE
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("This document contains proprietary information belonging to GoalCert Pte. Ltd.\nUnauthorized distribution or reproduction is strictly prohibited.")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

# ===================================================================
#  PAGE BREAK -> DOCUMENT CONTROL
# ===================================================================

doc.add_page_break()

doc.add_heading("Document Control", level=1)

# Document Control table - 8 rows
dc_table = doc.add_table(rows=9, cols=2)
dc_table.style = "Table Grid"
dc_table.alignment = WD_TABLE_ALIGNMENT.LEFT

dc_headers = ["Field", "Details"]
add_purple_header_row(dc_table, dc_headers)

dc_data = [
    ("Document Title", "GoalCert AutoMind \u2014 Data Flow Diagram (DFD)"),
    ("Document ID", "GC-AM-DFD-001"),
    ("Version", "1.0"),
    ("Status", "Draft"),
    ("Classification", "Confidential"),
    ("Author", "GoalCert Engineering"),
    ("Date", "23 June 2026"),
    ("Approved By", "Prem Kumar, CTO"),
]

for i, (field, detail) in enumerate(dc_data):
    row = dc_table.rows[i + 1]
    for j, val in enumerate([field, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in dc_table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(4.5)

doc.add_paragraph()

# Version History table
doc.add_heading("Version History", level=2)

vh_table = doc.add_table(rows=1, cols=4)
vh_table.style = "Table Grid"
vh_table.alignment = WD_TABLE_ALIGNMENT.LEFT

add_purple_header_row(vh_table, ["Version", "Date", "Author", "Changes"])

vh_data = [
    ("0.1", "18 Jun 2026", "GoalCert Engineering", "Initial draft \u2014 context diagram and Level 1 DFD structure"),
    ("0.2", "20 Jun 2026", "GoalCert Engineering", "Added process descriptions, data store details, and external entity catalog"),
    ("1.0", "23 Jun 2026", "GoalCert Engineering", "Finalised data flow table with 30 flows, added comprehensive descriptions"),
]

for i, row_data in enumerate(vh_data):
    row = vh_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")


# ===================================================================
#  TABLE OF CONTENTS placeholder
# ===================================================================

doc.add_page_break()

doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("1.", "Introduction", "4"),
    ("2.", "Notation", "5"),
    ("3.", "Context Diagram (Level 0)", "6"),
    ("4.", "Level 1 Data Flow Diagram", "8"),
    ("4.1", "Process Decomposition", "8"),
    ("4.2", "Data Stores", "8"),
    ("4.3", "Data Flow Table", "9"),
    ("5.", "Process Descriptions", "12"),
    ("6.", "Data Store Catalog", "16"),
    ("7.", "External Entity Catalog", "17"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}\t{title}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    if not "." in num or num.endswith("."):
        run.bold = True
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(0.6))


# ===================================================================
#  1. INTRODUCTION
# ===================================================================

doc.add_page_break()

doc.add_heading("1. Introduction", level=1)

doc.add_heading("1.1 Purpose", level=2)

add_body_paragraph(doc, (
    "This Data Flow Diagram (DFD) document describes how data moves through the GoalCert AutoMind platform, "
    "from the moment a user defines an AI agent and builds a workflow, through the execution engine that "
    "orchestrates node-by-node processing, to the delivery of results via dashboards, emails, Slack notifications, "
    "and real-time streaming consoles. The DFD serves as a technology-neutral blueprint that captures what data "
    "flows where across the system, rather than prescribing how the underlying implementation achieves these flows."
))

add_body_paragraph(doc, (
    "AutoMind is an agentic AI workforce platform that enables non-technical users to create, configure, and "
    "deploy autonomous AI agents through a visual workflow builder. Each agent encapsulates a workflow composed "
    "of interconnected nodes \u2014 triggers, AI actions, decision gates, integration dispatchers, web search modules, "
    "and code execution blocks \u2014 that the platform traverses using a breadth-first search (BFS) algorithm. "
    "Understanding the data flows between these components is essential for security auditing, system design "
    "validation, compliance assessment, and onboarding new engineers to the platform architecture."
))

doc.add_heading("1.2 Scope", level=2)

add_body_paragraph(doc, (
    "This document covers all data flows within the AutoMind platform as of version 1.0. It encompasses the "
    "complete lifecycle of data from user authentication through agent creation, workflow definition, execution "
    "orchestration, AI processing, external integration delivery, real-time event streaming, memory accumulation, "
    "and dashboard analytics aggregation. The DFD is structured in two levels: a Level 0 context diagram that "
    "shows the platform as a single process interacting with external entities, and a Level 1 diagram that "
    "decomposes the platform into seven internal processes and seven data stores."
))

doc.add_heading("1.3 Audience", level=2)

add_body_paragraph(doc, (
    "This document is intended for software architects, backend engineers, security auditors, product managers, "
    "and any stakeholder who needs to understand the data architecture of AutoMind. It assumes familiarity with "
    "DFD notation (Yourdon-DeMarco or Gane-Sarson conventions) and a general understanding of web application "
    "data patterns."
))

doc.add_heading("1.4 Conventions", level=2)

add_body_paragraph(doc, (
    "Data Flow Diagrams are technology-neutral by design. They describe what data flows where, not how the "
    "data is transported or transformed at the implementation level. Process identifiers use the Pn naming "
    "convention (P1, P2, ..., P7), data stores use DSn (DS1, ..., DS7), external entities use En (E1, ..., E6), "
    "and individual data flows use DF-nn (DF-01, ..., DF-30). Where this document references specific technologies "
    "(PostgreSQL, Redis, Celery), it does so only to clarify the nature of a data store or transport mechanism, "
    "not to constrain the design."
))


# ===================================================================
#  2. NOTATION
# ===================================================================

doc.add_page_break()

doc.add_heading("2. Notation", level=1)

add_body_paragraph(doc, (
    "The following table defines the DFD notation symbols used throughout this document. The notation follows "
    "the Yourdon-DeMarco convention, which is the most widely adopted standard for structured analysis and "
    "data flow modeling."
))

notation_table = doc.add_table(rows=1, cols=3)
notation_table.style = "Table Grid"
notation_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(notation_table, ["Symbol / Notation Style", "Element Type", "Meaning"])

notation_data = [
    ("Double-bordered rectangle", "External Entity", "An actor, user, or external system that exists outside the boundary of the AutoMind platform. External entities produce data that enters the system or consume data that the system outputs. They are not under the control of the platform."),
    ("Ellipse / rounded rectangle", "Process", "A transformation or processing step that receives input data, performs some operation, and produces output data. Processes are numbered (P1\u2013P7) and represent logical groupings of functionality within AutoMind."),
    ("Open-ended rectangle", "Data Store", "A repository where data is persisted or temporarily held. Data stores are numbered (DS1\u2013DS7) and correspond to database tables, caches, or message queues. Data can be read from or written to a data store."),
    ("Arrow (directed line)", "Data Flow", "The movement of data between elements in the diagram. Each arrow is labeled with a description of the data it carries and is assigned a unique identifier (DF-01 through DF-30). Arrows indicate direction of data movement."),
]

for i, row_data in enumerate(notation_data):
    row = notation_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in notation_table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(3.5)

add_body_paragraph(doc, (
    "Throughout this document, table-based representations are used to depict the DFD diagrams. In the context "
    "diagram (Level 0), external entities are positioned around a central process cell. In the Level 1 diagram, "
    "the complete set of data flows is cataloged in a tabular format that specifies the source, destination, and "
    "content of each flow. This tabular approach ensures precision and traceability that is sometimes lost in "
    "purely graphical representations."
))


# ===================================================================
#  3. CONTEXT DIAGRAM (Level 0)
# ===================================================================

doc.add_page_break()

doc.add_heading("3. Context Diagram (Level 0)", level=1)

add_body_paragraph(doc, (
    "The context diagram presents the GoalCert AutoMind platform as a single process, labeled P0, interacting "
    "with six external entities. This is the highest level of abstraction in the DFD hierarchy. At this level, "
    "all internal processing is hidden within P0, and the diagram focuses exclusively on the data that crosses "
    "the system boundary \u2014 what enters from external entities and what the platform sends back out."
))

add_body_paragraph(doc, (
    "P0 represents the entire AutoMind platform: the authentication layer, the agent and workflow management "
    "system, the execution engine that traverses workflow graphs, the AI processing subsystem that calls "
    "language model APIs, the integration dispatch layer that sends emails and Slack messages, the real-time "
    "streaming infrastructure, and the dashboard analytics engine. All of these are collapsed into a single "
    "process at Level 0 and will be decomposed into seven sub-processes at Level 1."
))

# Create the 5x5 context diagram table
ctx_table = doc.add_table(rows=5, cols=5)
ctx_table.style = "Table Grid"
ctx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set all cells to have minimal content initially
for row in ctx_table.rows:
    for cell in row.cells:
        p = cell.paragraphs[0]
        p.clear()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

def fill_cell(row, col, text, is_center=False, is_entity=False):
    cell = ctx_table.rows[row].cells[col]
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    if is_center:
        set_cell_shading(cell, "4902A2")
        run.font.color.rgb = WHITE
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif is_entity:
        set_cell_shading(cell, "F5F0FF")
        run.font.color.rgb = PURPLE
        run.font.size = Pt(8)
    else:
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

# Top-left: E1 Platform User
fill_cell(0, 0, "E1: Platform User\n\nSends: agent configs, workflow definitions, chat messages, execution triggers\n\nReceives: dashboards, reports, analytics, execution results, chat responses", is_entity=True)

# Top-center: arrow label
fill_cell(0, 2, "Prompts with context,\nmodel config\n\u2193\n\u2191\nGenerated text,\ntoken usage, cost", is_entity=False)

# Top-right: E2 OpenAI/Anthropic
fill_cell(0, 4, "E2: OpenAI / Anthropic\n\nReceives: API requests (messages, model, max_tokens, temperature)\n\nSends: generated text, token counts, cost metrics", is_entity=True)

# Left-center: E6 Cron Scheduler
fill_cell(2, 0, "E6: Cron Scheduler\n(RedBeat / Celery Beat)\n\nSends: scheduled trigger events (agent_id, cron expression)", is_entity=True)

# CENTER: P0
fill_cell(2, 2, "P0\nGoalCert\nAutoMind\nPlatform", is_center=True)

# Right side: arrow labels
fill_cell(2, 4, "", is_entity=False)

# Arrows from E1 to center
fill_cell(1, 1, "\u2199 Configs, triggers,\nchat messages\n\u2197 Results, dashboards,\nSSE streams")

# Arrows from E2 to center
fill_cell(1, 3, "\u2196 Prompts + context\n\u2198 AI responses + usage")

# Bottom-left: E3 Email
fill_cell(4, 0, "E3: Email Service\n(Resend API)\n\nReceives: email content (to, subject, HTML body, from address)", is_entity=True)

# Bottom-center: E4 Slack
fill_cell(4, 2, "E4: Slack\n(Webhook)\n\nReceives: webhook payload (channel, message text)", is_entity=True)

# Bottom-right: E5 Web
fill_cell(4, 4, "E5: Web\n(DuckDuckGo)\n\nReceives: search queries\nSends: search results (title, URL, snippet, rank)", is_entity=True)

# Arrow labels bottom
fill_cell(3, 0, "\u2198 Email delivery\ninstructions")
fill_cell(3, 2, "\u2193 Slack webhook\npayloads")
fill_cell(3, 4, "\u2199 Search queries\n\u2197 Search results")

# Set column widths
for row in ctx_table.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(1.2)
    row.cells[2].width = Inches(1.8)
    row.cells[3].width = Inches(1.2)
    row.cells[4].width = Inches(1.8)

add_figure_caption(doc, "Figure 1: Context Diagram (Level 0) \u2014 GoalCert AutoMind Platform and its external entities")

add_body_paragraph(doc, (
    "The context diagram reveals six distinct external entities that interact with AutoMind. The Platform User "
    "(E1) is the primary actor who creates agents, defines workflows, triggers executions, and consumes results "
    "through dashboards and real-time streaming consoles. OpenAI and Anthropic (E2) serve as the large language "
    "model providers, receiving structured prompts with context and returning generated text along with token "
    "usage and cost metrics. The Email Service (E3), powered by the Resend API, receives formatted email "
    "content for delivery to external recipients. Slack (E4) receives webhook payloads for team notifications. "
    "DuckDuckGo (E5) provides web search capabilities, accepting search queries and returning structured "
    "results. The Cron Scheduler (E6), backed by RedBeat and Celery Beat via Redis, sends scheduled trigger "
    "events to initiate automated workflow executions at user-defined intervals."
))

add_body_paragraph(doc, (
    "A critical observation from this diagram is that all data flowing into the system from E1 is "
    "authenticated \u2014 every API request carries a JWT token that binds the request to a specific user context. "
    "Data flowing to E2 is rate-limited and cost-tracked per execution. Data flowing to E3 and E4 is "
    "one-directional \u2014 the platform sends delivery instructions but does not process inbound email or Slack "
    "messages in this version. The interaction with E5 is stateless and requires no API key, as the platform "
    "scrapes DuckDuckGo's HTML search results directly."
))


# ===================================================================
#  4. LEVEL 1 DATA FLOW DIAGRAM
# ===================================================================

doc.add_page_break()

doc.add_heading("4. Level 1 Data Flow Diagram", level=1)

add_body_paragraph(doc, (
    "The Level 1 DFD decomposes the single P0 process from the context diagram into seven internal processes "
    "and seven data stores. This decomposition reveals the internal data routing that makes AutoMind function: "
    "how user requests are authenticated, how agents and workflows are persisted, how the execution engine "
    "traverses workflow graphs, how AI models are called, how integration messages are dispatched, how real-time "
    "events are streamed to connected clients, and how dashboard analytics are aggregated from execution history."
))

doc.add_heading("4.1 Process Decomposition", level=2)

add_body_paragraph(doc, (
    "The following table lists all seven processes that compose the AutoMind platform at Level 1. Each process "
    "represents a logically distinct area of functionality, though in the physical architecture some processes "
    "share the same FastAPI application server while others run in Celery worker processes."
))

proc_table = doc.add_table(rows=1, cols=3)
proc_table.style = "Table Grid"
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(proc_table, ["Process ID", "Process Name", "Description"])

proc_data = [
    ("P1", "Authentication & Session Management", "Handles user registration, login, password verification, JWT token generation, and session validation. Every authenticated API request passes through P1 for token verification before reaching any other process."),
    ("P2", "Agent & Workflow Management", "Manages the full lifecycle of agents (create, read, update, delete) and their associated workflows. Workflows are stored as JSONB definitions containing the React Flow graph of nodes and edges. Also handles template instantiation."),
    ("P3", "Execution Engine", "The core orchestration process. Loads a workflow definition, identifies the start node, performs BFS traversal across the graph, executes each node in order, handles decision branching, manages variable interpolation across nodes, and records execution logs."),
    ("P4", "AI Processing", "Handles all interactions with LLM providers. Receives prompts with variable context and agent memory, routes to the appropriate provider (OpenAI or Anthropic), processes responses, calculates token costs, and returns structured output with usage metrics."),
    ("P5", "Integration Dispatch", "Delivers messages to external services. Formats and sends emails via the Resend API, posts messages to Slack via incoming webhooks. Reads integration configuration (API keys, webhook URLs) from the data store."),
    ("P6", "Real-time Streaming", "Provides Server-Sent Events (SSE) endpoints for live execution consoles and agent chat. Subscribes to Redis pub/sub channels to receive execution log events and streams them to connected browser clients in real time."),
    ("P7", "Dashboard & Analytics Aggregation", "Aggregates execution data to produce dashboard statistics: total agents, active agents, tasks completed, estimated cost savings, average response time, and recent activity feeds."),
]

for i, row_data in enumerate(proc_data):
    row = proc_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in proc_table.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(4.2)

doc.add_paragraph()

doc.add_heading("4.2 Data Stores", level=2)

add_body_paragraph(doc, (
    "AutoMind uses two primary storage technologies: PostgreSQL for durable relational data and Redis for "
    "ephemeral messaging and task scheduling. The seven data stores capture all persistent and transient state "
    "within the platform."
))

ds_table = doc.add_table(rows=1, cols=3)
ds_table.style = "Table Grid"
ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(ds_table, ["Store ID", "Store Name", "Description"])

ds_data = [
    ("DS1", "Users (PostgreSQL)", "User accounts with email, hashed password, display name, and timestamps. Each user owns a set of agents and integration configurations. The users table is the root of the ownership hierarchy."),
    ("DS2", "Agents & Workflows (PostgreSQL + JSONB)", "Agent records (name, type, status, cron schedule, timezone) and their one-to-one workflow definitions stored as JSONB. The JSONB column holds the complete React Flow graph with nodes, edges, and node configurations."),
    ("DS3", "Executions & Node Logs (PostgreSQL)", "Execution records tracking status, timing, cost, trigger type, and error messages. Each execution has a set of node logs that record per-node input, output, timing, status, and LLM usage metrics."),
    ("DS4", "Agent Memory (PostgreSQL)", "Execution summaries generated by GPT-4o-mini after each workflow run. Stores the agent_id, execution_id, a natural-language summary of what the execution accomplished, key output data, and memory type classification."),
    ("DS5", "Templates (PostgreSQL)", "Pre-built agent templates with name, description, type, icon, color, feature list, and a complete workflow_definition that can be cloned to create new agents instantly."),
    ("DS6", "Integrations Config (PostgreSQL)", "Per-user integration settings keyed by service name (email, slack). Stores API keys, webhook URLs, and configuration as JSONB. Each user has a unique constraint on (user_id, service)."),
    ("DS7", "Task Queue & Pub/Sub (Redis)", "Celery task queue for asynchronous workflow execution, RedBeat schedule entries for cron-based triggers, and pub/sub channels for real-time execution log streaming to SSE endpoints."),
]

for i, row_data in enumerate(ds_data):
    row = ds_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in ds_table.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(2.2)
    row.cells[2].width = Inches(4.0)


# ===================================================================
#  4.3 DATA FLOW TABLE
# ===================================================================

doc.add_page_break()

doc.add_heading("4.3 Data Flow Table", level=2)

add_body_paragraph(doc, (
    "The following table catalogs every data flow in the Level 1 DFD. Each row represents a single directed "
    "flow of data between two elements (processes, data stores, or external entities). The flows are numbered "
    "sequentially and grouped by the primary process they involve. This table is the definitive reference for "
    "understanding what data moves where in AutoMind."
))

# Create the flow table
flow_table = doc.add_table(rows=1, cols=4)
flow_table.style = "Table Grid"
flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(flow_table, ["Flow ID", "From", "To", "Data Description"])

flow_data = [
    ("DF-01", "E1: Platform User", "P1: Auth & Session", "Login credentials (email address + plaintext password) or registration data (email + password + display name)"),
    ("DF-02", "P1: Auth & Session", "E1: Platform User", "JWT access token + user profile object (id, email, name, created_at)"),
    ("DF-03", "P1: Auth & Session", "DS1: Users", "New user record (email, bcrypt password hash, name) on registration; read query on login"),
    ("DF-04", "DS1: Users", "P1: Auth & Session", "Existing user record for password verification and token generation"),
    ("DF-05", "E1: Platform User", "P2: Agent & Workflow Mgmt", "Agent configuration (name, type, description, schedule_cron, schedule_timezone) and workflow definition (nodes + edges as JSON)"),
    ("DF-06", "P2: Agent & Workflow Mgmt", "DS2: Agents & Workflows", "Agent record insert/update + workflow JSONB definition upsert"),
    ("DF-07", "DS2: Agents & Workflows", "P2: Agent & Workflow Mgmt", "Agent list, agent detail, workflow definition for display in builder"),
    ("DF-08", "P2: Agent & Workflow Mgmt", "E1: Platform User", "Agent list, agent detail responses, workflow definition for visual builder rendering"),
    ("DF-09", "DS5: Templates", "P2: Agent & Workflow Mgmt", "Template workflow definition + metadata for agent creation from template"),
    ("DF-10", "E6: Cron Scheduler", "P3: Execution Engine", "Scheduled trigger event containing agent_id, dispatched via Celery task queue at cron intervals"),
    ("DF-11", "E1: Platform User", "P3: Execution Engine", "Manual execution trigger (agent_id) via POST /api/agents/{id}/execute endpoint"),
    ("DF-12", "P3: Execution Engine", "DS2: Agents & Workflows", "Read request for workflow definition and agent configuration to prepare execution"),
    ("DF-13", "DS2: Agents & Workflows", "P3: Execution Engine", "Workflow definition (nodes array + edges array) and agent metadata (name, type) for execution"),
    ("DF-14", "P3: Execution Engine", "DS3: Executions & Logs", "New execution record (status=pending, triggered_by, agent_id, workflow_id) and per-node log entries (node_id, type, label, status, input_data, output_data, timing, llm_usage)"),
    ("DF-15", "DS3: Executions & Logs", "P3: Execution Engine", "Execution record for status updates (pending to running to success/failed)"),
    ("DF-16", "P3: Execution Engine", "P4: AI Processing", "Prompt string with interpolated variables + model configuration (model name, max_tokens, temperature, system_prompt) + agent memory context from past executions"),
    ("DF-17", "P4: AI Processing", "E2: OpenAI / Anthropic", "API request payload: messages array (system + user roles), model identifier, max_tokens, temperature parameter"),
    ("DF-18", "E2: OpenAI / Anthropic", "P4: AI Processing", "API response: generated text content, input token count, output token count, finish reason"),
    ("DF-19", "P4: AI Processing", "P3: Execution Engine", "Structured output: parsed AI response (text or JSON), LLM usage metrics (model, input_tokens, output_tokens, calculated cost)"),
    ("DF-20", "P3: Execution Engine", "P5: Integration Dispatch", "Delivery instruction: service type (email or slack) + content payload (recipients/subject/body for email, message/webhook_url for Slack)"),
    ("DF-21", "P5: Integration Dispatch", "E3: Email (Resend)", "HTTP POST to Resend API: from address, to recipients list, subject line, HTML body content"),
    ("DF-22", "P5: Integration Dispatch", "E4: Slack", "HTTP POST to webhook URL: JSON payload with text field containing the message"),
    ("DF-23", "DS6: Integrations Config", "P5: Integration Dispatch", "API keys (Resend API key), webhook URLs (Slack webhook URL), and per-user service configuration"),
    ("DF-24", "P3: Execution Engine", "E5: Web (DuckDuckGo)", "Search query string (interpolated from node config) + max_results parameter"),
    ("DF-25", "E5: Web (DuckDuckGo)", "P3: Execution Engine", "Search results array: each result contains title, URL, snippet text, and rank position"),
    ("DF-26", "P3: Execution Engine", "DS7: Redis Pub/Sub", "Execution log events published to channel execution:{id}:logs containing timestamp, message, node_id, status"),
    ("DF-27", "DS7: Redis Pub/Sub", "P6: SSE Streaming", "Log events received via pub/sub subscription for active execution channels"),
    ("DF-28", "P6: SSE Streaming", "E1: Platform User", "Server-Sent Events: real-time execution console lines (node start, completion, errors, timing) and chat response tokens"),
    ("DF-29", "P3: Execution Engine", "DS4: Agent Memory", "Execution summary (GPT-generated natural language summary, key output variables, execution_id, agent_id, memory_type)"),
    ("DF-30", "DS4: Agent Memory", "P3: Execution Engine", "Past execution summaries formatted as context string for injection into AI node system prompts (most recent 10 memories)"),
    ("DF-31", "E1: Platform User", "P6: Chat Streaming", "Chat message text + conversation history array [{role, content}] + agent_id for context lookup"),
    ("DF-32", "DS4: Agent Memory", "P6: Chat Streaming", "Agent memory context (past 15 execution summaries) for system prompt construction in chat responses"),
    ("DF-33", "P7: Dashboard & Analytics", "DS3: Executions & Logs", "Aggregation queries: COUNT agents, COUNT active agents, COUNT successful executions, AVG duration_ms, recent 20 executions with agent metadata"),
    ("DF-34", "P7: Dashboard & Analytics", "E1: Platform User", "Dashboard statistics (total_agents, active_agents, tasks_completed, estimated_savings, avg_response_time) + activity feed (execution history with agent name, type, status, timing, cost)"),
    ("DF-35", "P2: Agent & Workflow Mgmt", "DS7: Redis (RedBeat)", "Schedule entry for cron-based agent triggers (agent_id, parsed crontab, task name) via RedBeat scheduler"),
]

for i, row_data in enumerate(flow_data):
    row = flow_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in flow_table.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(1.5)
    row.cells[3].width = Inches(3.5)

doc.add_paragraph()

add_figure_caption(doc, "Figure 2: Level 1 Data Flow Table \u2014 Complete catalog of 35 data flows within the AutoMind platform")

add_body_paragraph(doc, (
    "The data flow table above reveals several important architectural patterns. First, the Execution Engine (P3) "
    "is the most connected process, participating in 18 of the 35 flows. This is expected, as the execution engine "
    "is the central orchestrator that coordinates all other processes during a workflow run. Second, Redis (DS7) "
    "serves a dual role as both a Celery task queue for asynchronous execution dispatch and a pub/sub channel for "
    "real-time event streaming. Third, Agent Memory (DS4) creates a feedback loop where execution outputs from "
    "past runs are fed back into future AI prompts, enabling agents to develop contextual awareness over time."
))

add_body_paragraph(doc, (
    "Another notable pattern is the separation between the chat streaming path (DF-31, DF-32, DF-28) and the "
    "execution streaming path (DF-26, DF-27, DF-28). Both converge at P6 (SSE Streaming) and deliver events to "
    "the user via Server-Sent Events, but they originate from different triggers: chat is user-initiated with "
    "conversational context, while execution streaming is system-initiated as a byproduct of workflow traversal."
))


# ===================================================================
#  5. PROCESS DESCRIPTIONS
# ===================================================================

doc.add_page_break()

doc.add_heading("5. Process Descriptions", level=1)

add_body_paragraph(doc, (
    "This section provides detailed descriptions for each of the seven processes identified in the Level 1 DFD. "
    "Each process is documented with its inputs, processing logic, outputs, and the data stores it interacts with. "
    "These descriptions complement the data flow table by explaining the transformation logic within each process."
))

# P1
doc.add_heading("5.1 P1: Authentication & Session Management", level=2)

add_body_paragraph(doc, (
    "The authentication process is the gateway to the platform. Every API request except the login and "
    "registration endpoints must include a valid JWT token in the Authorization header. P1 handles three "
    "operations: user registration (creating a new account with bcrypt-hashed password), user login "
    "(verifying credentials and issuing a JWT), and token validation (decoding the JWT on every subsequent "
    "request to establish user identity). The JWT contains the user's ID and email, and is used by all "
    "downstream processes to enforce ownership-based access control."
))

p1_table = doc.add_table(rows=7, cols=2)
p1_table.style = "Table Grid"
add_purple_header_row(p1_table, ["Aspect", "Detail"])

p1_data = [
    ("Process ID", "P1"),
    ("Name", "Authentication & Session Management"),
    ("Input", "User credentials (email + password) for login; registration data (email + password + name) for signup; JWT token in Authorization header for session validation"),
    ("Processing", "Registration: validate email uniqueness against DS1, hash password with bcrypt, create user record, generate JWT. Login: query DS1 for user by email, verify password hash, generate JWT with user_id and email claims. Validation: decode JWT from header, extract user_id, load user from DS1."),
    ("Output", "JWT access token + user profile (id, email, name, created_at) on login/register; validated user context for downstream processes"),
    ("Data Stores Used", "DS1: Users (read for login/validation, write for registration)"),
]

for i, (aspect, detail) in enumerate(p1_data):
    row = p1_table.rows[i + 1]
    for j, val in enumerate([aspect, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in p1_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.5)

doc.add_paragraph()

# P2
doc.add_heading("5.2 P2: Agent & Workflow Management", level=2)

add_body_paragraph(doc, (
    "Agent and workflow management handles the complete CRUD lifecycle for agents and their associated "
    "workflows. An agent is a named entity with a type classification (e.g., content, research, outreach), "
    "an optional cron schedule for automated execution, and a one-to-one relationship with a workflow "
    "definition. The workflow definition is stored as a JSONB column containing the complete React Flow "
    "graph serialization: arrays of nodes (with type, position, and configuration data) and edges (with "
    "source, target, and optional branch labels). This process also handles template instantiation, where "
    "a pre-built workflow definition is cloned from DS5 to create a new agent with a ready-made workflow."
))

p2_table = doc.add_table(rows=7, cols=2)
p2_table.style = "Table Grid"
add_purple_header_row(p2_table, ["Aspect", "Detail"])

p2_data = [
    ("Process ID", "P2"),
    ("Name", "Agent & Workflow Management"),
    ("Input", "Agent configuration (name, type, description, schedule_cron, schedule_timezone) from E1; workflow definition (nodes + edges JSON) from visual builder; template selection from E1"),
    ("Processing", "Validates agent name uniqueness per user. Creates/updates agent record in DS2. Upserts workflow definition as JSONB. On schedule change, creates/updates RedBeat entry in DS7 for cron-based execution. On template selection, reads template from DS5 and clones workflow_definition to new agent."),
    ("Output", "Agent record with workflow definition returned to E1 for display; schedule entry written to DS7 for cron triggers"),
    ("Data Stores Used", "DS2: Agents & Workflows (read/write), DS5: Templates (read), DS7: Redis/RedBeat (write schedule entries)"),
]

for i, (aspect, detail) in enumerate(p2_data):
    row = p2_table.rows[i + 1]
    for j, val in enumerate([aspect, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in p2_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.5)

doc.add_paragraph()

# P3
doc.add_heading("5.3 P3: Execution Engine", level=2)

add_body_paragraph(doc, (
    "The Execution Engine is the heart of the AutoMind platform. It is responsible for loading a workflow "
    "definition, parsing the graph structure into nodes and edges, identifying the starting node (preferring "
    "trigger-type nodes with no incoming edges), and performing a breadth-first search (BFS) traversal of "
    "the graph. At each node, the engine identifies the node type, resolves the appropriate executor "
    "(TriggerNodeExecutor, AIActionNodeExecutor, DecisionNodeExecutor, IntegrationNodeExecutor, "
    "WebSearchNodeExecutor, CodeExecNodeExecutor, or EscalationNodeExecutor), interpolates template "
    "variables from the shared variable context, executes the node, merges the output variables back into "
    "the shared context, and determines the next nodes to visit. For decision nodes, the engine follows "
    "only the edge whose label matches the branch result (true or false). The engine runs inside a Celery "
    "worker process, managing its own database sessions for each operation."
))

p3_table = doc.add_table(rows=7, cols=2)
p3_table.style = "Table Grid"
add_purple_header_row(p3_table, ["Aspect", "Detail"])

p3_data = [
    ("Process ID", "P3"),
    ("Name", "Execution Engine"),
    ("Input", "Trigger event (scheduled from E6 or manual from E1), workflow definition from DS2, agent configuration, agent memory context from DS4"),
    ("Processing", "Creates execution record in DS3 (status=pending). Loads workflow definition, parses nodes and edges. Loads agent memory context from DS4 for AI prompt injection. Marks execution as running. Identifies root node (trigger type with no incoming edges). BFS traversal: for each node, resolves executor by type, interpolates variables using {{variable}} syntax, calls executor.execute(), merges output_variables into shared context, logs node result to DS3, publishes log event to DS7 Redis pub/sub. For decision nodes, evaluates condition and follows matching branch edge. On AI nodes, delegates to P4. On integration nodes, delegates to P5. On web search nodes, calls E5 directly. On completion, calculates total duration and cost, updates execution status, calls memory_service to generate and store execution summary in DS4."),
    ("Output", "Execution record with final status (success/failed), per-node logs with timing and output data, agent memory entry, SSE log events via DS7 pub/sub"),
    ("Data Stores Used", "DS2 (read workflow), DS3 (write execution + node logs), DS4 (read memory context, write execution summary), DS7 (publish real-time log events)"),
]

for i, (aspect, detail) in enumerate(p3_data):
    row = p3_table.rows[i + 1]
    for j, val in enumerate([aspect, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in p3_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.5)

doc.add_paragraph()

# P4
doc.add_heading("5.4 P4: AI Processing", level=2)

add_body_paragraph(doc, (
    "The AI Processing process encapsulates all interactions with external LLM providers. It receives a "
    "fully interpolated prompt string, model configuration parameters, and optionally an agent memory "
    "context string from P3. Based on the model name, it routes the request to either the OpenAI API "
    "(for GPT-4o, GPT-4o-mini, GPT-4.1, GPT-4.1-mini, GPT-4.1-nano, o3-mini models) or the Anthropic "
    "API (for Claude Sonnet 4, Claude Haiku 4.5 models). It constructs the appropriate API request format, "
    "executes the call asynchronously, parses the response, attempts JSON extraction from the response text, "
    "calculates the cost based on a per-model token pricing table, and returns structured output with usage "
    "metrics back to P3."
))

p4_table = doc.add_table(rows=7, cols=2)
p4_table.style = "Table Grid"
add_purple_header_row(p4_table, ["Aspect", "Detail"])

p4_data = [
    ("Process ID", "P4"),
    ("Name", "AI Processing"),
    ("Input", "Prompt string (with variables already interpolated), model identifier (e.g. gpt-4o-mini, claude-sonnet-4-20250514), max_tokens, temperature, optional system_prompt, optional output_variable name, agent memory context"),
    ("Processing", "Determines provider from model name prefix. Injects memory context into system prompt if available. For OpenAI: constructs messages array with system + user roles, calls AsyncOpenAI.chat.completions.create(). For Anthropic: constructs messages with user role and system parameter, calls AsyncAnthropic.messages.create(). Parses response text, attempts JSON extraction (handles code blocks and raw JSON). Calculates cost using per-model pricing table (input_rate, output_rate per million tokens). If no API key is configured, returns a mock response for development."),
    ("Output", "Structured output: output_variables dict (parsed JSON or {output: text}), llm_usage dict (model, input_tokens, output_tokens, cost)"),
    ("Data Stores Used", "None directly (API keys are loaded from environment configuration, not from a data store)"),
]

for i, (aspect, detail) in enumerate(p4_data):
    row = p4_table.rows[i + 1]
    for j, val in enumerate([aspect, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in p4_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.5)

doc.add_paragraph()

# P5
doc.add_heading("5.5 P5: Integration Dispatch", level=2)

add_body_paragraph(doc, (
    "The Integration Dispatch process handles the last-mile delivery of data to external communication "
    "services. Currently, it supports two integration types: email via the Resend API and Slack via "
    "incoming webhooks. When P3 encounters an integration node during workflow traversal, it delegates "
    "to P5 with the service type and content payload. P5 resolves recipient addresses (handling comma-"
    "separated strings, lists, and variable-resolved objects), formats the request, and makes the HTTP "
    "call to the external service. If API keys are not configured, P5 returns mock results for development "
    "and testing purposes."
))

p5_table = doc.add_table(rows=7, cols=2)
p5_table.style = "Table Grid"
add_purple_header_row(p5_table, ["Aspect", "Detail"])

p5_data = [
    ("Process ID", "P5"),
    ("Name", "Integration Dispatch"),
    ("Input", "Delivery instruction from P3: service type (email/slack), for email: recipients (string, list, or variable reference), subject, body (HTML), from address; for Slack: message text, webhook_url"),
    ("Processing", "Email: normalises recipients to a list (handles comma-separated strings, variable-resolved lists of objects with email fields, single addresses). Interpolates variables in subject and body. POSTs to Resend API at https://api.resend.com/emails with Bearer auth. Slack: interpolates variables in message. POSTs JSON {text: message} to webhook URL. Both: reads API keys from DS6 or environment config. Returns mock result if keys not configured."),
    ("Output", "Delivery result: emails_sent count + resend_id for email, message_sent boolean for Slack, error details on failure"),
    ("Data Stores Used", "DS6: Integrations Config (read API keys, webhook URLs)"),
]

for i, (aspect, detail) in enumerate(p5_data):
    row = p5_table.rows[i + 1]
    for j, val in enumerate([aspect, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in p5_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.5)

doc.add_paragraph()

# P6
doc.add_heading("5.6 P6: Real-time Streaming", level=2)

add_body_paragraph(doc, (
    "The Real-time Streaming process provides two SSE (Server-Sent Events) endpoints: one for live execution "
    "console streaming and one for agent chat. The execution streaming endpoint subscribes to a Redis pub/sub "
    "channel keyed by execution ID (execution:{id}:logs) and forwards every published event to the connected "
    "browser client as an SSE message. Events include node start, node completion, errors, timing information, "
    "and a terminal __STREAM_END__ sentinel that signals the client to close the connection. The chat endpoint "
    "constructs a system prompt from the agent's metadata and memory context, sends the conversation to "
    "OpenAI's streaming API, and forwards each token as an SSE event. Both endpoints use the sse-starlette "
    "library's EventSourceResponse for HTTP streaming."
))

p6_table = doc.add_table(rows=7, cols=2)
p6_table.style = "Table Grid"
add_purple_header_row(p6_table, ["Aspect", "Detail"])

p6_data = [
    ("Process ID", "P6"),
    ("Name", "Real-time Streaming (SSE)"),
    ("Input", "Execution streaming: execution_id from URL path, log events from DS7 pub/sub. Chat streaming: agent_id, user message, conversation history [{role, content}], agent memory context from DS4"),
    ("Processing", "Execution: creates Redis pub/sub subscription on channel execution:{execution_id}:logs. Listens for incoming messages. Decodes each message and yields as SSE data event. Closes on __STREAM_END__ sentinel or client disconnect. Chat: loads agent from DS2, fetches memory context (last 15 summaries) from DS4, constructs system prompt with agent identity and memory, appends conversation history (last 10 messages), calls OpenAI streaming API, yields each token chunk as SSE event with type=token, yields type=done on completion."),
    ("Output", "SSE event stream to E1: for execution: JSON events with timestamp, message, node_id, status; for chat: JSON events with type (token/done/error) and content"),
    ("Data Stores Used", "DS7: Redis Pub/Sub (subscribe to execution channels), DS4: Agent Memory (read for chat context), DS2: Agents (read agent metadata for chat)"),
]

for i, (aspect, detail) in enumerate(p6_data):
    row = p6_table.rows[i + 1]
    for j, val in enumerate([aspect, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in p6_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.5)

doc.add_paragraph()

# P7
doc.add_heading("5.7 P7: Dashboard & Analytics Aggregation", level=2)

add_body_paragraph(doc, (
    "The Dashboard & Analytics Aggregation process computes summary statistics and activity feeds from "
    "execution history. It serves two endpoints: /api/dashboard/stats for aggregate metrics and "
    "/api/dashboard/activity for a chronological feed of recent executions. The statistics include total "
    "agent count, active agent count, completed task count, estimated cost savings (calculated as "
    "tasks_completed * 0.5 hours * $50/hr), and average response time in milliseconds. The activity feed "
    "returns the 20 most recent executions with associated agent metadata (name, type), execution status, "
    "trigger method, timing, and cost information."
))

p7_table = doc.add_table(rows=7, cols=2)
p7_table.style = "Table Grid"
add_purple_header_row(p7_table, ["Aspect", "Detail"])

p7_data = [
    ("Process ID", "P7"),
    ("Name", "Dashboard & Analytics Aggregation"),
    ("Input", "Authenticated user context from P1; aggregation queries against DS3 filtered by user ownership"),
    ("Processing", "Stats: COUNT(agents) and COUNT(agents WHERE status=active) from agents table. COUNT(executions WHERE status=success) and AVG(duration_ms) from executions table, filtered to agents owned by current user. Calculates estimated_savings = tasks_completed * 0.5 * 50. Activity: SELECT executions JOIN agents WHERE agent.user_id = current_user ORDER BY created_at DESC LIMIT 20. Maps each row to ActivityEvent with execution_id, agent_name, agent_type, status, triggered_by, timing, cost."),
    ("Output", "DashboardStats (total_agents, active_agents, tasks_completed, estimated_savings, avg_response_time) and ActivityEvent list to E1"),
    ("Data Stores Used", "DS3: Executions & Logs (read for aggregation), DS2: Agents (read for ownership filtering and metadata join)"),
]

for i, (aspect, detail) in enumerate(p7_data):
    row = p7_table.rows[i + 1]
    for j, val in enumerate([aspect, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in p7_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.5)


# ===================================================================
#  6. DATA STORE CATALOG
# ===================================================================

doc.add_page_break()

doc.add_heading("6. Data Store Catalog", level=1)

add_body_paragraph(doc, (
    "This section provides a comprehensive catalog of all seven data stores used by the AutoMind platform. "
    "Each data store is documented with its technology, contents, and the processes that read from and write "
    "to it. The catalog serves as a reference for understanding data ownership, access patterns, and the "
    "physical storage technologies that back each logical store."
))

ds_cat_table = doc.add_table(rows=1, cols=6)
ds_cat_table.style = "Table Grid"
ds_cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(ds_cat_table, ["Store ID", "Name", "Technology", "Contents", "Read By", "Written By"])

ds_cat_data = [
    ("DS1", "Users", "PostgreSQL", "User accounts: id (UUID), email (unique), password_hash (bcrypt), name, created_at, updated_at", "P1 (login, token validation)", "P1 (registration)"),
    ("DS2", "Agents & Workflows", "PostgreSQL + JSONB", "Agents: id, user_id, name, type, status, schedule_cron, schedule_timezone, last_execution_at. Workflows: id, agent_id (unique FK), status, definition (JSONB with nodes + edges), deployed_at", "P2 (CRUD display), P3 (load workflow for execution), P6 (agent metadata for chat), P7 (ownership filtering)", "P2 (create/update agents and workflows), P3 (update last_execution_at)"),
    ("DS3", "Executions & Node Logs", "PostgreSQL", "Executions: id, agent_id, workflow_id, status, triggered_by, started_at, ended_at, duration_ms, error_message, variables (JSONB), total_cost. Node Logs: id, execution_id, node_id, node_type, node_label, status, started_at, ended_at, duration_ms, input_data (JSONB), output_data (JSONB), error_message, llm_usage (JSONB)", "P3 (status updates), P7 (aggregation queries)", "P3 (execution records, node logs)"),
    ("DS4", "Agent Memory", "PostgreSQL", "Memory entries: id, agent_id, execution_id, summary (text), key_outputs (JSONB), memory_type, created_at. Summaries generated by GPT-4o-mini.", "P3 (memory context for AI prompts, last 10), P6 (memory context for chat, last 15)", "P3 (save_execution_memory after each run)"),
    ("DS5", "Templates", "PostgreSQL", "Agent templates: id, name, description, type, workflow_definition (JSONB), icon, color, features (JSONB array), created_at. Pre-seeded catalog of reusable agent blueprints.", "P2 (template selection for agent creation)", "Admin seeding (not via application processes)"),
    ("DS6", "Integrations Config", "PostgreSQL", "Integration records: id, user_id, service (unique per user), config (JSONB with API keys, webhook URLs), status, created_at, updated_at", "P5 (read API keys and webhook URLs for delivery)", "P2 (user configures integrations via settings)"),
    ("DS7", "Task Queue & Pub/Sub", "Redis", "Celery task queue: pending workflow execution tasks. RedBeat schedule entries: cron-based agent triggers. Pub/sub channels: execution:{id}:logs for real-time event streaming.", "P3 (dequeue tasks), P6 (subscribe to pub/sub channels), E6 (RedBeat reads schedule entries)", "P2 (write schedule entries), P3 (publish log events, enqueue tasks), P3 via Celery (write task results)"),
]

for i, row_data in enumerate(ds_cat_data):
    row = ds_cat_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in ds_cat_table.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(0.9)
    row.cells[3].width = Inches(2.2)
    row.cells[4].width = Inches(1.4)
    row.cells[5].width = Inches(1.2)

add_body_paragraph(doc, (
    "The data store catalog highlights that PostgreSQL serves as the single source of truth for all durable "
    "platform state, while Redis provides the ephemeral messaging and scheduling infrastructure. The JSONB "
    "columns in DS2 (workflow definitions), DS3 (node I/O data and LLM usage), DS4 (key outputs), and DS6 "
    "(integration config) enable flexible, schema-on-read storage of complex nested structures without "
    "requiring rigid relational schemas for every data variant."
))

add_body_paragraph(doc, (
    "A key design decision is that DS7 (Redis) never stores data that cannot be reconstructed. Task queue "
    "entries are transient: if Redis is flushed, pending executions are lost but can be re-triggered. "
    "Pub/sub messages are fire-and-forget: if no subscriber is listening, events are dropped without "
    "consequence because the canonical execution state is persisted in DS3. RedBeat schedule entries "
    "can be reconstructed from the schedule_cron fields stored in DS2."
))


# ===================================================================
#  7. EXTERNAL ENTITY CATALOG
# ===================================================================

doc.add_page_break()

doc.add_heading("7. External Entity Catalog", level=1)

add_body_paragraph(doc, (
    "This section catalogs all six external entities that interact with the AutoMind platform. External "
    "entities exist outside the system boundary and are not under the direct control of the platform. They "
    "either produce data that enters the system, consume data that the system outputs, or both."
))

ee_table = doc.add_table(rows=1, cols=5)
ee_table.style = "Table Grid"
ee_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(ee_table, ["Entity ID", "Name", "Description", "Data Sent to Platform", "Data Received from Platform"])

ee_data = [
    ("E1", "Platform User", "A human user who interacts with the AutoMind web application via a browser. Users create agents, build workflows using the visual editor, trigger executions, monitor progress through real-time consoles, view dashboard analytics, and chat with agents about their execution history.", "Login/registration credentials, agent configurations (name, type, description, schedule), workflow definitions (nodes + edges from visual builder), manual execution triggers, chat messages with conversation history", "JWT authentication tokens, user profile data, agent lists and details, workflow definitions for builder rendering, execution results and status, real-time SSE events (console logs, chat tokens), dashboard statistics (agent counts, task completions, savings, response times), activity feeds"),
    ("E2", "OpenAI / Anthropic", "Cloud-hosted large language model API providers. OpenAI provides GPT-4o, GPT-4o-mini, GPT-4.1 family, and o3-mini models. Anthropic provides Claude Sonnet 4 and Claude Haiku 4.5 models. The platform selects the provider based on the model name configured in each AI action node.", "API requests containing: messages array (system prompt + user prompt), model identifier, max_tokens limit, temperature parameter. Sent via HTTPS REST API with Bearer token authentication.", "API responses containing: generated text content (in choices[0].message.content for OpenAI or content[0].text for Anthropic), input token count, output token count, finish reason. Used for cost calculation and output parsing."),
    ("E3", "Email Service (Resend)", "The Resend email delivery API, used to send transactional and automated emails from integration nodes. Resend provides reliable email delivery with tracking and analytics. Communication is outbound only.", "None (outbound-only integration)", "HTTP POST requests to https://api.resend.com/emails containing: from address (default automind@resend.dev), to recipients (list of email addresses), subject line, HTML body content. Authenticated with Bearer API key."),
    ("E4", "Slack", "Slack messaging platform, integrated via incoming webhook URLs. Used to deliver notifications, summaries, and automated messages from integration nodes to designated Slack channels. Communication is outbound only.", "None (outbound-only integration)", "HTTP POST requests to configured webhook URL containing JSON payload: {text: message_content}. No authentication required beyond the webhook URL itself."),
    ("E5", "Web (DuckDuckGo)", "DuckDuckGo web search engine, used as a web search data source by web_search nodes. The platform scrapes DuckDuckGo's HTML search results page directly (no API key required). Provides web intelligence to agents.", "Search results parsed from HTML response: array of result objects, each containing title, URL, snippet text, and rank position", "HTTP POST requests to https://html.duckduckgo.com/html/ with form data: query string and pagination parameter. Custom User-Agent header (Mozilla/5.0 compatible; AutoMind/1.0)."),
    ("E6", "Cron Scheduler", "The RedBeat scheduler running within Celery Beat, backed by Redis. Reads schedule entries created by P2 and dispatches Celery tasks at the configured cron intervals. Acts as an autonomous trigger source for scheduled agent executions.", "Scheduled trigger events dispatched as Celery tasks: execute_workflow_scheduled with args=[agent_id]. Triggered at intervals defined by the agent's cron expression and timezone.", "Schedule entries written by P2: RedBeat entries containing agent_id, parsed crontab schedule, and task name. Stored in Redis with the key pattern redbeat:agent:{agent_id}."),
]

for i, row_data in enumerate(ee_data):
    row = ee_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in ee_table.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(2.0)
    row.cells[3].width = Inches(2.0)
    row.cells[4].width = Inches(2.0)

add_body_paragraph(doc, (
    "The external entity catalog reveals a clear pattern in AutoMind's integration model. The platform has "
    "one bidirectional human actor (E1) who both sends and receives substantial amounts of data. It has one "
    "bidirectional machine actor (E2, the LLM providers) that processes prompts and returns generated content. "
    "It has two unidirectional outbound integrations (E3 Email and E4 Slack) that only receive delivery "
    "instructions. It has one bidirectional data source (E5 Web Search) that accepts queries and returns "
    "structured results. And it has one autonomous trigger source (E6 Cron Scheduler) that generates "
    "time-based execution events without human intervention."
))

add_body_paragraph(doc, (
    "This pattern reflects the platform's design philosophy: AutoMind is a workflow orchestration system that "
    "gathers intelligence (from AI models and web search), processes and transforms it (through the execution "
    "engine), delivers results (via email, Slack, and dashboards), and accumulates institutional knowledge "
    "(through agent memory). The external entities define the boundaries of this data lifecycle."
))


# ===================================================================
#  FOOTER / END
# ===================================================================

doc.add_page_break()

doc.add_heading("End of Document", level=1)

add_body_paragraph(doc, (
    "This Data Flow Diagram document provides a complete mapping of how data moves through the GoalCert "
    "AutoMind platform. The Level 0 context diagram establishes the system boundary and identifies six "
    "external entities. The Level 1 diagram decomposes the platform into seven processes and seven data "
    "stores, connected by 35 documented data flows. Process descriptions detail the transformation logic "
    "within each process. The data store catalog specifies the technology, contents, and access patterns "
    "of each storage component. The external entity catalog defines every actor and system that interacts "
    "with the platform."
))

add_body_paragraph(doc, (
    "This document should be read in conjunction with the High-Level Design (HLD), Low-Level Design (LLD), "
    "Business Requirements Document (BRD), and Functional Requirements Document (FRD) for a complete "
    "understanding of the AutoMind platform architecture."
))


# ===================================================================
#  SAVE
# ===================================================================

output_path = "/Users/narhen/automind/docs/GoalCert_AutoMind_DFD_v1_0.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Document saved to: {output_path}")
