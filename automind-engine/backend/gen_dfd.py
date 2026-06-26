"""
Generate GoalCert AutoMind Data Flow Diagrams (DFD) document as a professional .docx file.
Matches Flex Coach DFD format exactly.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# -- Color Constants (Flex Coach DFD palette) --
NAVY = RGBColor(0x1B, 0x2A, 0x4A)       # #1B2A4A - headings, header fills
BLUE = RGBColor(0x2E, 0x74, 0xB5)        # #2E74B5 - subtitle, heading 2
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
META_LEFT_BG = "E8EEF4"                   # metadata table left col
ALT_ROW_BG = "F0F4F8"                     # alternating table rows
HEADER_FILL = "1B2A4A"                    # table header fill
EXT_ENTITY_BG = "D6E4F0"                  # DFD external entity
PROCESS_BG = "E2EFDA"                     # DFD process
DATA_STORE_BG = "FFF2CC"                  # DFD data store
FONT_NAME = "Calibri"
OUTPUT_DIR = "/Users/narhen/automind/docs"


# ============================================================================
# Utility functions
# ============================================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def fmt_run(run, font=FONT_NAME, size=Pt(11), bold=False, color=BLACK, italic=False):
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_para(doc, text, size=Pt(11), bold=False, color=BLACK, align=None,
             space_before=0, space_after=Pt(6), italic=False, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    fmt_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.5 + indent_level * 1.0)
    for r in p.runs:
        r.clear()
    run = p.add_run(text)
    fmt_run(run, size=Pt(11))
    return p


# ============================================================================
# Cover page
# ============================================================================

def add_cover_page(doc):
    for _ in range(5):
        doc.add_paragraph()

    # "GoalCert AutoMind" - Calibri 32pt bold #1B2A4A
    add_para(doc, "GoalCert AutoMind", size=Pt(32), bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))

    # "Data Flow Diagrams (DFD)" - Calibri 20pt #2E74B5
    add_para(doc, "Data Flow Diagrams (DFD)", size=Pt(20), color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

    # Horizontal rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="1B2A4A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    doc.add_paragraph()

    add_para(doc, "Agentic AI Workforce Platform", size=Pt(14),
             color=RGBColor(100, 100, 100), align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=Pt(8), italic=True)

    add_para(doc, "Version 1.0  |  25 June 2026", size=Pt(11),
             color=RGBColor(120, 120, 120), align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=Pt(60))

    # Confidentiality
    add_para(doc, "CONFIDENTIAL", size=Pt(12), bold=True,
             color=RGBColor(180, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=Pt(4))
    add_para(doc,
             "This document contains proprietary information belonging to GoalCert Pte. Ltd. "
             "Unauthorised distribution, reproduction, or use of this document is strictly prohibited.",
             size=Pt(9), color=RGBColor(120, 120, 120),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))

    doc.add_page_break()


# ============================================================================
# Metadata table (left col fill #E8EEF4, right col white)
# ============================================================================

def add_metadata_table(doc):
    add_para(doc, "Document Information", size=Pt(20), bold=True, color=NAVY,
             space_after=Pt(12))

    rows_data = [
        ("Document Title", "Data Flow Diagrams (DFD)"),
        ("Product", "GoalCert AutoMind"),
        ("Version", "1.0"),
        ("Date", "25 June 2026"),
        ("Classification", "Confidential"),
        ("Prepared By", "AutoMind Engineering Team"),
        ("Reviewed By", "Prem Kumar"),
        ("Approved By", "Prem Kumar"),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, (key, value) in enumerate(rows_data):
        cell_key = table.rows[r_idx].cells[0]
        cell_val = table.rows[r_idx].cells[1]

        cell_key.text = ""
        p = cell_key.paragraphs[0]
        run = p.add_run(key)
        fmt_run(run, bold=True, size=Pt(10), color=NAVY)
        set_cell_shading(cell_key, META_LEFT_BG)

        cell_val.text = ""
        p = cell_val.paragraphs[0]
        run = p.add_run(value)
        fmt_run(run, size=Pt(10))

        cell_key.width = Cm(5)
        cell_val.width = Cm(12)

    doc.add_paragraph()


# ============================================================================
# Revision history (header fill #1B2A4A white text)
# ============================================================================

def add_revision_history(doc):
    add_para(doc, "Revision History", size=Pt(20), bold=True, color=NAVY,
             space_after=Pt(12))

    headers = ["Version", "Date", "Author", "Changes"]
    rows_data = [
        ["0.1", "12 June 2026", "Engineering Team", "Initial draft with Level 0 and Level 1 DFDs"],
        ["0.5", "18 June 2026", "Engineering Team", "Added Level 2 decompositions and data flow catalogs"],
        ["1.0", "25 June 2026", "Engineering Team", "Final review and release"],
    ]
    _create_std_table(doc, headers, rows_data)
    doc.add_page_break()


# ============================================================================
# Standard table helper (header fill #1B2A4A white bold, alt rows #F0F4F8)
# ============================================================================

def _create_std_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        fmt_run(run, bold=True, color=WHITE, size=Pt(10))
        set_cell_shading(cell, HEADER_FILL)

    # Data rows with alternating shading
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            fmt_run(run, size=Pt(10))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    doc.add_paragraph()
    return table


# ============================================================================
# DFD diagram table helper (colored cells for DFD elements)
# ============================================================================

def _create_dfd_table(doc, rows_data, col_count, col_widths=None):
    """Create a diagram table where each cell has a bg color and text.
    rows_data = list of lists of (text, bg_color_hex_or_None, text_color, bold, align)
    """
    table = doc.add_table(rows=len(rows_data), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_spec in enumerate(row):
            if c_idx >= col_count:
                break
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""

            text = cell_spec[0]
            bg = cell_spec[1] if len(cell_spec) > 1 else None
            txt_color = cell_spec[2] if len(cell_spec) > 2 else BLACK
            is_bold = cell_spec[3] if len(cell_spec) > 3 else False
            alignment = cell_spec[4] if len(cell_spec) > 4 else WD_ALIGN_PARAGRAPH.CENTER

            if bg:
                set_cell_shading(cell, bg)

            p = cell.paragraphs[0]
            p.alignment = alignment
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            if text:
                run = p.add_run(text)
                fmt_run(run, size=Pt(9), bold=is_bold, color=txt_color)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    doc.add_paragraph()
    return table


# ============================================================================
# Table of Contents (Manual Normal paragraphs)
# ============================================================================

def add_toc(doc):
    add_para(doc, "Table of Contents", size=Pt(20), bold=True, color=NAVY,
             space_after=Pt(16))

    toc_items = [
        (1, "1. Introduction"),
        (1, "2. DFD Notation Guide"),
        (1, "3. DFD Level 0 \u2013 Context Diagram"),
        (1, "4. DFD Level 1 \u2013 Process Decomposition"),
        (1, "5. DFD Level 2 \u2013 Sub-Process Detail"),
        (2, "5.1 Process 3.0: Workflow Execution (Detailed)"),
        (2, "5.2 Process 5.0: Chat & Memory (Detailed)"),
        (1, "6. Data Store Catalog"),
        (1, "7. Data Flow Catalog"),
    ]

    for level, text in toc_items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if level == 1:
            p.paragraph_format.left_indent = Cm(0)
            run = p.add_run(text)
            fmt_run(run, size=Pt(11), bold=True, color=NAVY)
        else:
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(text)
            fmt_run(run, size=Pt(11), bold=False, color=NAVY)

    doc.add_page_break()


# ============================================================================
# Heading helpers
# ============================================================================

def h1(doc, text):
    """Heading 1: 20pt bold #1B2A4A"""
    return add_para(doc, text, size=Pt(20), bold=True, color=NAVY,
                    space_before=Pt(18), space_after=Pt(10))

def h2(doc, text):
    """Heading 2: 15pt bold #2E74B5"""
    return add_para(doc, text, size=Pt(15), bold=True, color=BLUE,
                    space_before=Pt(12), space_after=Pt(6))

def body(doc, text, space_after=Pt(6)):
    """Body: Calibri 11pt"""
    return add_para(doc, text, size=Pt(11), color=BLACK, space_after=space_after)


# ============================================================================
# DFD Document Sections
# ============================================================================

def section_1_introduction(doc):
    h1(doc, "1. Introduction")

    body(doc,
         "This document presents the Data Flow Diagrams (DFDs) for the GoalCert AutoMind "
         "Agentic AI Platform. DFDs provide a structured visual representation of how data "
         "moves through the system, from external entities through internal processes to data "
         "stores. The diagrams are presented at three levels of abstraction: Context (Level 0), "
         "Process Decomposition (Level 1), and Sub-Process Detail (Level 2).")

    body(doc,
         "The purpose of this document is to communicate the data flow architecture of AutoMind "
         "to stakeholders, developers, and auditors. AutoMind enables users to create, configure, "
         "and deploy autonomous AI agents through a visual workflow builder. Each agent executes "
         "multi-step workflows comprising AI inference, web search, code execution, conditional "
         "logic, and external integrations \u2014 all orchestrated by a distributed task engine "
         "built on FastAPI, Celery, Redis, and PostgreSQL.")

    h2(doc, "1.1 Scope")
    body(doc,
         "This DFD document covers all data flows within the AutoMind platform, including user "
         "authentication, agent and workflow management, workflow execution orchestration, AI "
         "inference, web search, real-time streaming, chat with memory, scheduling, dashboard "
         "analytics, and integration services. It traces data from external entities (users, "
         "third-party APIs) through internal processes to persistent data stores.")

    h2(doc, "1.2 Intended Audience")
    add_bullet(doc, "Software architects and developers implementing or maintaining the platform")
    add_bullet(doc, "Quality assurance engineers validating data flow correctness")
    add_bullet(doc, "Security auditors reviewing data handling and access patterns")
    add_bullet(doc, "Project stakeholders requiring a high-level understanding of system data movement")

    doc.add_page_break()


def section_2_notation_guide(doc):
    h1(doc, "2. DFD Notation Guide")

    body(doc,
         "The following notation is used throughout this document to represent data flow diagram "
         "elements. Table-based diagrams use coloured cells to distinguish between element types.")

    # Notation table
    headers = ["Symbol", "Name", "Description", "Colour Code"]
    rows = [
        ["Rectangle", "External Entity",
         "An entity outside the system boundary that produces or consumes data (e.g., User, external API).",
         "#D6E4F0 (Light Blue)"],
        ["Rounded Rectangle", "Process",
         "A process or function within the system that transforms input data into output data.",
         "#E2EFDA (Light Green)"],
        ["Open-ended Rectangle", "Data Store",
         "A repository where data is stored persistently (e.g., PostgreSQL table, Redis key).",
         "#FFF2CC (Light Yellow)"],
        ["Arrow (\u2192)", "Data Flow",
         "The direction and path of data movement between elements. Labelled with the data being transferred.",
         "White cell with \u2192 text"],
    ]
    _create_std_table(doc, headers, rows)

    h2(doc, "2.1 Colour Legend for Diagram Tables")

    # Color legend as a small visual table
    legend_table = doc.add_table(rows=4, cols=2)
    legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    legend_items = [
        (EXT_ENTITY_BG, "External Entity"),
        (PROCESS_BG, "Process"),
        (DATA_STORE_BG, "Data Store"),
        (None, "Data Flow (arrow \u2192)"),
    ]
    for i, (bg, label) in enumerate(legend_items):
        cell_color = legend_table.rows[i].cells[0]
        cell_label = legend_table.rows[i].cells[1]

        cell_color.text = ""
        p = cell_color.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bg:
            set_cell_shading(cell_color, bg)
            run = p.add_run("  ")
        else:
            run = p.add_run("\u2192")
        fmt_run(run, size=Pt(10))

        cell_label.text = ""
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        fmt_run(run, size=Pt(10))

        cell_color.width = Cm(3)
        cell_label.width = Cm(10)

    doc.add_paragraph()
    doc.add_page_break()


def section_3_level0(doc):
    h1(doc, "3. DFD Level 0 \u2013 Context Diagram")

    body(doc,
         "The Context Diagram (Level 0) shows the AutoMind system as a single process "
         "interacting with its external entities. It establishes the system boundary and "
         "identifies all external data sources and sinks. The central process represents the "
         "entire AutoMind platform; external entities include the end user, OpenAI API, "
         "Anthropic API, DuckDuckGo search, Resend email service, and Slack webhook.")

    h2(doc, "3.1 Context Diagram")

    # Build the Level 0 diagram as a table
    # 5 columns: Entity | Arrow | SYSTEM | Arrow | Entity
    E = EXT_ENTITY_BG
    P = PROCESS_BG
    W = WHITE
    N = NAVY

    col_w = [Cm(3.2), Cm(2.2), Cm(4.5), Cm(2.2), Cm(3.2)]

    diagram_rows = [
        # Row 0: User -> system
        [("User\n(Web Browser)", E, BLACK, True),
         ("Credentials,\nAgent Config,\nWorkflow Def,\nChat Messages\n\u2192", None, BLACK, False),
         ("", None, WHITE, False),
         ("", None, WHITE, False),
         ("", None, WHITE, False)],
        # Row 1: Empty -> AUTOMIND SYSTEM -> Empty
        [("", None, WHITE, False),
         ("\u2192", None, BLACK, False),
         ("AutoMind\nSystem\n(0.0)", P, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("", None, WHITE, False)],
        # Row 2: OpenAI API
        [("OpenAI API", E, BLACK, True),
         ("Prompts \u2192\n\u2190 Completions", None, BLACK, False),
         ("", None, WHITE, False),
         ("JWT, Agent Data,\nExecution Results,\nSSE Streams\n\u2192", None, BLACK, False),
         ("User\n(Web Browser)", E, BLACK, True)],
        # Row 3: Anthropic API
        [("Anthropic API", E, BLACK, True),
         ("Prompts \u2192\n\u2190 Completions", None, BLACK, False),
         ("", None, WHITE, False),
         ("Emails \u2192", None, BLACK, False),
         ("Resend\n(Email API)", E, BLACK, True)],
        # Row 4: DuckDuckGo
        [("DuckDuckGo\n(Web Search)", E, BLACK, True),
         ("Queries \u2192\n\u2190 Results", None, BLACK, False),
         ("", None, WHITE, False),
         ("Messages \u2192", None, BLACK, False),
         ("Slack\n(Webhook)", E, BLACK, True)],
    ]

    _create_dfd_table(doc, diagram_rows, 5, col_w)

    h2(doc, "3.2 Level 0 Data Flow Catalog")

    headers = ["Flow ID", "Flow Name", "Source", "Destination", "Data Elements"]
    rows = [
        ["DF-0.1", "User Authentication", "User", "AutoMind System",
         "Email, password, registration details"],
        ["DF-0.2", "Auth Response", "AutoMind System", "User",
         "JWT token, user profile"],
        ["DF-0.3", "Agent Configuration", "User", "AutoMind System",
         "Agent name, type, description, schedule, template selection"],
        ["DF-0.4", "Workflow Definition", "User", "AutoMind System",
         "React Flow JSON (nodes, edges, viewport)"],
        ["DF-0.5", "Execution Trigger", "User", "AutoMind System",
         "Agent ID, manual trigger command"],
        ["DF-0.6", "Execution Results", "AutoMind System", "User",
         "Status, duration, variables, cost, node logs"],
        ["DF-0.7", "SSE Stream", "AutoMind System", "User",
         "Real-time execution log events (timestamp, message, node_id, status)"],
        ["DF-0.8", "Chat Messages", "User", "AutoMind System",
         "Message text, conversation history"],
        ["DF-0.9", "Chat Response Stream", "AutoMind System", "User",
         "SSE token stream (type: token/done/error)"],
        ["DF-0.10", "LLM Prompt", "AutoMind System", "OpenAI API",
         "System prompt, user prompt, model, max_tokens, temperature"],
        ["DF-0.11", "LLM Completion", "OpenAI API", "AutoMind System",
         "Response text, token counts (input/output)"],
        ["DF-0.12", "LLM Prompt (Anthropic)", "AutoMind System", "Anthropic API",
         "System prompt, messages, model, max_tokens, temperature"],
        ["DF-0.13", "LLM Completion (Anthropic)", "Anthropic API", "AutoMind System",
         "Response text blocks, token counts"],
        ["DF-0.14", "Search Query", "AutoMind System", "DuckDuckGo",
         "Search query string, max_results parameter"],
        ["DF-0.15", "Search Results", "DuckDuckGo", "AutoMind System",
         "HTML response parsed into title, URL, snippet, rank"],
        ["DF-0.16", "Email Request", "AutoMind System", "Resend",
         "From, to, subject, HTML body"],
        ["DF-0.17", "Slack Message", "AutoMind System", "Slack",
         "Webhook URL, message text payload"],
        ["DF-0.18", "Dashboard Request", "User", "AutoMind System",
         "Authenticated stats/activity query"],
        ["DF-0.19", "Dashboard Data", "AutoMind System", "User",
         "Total agents, active agents, tasks completed, savings, activity feed"],
    ]
    _create_std_table(doc, headers, rows)
    doc.add_page_break()


def section_4_level1(doc):
    h1(doc, "4. DFD Level 1 \u2013 Process Decomposition")

    body(doc,
         "Level 1 decomposes the single AutoMind System process from Level 0 into six major "
         "sub-processes. Each process has well-defined inputs and outputs, and communicates "
         "with specific data stores. The processes correspond to the major functional domains "
         "of the platform: authentication, agent management, workflow execution, scheduling, "
         "chat with memory, and analytics.")

    h2(doc, "4.1 Level 1 Diagram")

    # Level 1 diagram table - show processes and their connections
    E = EXT_ENTITY_BG
    P = PROCESS_BG
    D = DATA_STORE_BG

    col_w = [Cm(3.5), Cm(2.0), Cm(3.5), Cm(2.0), Cm(3.5)]

    diagram_rows = [
        # Row 0: Header row
        [("EXTERNAL ENTITIES", HEADER_FILL, WHITE, True),
         ("", None, WHITE, False),
         ("PROCESSES", HEADER_FILL, WHITE, True),
         ("", None, WHITE, False),
         ("DATA STORES", HEADER_FILL, WHITE, True)],
        # Row 1: User -> Auth
        [("User\n(Web Browser)", E, BLACK, True),
         ("Credentials\n\u2192", None, BLACK, False),
         ("1.0\nAuthentication", P, BLACK, True),
         ("\u2192\nUser Record", None, BLACK, False),
         ("D1: users\n(PostgreSQL)", D, BLACK, True)],
        # Row 2: User -> Agent Mgmt
        [("", None, WHITE, False),
         ("Agent Config\n\u2192", None, BLACK, False),
         ("2.0\nAgent\nManagement", P, BLACK, True),
         ("\u2192\nAgent, Workflow", None, BLACK, False),
         ("D2: agents\nD3: workflows\n(PostgreSQL)", D, BLACK, True)],
        # Row 3: OpenAI -> Workflow Exec
        [("OpenAI API\nAnthropic API\nDuckDuckGo", E, BLACK, True),
         ("Prompts,\nQueries\n\u2190 \u2192", None, BLACK, False),
         ("3.0\nWorkflow\nExecution", P, BLACK, True),
         ("\u2192\nExecution Logs,\nNode Logs", None, BLACK, False),
         ("D4: executions\nD5: exec_node_logs\n(PostgreSQL)", D, BLACK, True)],
        # Row 4: RedBeat -> Scheduling
        [("", None, WHITE, False),
         ("Cron Config\n\u2192", None, BLACK, False),
         ("4.0\nScheduling", P, BLACK, True),
         ("\u2192\nSchedule Entry", None, BLACK, False),
         ("D6: Redis\n(RedBeat)", D, BLACK, True)],
        # Row 5: User -> Chat
        [("", None, WHITE, False),
         ("Chat Message\n\u2192", None, BLACK, False),
         ("5.0\nChat &\nMemory", P, BLACK, True),
         ("\u2192\nMemory Record", None, BLACK, False),
         ("D7: agent_memory\n(PostgreSQL)", D, BLACK, True)],
        # Row 6: User -> Analytics
        [("", None, WHITE, False),
         ("Stats Query\n\u2192", None, BLACK, False),
         ("6.0\nAnalytics", P, BLACK, True),
         ("\u2190\nAggregated Data", None, BLACK, False),
         ("D4: executions\nD2: agents\n(PostgreSQL)", D, BLACK, True)],
    ]

    _create_dfd_table(doc, diagram_rows, 5, col_w)

    h2(doc, "4.2 Process Descriptions")

    process_descs = [
        ("1.0 Authentication",
         "Handles user registration, login, and JWT token issuance. Passwords are hashed "
         "with bcrypt (12 rounds) before storage. Validates credentials against the users "
         "table and returns a signed JWT (HS256, 24-hour expiry) containing user ID and email. "
         "The get_current_user dependency extracts and validates tokens on all protected endpoints."),
        ("2.0 Agent Management",
         "Manages the full lifecycle of AI agents: create, read, update, delete, pause, resume, "
         "and AI-powered generation from natural language descriptions. Each agent has a 1:1 "
         "relationship with a Workflow. Creating an agent also creates a default workflow. If "
         "a template_id is provided, the template workflow definition is copied. Agent names "
         "are unique per user. Pause/resume toggles the cron schedule via Process 4.0."),
        ("3.0 Workflow Execution",
         "The core orchestration engine. Receives a workflow definition and execution ID, "
         "performs BFS traversal of the node graph starting from the trigger node. Dispatches "
         "each node to the appropriate executor (ai_action, web_search, decision, integration, "
         "escalation, code_exec). Merges output variables into a shared context. Publishes "
         "real-time events via Redis pub/sub. Records per-node logs and cumulative LLM cost."),
        ("4.0 Scheduling",
         "Manages dynamic cron scheduling using RedBeat (Celery Beat backed by Redis). When "
         "an agent is deployed with a cron expression, creates a RedBeatSchedulerEntry that "
         "fires the execute_workflow_scheduled Celery task. Parses standard 5-field cron "
         "expressions. Unschedules entries on agent pause or delete."),
        ("5.0 Chat & Memory",
         "Enables conversational interaction with agents using their execution memory as "
         "context. Retrieves past execution summaries from agent_memory, constructs a system "
         "prompt with agent metadata and memory context, and streams GPT-4o-mini responses "
         "via SSE. Also stores execution memory summaries after each workflow run, optionally "
         "using AI summarisation."),
        ("6.0 Analytics",
         "Computes dashboard statistics by aggregating data from agents and executions tables. "
         "Provides total agents, active agents, completed tasks, estimated cost savings, and "
         "average response time. Returns the 20 most recent execution activity events with "
         "agent names, types, statuses, durations, and costs."),
    ]

    for title, desc in process_descs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{title}: ")
        fmt_run(run, bold=True, size=Pt(11), color=NAVY)
        run2 = p.add_run(desc)
        fmt_run(run2, size=Pt(11))

    doc.add_paragraph()

    h2(doc, "4.3 Level 1 Data Flow Catalog")

    headers = ["Flow ID", "Flow Name", "Source", "Destination", "Data Elements"]
    rows = [
        # Authentication flows
        ["DF-1.1", "Registration Request", "User", "1.0 Authentication",
         "Email, password, name"],
        ["DF-1.2", "Login Request", "User", "1.0 Authentication",
         "Email, password"],
        ["DF-1.3", "User Record Write", "1.0 Authentication", "D1: users",
         "User ID, email, password_hash, name, timestamps"],
        ["DF-1.4", "User Record Read", "D1: users", "1.0 Authentication",
         "Stored password_hash for verification"],
        ["DF-1.5", "Auth Response", "1.0 Authentication", "User",
         "JWT token, user profile (id, email, name)"],
        # Agent Management flows
        ["DF-2.1", "Agent CRUD", "User", "2.0 Agent Management",
         "Agent name, type, description, template_id, updates"],
        ["DF-2.2", "Agent Write", "2.0 Agent Management", "D2: agents",
         "Agent record (id, user_id, name, type, status, schedule)"],
        ["DF-2.3", "Workflow Write", "2.0 Agent Management", "D3: workflows",
         "Workflow definition JSON (nodes, edges, viewport)"],
        ["DF-2.4", "Template Read", "D8: agent_templates", "2.0 Agent Management",
         "Pre-built workflow definition, features, metadata"],
        ["DF-2.5", "Schedule Request", "2.0 Agent Management", "4.0 Scheduling",
         "Agent ID, cron expression, timezone"],
        ["DF-2.6", "Generation Prompt", "2.0 Agent Management", "OpenAI API",
         "Natural language description for agent generation"],
        ["DF-2.7", "Generated Workflow", "OpenAI API", "2.0 Agent Management",
         "Agent name, type, workflow definition JSON"],
        # Execution flows
        ["DF-3.1", "Execution Trigger", "User", "3.0 Workflow Execution",
         "Agent ID, trigger type (manual)"],
        ["DF-3.2", "Scheduled Trigger", "4.0 Scheduling", "3.0 Workflow Execution",
         "Agent ID, trigger type (schedule)"],
        ["DF-3.3", "Workflow Read", "D3: workflows", "3.0 Workflow Execution",
         "Workflow definition (nodes, edges)"],
        ["DF-3.4", "Execution Write", "3.0 Workflow Execution", "D4: executions",
         "Execution status, duration, cost, variables, error_message"],
        ["DF-3.5", "Node Log Write", "3.0 Workflow Execution", "D5: exec_node_logs",
         "Node ID, type, label, status, input/output data, LLM usage, duration"],
        ["DF-3.6", "SSE Events", "3.0 Workflow Execution", "D6: Redis",
         "Execution log events published to pub/sub channel"],
        ["DF-3.7", "SSE Stream", "D6: Redis", "User",
         "Real-time execution events streamed via SSE endpoint"],
        ["DF-3.8", "LLM Request", "3.0 Workflow Execution", "OpenAI/Anthropic API",
         "Prompt, system prompt, model, parameters"],
        ["DF-3.9", "LLM Response", "OpenAI/Anthropic API", "3.0 Workflow Execution",
         "Completion text, input/output token counts"],
        ["DF-3.10", "Search Request", "3.0 Workflow Execution", "DuckDuckGo",
         "Query string, max_results"],
        ["DF-3.11", "Search Response", "DuckDuckGo", "3.0 Workflow Execution",
         "Parsed results (title, URL, snippet, rank)"],
        ["DF-3.12", "Email Send", "3.0 Workflow Execution", "Resend API",
         "From, to, subject, HTML body"],
        ["DF-3.13", "Slack Send", "3.0 Workflow Execution", "Slack Webhook",
         "Webhook URL, message text"],
        ["DF-3.14", "Memory Save", "3.0 Workflow Execution", "5.0 Chat & Memory",
         "Execution summary, key outputs, agent ID"],
        # Scheduling flows
        ["DF-4.1", "Schedule Entry Write", "4.0 Scheduling", "D6: Redis",
         "RedBeat entry (task name, cron, args)"],
        ["DF-4.2", "Schedule Fire", "D6: Redis", "4.0 Scheduling",
         "Cron timer expiry notification"],
        # Chat & Memory flows
        ["DF-5.1", "Chat Request", "User", "5.0 Chat & Memory",
         "Message text, conversation history"],
        ["DF-5.2", "Memory Read", "D7: agent_memory", "5.0 Chat & Memory",
         "Past execution summaries (last 15)"],
        ["DF-5.3", "Memory Write", "5.0 Chat & Memory", "D7: agent_memory",
         "Execution summary, key outputs, memory_type"],
        ["DF-5.4", "Chat LLM Request", "5.0 Chat & Memory", "OpenAI API",
         "System prompt with memory, chat history, user message"],
        ["DF-5.5", "Chat SSE Stream", "5.0 Chat & Memory", "User",
         "Token-by-token response stream"],
        # Analytics flows
        ["DF-6.1", "Stats Query", "User", "6.0 Analytics",
         "Authenticated dashboard request"],
        ["DF-6.2", "Agent Data Read", "D2: agents", "6.0 Analytics",
         "Agent counts, statuses"],
        ["DF-6.3", "Execution Data Read", "D4: executions", "6.0 Analytics",
         "Success counts, durations, costs, activity events"],
        ["DF-6.4", "Dashboard Response", "6.0 Analytics", "User",
         "Stats (total, active, completed, savings, avg time), activity feed"],
    ]
    _create_std_table(doc, headers, rows)
    doc.add_page_break()


def section_5_level2(doc):
    h1(doc, "5. DFD Level 2 \u2013 Sub-Process Detail")

    body(doc,
         "Level 2 decomposes selected Level 1 processes into their constituent sub-processes. "
         "This level provides the detail needed for implementation and testing. Two key processes "
         "are decomposed: Process 3.0 (Workflow Execution) and Process 5.0 (Chat & Memory).")

    # ---- 5.1 Workflow Execution ----
    h2(doc, "5.1 Process 3.0: Workflow Execution (Detailed)")

    body(doc,
         "Process 3.0 is the most complex process in the system. It decomposes into six "
         "sub-processes that collectively parse, traverse, and execute a workflow graph. "
         "The WorkflowExecutor performs BFS traversal, dispatching each node to the appropriate "
         "sub-process based on node type. Output variables from each node are merged into a "
         "shared variable context accessible by all downstream nodes.")

    # Diagram for 5.1
    P = PROCESS_BG
    D = DATA_STORE_BG
    E = EXT_ENTITY_BG

    col_w = [Cm(3.2), Cm(2.0), Cm(3.2), Cm(2.0), Cm(3.2)]

    diagram_rows = [
        # Header
        [("INPUT", HEADER_FILL, WHITE, True),
         ("", None, WHITE, False),
         ("SUB-PROCESSES", HEADER_FILL, WHITE, True),
         ("", None, WHITE, False),
         ("OUTPUT / STORES", HEADER_FILL, WHITE, True)],
        # Parse Workflow
        [("Workflow\nDefinition\n(JSON)", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("3.1\nParse\nWorkflow", P, BLACK, True),
         ("\u2192\nNodes, Edges", None, BLACK, False),
         ("Node Map,\nEdge List\n(in memory)", D, BLACK, True)],
        # Execute Node
        [("Node Map\n(BFS Queue)", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("3.2\nExecute\nNode", P, BLACK, True),
         ("\u2192\nNode Log", None, BLACK, False),
         ("D5: exec_node_logs\n(PostgreSQL)", D, BLACK, True)],
        # AI Action
        [("Prompt Template\n+ Variables", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("3.3\nAI Action\nProcessing", P, BLACK, True),
         ("\u2192\nCompletion,\nToken Usage", None, BLACK, False),
         ("OpenAI API /\nAnthropic API", E, BLACK, True)],
        # Web Search
        [("Search Query\n+ Variables", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("3.4\nWeb Search", P, BLACK, True),
         ("\u2192\nSearch Results", None, BLACK, False),
         ("DuckDuckGo\n(HTML endpoint)", E, BLACK, True)],
        # Decision
        [("Condition Config\n(operands, op)", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("3.5\nDecision\nEvaluation", P, BLACK, True),
         ("\u2192\nbranch=true/false", None, BLACK, False),
         ("BFS Queue\n(filtered edges)", D, BLACK, True)],
        # Result Aggregation
        [("All Node Outputs\n+ Variables", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("3.6\nResult\nAggregation", P, BLACK, True),
         ("\u2192\nFinal Status,\nCost, Memory", None, BLACK, False),
         ("D4: executions\nD7: agent_memory", D, BLACK, True)],
    ]

    _create_dfd_table(doc, diagram_rows, 5, col_w)

    # Sub-process descriptions
    sub_procs = [
        ("3.1 Parse Workflow",
         "Receives the workflow definition (React Flow JSON) and extracts the nodes array and "
         "edges array. Builds a node_map (dict keyed by node ID) and identifies the start node "
         "by finding nodes with no incoming edges, preferring trigger-type nodes. Initialises "
         "the BFS queue with start node IDs."),
        ("3.2 Execute Node",
         "Dequeues a node from the BFS queue. Looks up the appropriate NodeExecutor from the "
         "executor registry based on node type. Creates a running node log entry in "
         "execution_node_logs. Calls the executor's execute() method with the node config, "
         "current variable context, trigger source, and memory context. On completion, updates "
         "the node log with status, output, duration, and LLM usage. Publishes SSE events via "
         "Redis pub/sub. Merges output variables into the shared context."),
        ("3.3 AI Action Processing",
         "Interpolates {variable} placeholders in the prompt and system_prompt templates. "
         "Prepends agent memory context to the system prompt. Routes to OpenAI or Anthropic "
         "based on model name prefix (gpt-*/o1*/o3*/o4* = OpenAI, otherwise Anthropic). "
         "Calls the provider API with configured parameters. Parses the response (attempts "
         "JSON extraction first, falls back to {output: raw_text}). Calculates cost from a "
         "per-model pricing table. Returns output_variables and llm_usage dicts."),
        ("3.4 Web Search",
         "Interpolates the query template against the variable context. Posts to DuckDuckGo's "
         "HTML endpoint (https://html.duckduckgo.com/html/) with the query. Parses the HTML "
         "response using regex to extract result links, titles, and snippets. Returns results "
         "as a list of dicts in the specified output_variable."),
        ("3.5 Decision Evaluation",
         "Interpolates left_operand and right_operand against the variable context. Converts "
         "to strings. Evaluates the operator (==, !=, >, <, >=, <=, contains). For numeric "
         "operators, attempts float conversion for numeric comparison with string fallback. "
         "Returns branch=true or branch=false. The get_next_nodes function filters outgoing "
         "edges by matching edge labels to the branch value, controlling BFS traversal."),
        ("3.6 Result Aggregation",
         "After BFS traversal completes (or on first node failure), aggregates results. "
         "Records final execution status (success/failed), total duration, cumulative LLM cost, "
         "error message (if any), and final variable context. Generates an execution memory "
         "summary (using GPT-4o-mini if available, otherwise a fallback template). Stores the "
         "memory in agent_memory. Updates agent.last_execution_at. Publishes __STREAM_END__ "
         "sentinel to Redis pub/sub to signal SSE stream closure."),
    ]

    for title, desc in sub_procs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{title}: ")
        fmt_run(run, bold=True, size=Pt(11), color=BLUE)
        run2 = p.add_run(desc)
        fmt_run(run2, size=Pt(11))

    doc.add_page_break()

    # ---- 5.2 Chat & Memory ----
    h2(doc, "5.2 Process 5.0: Chat & Memory (Detailed)")

    body(doc,
         "Process 5.0 decomposes into four sub-processes that handle memory retrieval, "
         "context construction, LLM response streaming, and memory storage. The chat feature "
         "allows users to converse with their agents, with responses grounded in the agent's "
         "actual execution history.")

    col_w2 = [Cm(3.2), Cm(2.0), Cm(3.2), Cm(2.0), Cm(3.2)]

    diagram_rows2 = [
        [("INPUT", HEADER_FILL, WHITE, True),
         ("", None, WHITE, False),
         ("SUB-PROCESSES", HEADER_FILL, WHITE, True),
         ("", None, WHITE, False),
         ("OUTPUT / STORES", HEADER_FILL, WHITE, True)],
        # Retrieve Memories
        [("Agent ID", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("5.1\nRetrieve\nMemories", P, BLACK, True),
         ("\u2190\nPast Summaries", None, BLACK, False),
         ("D7: agent_memory\n(PostgreSQL)", D, BLACK, True)],
        # Build Context
        [("Agent Metadata\n(name, type,\nstatus, schedule)", D, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("5.2\nBuild\nContext", P, BLACK, True),
         ("\u2192\nSystem Prompt\n+ Messages", None, BLACK, False),
         ("Constructed\nPrompt\n(in memory)", D, BLACK, True)],
        # Stream GPT Response
        [("User Message\n+ History", E, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("5.3\nStream GPT\nResponse", P, BLACK, True),
         ("\u2192\nToken Stream", None, BLACK, False),
         ("User\n(SSE Client)", E, BLACK, True)],
        # Store Summary
        [("Execution Results\n(from 3.6)", P, BLACK, True),
         ("\u2192", None, BLACK, False),
         ("5.4\nStore\nSummary", P, BLACK, True),
         ("\u2192\nMemory Record", None, BLACK, False),
         ("D7: agent_memory\n(PostgreSQL)", D, BLACK, True)],
    ]

    _create_dfd_table(doc, diagram_rows2, 5, col_w2)

    sub_procs2 = [
        ("5.1 Retrieve Memories",
         "Queries the agent_memory table for the specified agent_id, ordered by created_at "
         "descending, limited to 15 records (for chat) or 10 records (for execution context). "
         "Formats each memory as a timestamped line: [YYYY-MM-DD HH:MM] summary. Returns the "
         "formatted context string or empty string if no memories exist."),
        ("5.2 Build Context",
         "Constructs the system prompt by combining: agent name, type, description, status, "
         "schedule, and the memory context string from sub-process 5.1. Appends the last 10 "
         "messages from conversation history. Adds the new user message. The complete message "
         "list is ready for the LLM API call."),
        ("5.3 Stream GPT Response",
         "Calls OpenAI's chat.completions.create with model=gpt-4o-mini, stream=True, "
         "temperature=0.7, max_tokens=1024, and the constructed messages. Yields each content "
         "delta as an SSE event {type: 'token', content: text}. On completion, yields "
         "{type: 'done'}. On error, yields {type: 'error', content: message}. Checks "
         "request.is_disconnected() on each chunk to handle client disconnection."),
        ("5.4 Store Summary",
         "Called by Process 3.6 (Result Aggregation) after workflow execution. Receives "
         "agent_id, execution_id, agent_name, status, duration, node_count, and key_outputs. "
         "If an OpenAI API key is configured, generates a 2-3 sentence AI summary using "
         "GPT-4o-mini (temperature=0.3, max_tokens=200). Otherwise, creates a template-based "
         "summary. Stores the AgentMemory record with memory_type=execution_summary."),
    ]

    for title, desc in sub_procs2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{title}: ")
        fmt_run(run, bold=True, size=Pt(11), color=BLUE)
        run2 = p.add_run(desc)
        fmt_run(run2, size=Pt(11))

    doc.add_page_break()


def section_6_data_store_catalog(doc):
    h1(doc, "6. Data Store Catalog")

    body(doc,
         "The following table catalogs all data stores referenced in the DFD diagrams, including "
         "their descriptions and the processes that read from and write to each store.")

    headers = ["Store ID", "Name", "Description", "Read By", "Written By"]
    rows = [
        ["D1", "users",
         "User accounts with hashed passwords and profile information. PostgreSQL table with UUID PKs.",
         "1.0 Authentication (credential verification)",
         "1.0 Authentication (registration)"],
        ["D2", "agents",
         "AI agent definitions including name, type, status, schedule, and timestamps. "
         "Unique constraint on (user_id, name).",
         "2.0 Agent Management, 3.0 Workflow Execution (last_execution_at update), 6.0 Analytics",
         "2.0 Agent Management, 3.0 Workflow Execution (last_execution_at)"],
        ["D3", "workflows",
         "Workflow definitions stored as JSONB (React Flow nodes, edges, viewport). "
         "One-to-one relationship with agents.",
         "2.0 Agent Management, 3.0 Workflow Execution, 4.0 Scheduling",
         "2.0 Agent Management (create, update, deploy)"],
        ["D4", "executions",
         "Execution records tracking status, duration, cost, variables, and error messages. "
         "Linked to agent and workflow.",
         "3.0 Workflow Execution, 6.0 Analytics",
         "3.0 Workflow Execution"],
        ["D5", "execution_node_logs",
         "Per-node execution logs storing input/output data, duration, LLM usage, and status. "
         "Cascade-deleted with parent execution.",
         "3.0 Workflow Execution (node count aggregation), User (via API)",
         "3.0 Workflow Execution"],
        ["D6", "Redis",
         "Multi-purpose in-memory store: (a) Celery message broker for task dispatch, "
         "(b) pub/sub channels for real-time SSE streaming, (c) RedBeat storage for cron schedules.",
         "3.0 Workflow Execution (SSE subscription), 4.0 Scheduling (cron fire)",
         "3.0 Workflow Execution (SSE publish), 4.0 Scheduling (RedBeat entry)"],
        ["D7", "agent_memory",
         "Execution summaries and key outputs per agent. Used to build context for chat and "
         "AI action nodes. Stores AI-generated or template-based summaries.",
         "5.0 Chat & Memory (context retrieval), 3.0 Workflow Execution (memory context for AI nodes)",
         "5.0 Chat & Memory (store summary after execution)"],
        ["D8", "agent_templates",
         "Pre-built workflow templates with names, descriptions, types, workflow definitions, "
         "icons, colours, and feature tags. Read-only at runtime.",
         "2.0 Agent Management (template selection during agent creation)",
         "Seed scripts (initial data population)"],
        ["D9", "integrations",
         "User-configured integration credentials (Resend API keys, Slack webhook URLs). "
         "Sensitive fields masked in API responses. Unique constraint on (user_id, service).",
         "3.0 Workflow Execution (integration node config lookup)",
         "User (via Integrations API)"],
    ]
    _create_std_table(doc, headers, rows)
    doc.add_page_break()


def section_7_data_flow_catalog(doc):
    h1(doc, "7. Data Flow Catalog")

    body(doc,
         "The comprehensive data flow catalog below lists every data flow identified across "
         "all DFD levels. Each flow specifies its source, destination, data elements, and "
         "typical frequency of occurrence.")

    headers = ["Flow ID", "Name", "Source", "Destination", "Data Elements", "Frequency"]
    rows = [
        # Authentication
        ["DF-1.1", "Registration Request", "User", "1.0 Authentication",
         "Email, password, name", "On user sign-up"],
        ["DF-1.2", "Login Request", "User", "1.0 Authentication",
         "Email, password", "Per session start"],
        ["DF-1.3", "User Record Write", "1.0 Authentication", "D1: users",
         "ID, email, password_hash, name, timestamps", "On registration"],
        ["DF-1.4", "User Record Read", "D1: users", "1.0 Authentication",
         "Stored password_hash", "On login"],
        ["DF-1.5", "Auth Token Response", "1.0 Authentication", "User",
         "JWT (sub, email, exp), user profile", "On login/register"],

        # Agent Management
        ["DF-2.1", "Agent Create/Update", "User", "2.0 Agent Mgmt",
         "Name, type, description, template_id", "Per agent operation"],
        ["DF-2.2", "Agent Record Write", "2.0 Agent Mgmt", "D2: agents",
         "Agent fields (status, schedule, timestamps)", "Per CRUD operation"],
        ["DF-2.3", "Workflow Definition Write", "2.0 Agent Mgmt", "D3: workflows",
         "React Flow JSON (nodes, edges, viewport)", "Per workflow save/deploy"],
        ["DF-2.4", "Template Read", "D8: agent_templates", "2.0 Agent Mgmt",
         "Workflow definition, features", "On template-based creation"],
        ["DF-2.5", "Schedule Request", "2.0 Agent Mgmt", "4.0 Scheduling",
         "Agent ID, cron expression, timezone", "On deploy/pause/resume"],
        ["DF-2.6", "NL Generation Prompt", "2.0 Agent Mgmt", "OpenAI API",
         "Natural language description", "On /generate endpoint"],
        ["DF-2.7", "Generated Workflow", "OpenAI API", "2.0 Agent Mgmt",
         "Agent name, type, workflow JSON", "On /generate response"],

        # Workflow Execution
        ["DF-3.1", "Manual Trigger", "User", "3.0 Workflow Exec",
         "Agent ID, trigger=manual", "Per user-initiated run"],
        ["DF-3.2", "Scheduled Trigger", "4.0 Scheduling", "3.0 Workflow Exec",
         "Agent ID, trigger=schedule", "Per cron interval"],
        ["DF-3.3", "Workflow Definition Read", "D3: workflows", "3.0 Workflow Exec",
         "Nodes, edges, viewport JSON", "Per execution"],
        ["DF-3.4", "Execution Record Write", "3.0 Workflow Exec", "D4: executions",
         "Status, duration, cost, variables, error", "Per execution lifecycle"],
        ["DF-3.5", "Node Log Write", "3.0 Workflow Exec", "D5: exec_node_logs",
         "Node ID/type/label, status, I/O data, LLM usage", "Per node execution"],
        ["DF-3.6", "SSE Event Publish", "3.0 Workflow Exec", "D6: Redis",
         "JSON log: timestamp, message, node_id, status", "Per node state change"],
        ["DF-3.7", "SSE Event Stream", "D6: Redis", "User",
         "Execution events via EventSource", "Real-time during execution"],
        ["DF-3.8", "LLM Prompt", "3.3 AI Action", "OpenAI/Anthropic API",
         "System prompt, user prompt, model params", "Per AI action node"],
        ["DF-3.9", "LLM Completion", "OpenAI/Anthropic API", "3.3 AI Action",
         "Response text, input/output tokens", "Per AI action node"],
        ["DF-3.10", "Search Request", "3.4 Web Search", "DuckDuckGo",
         "Query string, max_results", "Per web search node"],
        ["DF-3.11", "Search Response", "DuckDuckGo", "3.4 Web Search",
         "HTML parsed to title, URL, snippet, rank", "Per web search node"],
        ["DF-3.12", "Email Dispatch", "3.0 Workflow Exec", "Resend API",
         "From, to[], subject, HTML body", "Per integration/escalation node"],
        ["DF-3.13", "Slack Dispatch", "3.0 Workflow Exec", "Slack Webhook",
         "Webhook URL, message text", "Per Slack integration node"],
        ["DF-3.14", "Memory Context Read", "D7: agent_memory", "3.0 Workflow Exec",
         "Past 10 execution summaries", "Per execution start"],
        ["DF-3.15", "Memory Write", "3.0 Workflow Exec", "D7: agent_memory",
         "Summary, key_outputs, execution_id", "Per execution completion"],
        ["DF-3.16", "Variable Merge", "3.2 Execute Node", "Shared Context",
         "Output variables from each node", "Per node completion"],
        ["DF-3.17", "Branch Decision", "3.5 Decision Eval", "3.2 Execute Node",
         "branch=true/false, filtered edge list", "Per decision node"],
        ["DF-3.18", "Code Exec I/O", "3.2 Execute Node", "Subprocess",
         "Python code (temp file), captured locals", "Per code_exec node"],

        # Scheduling
        ["DF-4.1", "RedBeat Entry Write", "4.0 Scheduling", "D6: Redis",
         "Task name, cron schedule, agent ID arg", "On agent deploy/resume"],
        ["DF-4.2", "RedBeat Entry Delete", "4.0 Scheduling", "D6: Redis",
         "Entry key removal", "On agent pause/delete"],
        ["DF-4.3", "Cron Fire", "D6: Redis (RedBeat)", "4.0 Scheduling",
         "Timer expiry, Celery task dispatch", "Per cron interval"],

        # Chat & Memory
        ["DF-5.1", "Chat Request", "User", "5.0 Chat & Memory",
         "Message text, history (role, content)[]", "Per chat message"],
        ["DF-5.2", "Memory Context Read", "D7: agent_memory", "5.1 Retrieve",
         "Last 15 summaries, timestamps", "Per chat request"],
        ["DF-5.3", "Context Assembly", "5.1 + Agent Metadata", "5.2 Build Context",
         "System prompt + memory + history + message", "Per chat request"],
        ["DF-5.4", "Chat LLM Request", "5.3 Stream GPT", "OpenAI API",
         "Messages[], model=gpt-4o-mini, stream=True", "Per chat request"],
        ["DF-5.5", "Chat Token Stream", "OpenAI API", "5.3 Stream GPT",
         "Content deltas (streaming chunks)", "Per chat response"],
        ["DF-5.6", "Chat SSE Events", "5.3 Stream GPT", "User",
         "SSE: {type: token/done/error, content}", "Per chat token"],
        ["DF-5.7", "Summary Generation", "5.4 Store", "OpenAI API",
         "Summarisation prompt (execution data)", "Per execution completion"],
        ["DF-5.8", "Memory Record Write", "5.4 Store", "D7: agent_memory",
         "AgentMemory record (summary, key_outputs)", "Per execution completion"],

        # Analytics
        ["DF-6.1", "Stats Request", "User", "6.0 Analytics",
         "Authenticated dashboard query", "Per dashboard load"],
        ["DF-6.2", "Agent Aggregation Read", "D2: agents", "6.0 Analytics",
         "Agent count, active count by user_id", "Per stats request"],
        ["DF-6.3", "Execution Aggregation Read", "D4: executions", "6.0 Analytics",
         "Success count, avg duration, costs", "Per stats request"],
        ["DF-6.4", "Activity Read", "D4: executions + D2: agents", "6.0 Analytics",
         "Last 20 executions with agent names/types", "Per activity request"],
        ["DF-6.5", "Dashboard Response", "6.0 Analytics", "User",
         "Stats object, activity event list", "Per dashboard load"],
    ]
    _create_std_table(doc, headers, rows)


# ============================================================================
# Main generator
# ============================================================================

def generate_dfd():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = BLACK

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build document
    add_cover_page(doc)
    add_metadata_table(doc)
    add_revision_history(doc)
    add_toc(doc)
    section_1_introduction(doc)
    section_2_notation_guide(doc)
    section_3_level0(doc)
    section_4_level1(doc)
    section_5_level2(doc)
    section_6_data_store_catalog(doc)
    section_7_data_flow_catalog(doc)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, "GoalCert_AutoMind_DFD_v1_0.docx")
    doc.save(path)
    print(f"DFD document saved to {path}")
    return path


if __name__ == "__main__":
    generate_dfd()
