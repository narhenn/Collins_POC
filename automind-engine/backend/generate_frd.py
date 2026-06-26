"""
Generate GoalCert AutoMind FRD (Functional Requirements Document)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──────────────────────────────────────────────────────
PURPLE = RGBColor(73, 2, 162)
DARK_GRAY = RGBColor(51, 51, 51)
MED_GRAY = RGBColor(102, 102, 102)
WHITE = RGBColor(255, 255, 255)
LIGHT_BG = "F5F0FF"
FONT = "Calibri"
BODY_SIZE = Pt(11)
H1_SIZE = Pt(16)
H2_SIZE = Pt(13)
H3_SIZE = Pt(11)
SAVE_PATH = "/Users/narhen/automind/docs/GoalCert_AutoMind_FRD_v1_0.docx"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "4902A2")

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_BG)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading1(doc, number, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.clear()
    run = p.add_run(f"{number}. {title}")
    run.font.name = FONT
    run.font.size = H1_SIZE
    run.font.color.rgb = PURPLE
    run.bold = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_heading2(doc, number, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.clear()
    run = p.add_run(f"{number} {title}")
    run.font.name = FONT
    run.font.size = H2_SIZE
    run.font.color.rgb = PURPLE
    run.bold = True
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading3(doc, number, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    p.clear()
    run = p.add_run(f"{number} {title}")
    run.font.name = FONT
    run.font.size = H3_SIZE
    run.font.color.rgb = PURPLE
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


def add_cover_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GoalCert")
    run.font.name = FONT
    run.font.size = Pt(36)
    run.font.color.rgb = PURPLE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AutoMind — Agentic AI Workforce Platform")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = MED_GRAY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = PURPLE
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Functional Requirements Document (FRD)")
    run.font.name = FONT
    run.font.size = Pt(22)
    run.font.color.rgb = PURPLE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Detailed functional specifications for the GoalCert AutoMind platform,\n"
        "defining module-level requirements, system behaviors, and traceability to business needs."
    )
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = MED_GRAY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(180, 0, 0)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This document contains proprietary information belonging to GoalCert Pte Ltd.\n"
        "Unauthorized distribution, reproduction, or use of this document is strictly prohibited."
    )
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = MED_GRAY

    doc.add_page_break()


def add_document_control(doc):
    p = doc.add_paragraph()
    run = p.add_run("Document Control")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = PURPLE
    run.bold = True

    doc.add_paragraph()

    rows = [
        ("Document Type", "Functional Requirements Document (FRD)"),
        ("Product", "GoalCert AutoMind"),
        ("Version", "1.0"),
        ("Date", "23 June 2026"),
        ("Classification", "CONFIDENTIAL"),
        ("Prepared By", "Engineering"),
        ("Reviewed By", "Product Management"),
        ("Approved By", "CTO"),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(rows):
        cell_l = table.rows[i].cells[0]
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell_l, "4902A2")
        cell_l.width = Cm(5)

        cell_r = table.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        cell_r.width = Cm(11)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Version History")
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = PURPLE
    run.bold = True

    doc.add_paragraph()

    add_styled_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ("0.1", "12 June 2026", "Engineering", "Initial draft — module decomposition and FR-AUTH, FR-AGT requirements"),
            ("0.5", "17 June 2026", "Engineering", "Added workflow builder, execution engine, memory, and chat modules"),
            ("0.9", "21 June 2026", "Engineering", "Incorporated review feedback, added traceability matrix and non-functional requirements"),
            ("1.0", "23 June 2026", "Engineering", "Final release — approved by CTO and Product Management"),
        ],
        col_widths=[2, 3, 3, 8],
    )

    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = PURPLE
    run.bold = True

    doc.add_paragraph()

    toc_items = [
        ("1.", "Introduction"),
        ("   1.1", "Purpose"),
        ("   1.2", "Conventions"),
        ("   1.3", "System Overview"),
        ("2.", "Actors & Role-Based Access"),
        ("3.", "Module M1 — Authentication & User Management"),
        ("4.", "Module M2 — Agent Management"),
        ("5.", "Module M3 — Workflow Builder"),
        ("   5.1", "Node Types"),
        ("   5.2", "Canvas & Configuration"),
        ("6.", "Module M4 — Execution Engine"),
        ("   6.1", "Execution Pipeline"),
        ("   6.2", "Variable Substitution"),
        ("7.", "Module M5 — Agent Memory & RAG"),
        ("8.", "Module M6 — Agent Chat Interface"),
        ("9.", "Module M7 — Dashboard & Analytics"),
        ("10.", "Module M8 — Reports"),
        ("11.", "Module M9 — Templates"),
        ("12.", "Module M10 — Integrations"),
        ("13.", "Module M11 — JILLA AI Concierge"),
        ("14.", "Non-Functional Requirements"),
        ("   14.1", "Performance"),
        ("   14.2", "Security"),
        ("   14.3", "Scalability"),
        ("   14.4", "Availability"),
        ("15.", "Integration Points"),
        ("16.", "Traceability Matrix"),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        run.font.name = FONT
        run.font.size = Pt(11) if not num.startswith("   ") else Pt(10)
        run.font.color.rgb = DARK_GRAY if not num.startswith("   ") else MED_GRAY
        run.bold = not num.startswith("   ")
        p.paragraph_format.space_after = Pt(2)
        if num.startswith("   "):
            p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()


def add_req_table(doc, reqs, col_widths=None):
    """Add a requirements table. Each req is (ID, Name, Priority, Description)."""
    if col_widths is None:
        col_widths = [2.5, 3, 1.5, 9]
    add_styled_table(doc,
        ["ID", "Requirement", "Priority", "Description"],
        reqs,
        col_widths=col_widths,
    )


# ═══════════════════════════════════════════════════════════════════
# SECTION BUILDERS
# ═══════════════════════════════════════════════════════════════════

def build_section_1(doc):
    add_heading1(doc, "1", "Introduction")

    add_heading2(doc, "1.1", "Purpose")
    add_body(doc,
        "This Functional Requirements Document (FRD) specifies the detailed functional behavior of the GoalCert "
        "AutoMind platform. It decomposes the business requirements defined in the companion Business Requirements "
        "Document (BRD v1.0) into module-level functional requirements that are implementable, testable, and "
        "traceable. The FRD serves as the primary contract between product management and engineering, ensuring "
        "that all parties share a common understanding of what the system must do."
    )
    add_body(doc,
        "The intended audience includes software engineers, QA engineers, UX designers, technical architects, and "
        "project managers who need precise specifications for implementation, testing, and project planning. Business "
        "stakeholders may reference this document to validate that their requirements are accurately captured."
    )

    add_heading2(doc, "1.2", "Conventions")
    add_body(doc,
        "Functional requirements in this document follow a consistent identification scheme: FR-<MODULE>-<N>, where "
        "<MODULE> is a short code identifying the functional module (e.g., AUTH for Authentication, WF for Workflow, "
        "EXEC for Execution Engine) and <N> is a sequential number within that module. Each requirement is assigned "
        "a MoSCoW priority: Must (required for launch), Should (high value, schedule permitting), Could (desirable, "
        "not critical), or Won't (explicitly excluded from this version)."
    )
    add_body(doc,
        "The keywords 'shall', 'must', and 'will' indicate mandatory behavior. The keyword 'should' indicates "
        "recommended behavior. The keyword 'may' indicates optional behavior. All API endpoints described assume "
        "a RESTful JSON API convention with JWT-based authentication unless otherwise stated."
    )

    add_heading2(doc, "1.3", "System Overview")
    add_body(doc,
        "GoalCert AutoMind is a web-based platform that enables users to create, configure, and deploy autonomous "
        "AI agents through a visual workflow builder. The system architecture consists of the following major components:"
    )
    components = [
        "React Frontend: Single-page application providing the workflow builder canvas, agent management dashboard, execution console, chat interface, and administration panels",
        "FastAPI Backend: RESTful API server handling authentication, agent CRUD operations, workflow persistence, execution orchestration, and real-time event streaming via SSE endpoints",
        "PostgreSQL Database: Primary data store for users, agents, workflows (stored as JSON DAGs), execution logs, memory summaries, and chat histories",
        "Redis: In-memory cache and message broker serving three roles — Celery message transport, RedBeat schedule storage, and ephemeral execution state caching for SSE",
        "Celery Workers: Distributed task execution workers that process agent workflow runs, executing nodes sequentially according to BFS traversal order and publishing progress events to Redis",
        "RedBeat Scheduler: Celery beat scheduler backed by Redis that manages cron-based agent execution schedules, handling timezone conversion and missed-execution detection",
    ]
    for c in components:
        add_bullet(doc, c)

    add_body(doc,
        "The execution pipeline operates as follows: when an agent is triggered (by schedule, manual action, or "
        "webhook), the API server dispatches a Celery task to the worker pool. The worker deserializes the workflow "
        "DAG, performs a BFS traversal to determine execution order, and processes each node sequentially. At each "
        "node, the worker resolves variable substitutions (replacing {variable_name} placeholders with outputs from "
        "previously executed nodes), executes the node's operation (web search, LLM call, integration call, etc.), "
        "records the result, publishes a progress event to Redis (consumed by the SSE endpoint), and advances to the "
        "next node. Upon completion, the worker generates an execution summary, updates the agent's memory store, and "
        "publishes a completion event."
    )


def build_section_2(doc):
    add_heading1(doc, "2", "Actors & Role-Based Access")

    add_body(doc,
        "The AutoMind platform defines three actor roles with hierarchical permissions. Role assignment is managed "
        "through the administration module and enforced at the API layer via JWT claims."
    )

    add_styled_table(doc,
        ["Role", "Description", "Key Permissions"],
        [
            ("User", "Standard platform user who creates and manages their own agents",
             "Create/edit/delete own agents; execute own agents; view own execution logs; access agent chat; browse template marketplace; view own cost data"),
            ("Admin", "Organization administrator with elevated privileges",
             "All User permissions; view and manage all agents within the organization; access organization-level analytics and reports; manage user accounts; configure organization settings; publish templates"),
            ("Super Admin", "Platform-wide administrator (GoalCert staff only)",
             "All Admin permissions; manage all organizations; access platform-wide analytics; configure system settings (AI providers, rate limits, feature flags); manage pricing and billing; JILLA configuration"),
        ],
        col_widths=[2.5, 5, 8.5],
    )

    add_body(doc,
        "All API endpoints enforce role-based access control through middleware that validates the requesting user's "
        "role against the required permission for the endpoint. Unauthenticated requests receive a 401 Unauthorized "
        "response. Authenticated requests lacking the required role receive a 403 Forbidden response with a descriptive "
        "error message."
    )


def build_section_3(doc):
    add_heading1(doc, "3", "Module M1 — Authentication & User Management")

    add_body(doc,
        "The Authentication and User Management module handles user registration, login, session management, profile "
        "management, and password security. This module establishes the identity foundation upon which all other "
        "modules enforce access control."
    )

    reqs = [
        ("FR-AUTH-1", "User Registration", "Must",
         "The system shall provide a registration endpoint (POST /api/auth/register) that accepts email, password, and display name. "
         "Passwords must be at least 8 characters with at least one uppercase letter, one lowercase letter, and one digit. "
         "Passwords shall be hashed using bcrypt with a work factor of 12 before storage. Upon successful registration, the system "
         "shall return a JWT access token (15-minute expiry) and a refresh token (7-day expiry). Duplicate email registrations shall "
         "return a 409 Conflict response."),

        ("FR-AUTH-2", "User Login", "Must",
         "The system shall provide a login endpoint (POST /api/auth/login) that accepts email and password. The system shall "
         "verify the password against the stored bcrypt hash and, upon success, return a JWT access token and refresh token with "
         "the user's ID, email, and role encoded in the token claims. Failed login attempts shall return a generic 401 Unauthorized "
         "response without revealing whether the email exists. After five consecutive failed login attempts from the same IP within "
         "a 15-minute window, the system shall impose a 5-minute rate limit on that IP."),

        ("FR-AUTH-3", "Token Refresh", "Must",
         "The system shall provide a token refresh endpoint (POST /api/auth/refresh) that accepts a valid refresh token and returns "
         "a new access token. Expired or revoked refresh tokens shall return a 401 Unauthorized response. The system shall maintain "
         "a refresh token revocation list in Redis to support immediate token invalidation on logout or password change."),

        ("FR-AUTH-4", "User Profile Management", "Must",
         "The system shall provide endpoints for reading (GET /api/users/me) and updating (PATCH /api/users/me) the authenticated "
         "user's profile. Updatable fields include display name, avatar URL, timezone preference, and notification preferences. "
         "Email changes shall require re-verification. Password changes shall require the current password and shall invalidate all "
         "existing refresh tokens."),

        ("FR-AUTH-5", "Password Reset", "Should",
         "The system shall provide a password reset flow consisting of a request endpoint (POST /api/auth/forgot-password) that sends "
         "a time-limited reset link (1-hour expiry) to the registered email, and a reset endpoint (POST /api/auth/reset-password) that "
         "accepts the reset token and new password. Used reset tokens shall be invalidated immediately. The reset email shall be rate "
         "limited to 3 requests per email per hour."),

        ("FR-AUTH-6", "Admin User Management", "Must",
         "Admin users shall have access to a user management endpoint (GET /api/admin/users) that lists all users within their "
         "organization with pagination and search. Admins shall be able to update user roles (POST /api/admin/users/{id}/role), "
         "deactivate user accounts (POST /api/admin/users/{id}/deactivate), and view user activity summaries. Role changes shall "
         "take effect on the user's next token refresh."),

        ("FR-AUTH-7", "SSO / SAML Integration", "Could",
         "The system shall support SAML 2.0 single sign-on for enterprise customers, enabling authentication through external "
         "identity providers (e.g., Okta, Azure AD). The SAML integration shall support SP-initiated SSO, just-in-time user "
         "provisioning, and attribute mapping for role assignment. This feature is scoped for enterprise-tier customers only."),
    ]

    add_req_table(doc, reqs)


def build_section_4(doc):
    add_heading1(doc, "4", "Module M2 — Agent Management")

    add_body(doc,
        "The Agent Management module provides CRUD operations for AI agents, schedule configuration, status management, "
        "and agent lifecycle operations. An agent is the top-level entity that encapsulates a workflow definition, "
        "execution schedule, memory store, and chat history."
    )

    reqs = [
        ("FR-AGT-1", "Create Agent", "Must",
         "The system shall provide an endpoint (POST /api/agents) that creates a new agent with the following fields: "
         "name (required, max 100 characters), description (optional, max 500 characters), system prompt (required, defines "
         "the agent's persona and instructions), AI model selection (required, one of the supported models), and status "
         "(defaults to 'draft'). The agent shall be associated with the authenticated user's ID. The workflow field shall "
         "be initialized as an empty DAG. The response shall return the created agent with a generated UUID."),

        ("FR-AGT-2", "List Agents", "Must",
         "The system shall provide an endpoint (GET /api/agents) that returns a paginated list of agents belonging to the "
         "authenticated user. The response shall include agent ID, name, status, last execution timestamp, next scheduled "
         "execution, execution count, and model. The endpoint shall support query parameters for pagination (page, per_page), "
         "sorting (sort_by, sort_order), and filtering (status, model). Admin users shall see all agents in their organization."),

        ("FR-AGT-3", "Update Agent", "Must",
         "The system shall provide an endpoint (PATCH /api/agents/{id}) that updates agent properties including name, description, "
         "system prompt, model, and status. Updates to the workflow field are handled separately by the Workflow Builder module. "
         "Status transitions shall follow the state machine: draft -> active, active -> paused, paused -> active, any -> archived. "
         "Invalid transitions shall return a 422 Unprocessable Entity response. Activating an agent with a scheduled trigger shall "
         "register the schedule with RedBeat."),

        ("FR-AGT-4", "Delete Agent", "Must",
         "The system shall provide an endpoint (DELETE /api/agents/{id}) that soft-deletes an agent by setting its status to "
         "'archived' and its deleted_at timestamp. Archived agents shall be excluded from list queries by default but recoverable "
         "by admin users. Deletion shall cancel any active RedBeat schedules, terminate any in-flight executions, and retain "
         "execution logs for the organization's configured retention period (default 90 days)."),

        ("FR-AGT-5", "Agent Schedule Configuration", "Must",
         "The system shall provide an endpoint (PUT /api/agents/{id}/schedule) that configures the agent's execution schedule. "
         "The endpoint shall accept a schedule type (cron, interval, or manual) and, for cron schedules, a cron expression and "
         "timezone. The system shall validate the cron expression and register it with RedBeat upon agent activation. The endpoint "
         "shall also support natural-language schedule input (e.g., 'every weekday at 9am SGT') which the backend parses into "
         "a cron expression. Webhook triggers are configured separately via FR-INT-5."),

        ("FR-AGT-6", "Agent Status Dashboard", "Must",
         "The system shall provide an endpoint (GET /api/agents/{id}/status) that returns the agent's current operational status "
         "including: current execution state (idle, running, failed, paused), last execution result (success/failure with summary), "
         "next scheduled execution timestamp, total execution count, success rate, average execution duration, and total cost "
         "accumulated. This data shall be refreshed on each API call, not cached."),

        ("FR-AGT-7", "Manual Agent Trigger", "Must",
         "The system shall provide an endpoint (POST /api/agents/{id}/execute) that triggers an immediate manual execution of the "
         "agent's workflow regardless of its schedule. The endpoint shall return an execution ID and the SSE stream URL for "
         "monitoring. If the agent is already running and concurrent execution is disabled (default), the request shall return "
         "a 409 Conflict response. Manual triggers shall be logged distinctly from scheduled triggers in execution history."),

        ("FR-AGT-8", "Agent Duplication", "Should",
         "The system shall provide an endpoint (POST /api/agents/{id}/duplicate) that creates a deep copy of an existing agent "
         "including its workflow definition, system prompt, model configuration, and schedule settings. The duplicated agent shall "
         "have a 'draft' status and a name suffixed with ' (Copy)'. Memory and execution history shall not be copied. This "
         "feature enables users to iterate on agent designs without modifying production agents."),
    ]

    add_req_table(doc, reqs)


def build_section_5(doc):
    add_heading1(doc, "5", "Module M3 — Workflow Builder")

    add_body(doc,
        "The Workflow Builder module provides the visual canvas for designing agent workflows. A workflow is represented "
        "as a directed acyclic graph (DAG) where nodes are processing units and edges define the data flow between them. "
        "The frontend renders an interactive canvas using a graph visualization library, and the backend persists the "
        "workflow as a JSON structure containing node definitions, edge connections, and node positions."
    )

    add_heading2(doc, "5.1", "Node Types")
    add_body(doc,
        "The platform supports seven node types, each with distinct behavior and configuration options. Every node "
        "shares a common schema (ID, type, label, position, configuration) and produces typed outputs that can be "
        "referenced by downstream nodes via variable substitution."
    )

    add_styled_table(doc,
        ["Node Type", "Purpose", "Inputs", "Outputs", "Configuration Fields"],
        [
            ("Trigger", "Initiates workflow execution based on a schedule, manual action, or webhook event",
             "None (entry point)", "trigger_data: payload from webhook or empty for cron/manual",
             "trigger_type (cron | manual | webhook), cron_expression, timezone, webhook_secret"),

            ("Web Search", "Performs web searches using configured search providers and returns structured results",
             "query: search query string (supports variable substitution)",
             "results: array of {title, url, snippet}; raw_text: concatenated snippet text",
             "query, max_results (1-20), search_provider (default: Google), include_raw_text (boolean)"),

            ("AI Action", "Sends a prompt to a large language model and returns the generated response",
             "prompt: text prompt with variable substitution; context: optional additional context",
             "response: LLM response text; tokens_used: {input, output, total}; cost: computed cost",
             "model (gpt-4, gpt-4o, claude-3.5-sonnet, claude-3-opus), prompt_template, temperature (0.0-2.0), max_tokens, system_prompt_override"),

            ("Decision", "Evaluates a condition against node output and routes execution to one of two branches",
             "value: the value to evaluate (from upstream node via variable substitution)",
             "branch: 'true' or 'false' indicating which path was taken",
             "condition_type (contains | equals | greater_than | less_than | regex_match), condition_value, case_sensitive (boolean)"),

            ("Integration", "Sends data to external services (email, Slack, webhooks) as a delivery action",
             "content: message body with variable substitution; recipient: target address/channel/URL",
             "status: delivery status (sent | failed); response: service response",
             "integration_type (email | slack | webhook), recipient, subject (email only), channel (slack only), webhook_url, headers (webhook only)"),

            ("Escalation", "Pauses workflow execution and notifies a human operator for review or intervention",
             "message: escalation reason with variable substitution",
             "resolution: human operator's response; resolved_by: operator ID; resolved_at: timestamp",
             "escalation_message, notify_channels (email, slack), timeout_hours (auto-fail if unresolved), assignee_email"),

            ("Code Execution", "Executes user-provided Python code in a sandboxed environment",
             "inputs: dictionary of variables from upstream nodes available in the code scope",
             "result: the return value of the executed code; stdout: captured print output",
             "code (Python source), timeout_seconds (max 30), allowed_imports (restricted list)"),
        ],
        col_widths=[2, 3, 3, 3, 5],
    )

    add_heading2(doc, "5.2", "Canvas & Configuration Requirements")

    reqs = [
        ("FR-WF-1", "Workflow Canvas Rendering", "Must",
         "The frontend shall render an interactive canvas that displays workflow nodes as styled cards with type-specific "
         "icons, labels, and connection ports. The canvas shall support pan, zoom (10% to 200%), grid snapping, and "
         "minimap navigation. Nodes shall be positioned using a force-directed layout algorithm with manual override. "
         "The canvas state (node positions, zoom level, pan offset) shall persist across sessions."),

        ("FR-WF-2", "Drag-and-Drop Node Addition", "Must",
         "The frontend shall provide a node palette panel listing all seven node types with descriptions. Users shall "
         "drag nodes from the palette onto the canvas to add them to the workflow. Upon placement, the node shall be "
         "assigned a unique ID and default configuration. The canvas shall prevent placing nodes outside the visible "
         "workflow area and shall auto-scroll when dragging near canvas edges."),

        ("FR-WF-3", "Edge Connection Management", "Must",
         "Users shall create edges by clicking and dragging from a node's output port to another node's input port. "
         "The system shall enforce DAG constraints: self-loops and cycles shall be rejected with a visual error indicator. "
         "Decision nodes shall have two output ports labeled 'True' and 'False'. All other nodes shall have a single output "
         "port. Edges shall be visually styled as curved Bezier lines with directional arrows. Users shall be able to "
         "delete edges by clicking on them and pressing Delete or via a context menu."),

        ("FR-WF-4", "Node Configuration Panel", "Must",
         "Selecting a node on the canvas shall open a configuration panel that displays fields specific to the node type "
         "(as defined in the Node Types table above). The panel shall validate inputs in real time (e.g., cron expression "
         "syntax, valid model selection, non-empty prompt templates). Variable substitution references shall be "
         "auto-suggested in a dropdown listing all available upstream node outputs. Changes shall auto-save after 500ms "
         "of inactivity."),

        ("FR-WF-5", "Workflow Validation", "Must",
         "The system shall validate the workflow graph before allowing agent activation. Validation rules include: "
         "(1) exactly one Trigger node must exist, (2) all nodes must be reachable from the Trigger node, (3) no "
         "cycles exist in the graph, (4) all required node configuration fields are populated, (5) all variable "
         "substitution references resolve to valid upstream node outputs. Validation errors shall be displayed inline "
         "on the canvas with node-level error indicators and a summary panel listing all issues."),

        ("FR-WF-6", "Workflow Persistence", "Must",
         "The backend shall provide an endpoint (PUT /api/agents/{id}/workflow) that persists the workflow DAG as a "
         "JSON structure containing: nodes (array of {id, type, label, position: {x, y}, config: {...}}), edges "
         "(array of {id, source_node_id, source_port, target_node_id, target_port}), and metadata (version, "
         "last_modified, modified_by). The endpoint shall validate the DAG structure server-side and return detailed "
         "validation errors if the workflow is invalid."),

        ("FR-WF-7", "Workflow Version History", "Should",
         "The system shall maintain a version history for each workflow, creating a new version snapshot each time the "
         "workflow is saved. Users shall be able to view past versions (GET /api/agents/{id}/workflow/versions), compare "
         "two versions side-by-side, and restore a previous version. The system shall retain the last 20 versions per agent."),

        ("FR-WF-8", "Undo/Redo Operations", "Should",
         "The canvas shall support undo (Ctrl+Z) and redo (Ctrl+Shift+Z) for all operations including node addition, "
         "deletion, movement, edge creation, edge deletion, and configuration changes. The undo stack shall persist "
         "for the current editing session with a maximum depth of 50 operations."),
    ]

    add_req_table(doc, reqs)


def build_section_6(doc):
    add_heading1(doc, "6", "Module M4 — Execution Engine")

    add_heading2(doc, "6.1", "Execution Pipeline")
    add_body(doc,
        "The Execution Engine is the core runtime component responsible for processing agent workflows. When an "
        "execution is triggered, the engine deserializes the workflow DAG, computes the BFS traversal order from the "
        "Trigger node, and processes nodes sequentially. Each node execution follows a standardized lifecycle: "
        "resolve variable substitutions in the node's configuration, execute the node's operation (calling external "
        "APIs, running code, evaluating conditions), record the node's output in the execution context, publish a "
        "progress event to the SSE stream, and advance to the next node in the traversal order."
    )
    add_body(doc,
        "For Decision nodes, the engine evaluates the configured condition and follows only the matching branch "
        "(True or False), skipping all nodes on the non-matching branch. The execution context maintains a key-value "
        "store where each node's output is stored under the node's ID, enabling downstream nodes to reference any "
        "upstream output via the {node_id.output_field} variable substitution syntax."
    )

    reqs = [
        ("FR-EXEC-1", "Execution Dispatch", "Must",
         "When an agent execution is triggered (via schedule, manual, or webhook), the API server shall create an execution "
         "record in the database with status 'pending', generate a unique execution ID, and dispatch a Celery task to the "
         "worker pool with the agent ID and execution ID as parameters. The API shall return the execution ID and SSE "
         "stream URL immediately without waiting for execution to begin. The Celery task shall update the execution status "
         "to 'running' upon worker pickup."),

        ("FR-EXEC-2", "BFS Workflow Traversal", "Must",
         "The execution engine shall compute the node execution order using Breadth-First Search starting from the Trigger "
         "node. Nodes at the same BFS depth level shall be executed sequentially in order of their node ID. The engine "
         "shall skip nodes on inactive Decision branches (the branch not taken). The traversal order shall be logged at "
         "the start of execution for debugging purposes."),

        ("FR-EXEC-3", "Variable Substitution", "Must",
         "Before executing each node, the engine shall resolve all variable substitution placeholders in the node's "
         "configuration fields. Placeholders follow the syntax {node_id.output_field} where node_id is the unique ID of "
         "an upstream node and output_field is a key in that node's output dictionary. Nested references (e.g., "
         "{node_1.results[0].title}) shall be supported using dot-notation and array indexing. Unresolvable references "
         "shall be replaced with an empty string and logged as a warning. The special variable {trigger_data} provides "
         "access to the trigger payload."),

        ("FR-EXEC-4", "Node Execution Lifecycle", "Must",
         "Each node execution shall follow this lifecycle: (1) status update to 'executing' published to SSE, (2) variable "
         "substitution resolution, (3) operation execution with timeout enforcement (configurable per node type, default "
         "60 seconds), (4) output recording in execution context, (5) cost calculation for AI Action nodes, (6) status "
         "update to 'completed' or 'failed' published to SSE. Failed nodes shall capture the error message and stack trace. "
         "The engine shall continue to the next node on failure unless the workflow is configured for fail-fast mode."),

        ("FR-EXEC-5", "SSE Execution Streaming", "Must",
         "The system shall provide an SSE endpoint (GET /api/agents/{id}/executions/{exec_id}/stream) that streams "
         "real-time execution events to connected clients. Event types include: execution_started, node_started, "
         "node_output (intermediate results), node_completed, node_failed, execution_completed, execution_failed. Each "
         "event shall include a timestamp, node ID (where applicable), and relevant payload data. The SSE connection shall "
         "support automatic reconnection with last-event-ID recovery. Events shall be published to Redis Pub/Sub and "
         "consumed by the API server's SSE endpoint."),

        ("FR-EXEC-6", "Execution History", "Must",
         "The system shall provide an endpoint (GET /api/agents/{id}/executions) that returns a paginated list of past "
         "executions with ID, trigger type, status, start time, end time, duration, total cost, and node count. A detail "
         "endpoint (GET /api/agents/{id}/executions/{exec_id}) shall return the full execution record including per-node "
         "results, timing, cost breakdown, and any error details. Execution records shall be retained for the configured "
         "retention period (default 90 days)."),

        ("FR-EXEC-7", "Cost Tracking Per Execution", "Must",
         "The execution engine shall track the cost of each AI Action node execution by recording input tokens, output "
         "tokens, model used, and computing the cost based on the provider's pricing table (maintained as a configuration "
         "file in the backend). The total execution cost shall be the sum of all AI Action node costs. Cost data shall be "
         "stored in the execution record and accessible via the execution detail endpoint. The pricing configuration shall "
         "be updatable by Super Admins without code deployment."),

        ("FR-EXEC-8", "Execution Timeout and Cancellation", "Must",
         "Each execution shall have a maximum total duration (configurable, default 10 minutes). If the timeout is reached, "
         "the engine shall abort the current node, set the execution status to 'timeout', and publish a timeout event. "
         "Users shall be able to cancel a running execution via an endpoint (POST /api/agents/{id}/executions/{exec_id}/cancel) "
         "which sends a revoke signal to the Celery worker. Cancelled executions shall record partial results for nodes "
         "that completed before cancellation."),

        ("FR-EXEC-9", "Concurrent Execution Control", "Should",
         "The system shall support a per-agent concurrency setting (default: 1) that limits the number of simultaneous "
         "executions for a single agent. When a trigger fires for an agent that has reached its concurrency limit, the "
         "system shall queue the execution (if queuing is enabled) or reject it with a 409 Conflict response. This prevents "
         "race conditions in agents that maintain state or interact with rate-limited external services."),

        ("FR-EXEC-10", "Webhook Trigger Processing", "Must",
         "The system shall provide a webhook endpoint (POST /api/webhooks/{agent_webhook_id}) that accepts external HTTP "
         "requests and triggers agent execution. The webhook payload shall be validated against an optional JSON schema "
         "configured by the user. The entire request body shall be available as {trigger_data} within the workflow. Webhook "
         "endpoints shall be secured with HMAC signature verification using a per-agent webhook secret."),
    ]

    add_heading2(doc, "6.2", "Variable Substitution System")
    add_body(doc,
        "The variable substitution system is a core mechanism that enables data flow between nodes in a workflow. "
        "When configuring a node, users can reference outputs from any upstream node using the {node_id.output_field} "
        "syntax. The execution engine resolves these references at runtime by looking up the referenced node's output "
        "in the execution context. This design allows users to build sophisticated data pipelines where each node "
        "processes and transforms the outputs of preceding nodes."
    )
    add_body(doc,
        "For example, a typical workflow might have a Web Search node (node_1) whose 'raw_text' output is referenced "
        "in an AI Action node's prompt template as: 'Analyze the following search results and identify key trends: "
        "{node_1.raw_text}'. The AI Action node's 'response' output might then be referenced in an Integration node's "
        "email body as: 'Here is your daily market briefing:\\n\\n{node_2.response}'. This chaining enables complex "
        "multi-step analysis workflows without requiring any code."
    )

    add_req_table(doc, reqs)


def build_section_7(doc):
    add_heading1(doc, "7", "Module M5 — Agent Memory & RAG")

    add_body(doc,
        "The Agent Memory module provides persistent memory that enables agents to learn from past executions and "
        "incorporate historical context into future runs. After each execution, the system generates a structured "
        "summary of what happened — key findings, decisions made, data analyzed, and anomalies detected — and stores "
        "it in the agent's memory store. On subsequent executions, relevant summaries are retrieved and injected into "
        "the AI Action nodes' context, implementing a Retrieval-Augmented Generation (RAG) pattern that gives agents "
        "continuity across runs."
    )

    reqs = [
        ("FR-MEM-1", "Execution Summary Generation", "Must",
         "Upon successful completion of an agent execution, the system shall generate a structured summary by sending "
         "the execution results (all node outputs, timing data, and any errors) to an LLM (GPT-4o-mini for cost efficiency) "
         "with a standardized summarization prompt. The summary shall include: key findings (bullet points), data sources "
         "consulted, decisions made by Decision nodes, outputs delivered via Integration nodes, anomalies or unexpected "
         "results, and a one-sentence executive summary. The summary shall be stored as a JSON document in the agent's "
         "memory table with the execution ID, timestamp, and word count."),

        ("FR-MEM-2", "Memory Injection into Executions", "Must",
         "At the start of each execution, before processing the first AI Action node, the engine shall retrieve the "
         "agent's most recent N memory summaries (configurable per agent, default 5) and inject them into the AI Action "
         "nodes' context. The injected context shall be formatted as a chronological list of past execution summaries "
         "with timestamps, enabling the LLM to reference historical data. The injection shall prepend the memory context "
         "to the node's system prompt, clearly delineated with a 'Past Execution Memory' header."),

        ("FR-MEM-3", "Memory Viewer", "Must",
         "The system shall provide an endpoint (GET /api/agents/{id}/memory) that returns a paginated list of memory "
         "summaries for the specified agent, sorted by timestamp descending. Each summary shall include the execution ID, "
         "timestamp, executive summary, and full structured summary. Users shall be able to expand individual summaries "
         "in the UI to read the full content. The memory viewer shall display the total memory count and storage size."),

        ("FR-MEM-4", "Memory Retention Policy", "Should",
         "The system shall support configurable memory retention policies per agent. The retention policy shall specify "
         "the maximum number of summaries to retain (default 50) and the maximum age in days (default 90). When the "
         "retention limit is exceeded, the oldest summaries shall be automatically purged. Admins shall be able to set "
         "organization-wide retention defaults. Individual memory entries shall be manually deletable by the agent owner."),

        ("FR-MEM-5", "Memory Search", "Should",
         "The system shall provide a search endpoint (GET /api/agents/{id}/memory/search?q=query) that performs "
         "full-text search across an agent's memory summaries and returns relevant matches ranked by relevance. This "
         "enables users and the chat interface to quickly locate specific past findings without scrolling through all "
         "summaries. The search shall support basic keyword matching with optional date range filtering."),

        ("FR-MEM-6", "Selective Memory Injection", "Could",
         "In addition to chronological retrieval (FR-MEM-2), the system shall support semantic memory retrieval where "
         "the current execution's trigger data and workflow configuration are used to retrieve the most contextually "
         "relevant past summaries rather than simply the most recent ones. This requires embedding memory summaries "
         "and performing vector similarity search at execution time. This feature enhances agent performance for "
         "use cases where relevance matters more than recency."),
    ]

    add_req_table(doc, reqs)


def build_section_8(doc):
    add_heading1(doc, "8", "Module M6 — Agent Chat Interface")

    add_body(doc,
        "The Agent Chat module provides a conversational interface where users can interact with their deployed agents "
        "in natural language. The chat interface serves as a debriefing tool, allowing users to ask agents about their "
        "latest execution results, request clarifications, and instruct agents to perform follow-up analyses. The agent "
        "responds using its memory context and workflow knowledge, providing a seamless extension of the automated "
        "execution experience."
    )

    reqs = [
        ("FR-CHAT-1", "Chat Session Management", "Must",
         "The system shall provide endpoints for creating (POST /api/agents/{id}/chat), listing (GET /api/agents/{id}/chat), "
         "and deleting (DELETE /api/agents/{id}/chat/{session_id}) chat sessions. Each session maintains an ordered list of "
         "messages with roles (user, assistant, system) and timestamps. A new chat session shall automatically include a "
         "system message containing the agent's system prompt, most recent execution summary, and recent memory context."),

        ("FR-CHAT-2", "Chat Message Streaming", "Must",
         "When a user sends a chat message (POST /api/agents/{id}/chat/{session_id}/message), the system shall stream "
         "the AI response in real time via Server-Sent Events. The SSE stream shall emit individual text chunks as they "
         "are generated by the LLM, enabling character-by-character rendering in the frontend. The stream shall conclude "
         "with a 'done' event containing the complete response and token usage. The frontend shall display a typing "
         "indicator during streaming and render markdown formatting in the response."),

        ("FR-CHAT-3", "Context-Aware Responses", "Must",
         "When generating chat responses, the system shall include the following context in the LLM prompt: the agent's "
         "system prompt, the most recent execution summary (including per-node outputs), the last N memory summaries "
         "(configurable, default 3), and the full chat session history (up to the model's context window limit). This "
         "context enables the agent to answer questions about its past performance, explain its reasoning, and reference "
         "specific data points from previous executions."),

        ("FR-CHAT-4", "Chat History Persistence", "Must",
         "All chat messages (user and assistant) shall be persisted in the database with the session ID, message role, "
         "content, timestamp, and token count. The system shall provide an endpoint (GET /api/agents/{id}/chat/{session_id}/messages) "
         "that returns the paginated message history for a session. Chat sessions shall be retained for the same period as "
         "execution logs (default 90 days)."),

        ("FR-CHAT-5", "Chat Cost Tracking", "Should",
         "Each chat message exchange shall track and display the token cost, consistent with the execution cost tracking "
         "system (FR-EXEC-7). The chat interface shall display the running cost of the current session and the total "
         "chat cost per agent in the agent's cost dashboard. Chat costs shall be separate from execution costs in "
         "analytics and billing."),

        ("FR-CHAT-6", "Chat Export", "Could",
         "The system shall provide an endpoint (GET /api/agents/{id}/chat/{session_id}/export) that exports a chat session "
         "as a formatted document (PDF or Markdown) including all messages, timestamps, and the context summary that was "
         "active during the session. This enables users to share agent conversations with stakeholders who do not have "
         "platform access."),
    ]

    add_req_table(doc, reqs)


def build_section_9(doc):
    add_heading1(doc, "9", "Module M7 — Dashboard & Analytics")

    add_body(doc,
        "The Dashboard and Analytics module provides real-time operational visibility and historical analytics across "
        "an individual user's agents (User role) or all agents within an organization (Admin role). The dashboard "
        "serves as the primary landing page after login and surfaces the most important metrics at a glance."
    )

    reqs = [
        ("FR-DASH-1", "User Dashboard Home", "Must",
         "The dashboard home page shall display a summary panel showing: total active agents, total executions (last 7 days), "
         "overall success rate (last 7 days), total cost (current billing period), and next upcoming scheduled execution. "
         "Below the summary, a list of the user's agents shall be displayed as cards showing agent name, status indicator "
         "(green=active, yellow=paused, gray=draft), last execution result, and a quick-action button for manual trigger."),

        ("FR-DASH-2", "Execution Timeline", "Must",
         "The dashboard shall include an execution timeline view showing all executions across all agents on a time-series "
         "chart. Each execution shall be plotted as a point colored by status (green=success, red=failed, yellow=timeout). "
         "Clicking an execution point shall navigate to the execution detail view. The timeline shall support date range "
         "selection with presets (today, last 7 days, last 30 days, custom range)."),

        ("FR-DASH-3", "Cost Analytics", "Must",
         "The dashboard shall include a cost analytics panel showing: cost by agent (bar chart), cost by model (pie chart), "
         "cost over time (line chart with daily granularity), and projected monthly cost based on trailing 7-day average. "
         "Each data point shall be drillable to the underlying execution records. Cost alerts (configurable thresholds) "
         "shall be displayed as warning banners when approaching or exceeding the configured limit."),

        ("FR-DASH-4", "Agent Performance Metrics", "Should",
         "For each agent, the dashboard shall provide a detailed performance view showing: execution success rate over time, "
         "average execution duration trend, cost per execution trend, most common failure reasons, and node-level performance "
         "heatmap (identifying bottleneck nodes by average duration). This data shall be accessible via GET /api/agents/{id}/analytics "
         "with configurable date range parameters."),

        ("FR-DASH-5", "Organization Analytics (Admin)", "Should",
         "Admin users shall have access to an organization-level analytics page showing aggregate metrics across all users "
         "and agents: total agents by status, total executions by status, cost by user, most active agents, platform adoption "
         "trend (new agents created per week), and model usage distribution. This view shall be restricted to Admin and "
         "Super Admin roles."),

        ("FR-DASH-6", "Real-Time Notifications", "Should",
         "The dashboard shall display real-time notifications for events such as: execution completed, execution failed, "
         "cost alert triggered, agent schedule missed, and escalation pending. Notifications shall appear as a bell icon "
         "badge with a dropdown panel. Users shall be able to configure notification preferences (in-app, email, Slack) "
         "per event type."),
    ]

    add_req_table(doc, reqs)


def build_section_10(doc):
    add_heading1(doc, "10", "Module M8 — Reports")

    add_body(doc,
        "The Reports module provides structured, exportable reports that aggregate execution data into business-friendly "
        "formats. Reports are distinct from real-time dashboard analytics in that they are generated on demand or on "
        "a schedule, formatted for stakeholder consumption, and designed to answer specific business questions."
    )

    reqs = [
        ("FR-RPT-1", "Report Generation", "Must",
         "The system shall provide an endpoint (POST /api/reports) that generates a report based on a specified template "
         "and date range. Available templates include: Agent Performance Summary, Cost Analysis Report, Execution Audit "
         "Trail, and Agent Utilization Report. The endpoint shall accept parameters for date range, agent filter (all or "
         "specific agents), and output format (JSON for in-app rendering, PDF for download, CSV for data export). Report "
         "generation for large date ranges shall be asynchronous with a callback notification upon completion."),

        ("FR-RPT-2", "Scheduled Reports", "Should",
         "Users shall be able to schedule recurring report generation on a daily, weekly, or monthly cadence. Scheduled "
         "reports shall be automatically generated and delivered via the user's preferred channel (email attachment or "
         "Slack message with download link). The scheduling system shall reuse the RedBeat infrastructure. Users shall "
         "be able to configure up to 10 scheduled reports per account."),

        ("FR-RPT-3", "Report History", "Must",
         "The system shall maintain a history of generated reports accessible via GET /api/reports. Each report record "
         "shall include the report type, date range, generation timestamp, file size, and download URL. Reports shall "
         "be retained for 30 days before automatic cleanup. Users shall be able to regenerate any historical report "
         "with the same parameters."),

        ("FR-RPT-4", "Custom Report Builder", "Could",
         "The system shall provide a report builder interface allowing users to create custom reports by selecting "
         "metrics (execution count, success rate, cost, duration), dimensions (agent, model, date, trigger type), "
         "and visualization types (table, bar chart, line chart, pie chart). Custom report definitions shall be "
         "savable and shareable within the organization."),

        ("FR-RPT-5", "Report Data API", "Should",
         "The system shall provide a data API endpoint (GET /api/reports/data) that returns raw analytics data in "
         "JSON format with support for aggregation (sum, average, count, min, max), grouping (by agent, model, day, "
         "week, month), and filtering (status, model, date range). This API enables frontend visualizations and "
         "third-party BI tool integrations."),
    ]

    add_req_table(doc, reqs)


def build_section_11(doc):
    add_heading1(doc, "11", "Module M9 — Templates")

    add_body(doc,
        "The Templates module provides a marketplace for sharing and discovering pre-built agent workflows. Templates "
        "reduce the time-to-value for new users by providing ready-made solutions for common use cases. The marketplace "
        "supports both GoalCert-curated official templates and community-contributed templates."
    )

    reqs = [
        ("FR-TPL-1", "Template Browsing", "Must",
         "The system shall provide an endpoint (GET /api/templates) that returns a paginated list of published templates. "
         "Each template record shall include: name, description, category, author, install count, average rating, node "
         "count, estimated execution cost, and a workflow preview thumbnail. The endpoint shall support filtering by "
         "category, sorting by popularity/rating/recency, and full-text search. Categories include: Sales & Marketing, "
         "Research & Analysis, Operations, Finance, HR, IT & DevOps, and Custom."),

        ("FR-TPL-2", "Template Installation", "Must",
         "The system shall provide an endpoint (POST /api/templates/{id}/install) that creates a new agent from a template. "
         "Installation shall deep-copy the template's workflow definition, system prompt, and default configuration into a "
         "new agent owned by the authenticated user. Users shall be prompted to customize agent name, model selection, and "
         "integration credentials (which are never included in templates) before activation. The template's install count "
         "shall be incremented upon successful installation."),

        ("FR-TPL-3", "Template Publishing", "Should",
         "Users with Pro or Admin roles shall be able to publish their agent's workflow as a template via POST "
         "/api/templates. The publishing flow shall: (1) strip all sensitive configuration values (API keys, email "
         "addresses, webhook URLs), (2) allow the author to write a description and select a category, (3) generate "
         "a workflow preview image, (4) submit for review. Published templates shall undergo an automated check for "
         "sensitive data leaks before becoming visible in the marketplace."),

        ("FR-TPL-4", "Template Ratings and Reviews", "Should",
         "Users who have installed a template shall be able to submit a rating (1-5 stars) and optional text review "
         "via POST /api/templates/{id}/reviews. The average rating shall be computed and displayed on the template card. "
         "Template authors shall be able to respond to reviews. Reviews shall be moderated for spam and abuse."),

        ("FR-TPL-5", "Official Template Curation", "Must",
         "The system shall support a distinction between official templates (curated by GoalCert, marked with a verified "
         "badge) and community templates. Official templates shall be seeded at launch covering at least 15 use cases "
         "across all categories. Official templates shall be maintained and updated by the GoalCert team to reflect "
         "platform changes and best practices."),

        ("FR-TPL-6", "Template Versioning", "Could",
         "Templates shall support versioning, allowing authors to publish updates without breaking existing installations. "
         "Users who installed a previous version shall receive a notification when an update is available. Updating shall "
         "create a new agent rather than modifying the existing one, preserving the user's customizations."),
    ]

    add_req_table(doc, reqs)


def build_section_12(doc):
    add_heading1(doc, "12", "Module M10 — Integrations")

    add_body(doc,
        "The Integrations module provides connectors that allow agents to interact with external services as part of "
        "their workflow execution. Integrations are exposed as Integration nodes in the workflow builder and as "
        "delivery channels for reports and notifications. The module is designed with an extensible connector "
        "architecture that allows new integrations to be added via the plugin SDK."
    )

    reqs = [
        ("FR-INT-1", "Email Integration (SMTP)", "Must",
         "The system shall provide an email integration that sends email messages via a configured SMTP server. "
         "The Integration node shall accept recipient email address(es), subject line, and HTML body — all supporting "
         "variable substitution. The system shall support both platform-managed SMTP (using GoalCert's SES configuration) "
         "and user-configured SMTP settings for custom domains. Email delivery status shall be tracked and reported in "
         "execution logs. Failed deliveries shall be retried up to 3 times with exponential backoff."),

        ("FR-INT-2", "Slack Integration", "Must",
         "The system shall provide a Slack integration that sends messages to specified Slack channels or users via "
         "the Slack Web API. Users shall authenticate their Slack workspace through an OAuth 2.0 flow managed by the "
         "platform. The Integration node shall accept channel/user selection, message text (with variable substitution "
         "and Slack markdown formatting), and optional attachment blocks. The integration shall support both posting "
         "new messages and threading replies to existing messages."),

        ("FR-INT-3", "Webhook Integration (Outgoing)", "Must",
         "The system shall provide a generic webhook integration that sends HTTP POST requests to user-specified URLs. "
         "The Integration node shall accept the target URL, HTTP headers (with variable substitution), and request body "
         "(JSON format with variable substitution). The system shall support custom authentication headers (Bearer token, "
         "API key, basic auth). Responses shall be captured in the node output. Timeout shall be configurable (default "
         "30 seconds) with retry support."),

        ("FR-INT-4", "Integration Credential Management", "Must",
         "The system shall provide a secure credential store (GET/POST /api/integrations/credentials) where users manage "
         "authentication credentials for their integrations. Credentials shall be encrypted at rest using AES-256 and "
         "shall never be included in API responses (only credential names and types are returned). Integration nodes "
         "shall reference credentials by ID rather than containing sensitive values directly. Credential deletion shall "
         "be blocked if any active agent references the credential."),

        ("FR-INT-5", "Webhook Trigger (Incoming)", "Must",
         "The system shall provide per-agent webhook endpoints (POST /api/webhooks/{agent_webhook_id}) that accept "
         "incoming HTTP requests and trigger agent execution. Each agent shall have a unique, non-guessable webhook URL "
         "generated upon creation. Webhook security shall be enforced via HMAC-SHA256 signature verification. The "
         "webhook payload shall be available to the workflow as {trigger_data}. Webhook URLs shall be regeneratable "
         "(invalidating the previous URL) via the agent management interface."),

        ("FR-INT-6", "Integration Health Monitoring", "Should",
         "The system shall periodically verify the health of configured integrations by performing connectivity tests "
         "(e.g., sending a test email, posting a test message to Slack, making a HEAD request to webhook URLs). "
         "Unhealthy integrations shall be flagged in the UI with the specific error. Users shall be able to manually "
         "trigger a health check for any integration."),

        ("FR-INT-7", "Integration Event Logging", "Must",
         "All integration interactions (outbound requests and inbound webhooks) shall be logged with timestamp, "
         "integration type, target, HTTP status code, response time, and success/failure status. Logs shall be "
         "accessible via the execution detail view and the integration management page. Sensitive data (credentials, "
         "email bodies) shall be redacted in logs by default, with a 'show sensitive' toggle for authorized users."),
    ]

    add_req_table(doc, reqs)


def build_section_13(doc):
    add_heading1(doc, "13", "Module M11 — JILLA AI Concierge")

    add_body(doc,
        "JILLA is an AI-powered concierge system embedded within the AutoMind platform. JILLA serves two primary "
        "functions: onboarding new users by guiding them through platform features and their first agent creation, "
        "and assisting sales teams during live client demonstrations by providing interactive, intelligent walkthroughs "
        "of platform capabilities. JILLA is accessible via a persistent chat widget on every page of the application."
    )

    reqs = [
        ("FR-JILLA-1", "Persistent Chat Widget", "Must",
         "The JILLA chat widget shall appear as a floating button in the bottom-right corner of every page in the "
         "application. Clicking the button shall open a chat panel that maintains conversation state across page "
         "navigations. The widget shall support minimize, maximize, and close actions. New users shall see an "
         "automatic greeting message from JILLA upon their first login, offering to guide them through the platform."),

        ("FR-JILLA-2", "Guided Onboarding Flow", "Must",
         "JILLA shall provide an interactive onboarding flow that walks new users through: (1) understanding the "
         "platform concept and terminology, (2) creating their first agent with a guided workflow builder tour, "
         "(3) configuring a simple two-node workflow (Trigger + AI Action), (4) running their first execution and "
         "viewing results, (5) exploring the template marketplace. Each step shall include contextual highlights "
         "on the relevant UI elements and progress indicators. Users shall be able to skip or restart the onboarding "
         "at any time."),

        ("FR-JILLA-3", "Feature Explanation", "Should",
         "JILLA shall be able to explain any platform feature in natural language when asked. JILLA's knowledge base "
         "shall include comprehensive documentation about all node types, workflow building, scheduling, memory, chat, "
         "integrations, templates, and analytics. Responses shall include links to relevant documentation and, where "
         "applicable, direct navigation links to the feature being discussed."),

        ("FR-JILLA-4", "Demo Mode for Sales", "Should",
         "JILLA shall support a 'demo mode' activated by Admin or Super Admin users that provides guided demonstration "
         "scripts for client presentations. Demo mode shall include pre-configured talking points, interactive workflow "
         "demonstrations with sample data, and competitive differentiation responses. JILLA shall adapt the demo flow "
         "based on the prospect's industry (e.g., finance, marketing, operations) by adjusting use case examples and "
         "template recommendations."),

        ("FR-JILLA-5", "Contextual Assistance", "Could",
         "JILLA shall offer proactive contextual assistance based on the user's current activity. For example, if a "
         "user is configuring an AI Action node and leaves the prompt template empty for more than 30 seconds, JILLA "
         "shall offer prompt writing tips. If a workflow validation fails, JILLA shall explain the errors and suggest "
         "fixes. Proactive suggestions shall be dismissable and respect a 'do not disturb' preference."),
    ]

    add_req_table(doc, reqs)


def build_section_14(doc):
    add_heading1(doc, "14", "Non-Functional Requirements")

    add_heading2(doc, "14.1", "Performance")
    add_body(doc,
        "Performance requirements define the acceptable response times and throughput for platform operations under "
        "normal and peak load conditions."
    )

    add_styled_table(doc,
        ["ID", "Requirement", "Target", "Measurement"],
        [
            ("NFR-PERF-1", "API response time for non-AI endpoints (CRUD, list, auth)", "P95 < 200ms", "Application Performance Monitoring (APM) under 100 concurrent users"),
            ("NFR-PERF-2", "Workflow canvas rendering for workflows up to 50 nodes", "< 2 seconds initial load", "Browser performance profiling (Chrome DevTools)"),
            ("NFR-PERF-3", "SSE event delivery latency from worker to client", "< 500ms", "End-to-end timing in execution logs"),
            ("NFR-PERF-4", "Dashboard page load time", "< 3 seconds on initial load, < 1 second on subsequent", "Real User Monitoring (RUM)"),
            ("NFR-PERF-5", "Concurrent agent executions supported", "500 simultaneous", "Load testing with simulated workflows"),
            ("NFR-PERF-6", "Database query performance for paginated list endpoints", "< 50ms per query", "PostgreSQL query plan analysis"),
        ],
        col_widths=[2.5, 5, 3.5, 5],
    )

    add_heading2(doc, "14.2", "Security")
    add_body(doc,
        "Security requirements define the controls necessary to protect user data, prevent unauthorized access, and "
        "maintain the confidentiality of agent workflows and execution results."
    )

    sec_reqs = [
        ("NFR-SEC-1", "All API communication shall be encrypted via TLS 1.2 or higher. HTTP requests shall be redirected to HTTPS."),
        ("NFR-SEC-2", "Passwords shall be hashed using bcrypt with a minimum work factor of 12. Plaintext passwords shall never be logged or stored."),
        ("NFR-SEC-3", "Integration credentials shall be encrypted at rest using AES-256 with keys managed via AWS KMS or equivalent."),
        ("NFR-SEC-4", "Code Execution nodes shall run in isolated sandboxed environments with no network access, no filesystem access beyond a temporary directory, and restricted Python imports."),
        ("NFR-SEC-5", "All user inputs shall be sanitized to prevent SQL injection, XSS, and command injection attacks. The ORM layer (SQLAlchemy) shall be used for all database queries."),
        ("NFR-SEC-6", "API rate limiting shall be enforced at 100 requests per minute per user for general endpoints and 10 requests per minute for authentication endpoints."),
        ("NFR-SEC-7", "Audit logs shall record all authentication events, role changes, agent activations, and data exports with user ID, IP address, timestamp, and action performed."),
        ("NFR-SEC-8", "Multi-tenant data isolation shall be enforced at the ORM query layer, ensuring all queries are scoped to the authenticated user's organization ID. No cross-tenant data access shall be possible."),
    ]

    for req_id, desc in sec_reqs:
        p = doc.add_paragraph()
        run = p.add_run(f"{req_id}: ")
        run.font.name = FONT
        run.font.size = BODY_SIZE
        run.font.color.rgb = PURPLE
        run.bold = True
        run2 = p.add_run(desc)
        run2.font.name = FONT
        run2.font.size = BODY_SIZE
        run2.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(4)

    add_heading2(doc, "14.3", "Scalability")
    add_body(doc,
        "The platform architecture shall support horizontal scaling of all stateless components. The following "
        "scalability requirements define the growth targets for the first 12 months."
    )

    add_styled_table(doc,
        ["ID", "Requirement", "Target"],
        [
            ("NFR-SCALE-1", "API server instances shall be horizontally scalable behind a load balancer", "Auto-scale 2-10 instances based on CPU/request count"),
            ("NFR-SCALE-2", "Celery worker pool shall scale independently based on queue depth", "Auto-scale 2-20 workers based on pending task count"),
            ("NFR-SCALE-3", "Database shall support the expected data volume for 12 months", "1,000 users, 5,000 agents, 500,000 execution records, 50GB storage"),
            ("NFR-SCALE-4", "Redis cluster shall handle concurrent pub/sub and caching load", "10,000 active SSE connections, 50,000 pub/sub messages per minute"),
            ("NFR-SCALE-5", "SSE streaming shall support concurrent connections without degrading API performance", "1,000 simultaneous SSE connections"),
        ],
        col_widths=[2.5, 7, 6.5],
    )

    add_heading2(doc, "14.4", "Availability")
    add_body(doc,
        "Availability requirements define the uptime targets and disaster recovery capabilities."
    )

    add_styled_table(doc,
        ["ID", "Requirement", "Target"],
        [
            ("NFR-AVAIL-1", "Platform uptime (excluding scheduled maintenance)", "99.9% (< 8.76 hours downtime per year)"),
            ("NFR-AVAIL-2", "Scheduled maintenance windows", "Maximum 2 hours per month, announced 48 hours in advance"),
            ("NFR-AVAIL-3", "Database backup frequency and retention", "Automated daily backups, 30-day retention, point-in-time recovery"),
            ("NFR-AVAIL-4", "Recovery Time Objective (RTO) for major incidents", "< 4 hours"),
            ("NFR-AVAIL-5", "Recovery Point Objective (RPO) for data loss", "< 1 hour (continuous WAL archiving)"),
            ("NFR-AVAIL-6", "Graceful degradation under partial outage", "If Redis is unavailable, scheduled execution queues but live monitoring degrades gracefully"),
        ],
        col_widths=[2.5, 7, 6.5],
    )


def build_section_15(doc):
    add_heading1(doc, "15", "Integration Points")

    add_body(doc,
        "This section documents the external system integration points that the AutoMind platform depends on or "
        "interacts with. Each integration point identifies the external system, the nature of the interaction, the "
        "protocol used, and the failure handling strategy."
    )

    add_styled_table(doc,
        ["External System", "Direction", "Protocol", "Purpose", "Failure Handling"],
        [
            ("OpenAI API", "Outbound", "HTTPS REST", "AI Action node execution (GPT-4, GPT-4o models)", "Retry 3x with exponential backoff; fall back to Anthropic if configured; circuit breaker after 5 consecutive failures"),
            ("Anthropic API", "Outbound", "HTTPS REST", "AI Action node execution (Claude models)", "Retry 3x with exponential backoff; fall back to OpenAI if configured; circuit breaker after 5 consecutive failures"),
            ("SMTP Server (AWS SES)", "Outbound", "SMTP/TLS", "Email delivery for Integration nodes and report distribution", "Retry 3x with 30-second intervals; queue failed emails for manual retry; alert user on permanent failure"),
            ("Slack Web API", "Outbound", "HTTPS REST", "Slack message delivery for Integration nodes and notifications", "Retry 3x; handle rate limiting (429 responses) with Retry-After header; alert on token revocation"),
            ("Stripe API", "Outbound", "HTTPS REST", "Subscription billing, payment processing, usage metering", "Retry with idempotency keys; webhook reconciliation for missed events; manual intervention for payment failures"),
            ("PostgreSQL", "Bidirectional", "TCP (psycopg2)", "Primary data persistence for all application data", "Connection pooling (SQLAlchemy); automatic reconnection; read replica failover for read-heavy endpoints"),
            ("Redis", "Bidirectional", "TCP (redis-py)", "Task queue, schedule storage, SSE pub/sub, caching", "Connection retry with backoff; graceful degradation (scheduled tasks queue, monitoring unavailable); sentinel failover"),
            ("Google Search API", "Outbound", "HTTPS REST", "Web Search node execution for internet research", "Retry 3x; cache recent results for identical queries (5-minute TTL); fallback to Bing API if configured"),
        ],
        col_widths=[2.5, 1.5, 2, 4.5, 5.5],
    )


def build_section_16(doc):
    add_heading1(doc, "16", "Traceability Matrix")

    add_body(doc,
        "The following traceability matrix maps each functional requirement to its originating business requirement(s) "
        "from the BRD. This ensures complete coverage of all business needs and provides a reference for impact analysis "
        "when requirements change."
    )

    matrix = [
        ("FR-AUTH-1 to FR-AUTH-7", "BR-12 (RBAC), BR-09 (Multi-Tenant)", "Authentication and access control foundation"),
        ("FR-AGT-1 to FR-AGT-8", "BR-01 (Visual Builder), BR-03 (Scheduling)", "Agent lifecycle management and scheduling"),
        ("FR-WF-1 to FR-WF-8", "BR-01 (Visual Builder)", "Workflow canvas, node types, and configuration"),
        ("FR-EXEC-1 to FR-EXEC-10", "BR-03 (Scheduling), BR-04 (Monitoring), BR-10 (Cost Tracking)", "Execution engine, streaming, and cost tracking"),
        ("FR-MEM-1 to FR-MEM-6", "BR-05 (Agent Memory)", "Memory generation, retrieval, and RAG injection"),
        ("FR-CHAT-1 to FR-CHAT-6", "BR-06 (Chat Interface)", "Conversational agent interface and streaming"),
        ("FR-DASH-1 to FR-DASH-6", "BR-04 (Monitoring), BR-10 (Cost Tracking), BR-11 (Reports)", "Dashboard and analytics visualizations"),
        ("FR-RPT-1 to FR-RPT-5", "BR-11 (Reports)", "Report generation, scheduling, and export"),
        ("FR-TPL-1 to FR-TPL-6", "BR-08 (Template Marketplace)", "Template browsing, installation, and publishing"),
        ("FR-INT-1 to FR-INT-7", "BR-07 (Integration Ecosystem)", "Email, Slack, webhook, and credential management"),
        ("FR-JILLA-1 to FR-JILLA-5", "BR-14 (JILLA Concierge)", "Onboarding, feature explanation, and demo mode"),
        ("NFR-PERF-1 to NFR-PERF-6", "BR-13 (Production Deployment)", "Performance targets"),
        ("NFR-SEC-1 to NFR-SEC-8", "BR-09 (Multi-Tenant), BR-12 (RBAC)", "Security controls"),
        ("NFR-SCALE-1 to NFR-SCALE-5", "BR-13 (Production Deployment)", "Scalability targets"),
        ("NFR-AVAIL-1 to NFR-AVAIL-6", "BR-13 (Production Deployment)", "Availability and disaster recovery"),
    ]

    add_styled_table(doc,
        ["Functional Requirements", "Business Requirements (BRD)", "Coverage Area"],
        matrix,
        col_widths=[4, 5, 7],
    )

    add_body(doc,
        "All fifteen business requirements (BR-01 through BR-15) are covered by one or more functional requirement "
        "groups. The traceability matrix should be updated whenever requirements are added, modified, or removed to "
        "maintain alignment between the BRD and FRD throughout the project lifecycle."
    )


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT
    font.size = BODY_SIZE
    font.color.rgb = DARK_GRAY

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_cover_page(doc)
    add_document_control(doc)
    add_toc(doc)

    build_section_1(doc)
    doc.add_page_break()
    build_section_2(doc)
    doc.add_page_break()
    build_section_3(doc)
    doc.add_page_break()
    build_section_4(doc)
    doc.add_page_break()
    build_section_5(doc)
    doc.add_page_break()
    build_section_6(doc)
    doc.add_page_break()
    build_section_7(doc)
    doc.add_page_break()
    build_section_8(doc)
    doc.add_page_break()
    build_section_9(doc)
    doc.add_page_break()
    build_section_10(doc)
    doc.add_page_break()
    build_section_11(doc)
    doc.add_page_break()
    build_section_12(doc)
    doc.add_page_break()
    build_section_13(doc)
    doc.add_page_break()
    build_section_14(doc)
    doc.add_page_break()
    build_section_15(doc)
    doc.add_page_break()
    build_section_16(doc)

    doc.save(SAVE_PATH)
    print(f"FRD saved to {SAVE_PATH}")


if __name__ == "__main__":
    main()
