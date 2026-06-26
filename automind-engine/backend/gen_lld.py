"""Generate GoalCert AutoMind Low-Level Design (LLD) document as .docx"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colours ────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "E8EEF4"
ALT_ROW = "F0F4F8"
HEADER_HEX = "1B2A4A"

OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__), "..", "docs", "GoalCert_AutoMind_LLD_v1_0.docx"
)


# ── Helper functions ───────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run(run, font_name="Calibri", size=11, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_styled_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=20, bold=True, color=NAVY)
        p.style = doc.styles["Heading 1"]
        for r in p.runs:
            set_run(r, size=20, bold=True, color=NAVY)
    elif level == 2:
        p.style = doc.styles["Heading 2"]
        for r in p.runs:
            set_run(r, size=15, bold=True, color=BLUE)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=10, bold=True, color=WHITE)
        set_cell_shading(cell, HEADER_HEX)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=10)
            if row_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.text = ""
    run = p.add_run(text)
    set_run(run, size=11)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=11)
    return p


# ── Document construction ─────────────────────────────────────

def build_document():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── COVER PAGE ──────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GoalCert AutoMind")
    set_run(run, size=32, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Low-Level Design (LLD)")
    set_run(run, size=20, color=BLUE)

    doc.add_page_break()

    # ── METADATA TABLE ──────────────────────────────────────
    meta_data = [
        ("Document Title", "GoalCert AutoMind - Low-Level Design (LLD)"),
        ("Version", "1.0"),
        ("Date", "25 June 2026"),
        ("Author", "Narhen Karthikeyan"),
        ("Reviewed By", "Prem (GoalCert)"),
        ("Status", "Draft"),
        ("Classification", "Internal"),
        ("Project", "GoalCert AutoMind Agentic AI Platform"),
        ("Technology Stack", "FastAPI, PostgreSQL, Redis, Celery, React"),
    ]
    table = doc.add_table(rows=len(meta_data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_data):
        left = table.rows[i].cells[0]
        right = table.rows[i].cells[1]
        left.text = ""
        right.text = ""
        run_l = left.paragraphs[0].add_run(label)
        set_run(run_l, size=11, bold=True)
        run_r = right.paragraphs[0].add_run(value)
        set_run(run_r, size=11)
        set_cell_shading(left, LIGHT_BG)

    doc.add_paragraph()

    # ── REVISION HISTORY ────────────────────────────────────
    add_styled_heading(doc, "Revision History", level=2)
    make_table(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ("1.0", "25 Jun 2026", "Narhen Karthikeyan", "Initial LLD based on implemented codebase"),
        ],
    )

    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    set_run(run, size=20, bold=True, color=NAVY)

    toc_l1 = [
        "1. Introduction",
        "2. Component Diagram",
        "3. API Endpoint Specification",
        "4. Database Schema",
        "5. Entity Relationships",
        "6. Sequence Diagrams",
        "7. Error Code Catalog",
        "8. Data Validation Rules",
    ]
    toc_l2 = {
        "3. API Endpoint Specification": [
            "3.1 Authentication APIs",
            "3.2 Agent APIs",
            "3.3 Execution APIs",
            "3.4 Chat APIs",
            "3.5 Template APIs",
            "3.6 Analytics APIs",
            "3.7 Report APIs",
            "3.8 Integration APIs",
        ],
        "4. Database Schema": [
            "4.1 users",
            "4.2 agents",
            "4.3 workflows",
            "4.4 executions",
            "4.5 execution_node_logs",
            "4.6 agent_memory",
            "4.7 agent_templates",
            "4.8 integrations",
        ],
        "6. Sequence Diagrams": [
            "6.1 Agent Execution Flow",
            "6.2 Chat with Memory",
            "6.3 Scheduled Execution",
        ],
    }

    for item in toc_l1:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run(run, size=11, bold=True, color=NAVY)
        if item in toc_l2:
            for sub in toc_l2[item]:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(1.5)
                run2 = p2.add_run(sub)
                set_run(run2, size=11, color=NAVY)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 1 - INTRODUCTION
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "1. Introduction", level=1)
    add_body(doc, (
        "This Low-Level Design document describes the internal architecture, API contracts, "
        "database schemas, data flows, error handling, and validation rules for the GoalCert "
        "AutoMind Agentic AI Platform. It is derived directly from the implemented codebase and "
        "serves as the single source of truth for developers, reviewers, and future maintainers."
    ))
    add_body(doc, (
        "AutoMind is a visual workflow-builder platform that lets non-technical users create, "
        "deploy, and monitor autonomous AI agents. Each agent is backed by a directed graph of "
        "executable nodes (LLM calls, web searches, code execution, integrations, decisions, and "
        "escalations) orchestrated by a Celery-based execution engine with real-time streaming "
        "via Redis pub/sub and Server-Sent Events."
    ))

    add_styled_heading(doc, "1.1 Scope", level=2)
    scope_items = [
        "REST API endpoint specifications with request/response schemas",
        "PostgreSQL database schema with column-level detail",
        "Entity relationship mappings",
        "Sequence diagrams for critical flows",
        "Error code catalog and HTTP status conventions",
        "Data validation and business rules",
    ]
    for s in scope_items:
        add_bullet(doc, s)

    add_styled_heading(doc, "1.2 Technology Stack", level=2)
    stack_items = [
        "Backend Framework: FastAPI 0.115+ (async, Python 3.12)",
        "ORM: SQLAlchemy 2.0 with asyncpg driver",
        "Database: PostgreSQL 15+",
        "Task Queue: Celery 5.x with RedBeat scheduler",
        "Message Broker / Cache: Redis 7+",
        "Auth: JWT (HS256) via PyJWT + bcrypt password hashing",
        "LLM Provider: OpenAI API (gpt-4o-mini)",
        "Frontend: React 18 with React Flow (workflow canvas)",
        "Real-time: SSE (sse-starlette) + Redis pub/sub",
    ]
    for s in stack_items:
        add_bullet(doc, s)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 2 - COMPONENT DIAGRAM
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "2. Component Diagram", level=1)
    add_body(doc, (
        "The AutoMind platform follows a layered architecture with clear separation between "
        "the API layer, business logic services, and the asynchronous execution engine."
    ))

    add_styled_heading(doc, "2.1 System Components", level=2)

    components = [
        ("React Frontend (SPA)", (
            "Visual workflow builder using React Flow. Handles agent CRUD, execution monitoring "
            "via SSE, chat interface, dashboard analytics, and integration management."
        )),
        ("FastAPI Application Server", (
            "Async Python HTTP server hosting 9 router modules: auth, agents, workflows, "
            "executions, templates, integrations, memory, chat, and dashboard. Handles JWT "
            "authentication, request validation, and database operations through SQLAlchemy."
        )),
        ("Celery Worker", (
            "Runs workflow executions asynchronously. Each execution creates a WorkflowExecutor "
            "that traverses the workflow graph using BFS, executing nodes sequentially. Supports "
            "both manual triggers and scheduled (cron) triggers via RedBeat."
        )),
        ("Workflow Execution Engine", (
            "Core execution logic in app/engine/. Parses workflow JSON (nodes + edges), "
            "resolves execution order, and dispatches to 7 node executors: trigger, ai_action, "
            "integration, decision, escalation, web_search, and code_exec."
        )),
        ("PostgreSQL Database", (
            "Persistent storage for 8 tables: users, agents, workflows, executions, "
            "execution_node_logs, agent_memory, agent_templates, and integrations. Uses UUIDs "
            "as primary keys with gen_random_uuid() server defaults."
        )),
        ("Redis", (
            "Serves three roles: (1) Celery message broker for task queuing, (2) RedBeat "
            "scheduler backend for cron-triggered agent runs, (3) pub/sub channel for real-time "
            "execution log streaming to the frontend via SSE."
        )),
        ("Memory Service", (
            "Generates and stores execution summaries after each workflow run. Optionally uses "
            "GPT-4o-mini to create concise summaries. Provides context injection for chat and "
            "future AI action nodes."
        )),
        ("Agent Generator Service", (
            "Converts natural language descriptions into complete workflow definitions using "
            "GPT-4o-mini. Returns structured JSON with nodes, edges, schedule, and metadata."
        )),
    ]

    for name, desc in components:
        p = doc.add_paragraph()
        run_name = p.add_run(name + ": ")
        set_run(run_name, size=11, bold=True)
        run_desc = p.add_run(desc)
        set_run(run_desc, size=11)

    add_styled_heading(doc, "2.2 Inter-Component Communication", level=2)
    comms = [
        "Frontend to API Server: HTTPS REST (JSON) + SSE for real-time streaming",
        "API Server to Database: SQLAlchemy async sessions (asyncpg connection pool)",
        "API Server to Celery: celery_app.send_task() over Redis broker",
        "Celery Worker to Database: Dedicated async engine per execution (pool_size=5)",
        "Celery Worker to Redis: Synchronous redis.publish() for execution log streaming",
        "Frontend to SSE: GET /api/executions/{id}/stream, subscribes to Redis pub/sub channel",
        "RedBeat Scheduler to Celery: Fires execute_workflow_scheduled task on cron match",
    ]
    for c in comms:
        add_bullet(doc, c)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 3 - API ENDPOINT SPECIFICATION
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "3. API Endpoint Specification", level=1)
    add_body(doc, (
        "All endpoints are prefixed with /api. Authentication is via Bearer JWT token in the "
        "Authorization header unless noted otherwise. Responses follow standard HTTP status codes "
        "with JSON error bodies containing a \"detail\" field."
    ))

    # 3.1 Authentication APIs
    add_styled_heading(doc, "3.1 Authentication APIs", level=2)
    add_body(doc, "Router prefix: /api/auth. No authentication required for register and login.")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("POST", "/api/auth/register",
             '{ email: str, password: str (min 6), name?: str }',
             '201: { user: UserResponse, token: str }',
             "No"),
            ("POST", "/api/auth/login",
             '{ email: str, password: str }',
             '200: { user: UserResponse, token: str }',
             "No"),
            ("GET", "/api/auth/me",
             "None",
             '200: { id, email, name, created_at }',
             "Yes"),
        ],
    )

    # 3.2 Agent APIs
    doc.add_paragraph()
    add_styled_heading(doc, "3.2 Agent APIs", level=2)
    add_body(doc, "Router prefix: /api/agents. All endpoints require authentication. Agents are scoped to the authenticated user.")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("GET", "/api/agents",
             "None",
             "200: AgentResponse[]",
             "Yes"),
            ("POST", "/api/agents",
             '{ name: str (2-50), type: sales|marketing|support|custom, description?: str, template_id?: str }',
             "201: AgentResponse",
             "Yes"),
            ("POST", "/api/agents/generate",
             '{ description: str }',
             "201: AgentResponse (AI-generated workflow)",
             "Yes"),
            ("GET", "/api/agents/{agent_id}",
             "None",
             "200: AgentResponse",
             "Yes"),
            ("PATCH", "/api/agents/{agent_id}",
             '{ name?, description?, schedule_cron?, schedule_timezone? }',
             "200: AgentResponse",
             "Yes"),
            ("DELETE", "/api/agents/{agent_id}",
             "None",
             "204: No Content",
             "Yes"),
            ("POST", "/api/agents/{agent_id}/pause",
             "None",
             "200: AgentResponse (status=paused)",
             "Yes"),
            ("POST", "/api/agents/{agent_id}/resume",
             "None",
             "200: AgentResponse (status=active)",
             "Yes"),
        ],
    )

    # 3.3 Execution APIs
    doc.add_paragraph()
    add_styled_heading(doc, "3.3 Execution APIs", level=2)
    add_body(doc, "Router prefix: /api. Handles workflow execution triggers, listing, detail views, and real-time streaming.")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("POST", "/api/agents/{agent_id}/execute",
             "None",
             "202: ExecutionResponse (queued to Celery)",
             "Yes"),
            ("GET", "/api/agents/{agent_id}/executions",
             "Query: limit (1-100, default 20), offset (>=0)",
             "200: ExecutionResponse[]",
             "Yes"),
            ("GET", "/api/executions/{execution_id}",
             "None",
             "200: { execution: ExecutionResponse, node_logs: NodeLogResponse[] }",
             "Yes"),
            ("GET", "/api/executions/{execution_id}/logs",
             "None",
             "200: NodeLogResponse[]",
             "Yes"),
            ("GET", "/api/executions/{execution_id}/stream",
             "None (SSE)",
             "SSE stream: { timestamp, message, node_id, status }",
             "No (UUID-based)"),
        ],
    )

    # 3.4 Chat APIs
    doc.add_paragraph()
    add_styled_heading(doc, "3.4 Chat APIs", level=2)
    add_body(doc, "Streaming chat with an agent using its execution memory as context. Uses OpenAI GPT-4o-mini for response generation.")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("POST", "/api/agents/{agent_id}/chat",
             '{ message: str, history: [{ role: user|assistant, content: str }] }',
             "SSE stream: { type: token|done|error, content?: str }",
             "Yes"),
        ],
    )

    # 3.5 Template APIs
    doc.add_paragraph()
    add_styled_heading(doc, "3.5 Template APIs", level=2)
    add_body(doc, "Router prefix: /api/templates. Public read-only endpoints for agent templates. No authentication required.")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("GET", "/api/templates",
             "None",
             "200: TemplateResponse[] (ordered by name)",
             "No"),
            ("GET", "/api/templates/{template_id}",
             "None",
             "200: TemplateResponse",
             "No"),
        ],
    )

    # 3.6 Analytics APIs
    doc.add_paragraph()
    add_styled_heading(doc, "3.6 Analytics APIs (Dashboard)", level=2)
    add_body(doc, "Router prefix: /api/dashboard. Provides aggregated statistics and recent activity for the authenticated user.")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("GET", "/api/dashboard/stats",
             "None",
             "200: { total_agents, active_agents, tasks_completed, estimated_savings, avg_response_time }",
             "Yes"),
            ("GET", "/api/dashboard/activity",
             "None",
             "200: ActivityEvent[] (last 20 executions with agent info)",
             "Yes"),
        ],
    )

    # 3.7 Report APIs (Workflow/Memory)
    doc.add_paragraph()
    add_styled_heading(doc, "3.7 Workflow and Memory APIs", level=2)
    add_body(doc, "Workflow management is nested under /api/agents/{agent_id}. Memory endpoints manage execution summaries stored for agent context.")

    add_body(doc, "Workflow Endpoints:")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("GET", "/api/agents/{agent_id}/workflow",
             "None",
             "200: WorkflowResponse",
             "Yes"),
            ("PUT", "/api/agents/{agent_id}/workflow",
             '{ definition: dict (React Flow JSON) }',
             "200: WorkflowResponse",
             "Yes"),
            ("POST", "/api/agents/{agent_id}/workflow/deploy",
             "None",
             "200: WorkflowResponse (status=active, deploys schedule)",
             "Yes"),
        ],
    )

    doc.add_paragraph()
    add_body(doc, "Memory Endpoints:")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("GET", "/api/agents/{agent_id}/memory",
             "Query: limit (default 20), offset (default 0)",
             "200: { memories: MemoryResponse[], total: int }",
             "Yes"),
            ("DELETE", "/api/agents/{agent_id}/memory",
             "None",
             "204: No Content (clears all memories)",
             "Yes"),
        ],
    )

    # 3.8 Integration APIs
    doc.add_paragraph()
    add_styled_heading(doc, "3.8 Integration APIs", level=2)
    add_body(doc, "Router prefix: /api/integrations. Manages third-party service connections (Resend email, Slack). Sensitive config values are masked in responses.")
    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Auth"],
        [
            ("GET", "/api/integrations",
             "None",
             "200: IntegrationResponse[] (config values masked)",
             "Yes"),
            ("POST", "/api/integrations",
             '{ service: "resend"|"slack", config: dict }',
             "201: IntegrationResponse (upserts if exists)",
             "Yes"),
            ("DELETE", "/api/integrations/{integration_id}",
             "None",
             "204: No Content",
             "Yes"),
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 4 - DATABASE SCHEMA
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "4. Database Schema", level=1)
    add_body(doc, (
        "All tables use PostgreSQL UUIDs as primary keys with gen_random_uuid() server defaults. "
        "Timestamps are timezone-aware (TIMESTAMP WITH TIME ZONE) with server_default=now(). "
        "The database uses asyncpg as the connection driver."
    ))

    # 4.1 users
    add_styled_heading(doc, "4.1 users", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique user identifier"),
            ("email", "VARCHAR(255)", "UNIQUE, NOT NULL", "User email address (login credential)"),
            ("password_hash", "VARCHAR(255)", "NOT NULL", "bcrypt-hashed password"),
            ("name", "VARCHAR(255)", "NULLABLE", "Display name"),
            ("created_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now()", "Account creation timestamp"),
            ("updated_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now(), ON UPDATE now()", "Last modification timestamp"),
        ],
    )

    # 4.2 agents
    doc.add_paragraph()
    add_styled_heading(doc, "4.2 agents", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique agent identifier"),
            ("user_id", "UUID", "FK -> users.id ON DELETE CASCADE, NOT NULL", "Owning user"),
            ("name", "VARCHAR(255)", "NOT NULL, UNIQUE(user_id, name)", "Agent name (unique per user)"),
            ("description", "TEXT", "NULLABLE", "Agent purpose description"),
            ("type", "VARCHAR(50)", "NOT NULL", "Agent category: sales, marketing, support, custom"),
            ("status", "VARCHAR(50)", "DEFAULT 'draft'", "Current state: draft, active, paused"),
            ("schedule_cron", "VARCHAR(100)", "NULLABLE", "Cron expression for scheduled runs"),
            ("schedule_timezone", "VARCHAR(50)", "DEFAULT 'UTC'", "IANA timezone for cron schedule"),
            ("created_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now()", "Creation timestamp"),
            ("updated_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now(), ON UPDATE now()", "Last modification timestamp"),
            ("last_execution_at", "TIMESTAMPTZ", "NULLABLE", "Timestamp of most recent execution"),
        ],
    )

    # 4.3 workflows
    doc.add_paragraph()
    add_styled_heading(doc, "4.3 workflows", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique workflow identifier"),
            ("agent_id", "UUID", "FK -> agents.id ON DELETE CASCADE, UNIQUE, NOT NULL", "One-to-one link to agent"),
            ("status", "VARCHAR(50)", "DEFAULT 'draft'", "Workflow state: draft, active"),
            ("definition", "JSONB", "NOT NULL, DEFAULT '{}'", "React Flow JSON: { nodes, edges, viewport }"),
            ("created_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now()", "Creation timestamp"),
            ("updated_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now(), ON UPDATE now()", "Last modification timestamp"),
            ("deployed_at", "TIMESTAMPTZ", "NULLABLE", "Timestamp of last deployment"),
        ],
    )

    # 4.4 executions
    doc.add_paragraph()
    add_styled_heading(doc, "4.4 executions", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique execution identifier"),
            ("agent_id", "UUID", "FK -> agents.id ON DELETE CASCADE, NOT NULL", "Agent that owns this execution"),
            ("workflow_id", "UUID", "FK -> workflows.id, NOT NULL", "Workflow version that was executed"),
            ("status", "VARCHAR(50)", "DEFAULT 'pending'", "Execution state: pending, running, success, failed"),
            ("triggered_by", "VARCHAR(50)", "NOT NULL", "Trigger source: manual, schedule"),
            ("started_at", "TIMESTAMPTZ", "NULLABLE", "Execution start time"),
            ("ended_at", "TIMESTAMPTZ", "NULLABLE", "Execution completion time"),
            ("duration_ms", "INTEGER", "NULLABLE", "Total execution duration in milliseconds"),
            ("error_message", "TEXT", "NULLABLE", "Error details if execution failed"),
            ("variables", "JSONB", "NOT NULL, DEFAULT '{}'", "Accumulated workflow variables (inputs + outputs)"),
            ("total_cost", "NUMERIC(10,6)", "NOT NULL, DEFAULT 0", "Total LLM API cost in USD"),
            ("created_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now()", "Record creation timestamp"),
        ],
    )

    # 4.5 execution_node_logs
    doc.add_paragraph()
    add_styled_heading(doc, "4.5 execution_node_logs", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique log entry identifier"),
            ("execution_id", "UUID", "FK -> executions.id ON DELETE CASCADE, NOT NULL", "Parent execution"),
            ("node_id", "VARCHAR(255)", "NOT NULL", "React Flow node ID from workflow definition"),
            ("node_type", "VARCHAR(50)", "NOT NULL", "Node type: trigger, ai_action, decision, etc."),
            ("node_label", "VARCHAR(255)", "NULLABLE", "Human-readable node label"),
            ("status", "VARCHAR(50)", "DEFAULT 'pending'", "Node state: pending, running, success, failed, skipped"),
            ("started_at", "TIMESTAMPTZ", "NULLABLE", "Node execution start time"),
            ("ended_at", "TIMESTAMPTZ", "NULLABLE", "Node execution end time"),
            ("duration_ms", "INTEGER", "NULLABLE", "Node execution duration in milliseconds"),
            ("input_data", "JSONB", "NOT NULL, DEFAULT '{}'", "Node configuration / input passed to executor"),
            ("output_data", "JSONB", "NOT NULL, DEFAULT '{}'", "Node execution results"),
            ("error_message", "TEXT", "NULLABLE", "Error message if node failed"),
            ("llm_usage", "JSONB", "NULLABLE", "LLM token usage and cost (ai_action nodes)"),
            ("created_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now()", "Record creation timestamp"),
        ],
    )

    # 4.6 agent_memory
    doc.add_paragraph()
    add_styled_heading(doc, "4.6 agent_memory", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique memory entry identifier"),
            ("agent_id", "UUID", "FK -> agents.id ON DELETE CASCADE, NOT NULL, INDEXED", "Owning agent"),
            ("execution_id", "UUID", "FK -> executions.id ON DELETE SET NULL, NULLABLE", "Source execution"),
            ("summary", "TEXT", "NOT NULL", "LLM-generated or fallback execution summary"),
            ("key_outputs", "JSONB", "DEFAULT '{}'", "Extracted key output variables from execution"),
            ("memory_type", "VARCHAR(50)", "DEFAULT 'execution_summary'", "Memory category"),
            ("created_at", "TIMESTAMPTZ", "DEFAULT now()", "Memory creation timestamp"),
        ],
    )

    # 4.7 agent_templates
    doc.add_paragraph()
    add_styled_heading(doc, "4.7 agent_templates", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique template identifier"),
            ("name", "VARCHAR(255)", "NOT NULL", "Template display name"),
            ("description", "TEXT", "NULLABLE", "Template description"),
            ("type", "VARCHAR(50)", "NOT NULL", "Template category: sales, marketing, support, custom"),
            ("workflow_definition", "JSONB", "NOT NULL", "Pre-built workflow JSON (nodes + edges + viewport)"),
            ("icon", "VARCHAR(50)", "NULLABLE", "Icon identifier for UI display"),
            ("color", "VARCHAR(50)", "NULLABLE", "Color identifier for UI display"),
            ("features", "JSONB", "NOT NULL, DEFAULT '[]'", "List of template feature descriptions"),
            ("created_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now()", "Template creation timestamp"),
        ],
    )

    # 4.8 integrations
    doc.add_paragraph()
    add_styled_heading(doc, "4.8 integrations", level=2)
    make_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ("id", "UUID", "PK, DEFAULT gen_random_uuid()", "Unique integration identifier"),
            ("user_id", "UUID", "FK -> users.id ON DELETE CASCADE, NOT NULL", "Owning user"),
            ("service", "VARCHAR(100)", "NOT NULL, UNIQUE(user_id, service)", "Service name: resend, slack"),
            ("config", "JSONB", "NOT NULL, DEFAULT '{}'", "Service configuration (API keys, webhooks)"),
            ("status", "VARCHAR(50)", "DEFAULT 'active'", "Integration state: active, inactive"),
            ("created_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now()", "Creation timestamp"),
            ("updated_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT now(), ON UPDATE now()", "Last modification timestamp"),
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 5 - ENTITY RELATIONSHIPS
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "5. Entity Relationships", level=1)
    add_body(doc, "The following entity relationships define the data model:")

    relationships = [
        ("User 1:N Agent", "A user owns zero or more agents. Deleting a user cascades to all their agents."),
        ("User 1:N Integration", "A user owns zero or more integrations. Unique constraint on (user_id, service) prevents duplicate service connections."),
        ("Agent 1:1 Workflow", "Each agent has exactly one workflow (one-to-one via unique FK). The workflow stores the React Flow graph definition."),
        ("Agent 1:N Execution", "Each agent can have many executions. Deleting an agent cascades to all its executions."),
        ("Execution 1:N ExecutionNodeLog", "Each execution produces one log entry per executed node. Cascades on execution delete."),
        ("Agent 1:N AgentMemory", "Each agent accumulates execution summaries as memory entries. Cascades on agent delete."),
        ("Execution 1:N AgentMemory", "Each memory entry optionally references its source execution (SET NULL on execution delete)."),
        ("Workflow N:1 Execution", "Each execution references the workflow ID that was run."),
        ("AgentTemplate (standalone)", "Templates are system-level resources not owned by any user. They provide pre-built workflow definitions for agent creation."),
    ]

    for title, desc in relationships:
        p = doc.add_paragraph()
        run_t = p.add_run(title + ": ")
        set_run(run_t, size=11, bold=True)
        run_d = p.add_run(desc)
        set_run(run_d, size=11)

    add_styled_heading(doc, "5.1 Cascade Delete Chain", level=2)
    add_body(doc, "When a user is deleted, the following cascade occurs:")
    cascade_steps = [
        "User deleted -> all Agents deleted (CASCADE)",
        "Agent deleted -> Workflow deleted (CASCADE)",
        "Agent deleted -> all Executions deleted (CASCADE)",
        "Agent deleted -> all AgentMemory entries deleted (CASCADE)",
        "Execution deleted -> all ExecutionNodeLogs deleted (CASCADE)",
        "Execution deleted -> AgentMemory.execution_id SET NULL",
        "User deleted -> all Integrations deleted (CASCADE)",
    ]
    for step in cascade_steps:
        add_bullet(doc, step)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 6 - SEQUENCE DIAGRAMS
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "6. Sequence Diagrams", level=1)

    # 6.1 Agent Execution Flow
    add_styled_heading(doc, "6.1 Agent Execution Flow (Manual Trigger)", level=2)
    add_body(doc, "This sequence describes the end-to-end flow when a user manually triggers an agent execution:")

    exec_steps = [
        "1. User clicks 'Run' in the frontend, sending POST /api/agents/{agent_id}/execute with JWT token.",
        "2. FastAPI validates the JWT, retrieves the agent and its workflow from PostgreSQL.",
        "3. API validates the workflow has nodes and is in active/draft status.",
        "4. API creates an Execution record with status='pending' and triggered_by='manual'.",
        "5. API dispatches a Celery task (execute_workflow) with the execution_id and workflow definition JSON.",
        "6. API returns 202 Accepted with the ExecutionResponse to the frontend immediately.",
        "7. Frontend opens an SSE connection to GET /api/executions/{execution_id}/stream.",
        "8. SSE handler subscribes to Redis pub/sub channel 'execution:{id}:logs'.",
        "9. Celery worker picks up the task, creates a WorkflowExecutor with a fresh asyncpg engine.",
        "10. Executor marks the execution as 'running' in PostgreSQL and loads agent memory context.",
        "11. Executor publishes 'Execution started' to Redis pub/sub (received by frontend via SSE).",
        "12. Executor parses the workflow JSON into nodes and edges, identifies the trigger (start) node.",
        "13. BFS traversal begins. For each node: create ExecutionNodeLog (status=running), publish log, execute node.",
        "14. For AI action nodes: call OpenAI API with config prompt + variable substitution + memory context.",
        "15. For decision nodes: evaluate condition and determine branch (true/false) for edge traversal.",
        "16. For integration nodes: send email via Resend or post to Slack webhook using user's integration config.",
        "17. After each node completes: update ExecutionNodeLog (status, duration, output), merge output_variables into workflow variables, publish completion log.",
        "18. After all nodes complete: update Execution (status=success, duration_ms, total_cost, variables). Update agent.last_execution_at.",
        "19. Memory Service generates an execution summary (via GPT-4o-mini or fallback template) and stores it as AgentMemory.",
        "20. Executor publishes '__STREAM_END__' to Redis. SSE handler closes the connection. Frontend shows completion.",
    ]
    for step in exec_steps:
        add_bullet(doc, step)

    # 6.2 Chat with Memory
    doc.add_paragraph()
    add_styled_heading(doc, "6.2 Chat with Memory", level=2)
    add_body(doc, "This sequence describes the agent chat flow with execution memory injection:")

    chat_steps = [
        "1. User sends a message via POST /api/agents/{agent_id}/chat with { message, history }.",
        "2. FastAPI validates JWT and confirms agent ownership.",
        "3. Memory Service queries the last 15 AgentMemory entries for this agent, ordered by created_at DESC.",
        "4. Memories are formatted into a text context block: 'Previous execution context (most recent first)' with timestamped summaries.",
        "5. A system prompt is constructed with agent name, type, description, status, schedule, and the memory context.",
        "6. The conversation history (last 10 messages) and new user message are assembled into the OpenAI messages array.",
        "7. An SSE EventSourceResponse is returned to the frontend.",
        "8. Inside the SSE generator: OpenAI chat.completions.create is called with model=gpt-4o-mini, stream=True.",
        "9. Each streaming chunk is forwarded as an SSE event: { type: 'token', content: delta_text }.",
        "10. When the stream completes, a final { type: 'done' } event is sent. On error, { type: 'error', content: error_message }.",
    ]
    for step in chat_steps:
        add_bullet(doc, step)

    # 6.3 Scheduled Execution
    doc.add_paragraph()
    add_styled_heading(doc, "6.3 Scheduled Execution", level=2)
    add_body(doc, "This sequence describes how cron-scheduled agent executions work:")

    sched_steps = [
        "1. When a workflow is deployed with a cron expression, the API calls schedule_agent() in the scheduler service.",
        "2. Scheduler creates a RedBeatSchedulerEntry with the agent's cron expression, storing it in Redis.",
        "3. The RedBeat scheduler (running as a Celery Beat process) continuously checks Redis for due entries.",
        "4. When a cron match occurs, RedBeat fires the execute_workflow_scheduled Celery task with the agent_id.",
        "5. The scheduled task handler queries PostgreSQL for the agent and verifies status='active'.",
        "6. It retrieves the agent's workflow and validates it has nodes.",
        "7. It creates a new Execution record with triggered_by='schedule'.",
        "8. It creates a WorkflowExecutor and calls execute(), following the same BFS traversal as manual execution.",
        "9. If the agent is paused or deleted, the scheduled task exits early without creating an execution.",
        "10. When an agent is paused or deleted, unschedule_agent() removes the RedBeat entry from Redis.",
    ]
    for step in sched_steps:
        add_bullet(doc, step)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 7 - ERROR CODE CATALOG
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "7. Error Code Catalog", level=1)
    add_body(doc, "The API uses standard HTTP status codes with a JSON body containing a \"detail\" field for error descriptions.")

    make_table(doc,
        ["HTTP Status", "Code", "Context", "Detail Message"],
        [
            ("400", "Bad Request", "POST /api/integrations", "Service must be one of: resend, slack"),
            ("400", "Bad Request", "POST /agents/{id}/execute", "Workflow is {status}, must be active or draft to execute"),
            ("400", "Bad Request", "POST /agents/{id}/execute", "Workflow has no nodes defined"),
            ("400", "Bad Request", "POST /agents/generate", "Description is required"),
            ("401", "Unauthorized", "POST /api/auth/login", "Invalid email or password"),
            ("401", "Unauthorized", "JWT validation", "Token has expired"),
            ("401", "Unauthorized", "JWT validation", "Invalid token"),
            ("401", "Unauthorized", "JWT validation", "Invalid token payload"),
            ("401", "Unauthorized", "JWT validation", "User not found"),
            ("404", "Not Found", "GET/PATCH/DELETE /agents/{id}", "Agent not found"),
            ("404", "Not Found", "GET /agents/{id}/workflow", "Workflow not found for this agent"),
            ("404", "Not Found", "GET /executions/{id}", "Execution not found"),
            ("404", "Not Found", "GET /templates/{id}", "Template not found"),
            ("404", "Not Found", "DELETE /integrations/{id}", "Integration not found"),
            ("404", "Not Found", "POST /agents/{id}/chat", "Agent not found"),
            ("409", "Conflict", "POST /api/auth/register", "A user with this email already exists"),
            ("409", "Conflict", "POST /api/agents", "An agent named '{name}' already exists"),
            ("422", "Unprocessable Entity", "POST /agents/generate", "AI generation validation error (dynamic)"),
            ("500", "Internal Server Error", "POST /agents/generate", "Failed to generate agent: {error}"),
            ("500", "Internal Server Error", "Global handler", "Internal server error (production) / {error} (debug)"),
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 8 - DATA VALIDATION RULES
    # ══════════════════════════════════════════════════════════
    add_styled_heading(doc, "8. Data Validation Rules", level=1)
    add_body(doc, "Validation is enforced at three layers: Pydantic schemas (request validation), SQLAlchemy models (database constraints), and business logic checks in route handlers.")

    add_styled_heading(doc, "8.1 Authentication Validation", level=2)
    make_table(doc,
        ["Field", "Rule", "Enforcement", "Error"],
        [
            ("email", "Must be a valid email format", "Pydantic EmailStr", "422 Validation Error"),
            ("password", "Minimum 6 characters", "Pydantic Field(min_length=6)", "422 Validation Error"),
            ("email", "Must be unique across users", "PostgreSQL UNIQUE constraint", "409 Conflict"),
            ("JWT token", "Must not be expired", "PyJWT decode (HS256)", "401 Token has expired"),
            ("JWT sub claim", "Must map to existing user", "DB lookup", "401 User not found"),
        ],
    )

    doc.add_paragraph()
    add_styled_heading(doc, "8.2 Agent Validation", level=2)
    make_table(doc,
        ["Field", "Rule", "Enforcement", "Error"],
        [
            ("name", "2-50 characters", "Pydantic Field(min_length=2, max_length=50)", "422 Validation Error"),
            ("type", "Must be one of: sales, marketing, support, custom", "Pydantic Literal type", "422 Validation Error"),
            ("name", "Unique per user", "PostgreSQL UNIQUE(user_id, name)", "409 Conflict"),
            ("agent_id", "Must belong to authenticated user", "DB query with user_id filter", "404 Agent not found"),
            ("description (generate)", "Must not be empty/whitespace", "Route handler strip() check", "400 Description is required"),
        ],
    )

    doc.add_paragraph()
    add_styled_heading(doc, "8.3 Workflow Validation", level=2)
    make_table(doc,
        ["Field", "Rule", "Enforcement", "Error"],
        [
            ("definition", "Must be a valid JSON dict", "Pydantic dict type", "422 Validation Error"),
            ("workflow.status", "Must be active or draft to execute", "Route handler check", "400 Bad Request"),
            ("workflow.definition.nodes", "Must have at least one node", "Route handler check", "400 Workflow has no nodes"),
            ("agent_id", "Workflow must exist for the agent", "DB lookup", "404 Workflow not found"),
        ],
    )

    doc.add_paragraph()
    add_styled_heading(doc, "8.4 Execution Validation", level=2)
    make_table(doc,
        ["Field", "Rule", "Enforcement", "Error"],
        [
            ("execution_id", "Must belong to an agent owned by current user", "JOIN query Agent.user_id", "404 Execution not found"),
            ("limit", "1-100 (default 20)", "Pydantic Query(ge=1, le=100)", "422 Validation Error"),
            ("offset", ">=0 (default 0)", "Pydantic Query(ge=0)", "422 Validation Error"),
        ],
    )

    doc.add_paragraph()
    add_styled_heading(doc, "8.5 Integration Validation", level=2)
    make_table(doc,
        ["Field", "Rule", "Enforcement", "Error"],
        [
            ("service", "Must be 'resend' or 'slack'", "Route handler ALLOWED_SERVICES check", "400 Bad Request"),
            ("service", "Unique per user", "PostgreSQL UNIQUE(user_id, service)", "Upsert (no error)"),
            ("config", "Must be a valid JSON dict", "Pydantic dict type", "422 Validation Error"),
            ("config response", "Sensitive keys (api_key, webhook_url, secret, token) are masked", "Route handler _mask_config()", "N/A (read-only masking)"),
        ],
    )

    doc.add_paragraph()
    add_styled_heading(doc, "8.6 Chat Validation", level=2)
    make_table(doc,
        ["Field", "Rule", "Enforcement", "Error"],
        [
            ("message", "Must be a non-empty string", "Pydantic str type", "422 Validation Error"),
            ("history", "Array of {role, content} objects", "Pydantic list[dict]", "422 Validation Error"),
            ("history", "Last 10 messages are used", "Route handler slicing [-10:]", "N/A (silent truncation)"),
            ("OPENAI_API_KEY", "Must be configured for chat to work", "Runtime check", "Returns fallback message"),
        ],
    )

    doc.add_paragraph()
    add_styled_heading(doc, "8.7 Cron Schedule Validation", level=2)
    make_table(doc,
        ["Field", "Rule", "Enforcement", "Error"],
        [
            ("schedule_cron", "Must be a valid 5-part cron expression", "scheduler_service.parse_cron()", "ValueError at scheduling time"),
            ("schedule_timezone", "Should be valid IANA timezone string", "Stored as-is, used by RedBeat", "Runtime scheduling error"),
        ],
    )

    # ── Save document ──────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"LLD saved to {os.path.abspath(OUTPUT_PATH)}")


if __name__ == "__main__":
    build_document()
