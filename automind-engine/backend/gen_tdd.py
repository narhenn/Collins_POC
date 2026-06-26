"""
Generate GoalCert AutoMind Technical Design Document (TDD) as .docx
Mirrors the exact Flex Coach TDD format.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
DARK_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_BLUE = RGBColor(0x2E, 0x74, 0xB5)
META_FILL = "E8EEF4"
HEADER_FILL = "1B2A4A"
ALT_ROW_FILL = "F0F4F8"
CODE_FILL = "F5F5F5"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    """Set background colour of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=Pt(11), font_name="Calibri",
                  alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Set spacing
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)


def add_table_with_headers(doc, headers, rows, col_widths=None):
    """Add a formatted table with header fill and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(10))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_FILL)
            set_cell_text(cell, str(val), size=Pt(10))

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")  # spacer
    return table


def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = DARK_NAVY
        run.font.name = "Calibri"


def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = ACCENT_BLUE
        run.font.name = "Calibri"


def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_code_block(doc, code_text):
    """Add a single-cell table with light gray fill containing monospace code."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("")  # spacer


def add_metadata_row(table, label, value, row_idx):
    """Add a row to the metadata table."""
    left = table.rows[row_idx].cells[0]
    right = table.rows[row_idx].cells[1]
    set_cell_shading(left, META_FILL)
    set_cell_text(left, label, bold=True, size=Pt(11))
    set_cell_text(right, value, size=Pt(11))


# ---------------------------------------------------------------------------
# Main Document Generation
# ---------------------------------------------------------------------------
def generate():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ======================================================================
    # COVER PAGE
    # ======================================================================
    for _ in range(6):
        doc.add_paragraph("")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("GoalCert AutoMind")
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = DARK_NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Technical Design Document (TDD)")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT_BLUE

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Metadata table
    meta_data = [
        ("Project", "GoalCert AutoMind - Agentic AI Platform"),
        ("Document Type", "Technical Design Document (TDD)"),
        ("Version", "1.0"),
        ("Date", "25 June 2026"),
        ("Author", "Narhen Karthikeyan"),
        ("Status", "Draft"),
        ("Classification", "Internal - Confidential"),
    ]
    meta_table = doc.add_table(rows=len(meta_data), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.style = "Table Grid"
    for i, (label, value) in enumerate(meta_data):
        add_metadata_row(meta_table, label, value, i)
    for row in meta_table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.0)

    doc.add_page_break()

    # ======================================================================
    # REVISION HISTORY
    # ======================================================================
    add_heading1(doc, "Revision History")

    rev_headers = ["Version", "Date", "Author", "Description"]
    rev_rows = [
        ["1.0", "25 June 2026", "Narhen Karthikeyan", "Initial release - complete TDD"],
    ]
    add_table_with_headers(doc, rev_headers, rev_rows, col_widths=[1.0, 1.5, 2.0, 2.5])

    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS (Manual Normal paragraphs)
    # ======================================================================
    add_heading1(doc, "Table of Contents")

    toc_entries = [
        (1, "1. Introduction"),
        (1, "2. System Context"),
        (1, "3. Technical Architecture"),
        (1, "4. Project Structure"),
        (1, "5. Authentication & Authorization Design"),
        (2, "5.1 JWT Token Structure"),
        (2, "5.2 Token Lifecycle"),
        (2, "5.3 Middleware Chain"),
        (1, "6. AI Agent Execution Engine Design"),
        (2, "6.1 Workflow Execution Pipeline"),
        (2, "6.2 Node Type Registry"),
        (2, "6.3 Variable Substitution System"),
        (1, "7. Notification System Design"),
        (1, "8. Database Design"),
        (2, "8.1 Schema Overview"),
        (2, "8.2 Indexing Strategy"),
        (2, "8.3 Migration Strategy"),
        (1, "9. Caching Strategy"),
        (1, "10. Logging & Monitoring"),
        (1, "11. Testing Strategy"),
        (1, "12. CI/CD Pipeline Design"),
        (1, "13. Environment Configuration"),
        (1, "14. Performance Benchmarks"),
    ]
    for level, text in toc_entries:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_NAVY
        if level == 1:
            run.font.bold = True
        else:
            run.font.bold = False
            p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # ======================================================================
    # 1. INTRODUCTION
    # ======================================================================
    add_heading1(doc, "1. Introduction")

    add_body(doc,
        "GoalCert AutoMind is an agentic AI platform that enables users to design, deploy, and "
        "manage autonomous AI workflows through a visual drag-and-drop interface. The platform "
        "allows non-technical users to compose multi-step AI pipelines by connecting configurable "
        "nodes such as LLM prompts, web searches, conditional branches, code execution, email/Slack "
        "integrations, and human escalation points."
    )
    add_body(doc,
        "This Technical Design Document provides a comprehensive specification of the system's "
        "architecture, data models, execution engine, authentication mechanisms, and infrastructure "
        "patterns. It serves as the authoritative reference for developers contributing to the "
        "platform."
    )

    add_heading2(doc, "1.1 Purpose")
    add_bullet(doc, "Define the technical architecture and design decisions for the AutoMind platform")
    add_bullet(doc, "Provide implementation guidance for backend services, frontend components, and infrastructure")
    add_bullet(doc, "Document the AI agent execution engine, node type registry, and variable substitution system")
    add_bullet(doc, "Specify database schema, caching strategy, and security model")

    add_heading2(doc, "1.2 Scope")
    add_body(doc,
        "This document covers the full-stack AutoMind platform including the FastAPI backend, "
        "React frontend, Celery task processing, Redis caching/pub-sub, PostgreSQL persistence, "
        "and integrations with OpenAI, Anthropic, DuckDuckGo, Resend (email), and Slack."
    )

    doc.add_page_break()

    # ======================================================================
    # 2. SYSTEM CONTEXT
    # ======================================================================
    add_heading1(doc, "2. System Context")

    add_body(doc,
        "The following table summarises the key technology choices and their roles within the "
        "GoalCert AutoMind platform."
    )

    ctx_headers = ["Aspect", "Detail"]
    ctx_rows = [
        ["Platform", "Web Application (SPA + REST API)"],
        ["Language", "Python 3.12 (backend), TypeScript 5.x (frontend)"],
        ["Backend Framework", "FastAPI 0.115 with Pydantic v2 validation"],
        ["Frontend Framework", "React 18 + Vite 6 + Tailwind CSS 4 + React Flow"],
        ["Database", "PostgreSQL 16 with SQLAlchemy 2.x async ORM"],
        ["Cache / Message Broker", "Redis 7 (Celery broker, pub/sub for SSE, RedBeat schedules)"],
        ["Task Queue", "Celery 5 with RedBeat scheduler for dynamic cron schedules"],
        ["AI Providers", "OpenAI (GPT-4o, GPT-4o-mini, GPT-4.1 family, o3-mini), Anthropic (Claude Sonnet 4, Claude Haiku 4.5)"],
        ["Web Search", "DuckDuckGo HTML API (no API key required)"],
        ["Email", "Resend API (transactional email)"],
        ["Messaging", "Slack Incoming Webhooks"],
        ["Authentication", "JWT (HS256) with bcrypt password hashing"],
        ["Hosting", "Render (backend), Vercel (frontend), Neon (PostgreSQL), Upstash (Redis)"],
    ]
    add_table_with_headers(doc, ctx_headers, ctx_rows, col_widths=[2.0, 5.0])

    doc.add_page_break()

    # ======================================================================
    # 3. TECHNICAL ARCHITECTURE
    # ======================================================================
    add_heading1(doc, "3. Technical Architecture")

    add_body(doc,
        "AutoMind follows a layered architecture pattern with clear separation of concerns. "
        "The frontend communicates exclusively through REST endpoints and SSE streams. "
        "Long-running workflow executions are offloaded to Celery workers to keep the API "
        "responsive. Real-time execution logs are delivered via Redis pub/sub to SSE connections."
    )

    arch_headers = ["Layer", "Technology", "Responsibility", "Key Patterns"]
    arch_rows = [
        [
            "Presentation",
            "React 18 + Vite 6 + Tailwind CSS 4 + TanStack Query v5",
            "Visual workflow builder, dashboard, execution monitoring, agent chat",
            "React Flow for DAG editor, Zustand for builder state, SSE for live console, TanStack Query for server-state caching",
        ],
        [
            "API Gateway",
            "FastAPI with Pydantic v2",
            "Request validation, JWT authentication, route handling, SSE streaming",
            "Dependency injection (Depends), async handlers, OAuth2PasswordBearer, EventSourceResponse",
        ],
        [
            "Business Logic",
            "Python service layer",
            "Agent generation, memory management, schedule management, workflow CRUD",
            "Service functions in app/services/, AI prompt chaining, OpenAI AsyncClient",
        ],
        [
            "Task Processing",
            "Celery 5 + Redis + RedBeat",
            "Asynchronous workflow execution, scheduled agent runs, task retry",
            "Shared tasks, asyncio.run() bridge, dynamic beat schedules via RedBeat entries",
        ],
        [
            "Data",
            "PostgreSQL 16 + SQLAlchemy 2.x async",
            "Persistent storage for users, agents, workflows, executions, templates, memory",
            "Mapped columns, JSONB for workflow definitions, async_sessionmaker, Alembic migrations",
        ],
        [
            "AI / ML",
            "OpenAI GPT-4o / GPT-4o-mini / GPT-4.1, Anthropic Claude, DuckDuckGo search",
            "LLM-driven node execution, agent generation from natural language, execution memory summarisation",
            "Dual-provider support (OpenAI + Anthropic), cost tracking per model, JSON output parsing",
        ],
    ]
    add_table_with_headers(doc, arch_headers, arch_rows, col_widths=[1.3, 1.7, 2.0, 2.5])

    doc.add_page_break()

    # ======================================================================
    # 4. PROJECT STRUCTURE
    # ======================================================================
    add_heading1(doc, "4. Project Structure")

    add_heading2(doc, "4.1 Backend (Python / FastAPI)")
    add_code_block(doc,
        "backend/\n"
        "  app/\n"
        "    main.py                    # FastAPI app, lifespan, middleware, routers\n"
        "    core/\n"
        "      config.py               # Pydantic Settings (env vars)\n"
        "      database.py             # SQLAlchemy async engine, session factory, Base\n"
        "      security.py             # JWT create/decode, bcrypt, get_current_user\n"
        "    models/\n"
        "      user.py                 # User model (id, email, password_hash, name)\n"
        "      agent.py                # Agent model (id, user_id, name, type, status, cron)\n"
        "      workflow.py             # Workflow model (id, agent_id, JSONB definition)\n"
        "      execution.py            # Execution + ExecutionNodeLog models\n"
        "      integration.py          # Integration model (service, JSONB config)\n"
        "      template.py             # AgentTemplate model (prebuilt workflows)\n"
        "      memory.py               # AgentMemory model (execution summaries)\n"
        "    schemas/\n"
        "      auth.py                 # UserRegister, UserLogin, AuthResponse\n"
        "      agent.py                # AgentCreate, AgentUpdate, AgentResponse\n"
        "      workflow.py             # WorkflowUpdate, WorkflowResponse\n"
        "      execution.py            # ExecutionResponse, ExecutionNodeLogResponse\n"
        "      integration.py          # IntegrationCreate, IntegrationResponse\n"
        "    routers/\n"
        "      auth.py                 # /api/auth/register, login, me\n"
        "      agents.py               # /api/agents CRUD + generate\n"
        "      workflows.py            # /api/agents/{id}/workflow\n"
        "      executions.py           # /api/agents/{id}/execute, /api/executions/{id}\n"
        "      templates.py            # /api/templates\n"
        "      dashboard.py            # /api/dashboard/stats, activity\n"
        "      integrations.py         # /api/integrations CRUD\n"
        "      memory.py               # /api/agents/{id}/memory\n"
        "      chat.py                 # /api/agents/{id}/chat (SSE streaming)\n"
        "    engine/\n"
        "      executor.py             # WorkflowExecutor - BFS graph traversal, DB logging\n"
        "      graph.py                # parse_workflow, topological_sort, get_next_nodes\n"
        "      variables.py            # {var} interpolation with dot-path resolution\n"
        "      nodes/\n"
        "        base.py               # BaseNodeExecutor ABC\n"
        "        trigger.py            # TriggerNodeExecutor\n"
        "        ai_action.py          # AIActionNodeExecutor (OpenAI + Anthropic)\n"
        "        web_search.py         # WebSearchNodeExecutor (DuckDuckGo)\n"
        "        code_exec.py          # CodeExecNodeExecutor (sandboxed subprocess)\n"
        "        decision.py           # DecisionNodeExecutor (conditional branching)\n"
        "        integration.py        # IntegrationNodeExecutor (email + Slack)\n"
        "        escalation.py         # EscalationNodeExecutor (human notification)\n"
        "    services/\n"
        "      agent_generator.py      # AI-powered workflow generation from NL\n"
        "      memory_service.py       # Execution memory save/retrieve\n"
        "      scheduler_service.py    # RedBeat schedule management\n"
        "    tasks/\n"
        "      celery_app.py           # Celery config (broker, RedBeat, serialisation)\n"
        "      workflow_tasks.py       # execute_workflow, execute_workflow_scheduled\n"
        "    integrations/\n"
        "      __init__.py             # Integration registry (extensible)\n"
        "  alembic/\n"
        "    versions/\n"
        "      16e837286122_initial_schema.py\n"
        "      0f2fc7c17eca_add_agent_memory_table.py\n"
    )

    add_heading2(doc, "4.2 Frontend (React / TypeScript)")
    add_code_block(doc,
        "frontend/src/\n"
        "  main.tsx                     # React entry point\n"
        "  App.tsx                      # Root router + providers\n"
        "  index.css                    # Tailwind CSS imports\n"
        "  api/\n"
        "    client.ts                  # Axios instance with JWT interceptor\n"
        "  stores/\n"
        "    authStore.ts               # Zustand auth store (token, user)\n"
        "    builderStore.ts            # Zustand workflow builder state\n"
        "  hooks/\n"
        "    useAgents.ts               # TanStack Query hooks for agents API\n"
        "    useWorkflow.ts             # TanStack Query hooks for workflow API\n"
        "    useExecutions.ts           # TanStack Query hooks for executions API\n"
        "    useIntegrations.ts         # TanStack Query hooks for integrations API\n"
        "    useMemory.ts               # TanStack Query hooks for memory API\n"
        "  pages/\n"
        "    LoginPage.tsx              # Login form\n"
        "    SignupPage.tsx             # Registration form\n"
        "    DashboardPage.tsx          # Stats cards + activity feed\n"
        "    AgentDetailPage.tsx        # Agent overview + execution history\n"
        "    WorkflowBuilderPage.tsx    # React Flow canvas + node config panel\n"
        "    ExecutionDetailPage.tsx    # Node log timeline + live console\n"
        "    TemplatesPage.tsx          # Prebuilt agent templates gallery\n"
        "    IntegrationsPage.tsx       # Email + Slack integration settings\n"
        "    AnalyticsPage.tsx          # Usage analytics\n"
        "    ReportsPage.tsx            # Execution reports\n"
        "    SettingsPage.tsx           # User settings + API keys\n"
        "  components/\n"
        "    layout/\n"
        "      AppLayout.tsx            # Sidebar + main content wrapper\n"
        "      Sidebar.tsx              # Navigation sidebar\n"
        "      ProtectedRoute.tsx       # JWT auth guard\n"
        "      GoalCertLogo.tsx         # Brand logo component\n"
        "      NotFound.tsx             # 404 page\n"
        "      ErrorBoundary.tsx        # React error boundary\n"
        "    builder/\n"
        "      WorkflowCanvas.tsx       # React Flow canvas\n"
        "      ComponentSidebar.tsx     # Draggable node palette\n"
        "      NodeConfigPanel.tsx      # Node configuration editor\n"
        "      nodes/\n"
        "        TriggerNode.tsx        # Trigger node visual component\n"
        "        AIActionNode.tsx       # AI Action node visual component\n"
        "        WebSearchNode.tsx      # Web Search node visual component\n"
        "        CodeExecNode.tsx       # Code Exec node visual component\n"
        "        DecisionNode.tsx       # Decision node visual component\n"
        "        IntegrationNode.tsx    # Integration node visual component\n"
        "        EscalationNode.tsx     # Escalation node visual component\n"
        "    executions/\n"
        "      LiveExecutionPanel.tsx   # Live execution monitoring\n"
        "      LiveConsole.tsx          # SSE-powered real-time log console\n"
        "      ExecutionTimeline.tsx    # Node execution timeline\n"
        "      NodeLogCard.tsx          # Individual node log display\n"
        "    agents/\n"
        "      CreateAgentModal.tsx     # Agent creation dialog\n"
        "      AgentStatusBadge.tsx     # Status indicator badge\n"
        "    dashboard/\n"
        "      StatsCards.tsx           # Dashboard statistics cards\n"
        "      AgentCard.tsx            # Agent summary card\n"
        "      ActivityFeed.tsx         # Recent execution activity\n"
        "    chat/\n"
        "      AgentChat.tsx            # Chat with agent (SSE streaming)\n"
        "  demo/\n"
        "    main-demo.tsx              # Demo mode entry point\n"
        "    demoAdapter.ts             # Demo mode API adapter\n"
        "    mockLayer.ts               # Mock data layer\n"
        "    mockFetch.ts               # Mock fetch interceptor\n"
        "    tour/                      # Guided demo tour components\n"
        "    data/                      # Static demo datasets\n"
        "  types/\n"
        "    index.ts                   # TypeScript type definitions\n"
        "  lib/\n"
        "    utils.ts                   # Utility functions\n"
    )

    doc.add_page_break()

    # ======================================================================
    # 5. AUTHENTICATION & AUTHORIZATION DESIGN
    # ======================================================================
    add_heading1(doc, "5. Authentication & Authorization Design")

    add_body(doc,
        "AutoMind uses stateless JWT authentication with bcrypt password hashing. "
        "All API endpoints (except /api/auth/register, /api/auth/login, /api/health, and SSE "
        "streams) require a valid Bearer token. The frontend stores the JWT in Zustand (memory) "
        "and attaches it to every API request via an Axios interceptor."
    )

    # 5.1 JWT Token Structure
    add_heading2(doc, "5.1 JWT Token Structure")
    add_body(doc, "The JWT payload contains the following claims:")
    add_code_block(doc,
        '{\n'
        '  "sub": "a3f1d8c2-7b4e-4a91-b5c6-9e2f3d8a1b0c",    // User UUID\n'
        '  "email": "user@example.com",                        // User email\n'
        '  "exp": 1750896000                                   // Expiry (UTC epoch)\n'
        '}'
    )
    add_body(doc,
        "Tokens are signed with HS256 using the JWT_SECRET environment variable. "
        "The signing key must be a cryptographically random string of at least 32 characters "
        "in production."
    )

    # 5.2 Token Lifecycle
    add_heading2(doc, "5.2 Token Lifecycle")
    tok_headers = ["Phase", "Action", "Detail"]
    tok_rows = [
        ["Registration", "POST /api/auth/register", "Creates user with bcrypt hash, returns JWT + user object"],
        ["Login", "POST /api/auth/login", "Validates credentials via bcrypt.checkpw, returns JWT + user object"],
        ["Token Use", "Authorization: Bearer <token>", "FastAPI OAuth2PasswordBearer extracts token from header"],
        ["Validation", "decode_access_token()", "Decodes with PyJWT, checks expiry, extracts sub (user_id)"],
        ["User Resolution", "get_current_user()", "Queries DB with user_id from token, returns User or 401"],
        ["Expiry", "JWT_EXPIRY_HOURS (default 24)", "Token becomes invalid; frontend redirects to login"],
        ["Refresh", "Re-authenticate", "No refresh token; user must log in again after expiry"],
    ]
    add_table_with_headers(doc, tok_headers, tok_rows, col_widths=[1.5, 2.5, 3.5])

    # 5.3 Middleware Chain
    add_heading2(doc, "5.3 Middleware Chain")
    add_body(doc,
        "FastAPI middleware is applied in registration order (outermost first):"
    )
    mw_headers = ["Order", "Middleware", "Purpose"]
    mw_rows = [
        ["1", "CORSMiddleware", "Allows cross-origin requests (all origins in DEBUG, FRONTEND_URL in production)"],
        ["2", "Global Exception Handler", "Catches unhandled exceptions, returns 500 JSON (detailed in DEBUG)"],
        ["3", "OAuth2PasswordBearer", "Extracts Bearer token from Authorization header per-endpoint via Depends()"],
        ["4", "get_current_user()", "Dependency injection - decodes JWT, loads user from DB, provides CurrentUser"],
    ]
    add_table_with_headers(doc, mw_headers, mw_rows, col_widths=[0.8, 2.2, 4.5])

    doc.add_page_break()

    # ======================================================================
    # 6. AI AGENT EXECUTION ENGINE DESIGN
    # ======================================================================
    add_heading1(doc, "6. AI Agent Execution Engine Design")

    add_body(doc,
        "The execution engine is the core of AutoMind. It takes a React Flow workflow definition "
        "(a directed acyclic graph of typed nodes and edges) and executes each node in topological "
        "order, passing variables between nodes via a shared context dictionary."
    )

    # 6.1 Workflow Execution Pipeline
    add_heading2(doc, "6.1 Workflow Execution Pipeline")

    add_body(doc,
        "Workflow execution is triggered via the /api/agents/{id}/execute endpoint or by a RedBeat "
        "scheduled task. The flow is as follows:"
    )

    add_bullet(doc, "1. API creates an Execution record (status=pending) and queues a Celery task with the execution_id and workflow_definition.")
    add_bullet(doc, "2. Celery worker picks up the task and calls asyncio.run() to bridge into the async executor (a new async DB engine is created per task to avoid cross-event-loop issues with asyncpg).")
    add_bullet(doc, "3. WorkflowExecutor.execute() marks the execution as running, loads the agent's memory context from past runs, and parses the workflow definition into nodes and edges.")
    add_bullet(doc, "4. The executor identifies the start node (a trigger node with no incoming edges) and begins BFS traversal of the graph.")
    add_bullet(doc, "5. For each node: a running log entry is created in the DB, the appropriate NodeExecutor is invoked, output variables are merged into the shared context, and the log is updated with duration/status/cost.")
    add_bullet(doc, "6. Decision nodes return a branch value (true/false) that determines which outgoing edge to follow.")
    add_bullet(doc, "7. On completion, the executor updates the execution record with final status, duration, total LLM cost, and saves an execution memory summary (optionally LLM-generated).")
    add_bullet(doc, "8. Throughout execution, real-time log messages are published to Redis pub/sub on channel execution:{id}:logs, which the frontend consumes via SSE.")

    # 6.2 Node Type Registry
    add_heading2(doc, "6.2 Node Type Registry")
    add_body(doc,
        "All node executors inherit from BaseNodeExecutor (ABC) and implement the async execute() "
        "method. The WorkflowExecutor maintains a registry mapping node type strings to executor "
        "instances."
    )

    node_headers = ["Node Type", "Executor Class", "Execution Logic", "Output"]
    node_rows = [
        [
            "trigger",
            "TriggerNodeExecutor",
            "Sets trigger_time (UTC ISO) and triggered_by (manual/schedule) variables. Entry point for every workflow.",
            "trigger_time, triggered_by",
        ],
        [
            "ai_action",
            "AIActionNodeExecutor",
            "Interpolates prompt template with current variables. Injects memory context as system prompt prefix. Calls OpenAI or Anthropic based on model prefix (gpt-*/o* = OpenAI, claude-* = Anthropic). Parses JSON from response. Tracks token usage and cost.",
            "output_variable (configurable), llm_usage (model, tokens, cost)",
        ],
        [
            "web_search",
            "WebSearchNodeExecutor",
            "Interpolates query template, POSTs to DuckDuckGo HTML endpoint, parses result links and snippets via regex. Returns structured results with title, URL, snippet, rank.",
            "output_variable (configurable) = list of search results",
        ],
        [
            "code_exec",
            "CodeExecNodeExecutor",
            "Interpolates code template. Validates against forbidden patterns (os, subprocess, socket, eval, exec, open). Writes to temp file, runs in subprocess with timeout (max 30s). Captures local variables of simple types as output.",
            "output_variable (configurable) = dict of captured locals",
        ],
        [
            "decision",
            "DecisionNodeExecutor",
            "Interpolates left and right operands. Evaluates condition using operator (==, !=, >, <, >=, <=, contains). Returns branch (true/false) to control graph traversal.",
            "branch (true/false), condition_result (bool)",
        ],
        [
            "integration",
            "IntegrationNodeExecutor",
            "Email: interpolates recipients/subject/body, sends via Resend API. Slack: interpolates message, POSTs to webhook URL. Falls back to mock responses when API keys are missing.",
            "emails_sent / message_sent, service, recipients/message",
        ],
        [
            "escalation",
            "EscalationNodeExecutor",
            "Interpolates recipient_email and message_template. Sends escalation alert email via Resend API. Used for human-in-the-loop breakpoints.",
            "escalation_sent (bool), recipient, subject",
        ],
    ]
    add_table_with_headers(doc, node_headers, node_rows, col_widths=[1.2, 1.8, 2.8, 1.7])

    # 6.3 Variable Substitution System
    add_heading2(doc, "6.3 Variable Substitution System")
    add_body(doc,
        "The variable substitution engine (app/engine/variables.py) provides template interpolation "
        "across all node configurations. It supports simple placeholders, nested dot-path access, and "
        "type preservation."
    )
    add_bullet(doc, "Simple: {company} resolves to variables['company']")
    add_bullet(doc, "Nested: {leads.0.name} resolves to variables['leads'][0]['name']")
    add_bullet(doc, "Type preservation: If an entire string is a single placeholder (e.g. '{data}'), the raw Python value is returned (dict, list, int), not its string representation.")
    add_bullet(doc, "Unresolved placeholders are left as-is (no error thrown).")
    add_bullet(doc, "Recursive interpolation: dicts and lists are recursed into, applying interpolation to all nested values.")

    add_code_block(doc,
        '# Example usage in the executor:\n'
        'prompt = interpolate("{company} quarterly report for {quarter}", variables)\n'
        '# With variables = {"company": "Acme Corp", "quarter": "Q2 2026"}\n'
        '# Result: "Acme Corp quarterly report for Q2 2026"'
    )

    doc.add_page_break()

    # ======================================================================
    # 7. NOTIFICATION SYSTEM DESIGN
    # ======================================================================
    add_heading1(doc, "7. Notification System Design")

    add_body(doc,
        "AutoMind provides real-time execution monitoring through Server-Sent Events (SSE) powered "
        "by Redis pub/sub. This enables the frontend LiveConsole component to display log messages "
        "as nodes execute, without polling."
    )

    add_heading2(doc, "7.1 SSE Streaming Architecture")
    add_bullet(doc, "Publisher: WorkflowExecutor._publish_log() writes JSON messages to Redis channel execution:{execution_id}:logs using a synchronous Redis client (runs inside the Celery worker).")
    add_bullet(doc, "Subscriber: GET /api/executions/{execution_id}/stream opens an SSE connection. An async Redis client subscribes to the same channel and yields messages as SSE events.")
    add_bullet(doc, "Stream lifecycle: The stream closes when the executor publishes a __STREAM_END__ sentinel message, or when the client disconnects.")
    add_bullet(doc, "No authentication: SSE stream endpoints do not require JWT authentication because execution IDs are UUIDs (128-bit, unguessable). This simplifies EventSource usage in the browser.")

    add_heading2(doc, "7.2 Log Message Format")
    add_code_block(doc,
        '{\n'
        '  "timestamp": "2026-06-25T10:30:00.000000+00:00",\n'
        '  "message": "Running node: Generate Report",\n'
        '  "node_id": "node-3",\n'
        '  "status": "running"   // running | success | error | info | done\n'
        '}'
    )

    add_heading2(doc, "7.3 Agent Chat (SSE)")
    add_body(doc,
        "The /api/agents/{id}/chat endpoint provides conversational interaction with agents. "
        "It injects the agent's execution memory as system context and streams GPT-4o-mini "
        "responses token-by-token via SSE. The frontend AgentChat component renders tokens "
        "incrementally for a real-time chat experience."
    )

    doc.add_page_break()

    # ======================================================================
    # 8. DATABASE DESIGN
    # ======================================================================
    add_heading1(doc, "8. Database Design")

    add_body(doc,
        "AutoMind uses PostgreSQL with SQLAlchemy 2.x async ORM. All tables use UUID primary keys "
        "generated by PostgreSQL's gen_random_uuid(). JSONB columns store workflow definitions, "
        "execution variables, node log I/O, and integration configs."
    )

    # 8.1 Schema Overview
    add_heading2(doc, "8.1 Schema Overview")

    # Users table
    add_body(doc, "Table: users")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "User identifier"],
            ["email", "VARCHAR(255)", "UNIQUE, NOT NULL", "Login email"],
            ["password_hash", "VARCHAR(255)", "NOT NULL", "bcrypt hash"],
            ["name", "VARCHAR(255)", "NULLABLE", "Display name"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Registration timestamp"],
            ["updated_at", "TIMESTAMPTZ", "DEFAULT now(), ON UPDATE", "Last modification"],
        ],
        col_widths=[1.5, 1.5, 2.0, 2.5],
    )

    # Agents table
    add_body(doc, "Table: agents")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "Agent identifier"],
            ["user_id", "UUID", "FK users.id CASCADE, NOT NULL", "Owner reference"],
            ["name", "VARCHAR(255)", "NOT NULL, UNIQUE(user_id, name)", "Agent name"],
            ["description", "TEXT", "NULLABLE", "Agent description"],
            ["type", "VARCHAR(50)", "NOT NULL", "sales / marketing / support / custom"],
            ["status", "VARCHAR(50)", "DEFAULT 'draft'", "draft / active / paused / archived"],
            ["schedule_cron", "VARCHAR(100)", "NULLABLE", "Cron expression for scheduled runs"],
            ["schedule_timezone", "VARCHAR(50)", "DEFAULT 'UTC'", "Timezone for cron schedule"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Creation timestamp"],
            ["updated_at", "TIMESTAMPTZ", "DEFAULT now(), ON UPDATE", "Last modification"],
            ["last_execution_at", "TIMESTAMPTZ", "NULLABLE", "Most recent execution end time"],
        ],
        col_widths=[1.6, 1.3, 2.2, 2.4],
    )

    # Workflows table
    add_body(doc, "Table: workflows")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "Workflow identifier"],
            ["agent_id", "UUID", "FK agents.id CASCADE, UNIQUE", "Owning agent (1:1)"],
            ["status", "VARCHAR(50)", "DEFAULT 'draft'", "draft / active"],
            ["definition", "JSONB", "DEFAULT {}, NOT NULL", "React Flow nodes + edges + viewport"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Creation timestamp"],
            ["updated_at", "TIMESTAMPTZ", "DEFAULT now(), ON UPDATE", "Last modification"],
            ["deployed_at", "TIMESTAMPTZ", "NULLABLE", "Last deployment timestamp"],
        ],
        col_widths=[1.5, 1.3, 2.2, 2.5],
    )

    # Executions table
    add_body(doc, "Table: executions")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "Execution identifier"],
            ["agent_id", "UUID", "FK agents.id CASCADE", "Parent agent"],
            ["workflow_id", "UUID", "FK workflows.id", "Workflow snapshot reference"],
            ["status", "VARCHAR(50)", "DEFAULT 'pending'", "pending / running / success / failed"],
            ["triggered_by", "VARCHAR(50)", "NOT NULL", "manual / schedule"],
            ["started_at", "TIMESTAMPTZ", "NULLABLE", "Execution start time"],
            ["ended_at", "TIMESTAMPTZ", "NULLABLE", "Execution end time"],
            ["duration_ms", "INTEGER", "NULLABLE", "Total execution duration in milliseconds"],
            ["error_message", "TEXT", "NULLABLE", "Error details if failed"],
            ["variables", "JSONB", "DEFAULT {}", "Final variable context after execution"],
            ["total_cost", "NUMERIC(10,6)", "DEFAULT 0", "Total LLM API cost in USD"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Record creation timestamp"],
        ],
        col_widths=[1.5, 1.3, 2.0, 2.7],
    )

    # ExecutionNodeLogs table
    add_body(doc, "Table: execution_node_logs")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "Log entry identifier"],
            ["execution_id", "UUID", "FK executions.id CASCADE", "Parent execution"],
            ["node_id", "VARCHAR(255)", "NOT NULL", "React Flow node ID"],
            ["node_type", "VARCHAR(50)", "NOT NULL", "Node type (trigger, ai_action, etc.)"],
            ["node_label", "VARCHAR(255)", "NULLABLE", "Human-readable node label"],
            ["status", "VARCHAR(50)", "DEFAULT 'pending'", "pending / running / success / failed / skipped"],
            ["started_at", "TIMESTAMPTZ", "NULLABLE", "Node execution start"],
            ["ended_at", "TIMESTAMPTZ", "NULLABLE", "Node execution end"],
            ["duration_ms", "INTEGER", "NULLABLE", "Node execution duration"],
            ["input_data", "JSONB", "DEFAULT {}", "Node configuration / input"],
            ["output_data", "JSONB", "DEFAULT {}", "Node execution output"],
            ["error_message", "TEXT", "NULLABLE", "Error details if failed"],
            ["llm_usage", "JSONB", "NULLABLE", "LLM token usage and cost breakdown"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Record creation timestamp"],
        ],
        col_widths=[1.5, 1.3, 2.0, 2.7],
    )

    # Integrations table
    add_body(doc, "Table: integrations")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "Integration identifier"],
            ["user_id", "UUID", "FK users.id CASCADE", "Owner reference"],
            ["service", "VARCHAR(100)", "NOT NULL, UNIQUE(user_id, service)", "Service name (email, slack)"],
            ["config", "JSONB", "DEFAULT {}", "Service-specific configuration"],
            ["status", "VARCHAR(50)", "DEFAULT 'active'", "active / inactive"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Creation timestamp"],
            ["updated_at", "TIMESTAMPTZ", "DEFAULT now(), ON UPDATE", "Last modification"],
        ],
        col_widths=[1.5, 1.3, 2.2, 2.5],
    )

    # AgentTemplates table
    add_body(doc, "Table: agent_templates")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "Template identifier"],
            ["name", "VARCHAR(255)", "NOT NULL", "Template name"],
            ["description", "TEXT", "NULLABLE", "Template description"],
            ["type", "VARCHAR(50)", "NOT NULL", "Agent type category"],
            ["workflow_definition", "JSONB", "NOT NULL", "Prebuilt workflow (nodes + edges)"],
            ["icon", "VARCHAR(50)", "NULLABLE", "Display icon identifier"],
            ["color", "VARCHAR(50)", "NULLABLE", "Display colour code"],
            ["features", "JSONB", "DEFAULT []", "Feature tags list"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Creation timestamp"],
        ],
        col_widths=[1.8, 1.3, 2.0, 2.4],
    )

    # AgentMemory table
    add_body(doc, "Table: agent_memory")
    add_table_with_headers(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, gen_random_uuid()", "Memory entry identifier"],
            ["agent_id", "UUID", "FK agents.id CASCADE, INDEX", "Parent agent"],
            ["execution_id", "UUID", "FK executions.id SET NULL, NULLABLE", "Source execution"],
            ["summary", "TEXT", "NOT NULL", "LLM-generated execution summary"],
            ["key_outputs", "JSONB", "DEFAULT {}", "Structured key outputs from execution"],
            ["memory_type", "VARCHAR(50)", "DEFAULT 'execution_summary'", "Memory category"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Creation timestamp"],
        ],
        col_widths=[1.5, 1.3, 2.2, 2.5],
    )

    # 8.2 Indexing Strategy
    add_heading2(doc, "8.2 Indexing Strategy")
    idx_headers = ["Table", "Index", "Columns", "Purpose"]
    idx_rows = [
        ["users", "uq_users_email", "email (UNIQUE)", "Fast email lookup during login"],
        ["agents", "uq_agent_user_name", "(user_id, name) UNIQUE", "Prevent duplicate agent names per user"],
        ["agents", "ix_agents_user_id", "user_id (FK)", "List agents by user"],
        ["workflows", "uq_workflows_agent_id", "agent_id (UNIQUE)", "Enforce 1:1 agent-workflow relationship"],
        ["executions", "ix_executions_agent_id", "agent_id (FK)", "List executions by agent"],
        ["execution_node_logs", "ix_enl_execution_id", "execution_id (FK)", "List node logs by execution"],
        ["integrations", "uq_integration_user_service", "(user_id, service) UNIQUE", "One integration per service per user"],
        ["agent_memory", "ix_agent_memory_agent_id", "agent_id (INDEX)", "Retrieve memory by agent"],
    ]
    add_table_with_headers(doc, idx_headers, idx_rows, col_widths=[1.8, 2.0, 1.8, 2.0])

    # 8.3 Migration Strategy
    add_heading2(doc, "8.3 Migration Strategy")
    add_bullet(doc, "Alembic is used for database migrations with async driver support (asyncpg).")
    add_bullet(doc, "Migration files are stored in alembic/versions/ with auto-generated revision IDs.")
    add_bullet(doc, "Initial schema migration (16e837286122) creates all core tables: users, agents, workflows, executions, execution_node_logs, integrations, agent_templates.")
    add_bullet(doc, "Agent memory table added in migration 0f2fc7c17eca.")
    add_bullet(doc, "Migrations are applied via 'alembic upgrade head' during deployment.")
    add_bullet(doc, "In development, Base.metadata.create_all() is called during FastAPI lifespan startup for rapid iteration.")

    doc.add_page_break()

    # ======================================================================
    # 9. CACHING STRATEGY
    # ======================================================================
    add_heading1(doc, "9. Caching Strategy")

    add_body(doc,
        "Redis serves multiple roles in the AutoMind architecture beyond traditional caching:"
    )

    cache_headers = ["Use Case", "Redis Feature", "Configuration", "Details"]
    cache_rows = [
        ["Celery Broker", "Message Queue", "REDIS_URL (default db 0)", "Task messages for execute_workflow and execute_workflow_scheduled tasks. JSON serialisation."],
        ["Celery Result Backend", "Key-Value Store", "REDIS_URL (default db 0)", "Task result storage for Celery. JSON serialised results with configurable TTL."],
        ["RedBeat Schedules", "Sorted Sets + Hashes", "redbeat_redis_url = REDIS_URL", "Dynamic Celery Beat schedules stored as RedBeat entries. Keys: redbeat:agent:{agent_id}. Supports create/update/delete of cron schedules at runtime."],
        ["Execution Logs (Pub/Sub)", "Pub/Sub Channels", "Channel: execution:{id}:logs", "Real-time log broadcasting from Celery workers to SSE endpoints. Synchronous publisher (redis-py), async subscriber (redis.asyncio)."],
        ["Worker Configuration", "Celery Settings", "worker_prefetch_multiplier=1", "Ensures fair task distribution. task_acks_late=True for reliability. task_reject_on_worker_lost=True."],
    ]
    add_table_with_headers(doc, cache_headers, cache_rows, col_widths=[1.5, 1.5, 1.8, 2.7])

    doc.add_page_break()

    # ======================================================================
    # 10. LOGGING & MONITORING
    # ======================================================================
    add_heading1(doc, "10. Logging & Monitoring")

    add_body(doc,
        "AutoMind uses Python's built-in logging module throughout the backend. "
        "Key subsystems have dedicated loggers:"
    )
    add_bullet(doc, "app.engine.executor - Workflow execution lifecycle, node failures, completion")
    add_bullet(doc, "app.engine.nodes.ai_action - LLM API calls, token usage, cost tracking, errors")
    add_bullet(doc, "app.engine.nodes.web_search - DuckDuckGo search requests and parsing")
    add_bullet(doc, "app.engine.nodes.code_exec - Sandboxed code execution, subprocess management")
    add_bullet(doc, "app.engine.nodes.integration - Email (Resend) and Slack webhook delivery")
    add_bullet(doc, "app.engine.nodes.escalation - Escalation email delivery")
    add_bullet(doc, "app.services.memory_service - Execution memory generation and retrieval")
    add_bullet(doc, "app.services.scheduler_service - RedBeat schedule create/delete operations")
    add_bullet(doc, "app.tasks.workflow_tasks - Celery task lifecycle, failure handling")
    add_bullet(doc, "app.routers.chat - Agent chat streaming errors")

    add_heading2(doc, "10.1 Execution Audit Trail")
    add_body(doc,
        "Every workflow execution produces a complete audit trail in the database:"
    )
    add_bullet(doc, "Execution record: status, triggered_by, duration_ms, total_cost, final variables, error_message")
    add_bullet(doc, "Node logs: per-node status, duration, input config, output data, LLM usage breakdown")
    add_bullet(doc, "Agent memory: LLM-generated natural language summary of each execution for cross-run context")

    add_heading2(doc, "10.2 Health Check")
    add_body(doc,
        "GET /api/health returns {\"status\": \"ok\", \"service\": \"automind-api\"} for uptime monitoring "
        "and load balancer health probes."
    )

    doc.add_page_break()

    # ======================================================================
    # 11. TESTING STRATEGY
    # ======================================================================
    add_heading1(doc, "11. Testing Strategy")

    test_headers = ["Layer", "Framework", "Scope", "Approach"]
    test_rows = [
        ["Unit Tests", "pytest + pytest-asyncio", "Node executors, variable interpolation, graph utilities, security helpers", "Mock external APIs (OpenAI, Anthropic, DuckDuckGo, Resend, Slack). Test each node executor in isolation with controlled config/variables."],
        ["Integration Tests", "pytest + httpx (AsyncClient)", "API endpoints, DB operations, authentication flow", "Use TestClient against FastAPI app with test database. Verify CRUD operations, JWT flow, execution creation."],
        ["Engine Tests", "pytest-asyncio", "WorkflowExecutor end-to-end, decision branching, variable propagation", "Build test workflow definitions, execute via WorkflowExecutor with mocked node executors, verify variable flow and node ordering."],
        ["Frontend Unit", "Vitest + React Testing Library", "Component rendering, hook behaviour, store logic", "Test React components with mocked API responses. Verify Zustand store state transitions."],
        ["E2E Tests", "Playwright", "Full user flows: login, create agent, build workflow, execute, view results", "Run against deployed preview environment. Validate critical paths end-to-end."],
        ["Security Tests", "pytest", "JWT expiry, invalid tokens, CORS, SQL injection prevention", "Test boundary conditions: expired tokens, malformed payloads, unauthorized access attempts."],
    ]
    add_table_with_headers(doc, test_headers, test_rows, col_widths=[1.3, 1.5, 2.0, 2.7])

    doc.add_page_break()

    # ======================================================================
    # 12. CI/CD PIPELINE DESIGN
    # ======================================================================
    add_heading1(doc, "12. CI/CD Pipeline Design")

    cicd_headers = ["Stage", "Tool", "Trigger", "Actions"]
    cicd_rows = [
        ["Lint & Format", "Ruff, ESLint, Prettier", "Every push / PR", "Python: ruff check + ruff format. TypeScript: eslint + prettier. Fail pipeline on violations."],
        ["Type Check", "mypy, tsc", "Every push / PR", "Backend: mypy --strict on core modules. Frontend: tsc --noEmit."],
        ["Unit Tests", "pytest, Vitest", "Every push / PR", "Run pytest with coverage (target 80%). Run Vitest with coverage."],
        ["Build", "Vite, Docker", "Every push to main", "Frontend: vite build (output to dist/). Backend: Docker image build."],
        ["Database Migration", "Alembic", "Pre-deployment", "Run alembic upgrade head against target database."],
        ["Deploy Backend", "Render", "Push to main", "Auto-deploy from Git. Render runs gunicorn + uvicorn workers. Health check on /api/health."],
        ["Deploy Frontend", "Vercel", "Push to main", "Auto-deploy from Git. Vercel builds Vite app and deploys to edge CDN."],
        ["Post-Deploy Smoke", "curl / httpx", "After deployment", "Hit /api/health, verify 200. Check frontend loads. Verify SSE stream connectivity."],
    ]
    add_table_with_headers(doc, cicd_headers, cicd_rows, col_widths=[1.3, 1.3, 1.3, 3.6])

    doc.add_page_break()

    # ======================================================================
    # 13. ENVIRONMENT CONFIGURATION
    # ======================================================================
    add_heading1(doc, "13. Environment Configuration")

    add_body(doc,
        "All configuration is loaded via Pydantic Settings from environment variables "
        "(with .env file fallback in development). The following variables control the platform:"
    )

    env_headers = ["Variable", "Type", "Default", "Description"]
    env_rows = [
        ["DATABASE_URL", "str", "postgresql+asyncpg://postgres:postgres@localhost:5432/automind", "PostgreSQL connection string (asyncpg driver)"],
        ["REDIS_URL", "str", "redis://localhost:6379/0", "Redis connection URL for Celery broker, result backend, pub/sub, and RedBeat"],
        ["JWT_SECRET", "str", "dev-secret-change-me", "Secret key for JWT HS256 signing (must be changed in production)"],
        ["JWT_ALGORITHM", "str", "HS256", "JWT signing algorithm"],
        ["JWT_EXPIRY_HOURS", "int", "24", "JWT token validity period in hours"],
        ["OPENAI_API_KEY", "str", "(empty)", "OpenAI API key for GPT-4o, GPT-4o-mini, GPT-4.1 models"],
        ["ANTHROPIC_API_KEY", "str", "(empty)", "Anthropic API key for Claude Sonnet 4, Claude Haiku 4.5 models"],
        ["RESEND_API_KEY", "str", "(empty)", "Resend API key for transactional email delivery"],
        ["SLACK_WEBHOOK_URL", "str", "(empty)", "Default Slack incoming webhook URL for Slack integration nodes"],
        ["FRONTEND_URL", "str", "http://localhost:5173", "Frontend origin URL for CORS (used when DEBUG=False)"],
        ["DEBUG", "bool", "True", "Enable debug mode: verbose SQL logging, detailed error responses, permissive CORS"],
    ]
    add_table_with_headers(doc, env_headers, env_rows, col_widths=[1.8, 0.6, 1.5, 3.6])

    doc.add_page_break()

    # ======================================================================
    # 14. PERFORMANCE BENCHMARKS
    # ======================================================================
    add_heading1(doc, "14. Performance Benchmarks")

    add_body(doc,
        "The following table defines target performance benchmarks for key system operations. "
        "Actual latency depends on external API response times (OpenAI, Anthropic, Resend, Slack) "
        "and database load."
    )

    perf_headers = ["Operation", "Endpoint / Component", "Target Latency (p95)", "Notes"]
    perf_rows = [
        ["Health Check", "GET /api/health", "< 10ms", "No DB call, static response"],
        ["User Login", "POST /api/auth/login", "< 200ms", "bcrypt verify + JWT sign + 1 DB query"],
        ["List Agents", "GET /api/agents", "< 100ms", "Paginated query with user_id filter"],
        ["Get Dashboard Stats", "GET /api/dashboard/stats", "< 150ms", "Aggregate queries (COUNT, AVG) on agents + executions"],
        ["Save Workflow", "PUT /api/agents/{id}/workflow", "< 100ms", "JSONB upsert, single row"],
        ["Trigger Execution", "POST /api/agents/{id}/execute", "< 200ms", "Creates Execution row + queues Celery task (async)"],
        ["Celery Task Pickup", "execute_workflow task", "< 500ms", "Time from queue to worker start"],
        ["Trigger Node", "TriggerNodeExecutor", "< 5ms", "Sets timestamp variables only"],
        ["AI Action Node (GPT-4o-mini)", "AIActionNodeExecutor", "1-5s", "Depends on prompt length and max_tokens"],
        ["AI Action Node (GPT-4o)", "AIActionNodeExecutor", "2-10s", "Depends on prompt length and max_tokens"],
        ["Web Search Node", "WebSearchNodeExecutor", "500ms-2s", "DuckDuckGo HTML parse, network dependent"],
        ["Code Exec Node", "CodeExecNodeExecutor", "< 30s", "Subprocess with configurable timeout (max 30s)"],
        ["Decision Node", "DecisionNodeExecutor", "< 5ms", "In-memory comparison, no I/O"],
        ["Integration Node (Email)", "IntegrationNodeExecutor", "500ms-2s", "Resend API call"],
        ["Integration Node (Slack)", "IntegrationNodeExecutor", "200ms-1s", "Slack webhook POST"],
        ["SSE Stream Setup", "GET /executions/{id}/stream", "< 100ms", "Redis pub/sub subscribe"],
        ["Full Workflow (5 nodes)", "WorkflowExecutor", "5-30s", "Depends on node types; AI nodes dominate"],
        ["Agent Generation", "generate_agent_from_description", "3-8s", "GPT-4o-mini with 4096 max_tokens"],
        ["DB Connection Pool", "SQLAlchemy async engine", "pool_size=20", "20 concurrent async connections"],
    ]
    add_table_with_headers(doc, perf_headers, perf_rows, col_widths=[1.8, 2.2, 1.3, 2.2])

    # ======================================================================
    # SAVE
    # ======================================================================
    output_dir = "/Users/narhen/automind/docs"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "GoalCert_AutoMind_TDD_v1_0.docx")
    doc.save(output_path)
    print(f"TDD saved to {output_path}")


if __name__ == "__main__":
    generate()
