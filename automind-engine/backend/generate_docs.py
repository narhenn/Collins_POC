"""
Generate GoalCert AutoMind HLD and LLD documents as professional .docx files.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# -- Constants --
PRIMARY_COLOR = RGBColor(73, 2, 162)
DARK_TEXT = RGBColor(33, 33, 33)
WHITE = RGBColor(255, 255, 255)
LIGHT_BG = RGBColor(245, 243, 250)
FONT_NAME = "Calibri"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(18)
HEADING2_SIZE = Pt(14)
HEADING3_SIZE = Pt(12)
OUTPUT_DIR = "/Users/narhen/automind/docs"


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def format_run(run, font_name=FONT_NAME, size=BODY_SIZE, bold=False, color=DARK_TEXT, italic=False):
    """Apply formatting to a run."""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_formatted_paragraph(doc, text, style=None, font_name=FONT_NAME, size=BODY_SIZE,
                            bold=False, color=DARK_TEXT, alignment=None, space_before=0,
                            space_after=Pt(6), italic=False):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    format_run(run, font_name, size, bold, color, italic)
    return p


def create_professional_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        format_run(run, bold=True, color=WHITE, size=Pt(10))
        set_cell_shading(cell, "4902A2")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            format_run(run, size=Pt(10))
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F3FA")

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    # Add spacing after table
    doc.add_paragraph()
    return table


def create_two_col_table(doc, rows):
    """Create a two-column key-value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, (key, value) in enumerate(rows):
        cell_key = table.rows[r_idx].cells[0]
        cell_val = table.rows[r_idx].cells[1]

        cell_key.text = ""
        p = cell_key.paragraphs[0]
        run = p.add_run(key)
        format_run(run, bold=True, size=Pt(10))
        set_cell_shading(cell_key, "4902A2")
        run.font.color.rgb = WHITE

        cell_val.text = ""
        p = cell_val.paragraphs[0]
        run = p.add_run(str(value))
        format_run(run, size=Pt(10))
        if r_idx % 2 == 1:
            set_cell_shading(cell_val, "F5F3FA")

        cell_key.width = Cm(5)
        cell_val.width = Cm(12)

    doc.add_paragraph()
    return table


def add_cover_page(doc, doc_title, description):
    """Add a professional cover page."""
    # Add spacing before header
    for _ in range(4):
        doc.add_paragraph()

    # GoalCert header
    add_formatted_paragraph(doc, "GoalCert", size=Pt(36), bold=True,
                            color=PRIMARY_COLOR, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=Pt(4))

    # Subtitle
    add_formatted_paragraph(doc, "AutoMind \u2014 Agentic AI Workforce Platform",
                            size=Pt(16), color=RGBColor(100, 100, 100),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

    # Horizontal rule (via paragraph border)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="4902A2"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    doc.add_paragraph()

    # Document title
    add_formatted_paragraph(doc, doc_title, size=Pt(24), bold=True,
                            color=DARK_TEXT, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=Pt(12))

    # Description
    add_formatted_paragraph(doc, description, size=Pt(12),
                            color=RGBColor(100, 100, 100),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=Pt(36), italic=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # Confidentiality notice
    add_formatted_paragraph(
        doc,
        "CONFIDENTIAL",
        size=Pt(12), bold=True, color=RGBColor(180, 0, 0),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4)
    )
    add_formatted_paragraph(
        doc,
        "This document contains proprietary information belonging to GoalCert Pte. Ltd. "
        "Unauthorised distribution, reproduction, or use of this document is strictly prohibited.",
        size=Pt(9), color=RGBColor(120, 120, 120),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0)
    )

    doc.add_page_break()


def add_document_control(doc, doc_type, version="1.0", date="23 June 2026"):
    """Add Document Control table."""
    add_formatted_paragraph(doc, "Document Control", size=HEADING1_SIZE, bold=True,
                            color=PRIMARY_COLOR, space_after=Pt(12))

    rows = [
        ("Document Type", doc_type),
        ("Product", "GoalCert AutoMind"),
        ("Version", version),
        ("Date", date),
        ("Classification", "Confidential"),
        ("Prepared By", "AutoMind Engineering Team"),
        ("Reviewed By", "Prem Kumar"),
        ("Approved By", "Prem Kumar"),
    ]
    create_two_col_table(doc, rows)


def add_version_history(doc):
    """Add Version History table."""
    add_formatted_paragraph(doc, "Version History", size=HEADING1_SIZE, bold=True,
                            color=PRIMARY_COLOR, space_after=Pt(12))

    create_professional_table(
        doc,
        headers=["Version", "Date", "Author", "Changes"],
        rows=[
            ["0.1", "10 June 2026", "Engineering Team", "Initial draft"],
            ["0.5", "16 June 2026", "Engineering Team", "Architecture review feedback incorporated"],
            ["1.0", "23 June 2026", "Engineering Team", "Release candidate"],
        ]
    )
    doc.add_page_break()


def add_heading1(doc, text, numbered=True, number=""):
    """Add a Heading 1 style paragraph."""
    prefix = f"{number}  " if numbered and number else ""
    p = add_formatted_paragraph(doc, f"{prefix}{text}", size=HEADING1_SIZE, bold=True,
                                color=PRIMARY_COLOR, space_before=Pt(18), space_after=Pt(10))
    return p


def add_heading2(doc, text, number=""):
    """Add a Heading 2 style paragraph."""
    prefix = f"{number}  " if number else ""
    p = add_formatted_paragraph(doc, f"{prefix}{text}", size=HEADING2_SIZE, bold=True,
                                color=PRIMARY_COLOR, space_before=Pt(12), space_after=Pt(6))
    return p


def add_heading3(doc, text, number=""):
    """Add a Heading 3 style paragraph."""
    prefix = f"{number}  " if number else ""
    p = add_formatted_paragraph(doc, f"{prefix}{text}", size=HEADING3_SIZE, bold=True,
                                color=PRIMARY_COLOR, space_before=Pt(8), space_after=Pt(4))
    return p


def add_body(doc, text, space_after=Pt(6)):
    """Add body text."""
    return add_formatted_paragraph(doc, text, size=BODY_SIZE, color=DARK_TEXT, space_after=space_after)


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    # Clear default text and add formatted run
    for r in p.runs:
        r.clear()
    run = p.add_run(text)
    format_run(run, size=Pt(10))
    return p


def add_architecture_diagram(doc):
    """Create a text-based architecture diagram as a formatted table."""
    add_heading2(doc, "Architecture Overview Diagram", "3.1")
    add_body(doc, "The following diagram illustrates the three-tier architecture of the AutoMind platform.")
    doc.add_paragraph()

    # Create outer table with 3 rows for the 3 tiers
    table = doc.add_table(rows=5, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style all cells
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(16)

    # Row 0: PRESENTATION LAYER header
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, "4902A2")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRESENTATION LAYER")
    format_run(run, bold=True, color=WHITE, size=Pt(12))

    # Row 1: Presentation details
    cell = table.rows[1].cells[0]
    cell.text = ""
    set_cell_shading(cell, "EDE8F5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("React 19  +  Vite  +  TypeScript  +  TanStack Query\n")
    format_run(run, size=Pt(10), color=DARK_TEXT)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    components = ["Dashboard", "Agent Detail", "Workflow Builder", "Reports", "Settings"]
    run2 = p2.add_run("    ".join(f"[ {c} ]" for c in components))
    format_run(run2, size=Pt(9), color=RGBColor(73, 2, 162), bold=True)

    # Row 2: APPLICATION LAYER header
    cell = table.rows[2].cells[0]
    cell.text = ""
    set_cell_shading(cell, "4902A2")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APPLICATION LAYER")
    format_run(run, bold=True, color=WHITE, size=Pt(12))

    # Row 3: Application details
    cell = table.rows[3].cells[0]
    cell.text = ""
    set_cell_shading(cell, "EDE8F5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FastAPI  +  Celery  +  Redis  +  SQLAlchemy 2.0\n")
    format_run(run, size=Pt(10), color=DARK_TEXT)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    services = ["Agent Mgmt", "Exec Engine", "AI Layer", "SSE/Chat", "Scheduler"]
    run2 = p2.add_run("    ".join(f"[ {s} ]" for s in services))
    format_run(run2, size=Pt(9), color=RGBColor(73, 2, 162), bold=True)

    # Row 4: DATA LAYER as nested table
    cell = table.rows[4].cells[0]
    cell.text = ""
    set_cell_shading(cell, "4902A2")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DATA LAYER\n")
    format_run(run, bold=True, color=WHITE, size=Pt(12))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    datastores = ["PostgreSQL 14+", "Redis 7.x", "pgvector (planned)"]
    run2 = p2.add_run("    ".join(f"[ {d} ]" for d in datastores))
    format_run(run2, size=Pt(10), color=WHITE, bold=True)

    doc.add_paragraph()


def add_component_table(doc, component_name, language, input_desc, output_desc, error_handling, dependencies):
    """Add a component specification table."""
    rows = [
        ("Component", component_name),
        ("Language / Runtime", language),
        ("Input", input_desc),
        ("Output", output_desc),
        ("Error Handling", error_handling),
        ("Dependencies", dependencies),
    ]
    create_two_col_table(doc, rows)


def add_api_table(doc, endpoints):
    """Add an API specification table. endpoints = list of (method, path, auth, request_body, response, status_codes)."""
    create_professional_table(
        doc,
        headers=["Method", "Path", "Auth", "Request Body", "Response", "Status Codes"],
        rows=endpoints,
    )


# =============================================================================
# HLD DOCUMENT
# =============================================================================

def generate_hld():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE
    font.color.rgb = DARK_TEXT

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Cover Page --
    add_cover_page(
        doc,
        "High-Level Design Document",
        "Architecture and system design for the AutoMind agentic AI workflow platform"
    )

    # -- Document Control --
    add_document_control(doc, "High-Level Design (HLD)")
    add_version_history(doc)

    # -- Table of Contents placeholder --
    add_heading1(doc, "Table of Contents", number="")
    add_body(doc, "[Table of contents to be generated from headings]")
    doc.add_page_break()

    # =====================================================================
    # 1. Introduction
    # =====================================================================
    add_heading1(doc, "Introduction", number="1")

    add_heading2(doc, "Purpose", number="1.1")
    add_body(doc,
        "This High-Level Design document describes the architecture, design principles, and technology "
        "decisions for the GoalCert AutoMind platform. AutoMind is an agentic AI workforce platform that "
        "enables users to create, configure, and deploy autonomous AI agents through a visual workflow "
        "builder. Each agent executes multi-step workflows comprising AI inference, web search, code "
        "execution, conditional logic, and external integrations, all orchestrated by a distributed "
        "task engine."
    )

    add_heading2(doc, "Scope", number="1.2")
    add_body(doc,
        "This document covers the end-to-end system architecture from the browser-based React frontend "
        "through the FastAPI application layer to the PostgreSQL and Redis data tier. It addresses "
        "deployment models, security architecture, scalability considerations, and cross-cutting concerns "
        "such as multi-tenancy and cost tracking."
    )

    add_heading2(doc, "References", number="1.3")
    create_professional_table(
        doc,
        headers=["Reference", "Description"],
        rows=[
            ["GoalCert AutoMind LLD v1.0", "Low-Level Design document for detailed component specifications"],
            ["React Flow Documentation", "https://reactflow.dev"],
            ["FastAPI Documentation", "https://fastapi.tiangolo.com"],
            ["Celery Documentation", "https://docs.celeryq.dev"],
            ["SQLAlchemy 2.0 Documentation", "https://docs.sqlalchemy.org/en/20/"],
            ["RedBeat Documentation", "https://github.com/sibson/redbeat"],
        ]
    )

    # =====================================================================
    # 2. Design Goals & Principles
    # =====================================================================
    add_heading1(doc, "Design Goals & Principles", number="2")
    add_body(doc,
        "The following architectural principles guide all design decisions in the AutoMind platform."
    )

    principles = [
        ("Agent-First Design", "Every feature is designed around the agent lifecycle: create, configure, deploy, execute, monitor, and iterate."),
        ("Visual Workflow Composition", "Complex agent behaviours are composed through a drag-and-drop canvas rather than code, lowering the barrier to automation."),
        ("Async-Native Execution", "All I/O-bound operations use Python async/await. The FastAPI backend, SQLAlchemy 2.0 async sessions, and httpx async HTTP clients ensure non-blocking throughput."),
        ("Provider-Agnostic AI", "The AI integration layer abstracts over multiple LLM providers (OpenAI GPT-4o, Anthropic Claude Sonnet 4) so agents are not locked into a single vendor."),
        ("Event-Driven Real-Time Feedback", "Workflow execution progress is streamed to the frontend via Server-Sent Events (SSE) backed by Redis pub/sub, giving users immediate visibility into agent behaviour."),
        ("Multi-Tenant Isolation", "Every database query is filtered by user_id. Agents, workflows, executions, memories, and integrations are strictly scoped to the authenticated user."),
        ("Graceful Degradation", "Missing API keys, unavailable integrations, and transient failures produce mock or fallback responses rather than crashing the workflow."),
        ("Cost Transparency", "Every LLM call records input/output token counts and calculates cost using a per-model pricing table, enabling users to track spend per agent."),
        ("Separation of Concerns", "The system is decomposed into clearly bounded services: authentication, agent management, workflow execution, AI inference, scheduling, and integrations."),
        ("Extensibility via Node Types", "New capabilities are added by implementing the BaseNodeExecutor interface and registering the new type in the executor registry, with zero changes to the core engine."),
    ]

    for title, desc in principles:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"\u2022  {title}: ")
        format_run(run, bold=True, size=Pt(10), color=PRIMARY_COLOR)
        run2 = p.add_run(desc)
        format_run(run2, size=Pt(10))

    # =====================================================================
    # 3. High-Level Architecture
    # =====================================================================
    add_heading1(doc, "High-Level Architecture", number="3")
    add_body(doc,
        "AutoMind follows a three-tier architecture: Presentation (React SPA), Application (FastAPI + Celery), "
        "and Data (PostgreSQL + Redis). The tiers communicate exclusively through well-defined REST APIs and "
        "Server-Sent Events."
    )

    add_architecture_diagram(doc)

    add_body(doc,
        "Communication flow: the React SPA issues REST API calls to FastAPI endpoints. Long-running "
        "workflow executions are offloaded to Celery workers. Real-time execution progress is published "
        "to Redis pub/sub channels and consumed by the frontend via SSE endpoints."
    )

    doc.add_page_break()

    # =====================================================================
    # 4. Tier 1 - Presentation Layer
    # =====================================================================
    add_heading1(doc, "Tier 1 \u2014 Presentation Layer", number="4")

    add_heading2(doc, "React SPA (Vite, TypeScript, React 19)", number="4.1")
    add_body(doc,
        "The frontend is a single-page application bootstrapped with Vite and written in TypeScript. "
        "React 19 provides the component model. The application is styled with Tailwind CSS and uses "
        "Lucide React for iconography. All API communication uses TanStack Query for automatic caching, "
        "deduplication, and background refetching."
    )

    add_heading2(doc, "Component Architecture", number="4.2")
    add_body(doc, "The frontend is organised into the following directory structure:")
    create_professional_table(
        doc,
        headers=["Directory", "Purpose", "Examples"],
        rows=[
            ["pages/", "Top-level route components", "Dashboard, AgentDetail, WorkflowBuilder, Settings"],
            ["components/", "Reusable UI components", "AgentCard, ExecutionTimeline, StatsCard, ChatPanel"],
            ["hooks/", "Custom React hooks", "useAgents, useExecutions, useWorkflow, useSSE"],
            ["stores/", "Zustand state stores", "authStore (JWT token, user), builderStore (canvas state)"],
            ["lib/", "Utility functions", "API client (axios/fetch), formatters, constants"],
        ]
    )

    add_heading2(doc, "State Management (Zustand + TanStack Query)", number="4.3")
    add_body(doc,
        "Client-side state (authentication tokens, UI preferences, workflow builder canvas state) is "
        "managed with Zustand. Server-side state (agents, executions, templates) is managed by TanStack "
        "Query, which provides automatic cache invalidation, optimistic updates, and infinite scroll "
        "pagination. This dual approach avoids the complexity of a single global store while keeping "
        "server data fresh."
    )

    add_heading2(doc, "Workflow Builder (React Flow Canvas)", number="4.4")
    add_body(doc,
        "The visual workflow builder is powered by React Flow 12.x. Users drag node types onto a canvas, "
        "connect them with edges, and configure each node through a side panel. The canvas state (nodes, "
        "edges, viewport) is serialised to JSON and persisted via the PUT /agents/:id/workflow endpoint. "
        "Supported node types include: trigger, ai_action, web_search, decision, integration, escalation, "
        "and code_exec. Decision nodes support conditional branching with true/false output handles."
    )

    doc.add_page_break()

    # =====================================================================
    # 5. Tier 2 - Application Layer
    # =====================================================================
    add_heading1(doc, "Tier 2 \u2014 Application Layer", number="5")

    add_heading2(doc, "FastAPI Backend (async, SQLAlchemy 2.0)", number="5.1")
    add_body(doc,
        "The backend is a FastAPI application running on Uvicorn. It uses SQLAlchemy 2.0 with asyncpg for "
        "fully async database access. Pydantic v2 models validate all request and response payloads. "
        "The application is structured into routers (one per domain), services (business logic), and "
        "models (SQLAlchemy ORM definitions)."
    )

    add_heading2(doc, "Authentication & Authorisation (JWT, bcrypt)", number="5.2")
    add_body(doc,
        "User passwords are hashed with bcrypt (12 rounds). On login or registration, the server issues "
        "a JWT (HS256, 24-hour expiry) containing the user ID and email. All protected endpoints extract "
        "the current user via a FastAPI dependency that decodes and validates the token, then fetches the "
        "user record from PostgreSQL."
    )

    add_heading2(doc, "Agent Management Service", number="5.3")
    add_body(doc,
        "Agents are the top-level domain object. Each agent has a name, type (sales, marketing, support, "
        "custom), status (draft, active, paused), optional cron schedule, and a one-to-one relationship "
        "with a Workflow. The service supports CRUD operations, pause/resume (which registers or removes "
        "the cron schedule), and AI-powered agent generation from natural language descriptions. Agent "
        "names are unique per user, enforced by a composite unique constraint."
    )

    add_heading2(doc, "Workflow Execution Engine", number="5.4")
    add_body(doc,
        "The WorkflowExecutor is the core of the system. Given an execution ID and workflow definition, "
        "it performs a BFS traversal of the workflow graph starting from the trigger node. Each node is "
        "dispatched to the appropriate NodeExecutor. Output variables from each node are merged into a "
        "shared variable context that subsequent nodes can reference via {variable} placeholders. Decision "
        "nodes route execution to the true or false branch based on configurable conditions."
    )
    add_body(doc,
        "After all nodes complete, the executor records the final status, total duration, cumulative LLM "
        "cost, and output variables. It also generates an execution memory summary (optionally using GPT-4o-mini) "
        "and stores it for future agent context."
    )

    add_heading2(doc, "Task Queue (Celery + Redis + RedBeat)", number="5.5")
    add_body(doc,
        "Workflow executions are offloaded to Celery workers to avoid blocking the API server. The API "
        "creates an Execution record in pending state, then dispatches a Celery task. The worker runs "
        "the WorkflowExecutor inside asyncio.run() with a fresh database engine (asyncpg engines are "
        "not reusable across event loops). RedBeat provides dynamic cron scheduling backed by Redis, "
        "allowing agents to be scheduled and unscheduled at runtime without restarting workers."
    )

    add_heading2(doc, "Real-time Communication (SSE via Redis pub/sub)", number="5.6")
    add_body(doc,
        "During execution, the WorkflowExecutor publishes structured log events to a Redis pub/sub "
        "channel keyed by execution ID. The frontend connects to a SSE endpoint that subscribes to this "
        "channel and streams events as they arrive. A special __STREAM_END__ sentinel signals completion. "
        "The chat feature also uses SSE to stream LLM token-by-token responses."
    )

    add_heading2(doc, "AI Integration Layer (OpenAI, Anthropic APIs)", number="5.7")
    add_body(doc,
        "The AIActionNodeExecutor supports two LLM providers: OpenAI (GPT-4o, GPT-4o-mini, GPT-4.1, "
        "o3-mini) and Anthropic (Claude Sonnet 4, Claude Haiku 4.5). Provider selection is determined "
        "by the model name prefix. Each call records input tokens, output tokens, and calculates cost "
        "from a per-model pricing table. If no API key is configured, the executor returns a mock response "
        "so the workflow does not fail."
    )

    add_heading2(doc, "Web Search Service (DuckDuckGo Scraping)", number="5.8")
    add_body(doc,
        "The WebSearchNodeExecutor provides zero-API-key web search by posting to DuckDuckGo's HTML "
        "endpoint and parsing results with regex. Each result includes title, URL, snippet, and rank. "
        "The output is stored in the workflow variable context for downstream AI nodes to reason over."
    )

    doc.add_page_break()

    # =====================================================================
    # 6. Tier 3 - Data Layer
    # =====================================================================
    add_heading1(doc, "Tier 3 \u2014 Data Layer", number="6")

    add_heading2(doc, "PostgreSQL (Primary Datastore)", number="6.1")
    add_body(doc,
        "PostgreSQL 14+ is the primary relational database. All tables use UUID primary keys generated "
        "by gen_random_uuid(). JSONB columns store flexible data: workflow definitions, execution variables, "
        "node log I/O, integration configurations, and template features. The schema is managed via "
        "Alembic migrations."
    )
    add_body(doc,
        "Key tables: users, agents, workflows, executions, execution_node_logs, agent_memory, "
        "agent_templates, and integrations."
    )

    add_heading2(doc, "Redis (Cache, Message Broker, Pub/Sub)", number="6.2")
    add_body(doc,
        "Redis 7.x serves three roles: (1) Celery message broker for task dispatch, (2) real-time "
        "pub/sub backend for SSE execution streaming, and (3) RedBeat storage for dynamic cron schedules. "
        "Redis is a single-node deployment in the current architecture."
    )

    add_heading2(doc, "Future: pgvector for Embeddings", number="6.3")
    add_body(doc,
        "A planned enhancement will add a agent_memory_embeddings table with a vector(1536) column "
        "using the pgvector extension. This will enable semantic similarity search over agent memory, "
        "allowing agents to retrieve contextually relevant past execution summaries rather than relying "
        "on a simple chronological window."
    )

    doc.add_page_break()

    # =====================================================================
    # 7. Cross-Cutting Concerns
    # =====================================================================
    add_heading1(doc, "Cross-Cutting Concerns", number="7")

    add_heading2(doc, "Multi-Tenancy", number="7.1")
    add_body(doc,
        "AutoMind uses a shared-database, shared-schema multi-tenancy model. Every table that stores "
        "user-owned data includes a user_id foreign key. All queries in the API layer are filtered by "
        "the authenticated user's ID, extracted from the JWT token via a FastAPI dependency. Agent names "
        "are unique per user (composite unique constraint on user_id + name). Integration configs are "
        "likewise scoped per user with a composite unique constraint on user_id + service."
    )

    add_heading2(doc, "Cost Tracking", number="7.2")
    add_body(doc,
        "Every AI action node records the LLM model used, input token count, output token count, and "
        "calculated cost (USD). Costs are computed from a hardcoded per-model pricing table covering "
        "OpenAI and Anthropic models. The cumulative cost is stored on the execution record and surfaced "
        "in the dashboard. The pricing table includes GPT-4o ($2.50/$10.00 per M tokens), GPT-4o-mini "
        "($0.15/$0.60), Claude Sonnet 4 ($3.00/$15.00), and others."
    )

    add_heading2(doc, "Error Handling & Recovery", number="7.3")
    add_body(doc,
        "The system implements layered error handling. At the node level, exceptions are caught and "
        "recorded in the execution_node_logs table with status=failed. At the workflow level, a node "
        "failure terminates BFS traversal and marks the execution as failed. At the API level, a global "
        "exception handler returns a 500 response with the error detail in debug mode. The Celery task "
        "wrapper catches exceptions and attempts to mark the execution as failed in the database even "
        "if the executor crashes."
    )

    add_heading2(doc, "Logging & Observability", number="7.4")
    add_body(doc,
        "Python's standard logging module is used throughout. Each service, router, and executor module "
        "defines a module-level logger. Execution events are published in real-time to Redis pub/sub, "
        "providing a live audit trail. Execution node logs persist every node's input, output, duration, "
        "LLM usage, and error messages to PostgreSQL for post-hoc analysis."
    )

    doc.add_page_break()

    # =====================================================================
    # 8. Technology Stack
    # =====================================================================
    add_heading1(doc, "Technology Stack", number="8")
    add_body(doc, "The following table enumerates all technologies used across the platform.")

    create_professional_table(
        doc,
        headers=["Layer", "Technology", "Version", "Purpose"],
        rows=[
            ["Frontend", "React", "19.2", "UI framework"],
            ["Frontend", "Vite", "8.0", "Build tool and dev server"],
            ["Frontend", "TypeScript", "6.0", "Type safety"],
            ["Frontend", "TanStack Query", "5.x", "Server state management"],
            ["Frontend", "React Flow", "12.x", "Workflow canvas"],
            ["Frontend", "Zustand", "5.x", "Client state management"],
            ["Frontend", "Tailwind CSS", "3.4", "Utility-first styling"],
            ["Frontend", "Lucide React", "1.x", "Icon library"],
            ["Backend", "FastAPI", "0.115+", "Async API framework"],
            ["Backend", "SQLAlchemy", "2.0", "Async ORM"],
            ["Backend", "Pydantic", "2.x", "Request/response validation"],
            ["Backend", "Celery", "5.x", "Distributed task queue"],
            ["Backend", "RedBeat", "2.x", "Dynamic cron scheduling"],
            ["Backend", "Redis", "7.x", "Cache, broker, pub/sub"],
            ["Database", "PostgreSQL", "14+", "Primary relational datastore"],
            ["Database", "pgvector", "0.7+", "Vector embeddings (planned)"],
            ["AI", "OpenAI API", "GPT-4o, 4o-mini", "LLM inference"],
            ["AI", "Anthropic API", "Claude Sonnet 4", "LLM inference"],
            ["Auth", "JWT + bcrypt", "\u2014", "Authentication and password hashing"],
            ["Email", "Resend API", "\u2014", "Transactional email delivery"],
            ["Deploy", "Railway / Render", "\u2014", "Cloud hosting (PaaS)"],
        ]
    )

    doc.add_page_break()

    # =====================================================================
    # 9. Deployment Architecture
    # =====================================================================
    add_heading1(doc, "Deployment Architecture", number="9")

    add_heading2(doc, "Development (localhost)", number="9.1")
    add_body(doc,
        "During development, the system runs locally with the following processes: "
        "(1) Vite dev server on port 5173 with hot module replacement, "
        "(2) Uvicorn running the FastAPI app on port 8000, "
        "(3) A local PostgreSQL instance, "
        "(4) A local Redis instance, and "
        "(5) A Celery worker process. "
        "Environment variables are loaded from a .env file via Pydantic Settings."
    )

    add_heading2(doc, "Production (Railway / Render)", number="9.2")
    add_body(doc,
        "In production, the backend is containerised using two Dockerfiles: Dockerfile (API server) and "
        "Dockerfile.worker (Celery worker). Both images share the same codebase and requirements.txt. "
        "Railway or Render provisions managed PostgreSQL and Redis add-ons. The frontend is built as "
        "static assets and served via a CDN or the API server. Environment variables (DATABASE_URL, "
        "REDIS_URL, JWT_SECRET, API keys) are configured through the hosting platform's dashboard."
    )

    add_heading2(doc, "Demo Mode (Single-File HTML)", number="9.3")
    add_body(doc,
        "For demonstrations and stakeholder reviews, a standalone single-file HTML demo can be generated "
        "that embeds the full UI with mock data. This allows showcasing the platform without deploying "
        "any backend infrastructure."
    )

    doc.add_page_break()

    # =====================================================================
    # 10. Security Architecture
    # =====================================================================
    add_heading1(doc, "Security Architecture", number="10")

    add_body(doc, "Security is implemented at multiple layers throughout the platform.")

    create_professional_table(
        doc,
        headers=["Layer", "Mechanism", "Detail"],
        rows=[
            ["Authentication", "JWT (HS256)", "24-hour expiry, user ID and email in payload"],
            ["Password Storage", "bcrypt", "12-round hashing, constant-time comparison"],
            ["Authorisation", "User-scoped queries", "All DB queries filtered by user_id from JWT"],
            ["Transport", "HTTPS", "TLS termination at load balancer / PaaS"],
            ["CORS", "Origin allowlist", "Strict origins in production, wildcard in debug only"],
            ["Code Sandbox", "Subprocess isolation", "Forbidden imports (os, sys, subprocess, socket, requests), 30s timeout"],
            ["API Keys", "Server-side only", "OpenAI, Anthropic, Resend keys never sent to client"],
            ["Integration Secrets", "Masked responses", "API keys and webhook URLs are masked before returning to client"],
            ["Input Validation", "Pydantic v2", "All request bodies validated with strict type checking"],
            ["SQL Injection", "SQLAlchemy ORM", "Parameterised queries via ORM, no raw SQL"],
        ]
    )

    doc.add_page_break()

    # =====================================================================
    # 11. Integration Architecture
    # =====================================================================
    add_heading1(doc, "Integration Architecture", number="11")

    add_body(doc,
        "AutoMind integrates with external services through a pluggable integration framework. "
        "Users configure integrations via the Settings page, and integration nodes in workflows "
        "consume these configurations at execution time."
    )

    create_professional_table(
        doc,
        headers=["Service", "Protocol", "Node Type", "Capabilities"],
        rows=[
            ["OpenAI", "REST API (HTTPS)", "ai_action", "Chat completions (GPT-4o, 4o-mini, 4.1 family, o3-mini)"],
            ["Anthropic", "REST API (HTTPS)", "ai_action", "Messages API (Claude Sonnet 4, Haiku 4.5)"],
            ["DuckDuckGo", "HTML scraping (HTTPS)", "web_search", "Web search without API key"],
            ["Resend", "REST API (HTTPS)", "integration / escalation", "Transactional email delivery"],
            ["Slack", "Webhook (HTTPS)", "integration", "Channel message posting via incoming webhook"],
        ]
    )

    add_body(doc,
        "New integrations can be added by implementing the IntegrationNodeExecutor._send_<service> "
        "pattern and registering the service name in the ALLOWED_SERVICES set."
    )

    # =====================================================================
    # 12. Scalability Considerations
    # =====================================================================
    add_heading1(doc, "Scalability Considerations", number="12")

    add_body(doc,
        "The AutoMind architecture is designed to scale horizontally at each tier."
    )

    create_professional_table(
        doc,
        headers=["Concern", "Current Design", "Scale Path"],
        rows=[
            ["API throughput", "Single Uvicorn process, async handlers", "Multiple Uvicorn workers behind load balancer"],
            ["Workflow execution", "Single Celery worker", "Multiple Celery workers with configurable concurrency"],
            ["Database connections", "asyncpg connection pool (5 conns)", "PgBouncer connection pooler, read replicas"],
            ["Redis", "Single Redis node", "Redis Sentinel or Redis Cluster"],
            ["Scheduled agents", "RedBeat on single worker", "RedBeat with dedicated beat worker"],
            ["LLM calls", "Sequential per workflow", "Parallel node execution for independent branches"],
            ["Memory / Embeddings", "Chronological window", "pgvector ANN index for semantic retrieval"],
            ["Frontend", "Static SPA", "CDN distribution, code splitting, lazy loading"],
        ]
    )

    add_body(doc,
        "The primary bottleneck at scale will be LLM API latency, which is external and cannot be "
        "reduced architecturally. The system mitigates this by offloading all LLM calls to Celery "
        "workers, keeping the API server responsive."
    )

    # Save
    path = os.path.join(OUTPUT_DIR, "GoalCert_AutoMind_HLD_v1_0.docx")
    doc.save(path)
    print(f"HLD saved to {path}")
    return path


# =============================================================================
# LLD DOCUMENT
# =============================================================================

def generate_lld():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE
    font.color.rgb = DARK_TEXT

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Cover Page --
    add_cover_page(
        doc,
        "Low-Level Design Document",
        "Detailed component specifications, data models, and API contracts "
        "for the AutoMind agentic AI workflow platform"
    )

    # -- Document Control --
    add_document_control(doc, "Low-Level Design (LLD)")
    add_version_history(doc)

    # -- Table of Contents placeholder --
    add_heading1(doc, "Table of Contents", number="")
    add_body(doc, "[Table of contents to be generated from headings]")
    doc.add_page_break()

    # =====================================================================
    # 1. Introduction
    # =====================================================================
    add_heading1(doc, "Introduction", number="1")
    add_body(doc,
        "This Low-Level Design document provides detailed specifications for every component, data model, "
        "and API endpoint in the GoalCert AutoMind platform. It is intended for developers implementing "
        "or maintaining the system and should be read alongside the High-Level Design document."
    )
    add_body(doc,
        "The document is structured around three pillars: Component Design (section 2), Data Model "
        "(section 3), and API Specification (section 4). Sections 5 through 7 cover sequence flows, "
        "security implementation, and error handling."
    )

    doc.add_page_break()

    # =====================================================================
    # 2. Component / Module Design
    # =====================================================================
    add_heading1(doc, "Component / Module Design", number="2")
    add_body(doc,
        "Each component below is described with its responsibilities, inputs, outputs, error handling "
        "strategy, and dependencies."
    )

    # 2.1 Authentication Service
    add_heading2(doc, "Authentication Service", number="2.1")
    add_body(doc,
        "Handles user registration, login, JWT generation, and token validation. Passwords are hashed "
        "with bcrypt before storage. The service issues HS256 JWTs with a 24-hour expiry. A FastAPI "
        "dependency (get_current_user) extracts the token from the Authorization header, decodes it, "
        "and fetches the corresponding User record."
    )
    add_component_table(doc,
        "Authentication Service",
        "Python 3.12 / FastAPI",
        "UserRegister (email, password, name) or UserLogin (email, password)",
        "AuthResponse (user object + JWT token string)",
        "IntegrityError on duplicate email returns 409. Invalid credentials return 401. Expired or "
        "malformed tokens return 401.",
        "bcrypt, PyJWT, SQLAlchemy (User model), FastAPI OAuth2PasswordBearer"
    )

    # 2.2 Agent CRUD Service
    add_heading2(doc, "Agent CRUD Service", number="2.2")
    add_body(doc,
        "Manages the lifecycle of AI agents. Creating an agent also creates a default empty Workflow. "
        "If a template_id is provided, the template's workflow_definition is copied to the new workflow. "
        "The service computes per-agent statistics (total executions, success rate) from the executions "
        "table using SQL aggregation. Pause/resume operations toggle the agent status and register or "
        "remove the RedBeat cron schedule."
    )
    add_component_table(doc,
        "Agent CRUD Service",
        "Python 3.12 / FastAPI",
        "AgentCreate (name, type, description, template_id), AgentUpdate (partial fields), agent_id path param",
        "AgentResponse (id, name, type, status, schedule, stats, timestamps)",
        "Duplicate agent name returns 409 (IntegrityError). Agent not found returns 404. All queries "
        "filtered by user_id for multi-tenant isolation.",
        "SQLAlchemy (Agent, Workflow, Execution, AgentTemplate models), SchedulerService"
    )

    # 2.3 Workflow Engine
    add_heading2(doc, "Workflow Engine", number="2.3")
    add_body(doc,
        "The WorkflowExecutor class is the central orchestrator. It is instantiated with an execution_id "
        "and workflow_definition (React Flow JSON). The engine parses the definition into nodes and edges, "
        "identifies the start node (trigger type preferred), and performs a BFS traversal. Each node is "
        "dispatched to the appropriate executor from a registry dict. Output variables are merged into a "
        "shared context dict. Decision nodes return a branch value (true/false) that filters which outgoing "
        "edges are followed."
    )
    add_component_table(doc,
        "Workflow Engine (WorkflowExecutor)",
        "Python 3.12 / asyncio",
        "execution_id (UUID string), workflow_definition (dict with nodes, edges, viewport), "
        "optional initial variables dict",
        "Final variables dict containing all outputs from all executed nodes",
        "Node-level exceptions are caught, logged to execution_node_logs, and terminate traversal. "
        "The execution is marked as failed with the error message. A __STREAM_END__ event is always "
        "published on completion (success or failure).",
        "graph.py (parse_workflow, get_next_nodes), All NodeExecutor classes, Redis (pub/sub), "
        "SQLAlchemy (Execution, ExecutionNodeLog), MemoryService"
    )

    # 2.4 Node Executors
    add_heading2(doc, "Node Executors", number="2.4")
    add_body(doc,
        "Each node type has a dedicated executor class implementing the BaseNodeExecutor interface. "
        "The abstract method signature is: execute(config: dict, variables: dict, **kwargs) -> dict."
    )

    node_types = [
        ("trigger", "TriggerNodeExecutor", "Sets trigger_time and triggered_by in the variable context. "
         "This is always the first node in a workflow."),
        ("web_search", "WebSearchNodeExecutor", "Posts to DuckDuckGo's HTML endpoint, parses results "
         "with regex, and stores them in the specified output_variable."),
        ("ai_action", "AIActionNodeExecutor", "Interpolates variables into the prompt template, calls "
         "the appropriate LLM provider (OpenAI or Anthropic), parses the response (attempts JSON first), "
         "calculates cost, and stores the output in the specified variable."),
        ("decision", "DecisionNodeExecutor", "Evaluates a condition (left_operand operator right_operand) "
         "with support for ==, !=, >, <, >=, <=, and contains. Returns branch=true or branch=false."),
        ("integration", "IntegrationNodeExecutor", "Dispatches to email (Resend API) or Slack (webhook) "
         "based on the service field. Supports variable interpolation in recipients, subject, and body."),
        ("escalation", "EscalationNodeExecutor", "Sends an escalation email via Resend API to a specified "
         "recipient. Returns mock response if no API key is configured."),
        ("code_exec", "CodeExecNodeExecutor", "Executes user-provided Python code in a subprocess with "
         "forbidden import checking and a configurable timeout (max 30 seconds). Captures local variables "
         "as output."),
    ]

    create_professional_table(
        doc,
        headers=["Node Type", "Executor Class", "Description"],
        rows=[[t, c, d] for t, c, d in node_types],
    )

    # 2.5 AI Integration Service
    add_heading2(doc, "AI Integration Service", number="2.5")
    add_body(doc,
        "The AI integration layer within AIActionNodeExecutor abstracts over two LLM providers. "
        "Model routing is determined by the model name prefix: gpt-*, o1*, o3*, o4* route to OpenAI; "
        "all others (claude-*) route to Anthropic."
    )
    add_body(doc, "Token cost calculation uses the following pricing table (USD per million tokens):")
    create_professional_table(
        doc,
        headers=["Model", "Input Cost / M tokens", "Output Cost / M tokens"],
        rows=[
            ["gpt-4o", "$2.50", "$10.00"],
            ["gpt-4o-mini", "$0.15", "$0.60"],
            ["gpt-4.1", "$2.00", "$8.00"],
            ["gpt-4.1-mini", "$0.40", "$1.60"],
            ["gpt-4.1-nano", "$0.10", "$0.40"],
            ["o3-mini", "$1.10", "$4.40"],
            ["claude-sonnet-4-20250514", "$3.00", "$15.00"],
            ["claude-haiku-4-5-20251001", "$0.80", "$4.00"],
        ]
    )

    # 2.6 Web Search Service
    add_heading2(doc, "Web Search Service", number="2.6")
    add_body(doc,
        "The web search service requires no API key. It submits a POST request to DuckDuckGo's HTML "
        "endpoint (https://html.duckduckgo.com/html/) with the search query. The HTML response is "
        "parsed using regex to extract result links, titles, and snippets. Results are returned as a "
        "list of dicts with keys: title, url, snippet, rank."
    )
    add_component_table(doc,
        "Web Search Service",
        "Python 3.12 / httpx (async)",
        "query (str, interpolated from template), max_results (int, default 5), output_variable (str)",
        "Dict with output_variables: {output_variable: [{title, url, snippet, rank}, ...]}",
        "HTTP errors and timeouts are caught; an empty results list and error message are returned.",
        "httpx, re (regex), variables.interpolate"
    )

    # 2.7 SSE Streaming Service
    add_heading2(doc, "SSE Streaming Service", number="2.7")
    add_body(doc,
        "Real-time execution monitoring is implemented via Server-Sent Events. The executor publishes "
        "JSON-encoded log entries to a Redis pub/sub channel named execution:{execution_id}:logs. "
        "Each log entry contains: timestamp (ISO 8601), message, node_id, and status (info, running, "
        "success, error, done). The SSE endpoint subscribes to this channel using aioredis and yields "
        "events to the client. The stream terminates when a __STREAM_END__ message is received or the "
        "client disconnects."
    )
    add_component_table(doc,
        "SSE Streaming Service",
        "Python 3.12 / FastAPI + sse-starlette",
        "execution_id (UUID path parameter)",
        "SSE event stream: JSON objects with timestamp, message, node_id, status fields",
        "Client disconnect is detected via request.is_disconnected(). Redis subscription is always "
        "cleaned up in a finally block.",
        "redis.asyncio (aioredis), sse-starlette.EventSourceResponse"
    )

    # 2.8 Chat Service
    add_heading2(doc, "Chat Service", number="2.8")
    add_body(doc,
        "The chat service enables conversational interaction with an agent using its execution memory "
        "as context. The system prompt includes the agent's name, type, description, schedule, and "
        "formatted memory context (last 15 execution summaries). User messages and conversation history "
        "(last 10 turns) are appended. The response is streamed token-by-token via SSE using OpenAI's "
        "streaming API (gpt-4o-mini)."
    )
    add_component_table(doc,
        "Chat Service",
        "Python 3.12 / FastAPI + OpenAI Streaming",
        "ChatRequest (message: str, history: list of role/content dicts), agent_id path param",
        "SSE stream of {type: 'token', content: str} events, terminated by {type: 'done'}",
        "Streaming errors emit a {type: 'error', content: str} event. Missing OpenAI API key returns "
        "a non-streaming fallback response.",
        "OpenAI AsyncOpenAI, MemoryService.get_agent_context, sse-starlette"
    )

    # 2.9 Scheduling Service
    add_heading2(doc, "Scheduling Service", number="2.9")
    add_body(doc,
        "The SchedulerService manages dynamic cron scheduling using RedBeat, a Celery Beat scheduler "
        "backed by Redis. When an agent is deployed with a cron expression, a RedBeatSchedulerEntry is "
        "created that fires the execute_workflow_scheduled Celery task. When an agent is paused or deleted, "
        "the entry is removed. Cron expressions follow the standard 5-field format (minute, hour, "
        "day_of_month, month, day_of_week)."
    )
    add_component_table(doc,
        "Scheduling Service",
        "Python 3.12 / Celery + RedBeat",
        "agent_id (UUID string), cron_str (5-field cron expression), timezone (str, default UTC)",
        "None (side effect: RedBeat entry created/deleted in Redis)",
        "Invalid cron expression raises ValueError. Unschedule silently succeeds if no entry exists.",
        "RedBeatSchedulerEntry, celery.schedules.crontab, celery_app"
    )

    # 2.10 Code Execution Sandbox
    add_heading2(doc, "Code Execution Sandbox", number="2.10")
    add_body(doc,
        "The code execution sandbox runs user-provided Python code in a subprocess with security "
        "restrictions. Before execution, the code is scanned for forbidden patterns: import os, "
        "import sys, import subprocess, import shutil, __import__, eval(, exec(, open(, import socket, "
        "import requests, import urllib. The code is wrapped in a try/except block that captures local "
        "variables of simple types (str, int, float, bool, list, dict) as JSON output."
    )
    add_component_table(doc,
        "Code Execution Sandbox",
        "Python 3.12 / asyncio.create_subprocess_exec",
        "code (str, interpolated), timeout (int, max 30s), output_variable (str)",
        "Dict with output_variables: {output_variable: {captured local variables}}",
        "Forbidden import detection returns an error dict without executing. Subprocess timeout raises "
        "asyncio.TimeoutError. Non-zero exit codes return stderr output.",
        "asyncio, tempfile, json"
    )

    doc.add_page_break()

    # =====================================================================
    # 3. Data Model
    # =====================================================================
    add_heading1(doc, "Data Model", number="3")

    add_heading2(doc, "Entity Relationship Overview", number="3.1")
    add_body(doc,
        "The AutoMind data model consists of 8 tables with the following relationships. All primary keys "
        "are UUIDs generated server-side. All timestamps use timezone-aware datetime."
    )

    create_professional_table(
        doc,
        headers=["Entity", "Relationship", "Target", "Cardinality", "On Delete"],
        rows=[
            ["users", "has many", "agents", "1:N", "CASCADE"],
            ["users", "has many", "integrations", "1:N", "CASCADE"],
            ["agents", "has one", "workflows", "1:1", "CASCADE"],
            ["agents", "has many", "executions", "1:N", "CASCADE"],
            ["workflows", "belongs to", "agents", "N:1", "\u2014"],
            ["executions", "has many", "execution_node_logs", "1:N", "CASCADE"],
            ["executions", "belongs to", "agents", "N:1", "\u2014"],
            ["executions", "belongs to", "workflows", "N:1", "\u2014"],
            ["agent_memory", "belongs to", "agents", "N:1", "CASCADE"],
            ["agent_memory", "references", "executions", "N:1", "SET NULL"],
        ]
    )

    # 3.2 users
    add_heading2(doc, "Table: users", number="3.2")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique user identifier"],
            ["email", "VARCHAR(255)", "UNIQUE, NOT NULL", "User email address"],
            ["password_hash", "VARCHAR(255)", "NOT NULL", "bcrypt-hashed password"],
            ["name", "VARCHAR(255)", "NULLABLE", "Display name"],
            ["created_at", "TIMESTAMPTZ", "NOT NULL, default now()", "Account creation timestamp"],
            ["updated_at", "TIMESTAMPTZ", "NOT NULL, default now(), on update now()", "Last modification timestamp"],
        ]
    )

    # 3.3 agents
    add_heading2(doc, "Table: agents", number="3.3")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique agent identifier"],
            ["user_id", "UUID", "FK -> users.id, ON DELETE CASCADE, NOT NULL", "Owning user"],
            ["name", "VARCHAR(255)", "NOT NULL, UNIQUE(user_id, name)", "Agent name (unique per user)"],
            ["description", "TEXT", "NULLABLE", "Free-text description"],
            ["type", "VARCHAR(50)", "NOT NULL", "Agent type: sales, marketing, support, custom"],
            ["status", "VARCHAR(50)", "default 'draft'", "Lifecycle status: draft, active, paused"],
            ["schedule_cron", "VARCHAR(100)", "NULLABLE", "5-field cron expression"],
            ["schedule_timezone", "VARCHAR(50)", "default 'UTC'", "IANA timezone for cron"],
            ["created_at", "TIMESTAMPTZ", "NOT NULL, default now()", "Creation timestamp"],
            ["updated_at", "TIMESTAMPTZ", "NOT NULL, default now(), on update", "Last update"],
            ["last_execution_at", "TIMESTAMPTZ", "NULLABLE", "Timestamp of most recent execution"],
        ]
    )

    # 3.4 workflows
    add_heading2(doc, "Table: workflows", number="3.4")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique workflow identifier"],
            ["agent_id", "UUID", "FK -> agents.id, UNIQUE, ON DELETE CASCADE", "Parent agent (1:1)"],
            ["status", "VARCHAR(50)", "default 'draft'", "Workflow status: draft, active"],
            ["definition", "JSONB", "NOT NULL, default '{}'", "React Flow JSON (nodes, edges, viewport)"],
            ["created_at", "TIMESTAMPTZ", "NOT NULL, default now()", "Creation timestamp"],
            ["updated_at", "TIMESTAMPTZ", "NOT NULL, default now(), on update", "Last update"],
            ["deployed_at", "TIMESTAMPTZ", "NULLABLE", "Deployment timestamp"],
        ]
    )

    # 3.5 executions
    add_heading2(doc, "Table: executions", number="3.5")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique execution identifier"],
            ["agent_id", "UUID", "FK -> agents.id, ON DELETE CASCADE", "Parent agent"],
            ["workflow_id", "UUID", "FK -> workflows.id", "Executed workflow snapshot"],
            ["status", "VARCHAR(50)", "default 'pending'", "pending, running, success, failed"],
            ["triggered_by", "VARCHAR(50)", "NOT NULL", "manual, schedule"],
            ["started_at", "TIMESTAMPTZ", "NULLABLE", "Execution start time"],
            ["ended_at", "TIMESTAMPTZ", "NULLABLE", "Execution end time"],
            ["duration_ms", "INTEGER", "NULLABLE", "Total execution duration in milliseconds"],
            ["error_message", "TEXT", "NULLABLE", "Error details if status=failed"],
            ["variables", "JSONB", "NOT NULL, default '{}'", "Final variable context after execution"],
            ["total_cost", "NUMERIC(10,6)", "NOT NULL, default 0", "Cumulative LLM cost in USD"],
            ["created_at", "TIMESTAMPTZ", "NOT NULL, default now()", "Record creation timestamp"],
        ]
    )

    # 3.6 execution_node_logs
    add_heading2(doc, "Table: execution_node_logs", number="3.6")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique log entry identifier"],
            ["execution_id", "UUID", "FK -> executions.id, ON DELETE CASCADE", "Parent execution"],
            ["node_id", "VARCHAR(255)", "NOT NULL", "React Flow node ID"],
            ["node_type", "VARCHAR(50)", "NOT NULL", "Node type (trigger, ai_action, etc.)"],
            ["node_label", "VARCHAR(255)", "NULLABLE", "Human-readable node label"],
            ["status", "VARCHAR(50)", "default 'pending'", "pending, running, success, failed, skipped"],
            ["started_at", "TIMESTAMPTZ", "NULLABLE", "Node execution start"],
            ["ended_at", "TIMESTAMPTZ", "NULLABLE", "Node execution end"],
            ["duration_ms", "INTEGER", "NULLABLE", "Node execution duration"],
            ["input_data", "JSONB", "NOT NULL, default '{}'", "Node configuration at execution time"],
            ["output_data", "JSONB", "NOT NULL, default '{}'", "Node output / result data"],
            ["error_message", "TEXT", "NULLABLE", "Error details if failed"],
            ["llm_usage", "JSONB", "NULLABLE", "LLM token counts and cost (ai_action nodes)"],
            ["created_at", "TIMESTAMPTZ", "NOT NULL, default now()", "Record creation timestamp"],
        ]
    )

    # 3.7 agent_memory
    add_heading2(doc, "Table: agent_memory", number="3.7")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique memory identifier"],
            ["agent_id", "UUID", "FK -> agents.id, ON DELETE CASCADE, INDEX", "Parent agent"],
            ["execution_id", "UUID", "FK -> executions.id, ON DELETE SET NULL, NULLABLE", "Source execution"],
            ["summary", "TEXT", "NOT NULL", "Execution summary (AI-generated or fallback)"],
            ["key_outputs", "JSONB", "default '{}'", "Structured key outputs from execution"],
            ["memory_type", "VARCHAR(50)", "default 'execution_summary'", "Memory category"],
            ["created_at", "TIMESTAMPTZ", "default now()", "Creation timestamp"],
        ]
    )

    # 3.8 agent_templates
    add_heading2(doc, "Table: agent_templates", number="3.8")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique template identifier"],
            ["name", "VARCHAR(255)", "NOT NULL", "Template display name"],
            ["description", "TEXT", "NULLABLE", "Template description"],
            ["type", "VARCHAR(50)", "NOT NULL", "Agent type category"],
            ["workflow_definition", "JSONB", "NOT NULL", "Pre-built workflow (nodes, edges, viewport)"],
            ["icon", "VARCHAR(50)", "NULLABLE", "Lucide icon name"],
            ["color", "VARCHAR(50)", "NULLABLE", "Theme colour"],
            ["features", "JSONB", "default '[]'", "Feature tag list"],
            ["created_at", "TIMESTAMPTZ", "default now()", "Creation timestamp"],
        ]
    )

    # 3.9 integrations
    add_heading2(doc, "Table: integrations", number="3.9")
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK, default gen_random_uuid()", "Unique integration identifier"],
            ["user_id", "UUID", "FK -> users.id, ON DELETE CASCADE", "Owning user"],
            ["service", "VARCHAR(100)", "NOT NULL, UNIQUE(user_id, service)", "Service name (resend, slack)"],
            ["config", "JSONB", "default '{}'", "Service-specific config (API keys, webhooks)"],
            ["status", "VARCHAR(50)", "default 'active'", "Integration status"],
            ["created_at", "TIMESTAMPTZ", "default now()", "Creation timestamp"],
            ["updated_at", "TIMESTAMPTZ", "default now(), on update", "Last update"],
        ]
    )

    # 3.10 Future: agent_memory_embeddings
    add_heading2(doc, "Future: agent_memory_embeddings", number="3.10")
    add_body(doc,
        "A planned table for vector-based semantic memory retrieval. Requires the pgvector extension."
    )
    create_professional_table(
        doc,
        headers=["Column", "Type", "Constraints", "Description"],
        rows=[
            ["id", "UUID", "PK", "Unique embedding identifier"],
            ["agent_id", "UUID", "FK -> agents.id, ON DELETE CASCADE", "Parent agent"],
            ["content", "TEXT", "NOT NULL", "Source text that was embedded"],
            ["embedding", "vector(1536)", "NOT NULL", "OpenAI text-embedding-3-small vector"],
            ["created_at", "TIMESTAMPTZ", "default now()", "Creation timestamp"],
        ]
    )
    add_body(doc, "An IVFFlat or HNSW index on the embedding column will enable efficient approximate nearest neighbour search.")

    doc.add_page_break()

    # =====================================================================
    # 4. API Specification
    # =====================================================================
    add_heading1(doc, "API Specification", number="4")
    add_body(doc,
        "All API endpoints are prefixed with /api. Authentication is via Bearer token in the "
        "Authorization header. Responses use JSON. Timestamps are ISO 8601 with timezone."
    )

    # 4.1 Authentication
    add_heading2(doc, "REST API \u2014 Authentication", number="4.1")
    add_api_table(doc, [
        ["POST", "/api/auth/register", "None", '{"email", "password", "name?"}', '{"user": {...}, "token": "jwt"}', "201, 409"],
        ["POST", "/api/auth/login", "None", '{"email", "password"}', '{"user": {...}, "token": "jwt"}', "200, 401"],
        ["GET", "/api/auth/me", "Bearer JWT", "\u2014", '{"id", "email", "name", "created_at"}', "200, 401"],
    ])

    # 4.2 Agents
    add_heading2(doc, "REST API \u2014 Agents", number="4.2")
    add_api_table(doc, [
        ["GET", "/api/agents", "Bearer JWT", "\u2014", "AgentResponse[]", "200"],
        ["POST", "/api/agents", "Bearer JWT", '{"name", "type", "description?", "template_id?"}', "AgentResponse", "201, 409"],
        ["POST", "/api/agents/generate", "Bearer JWT", '{"description"}', "AgentResponse", "201, 422, 500"],
        ["GET", "/api/agents/:id", "Bearer JWT", "\u2014", "AgentResponse", "200, 404"],
        ["PATCH", "/api/agents/:id", "Bearer JWT", "Partial AgentUpdate", "AgentResponse", "200, 404"],
        ["DELETE", "/api/agents/:id", "Bearer JWT", "\u2014", "\u2014", "204, 404"],
        ["POST", "/api/agents/:id/pause", "Bearer JWT", "\u2014", "AgentResponse", "200, 404"],
        ["POST", "/api/agents/:id/resume", "Bearer JWT", "\u2014", "AgentResponse", "200, 404"],
    ])

    # 4.3 Workflows
    add_heading2(doc, "REST API \u2014 Workflows", number="4.3")
    add_api_table(doc, [
        ["GET", "/api/agents/:id/workflow", "Bearer JWT", "\u2014", "WorkflowResponse", "200, 404"],
        ["PUT", "/api/agents/:id/workflow", "Bearer JWT", '{"definition": {nodes, edges, viewport}}', "WorkflowResponse", "200, 404"],
        ["POST", "/api/agents/:id/workflow/deploy", "Bearer JWT", "\u2014", "WorkflowResponse", "200, 404"],
    ])

    # 4.4 Executions
    add_heading2(doc, "REST API \u2014 Executions", number="4.4")
    add_api_table(doc, [
        ["POST", "/api/agents/:id/execute", "Bearer JWT", "\u2014", "ExecutionResponse", "202, 400, 404"],
        ["GET", "/api/agents/:id/executions", "Bearer JWT", "?limit=20&offset=0", "ExecutionResponse[]", "200, 404"],
        ["GET", "/api/executions/:id", "Bearer JWT", "\u2014", "ExecutionDetailResponse", "200, 404"],
        ["GET", "/api/executions/:id/logs", "Bearer JWT", "\u2014", "ExecutionNodeLogResponse[]", "200, 404"],
        ["GET", "/api/executions/:id/stream", "None*", "\u2014", "SSE event stream", "200"],
    ])
    add_body(doc, "* The SSE stream endpoint does not require authentication as execution IDs are UUIDs (unguessable).", space_after=Pt(2))

    # 4.5 Memory
    add_heading2(doc, "REST API \u2014 Memory", number="4.5")
    add_api_table(doc, [
        ["GET", "/api/agents/:id/memory", "Bearer JWT", "?limit=20&offset=0", '{"memories": [...], "total": int}', "200, 404"],
        ["DELETE", "/api/agents/:id/memory", "Bearer JWT", "\u2014", "\u2014", "204, 404"],
    ])

    # 4.6 Chat
    add_heading2(doc, "REST API \u2014 Chat", number="4.6")
    add_api_table(doc, [
        ["POST", "/api/agents/:id/chat", "Bearer JWT", '{"message", "history": [{role, content}]}', "SSE stream: {type, content}", "200, 404"],
    ])

    # 4.7 Dashboard
    add_heading2(doc, "REST API \u2014 Dashboard", number="4.7")
    add_api_table(doc, [
        ["GET", "/api/dashboard/stats", "Bearer JWT", "\u2014", '{"total_agents", "active_agents", "tasks_completed", "estimated_savings", "avg_response_time"}', "200"],
        ["GET", "/api/dashboard/activity", "Bearer JWT", "\u2014", "ActivityEvent[] (last 20)", "200"],
    ])

    # 4.8 Templates
    add_heading2(doc, "REST API \u2014 Templates", number="4.8")
    add_api_table(doc, [
        ["GET", "/api/templates", "None", "\u2014", "TemplateResponse[]", "200"],
        ["GET", "/api/templates/:id", "None", "\u2014", "TemplateResponse", "200, 404"],
    ])

    # 4.9 Integrations
    add_heading2(doc, "REST API \u2014 Integrations", number="4.9")
    add_api_table(doc, [
        ["GET", "/api/integrations", "Bearer JWT", "\u2014", "IntegrationResponse[] (masked)", "200"],
        ["POST", "/api/integrations", "Bearer JWT", '{"service", "config": {}}', "IntegrationResponse", "201, 400"],
        ["DELETE", "/api/integrations/:id", "Bearer JWT", "\u2014", "\u2014", "204, 404"],
    ])
    add_body(doc,
        "Allowed services: resend, slack. Sensitive config keys (api_key, webhook_url, secret, token) "
        "are masked in GET responses. POST performs an upsert: if the user already has an integration "
        "for the service, the config is updated."
    )

    doc.add_page_break()

    # =====================================================================
    # 5. Key Sequence Flows
    # =====================================================================
    add_heading1(doc, "Key Sequence Flows", number="5")

    # 5.1 Agent Execution Flow
    add_heading2(doc, "Agent Execution Flow", number="5.1")
    add_body(doc, "The following steps describe a complete manual agent execution from trigger to completion.")
    steps = [
        "User clicks 'Run' on the agent detail page.",
        "Frontend sends POST /api/agents/:id/execute with Bearer JWT.",
        "API validates agent ownership and workflow readiness (active or draft, has nodes).",
        "API creates an Execution record with status=pending and triggered_by=manual.",
        "API dispatches a Celery task (execute_workflow) with execution_id and workflow_definition.",
        "API returns 202 Accepted with the ExecutionResponse.",
        "Frontend opens an EventSource to GET /api/executions/:id/stream.",
        "Celery worker picks up the task, creates a fresh async DB engine.",
        "WorkflowExecutor marks execution as running, loads agent memory context.",
        "Executor publishes 'Execution started' to Redis pub/sub channel.",
        "Executor parses workflow into nodes and edges, identifies start node (trigger).",
        "BFS traversal begins. For each node: dispatch to NodeExecutor, merge output variables.",
        "Each node's start and completion are published to Redis and logged to execution_node_logs.",
        "Decision nodes evaluate conditions and return branch=true or branch=false.",
        "Only edges matching the branch label are followed.",
        "AI action nodes call OpenAI or Anthropic, record token usage and cost.",
        "On BFS completion (or first node failure), executor marks execution as success or failed.",
        "Executor saves execution memory summary (AI-generated if API key available).",
        "Executor updates agent.last_execution_at and publishes __STREAM_END__.",
        "Frontend receives stream end, refetches execution detail to show final results.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{i}. ")
        format_run(run, bold=True, size=Pt(10), color=PRIMARY_COLOR)
        run2 = p.add_run(step)
        format_run(run2, size=Pt(10))

    doc.add_paragraph()

    # 5.2 AI Action Node Execution
    add_heading2(doc, "AI Action Node Execution", number="5.2")
    steps = [
        "Executor receives node config containing prompt template, model, max_tokens, temperature, system_prompt, and output_variable.",
        "Variable interpolation replaces all {variable} placeholders in prompt and system_prompt with current variable context values.",
        "Agent memory context (past execution summaries) is prepended to the system prompt.",
        "Provider routing: if model starts with gpt-/o1/o3/o4, use OpenAI; otherwise use Anthropic.",
        "If the required API key is not configured, return a mock response (workflow continues).",
        "Call the provider API with the constructed messages, model, max_tokens, and temperature.",
        "Parse the response: attempt JSON parsing first (handling code fence wrapping), fall back to {output: raw_text}.",
        "Calculate cost: (input_tokens * input_rate + output_tokens * output_rate) / 1,000,000.",
        "If output_variable is specified, wrap the parsed output under that key.",
        "Return output_variables dict and llm_usage dict (model, input_tokens, output_tokens, cost).",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{i}. ")
        format_run(run, bold=True, size=Pt(10), color=PRIMARY_COLOR)
        run2 = p.add_run(step)
        format_run(run2, size=Pt(10))

    doc.add_paragraph()

    # 5.3 Decision Node Execution
    add_heading2(doc, "Decision Node Execution", number="5.3")
    steps = [
        "Executor receives config with left_operand, operator, and right_operand.",
        "Both operands are interpolated against the current variable context.",
        "Operands are converted to strings for comparison.",
        "For ==, !=, and contains: string comparison is performed directly.",
        "For >, <, >=, <=: numeric conversion is attempted (float). If both convert, numeric comparison is used; otherwise, string comparison.",
        "Result is a boolean. branch is set to 'true' or 'false'.",
        "The get_next_nodes function filters outgoing edges by matching the edge label/sourceHandle to the branch value.",
        "Only matching downstream nodes are enqueued in the BFS.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{i}. ")
        format_run(run, bold=True, size=Pt(10), color=PRIMARY_COLOR)
        run2 = p.add_run(step)
        format_run(run2, size=Pt(10))

    doc.add_paragraph()

    # 5.4 Chat Flow
    add_heading2(doc, "Chat Flow", number="5.4")
    steps = [
        "User opens the chat panel for an agent and sends a message.",
        "Frontend sends POST /api/agents/:id/chat with message text and conversation history.",
        "API verifies agent ownership.",
        "API calls MemoryService.get_agent_context to load the last 15 execution summaries.",
        "A system prompt is constructed containing agent name, type, description, status, schedule, and memory context.",
        "The last 10 messages from conversation history are appended, followed by the new user message.",
        "OpenAI streaming API is called with gpt-4o-mini, temperature=0.7, max_tokens=1024.",
        "Tokens are streamed to the client as SSE events: {type: 'token', content: 'word'}.",
        "On completion, a {type: 'done'} event is sent.",
        "On error, a {type: 'error', content: 'message'} event is sent.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{i}. ")
        format_run(run, bold=True, size=Pt(10), color=PRIMARY_COLOR)
        run2 = p.add_run(step)
        format_run(run2, size=Pt(10))

    doc.add_paragraph()

    # 5.5 Scheduled Execution
    add_heading2(doc, "Scheduled Execution", number="5.5")
    steps = [
        "User deploys an agent with a cron schedule (e.g., '0 9 * * 1-5' for weekdays at 9 AM).",
        "SchedulerService.schedule_agent creates a RedBeatSchedulerEntry in Redis.",
        "RedBeat fires the execute_workflow_scheduled Celery task at the scheduled time.",
        "The task loads the agent and workflow from PostgreSQL.",
        "If the agent is not active or has no valid workflow, the task exits silently.",
        "An Execution record is created with triggered_by=schedule.",
        "WorkflowExecutor runs the workflow (same path as manual execution).",
        "On agent pause or delete, SchedulerService.unschedule_agent removes the RedBeat entry.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{i}. ")
        format_run(run, bold=True, size=Pt(10), color=PRIMARY_COLOR)
        run2 = p.add_run(step)
        format_run(run2, size=Pt(10))

    doc.add_page_break()

    # =====================================================================
    # 6. Security Implementation
    # =====================================================================
    add_heading1(doc, "Security Implementation", number="6")

    add_heading2(doc, "Password Hashing (bcrypt, 12 rounds)", number="6.1")
    add_body(doc,
        "User passwords are hashed using bcrypt with a work factor of 12 rounds (auto-generated salt). "
        "Verification uses bcrypt.checkpw which performs constant-time comparison to prevent timing "
        "attacks. Passwords are encoded as UTF-8 before hashing."
    )

    add_heading2(doc, "JWT Tokens (HS256, 24h expiry)", number="6.2")
    add_body(doc,
        "Access tokens are signed with HMAC-SHA256 using a server-side secret (JWT_SECRET environment "
        "variable). The payload contains: sub (user UUID), email, and exp (expiry timestamp, 24 hours "
        "from issuance). The PyJWT library handles encoding and decoding. Expired tokens return 401 "
        "with 'Token has expired'. Invalid tokens return 401 with 'Invalid token'."
    )

    add_heading2(doc, "Multi-Tenant Isolation (user_id filtering)", number="6.3")
    add_body(doc,
        "Every endpoint that accesses user-owned resources includes a WHERE user_id = :current_user_id "
        "clause. The current user is resolved from the JWT token via a FastAPI dependency. There is no "
        "admin or superuser role; all queries are strictly scoped. Foreign key constraints with CASCADE "
        "delete ensure orphaned records cannot persist."
    )

    add_heading2(doc, "Code Sandbox (blocked imports, 30s timeout, subprocess isolation)", number="6.4")
    add_body(doc,
        "User-submitted Python code is executed in a separate subprocess (not the API or worker process). "
        "Before execution, a blocklist scan rejects code containing: import os, import sys, import subprocess, "
        "import shutil, __import__, eval(, exec(, open(, import socket, import requests, import urllib. "
        "The subprocess has a hard timeout of 30 seconds (configurable, clamped to max 30). The code runs "
        "in a temporary file that is deleted after execution."
    )

    add_heading2(doc, "CORS Configuration", number="6.5")
    add_body(doc,
        "In production (DEBUG=False), CORS is restricted to the configured FRONTEND_URL origin. In "
        "development (DEBUG=True), all origins are allowed for convenience. Credentials are always "
        "allowed. All HTTP methods and headers are permitted."
    )

    doc.add_page_break()

    # =====================================================================
    # 7. Error Handling Strategy
    # =====================================================================
    add_heading1(doc, "Error Handling Strategy", number="7")

    add_body(doc, "The platform implements a multi-layered error handling strategy.")

    create_professional_table(
        doc,
        headers=["Layer", "Mechanism", "Behaviour"],
        rows=[
            ["Node Executor", "try/except in _execute_node", "Exception logged to execution_node_logs with status=failed. Error published via SSE. BFS terminates."],
            ["Workflow Engine", "Status tracking in _finish_execution", "Execution marked as success or failed. Duration, cost, variables, and error message persisted. __STREAM_END__ always sent."],
            ["Celery Task", "try/except in execute_workflow_task", "If executor crashes, _mark_failed attempts to set execution.status=failed in DB."],
            ["API Router", "HTTPException raises", "Standard FastAPI exception handling: 400, 401, 404, 409, 422 with detail messages."],
            ["Global Handler", "app.exception_handler(Exception)", "Uncaught exceptions return 500. In DEBUG mode, error detail is included; in production, a generic message is returned."],
            ["AI Integration", "Graceful degradation", "Missing API keys produce mock responses. API errors are caught and returned as error fields without crashing the workflow."],
            ["Email / Slack", "Mock fallback", "Missing Resend or Slack credentials return mock success responses with mock=True flag."],
            ["SSE Stream", "finally block cleanup", "Redis pub/sub subscription is always unsubscribed and connection closed, even on client disconnect."],
        ]
    )

    add_body(doc,
        "This layered approach ensures that a failure at any level is captured, logged, and "
        "communicated to the user without leaving the system in an inconsistent state."
    )

    # Save
    path = os.path.join(OUTPUT_DIR, "GoalCert_AutoMind_LLD_v1_0.docx")
    doc.save(path)
    print(f"LLD saved to {path}")
    return path


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_hld()
    generate_lld()
    print("\nBoth documents generated successfully.")
