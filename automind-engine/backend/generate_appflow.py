"""Generate GoalCert AutoMind Application Flow document as a professional .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# -- Colour constants (matching DFD document) --
PURPLE = RGBColor(73, 2, 162)
WHITE = RGBColor(255, 255, 255)
GRAY = RGBColor(128, 128, 128)
LIGHT_GRAY = RGBColor(245, 245, 245)
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(51, 51, 51)
LIGHT_PURPLE = RGBColor(240, 230, 255)

doc = Document()

# ---------- default style ----------
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = DARK_GRAY
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.space_before = Pt(0)
paragraph_format.line_spacing = 1.15

# ---------- heading styles ----------
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hf = hs.font
    hf.name = "Calibri"
    hf.color.rgb = PURPLE
    hf.bold = True
    if level == 1:
        hf.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hf.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hf.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


# ===================================================================
#  Helper functions
# ===================================================================

def set_cell_shading(cell, hex_color):
    """Apply background shading to a cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_purple_header_row(table, headers):
    """Style the first row of a table as a purple header."""
    row = table.rows[0]
    for idx, cell in enumerate(row.cells):
        set_cell_shading(cell, "4902A2")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(headers[idx])
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)


def add_body_paragraph(doc, text, bold=False, italic=False):
    """Add a body paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(8)
    return p


def add_figure_caption(doc, text):
    """Add a figure caption in italic, centered."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)


def create_flow_diagram(doc, steps, caption):
    """Create a block-diagram table showing flow steps with colored cells and arrows.

    steps: list of strings representing sequential steps in the flow.
    """
    # Build a single-row table with alternating step cells and arrow cells
    num_cols = len(steps) * 2 - 1  # steps + arrows between them

    # If flow is too wide (more than 5 steps), split into rows
    if len(steps) <= 5:
        _create_flow_row(doc, steps)
    else:
        # Split into rows of up to 4 steps
        row_size = 4
        for i in range(0, len(steps), row_size):
            chunk = steps[i:i + row_size]
            is_first = (i == 0)
            is_last = (i + row_size >= len(steps))
            _create_flow_row(doc, chunk, is_first_row=is_first, is_last_row=is_last)
            if not is_last:
                # Add a down-arrow connector between rows
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run("\u2193")
                run.font.size = Pt(16)
                run.font.color.rgb = PURPLE
                run.bold = True
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

    add_figure_caption(doc, caption)


def _create_flow_row(doc, steps, is_first_row=True, is_last_row=True):
    """Create a single row of the flow diagram."""
    num_cols = len(steps) * 2 - 1
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row = table.rows[0]
    step_idx = 0
    for col_idx in range(num_cols):
        cell = row.cells[col_idx]
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        if col_idx % 2 == 0:
            # Step cell
            is_terminal = (is_first_row and col_idx == 0) or (is_last_row and col_idx == num_cols - 1)
            run = p.add_run(steps[step_idx])
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.bold = True

            if is_terminal:
                # Purple background for start/end
                set_cell_shading(cell, "4902A2")
                run.font.color.rgb = WHITE
            else:
                # Light purple for intermediate steps
                set_cell_shading(cell, "F0E6FF")
                run.font.color.rgb = PURPLE

            step_idx += 1
        else:
            # Arrow cell
            run = p.add_run("\u2192")
            run.font.size = Pt(14)
            run.font.color.rgb = PURPLE
            run.bold = True
            # No background - white cell with no border appearance
            set_cell_shading(cell, "FFFFFF")
            # Make arrow cells narrow
            cell.width = Inches(0.4)

    # Set step cell widths
    for col_idx in range(num_cols):
        if col_idx % 2 == 0:
            row.cells[col_idx].width = Inches(1.4)
        else:
            row.cells[col_idx].width = Inches(0.4)


def create_steps_table(doc, steps_data):
    """Create a detailed steps table.

    steps_data: list of (step_number, step_name, description, component) tuples
    """
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_purple_header_row(table, ["Step", "Action", "Description", "Component"])

    for i, (step_num, action, description, component) in enumerate(steps_data):
        row = table.add_row()
        values = [step_num, action, description, component]
        for j, val in enumerate(row.cells):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(values[j]))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY
            if j == 0:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        if i % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "F5F0FF")

    for row in table.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(3.5)
        row.cells[3].width = Inches(1.5)

    return table


# ===================================================================
#  COVER PAGE
# ===================================================================

for _ in range(6):
    doc.add_paragraph()

# GoalCert
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GoalCert")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = PURPLE
run.font.name = "Calibri"

# AutoMind subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AutoMind \u2014 Agentic AI Workforce Platform")
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

# Spacing
doc.add_paragraph()

# Document title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Application Flow Document")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = DARK_GRAY
run.font.name = "Calibri"

# Description
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "A comprehensive mapping of user journeys and system flows\n"
    "across the AutoMind platform, from authentication through\n"
    "agent creation, execution, and monitoring."
)
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

for _ in range(4):
    doc.add_paragraph()

# Confidentiality
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = PURPLE
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This document contains proprietary information belonging to GoalCert Pte. Ltd.\n"
    "Unauthorized distribution or reproduction is strictly prohibited."
)
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = "Calibri"


# ===================================================================
#  DOCUMENT CONTROL
# ===================================================================

doc.add_page_break()

doc.add_heading("Document Control", level=1)

dc_table = doc.add_table(rows=9, cols=2)
dc_table.style = "Table Grid"
dc_table.alignment = WD_TABLE_ALIGNMENT.LEFT

add_purple_header_row(dc_table, ["Field", "Details"])

dc_data = [
    ("Document Title", "GoalCert AutoMind \u2014 Application Flow Document"),
    ("Document ID", "GC-AM-AF-001"),
    ("Version", "1.0"),
    ("Status", "Draft"),
    ("Classification", "Confidential"),
    ("Author", "GoalCert Engineering"),
    ("Date", "25 June 2026"),
    ("Approved By", "Prem Kumar, CTO"),
]

for i, (field, detail) in enumerate(dc_data):
    row = dc_table.rows[i + 1]
    for j, val in enumerate([field, detail]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in dc_table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(4.5)

doc.add_paragraph()

# Version History
doc.add_heading("Version History", level=2)

vh_table = doc.add_table(rows=1, cols=4)
vh_table.style = "Table Grid"
vh_table.alignment = WD_TABLE_ALIGNMENT.LEFT

add_purple_header_row(vh_table, ["Version", "Date", "Author", "Changes"])

vh_data = [
    ("0.1", "20 Jun 2026", "GoalCert Engineering", "Initial draft \u2014 authentication and agent creation flows"),
    ("0.2", "23 Jun 2026", "GoalCert Engineering", "Added execution, workflow builder, chat, and monitoring flows"),
    ("1.0", "25 Jun 2026", "GoalCert Engineering", "Finalised all eight application flows with detailed step tables"),
]

for i, row_data in enumerate(vh_data):
    row = vh_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")


# ===================================================================
#  TABLE OF CONTENTS
# ===================================================================

doc.add_page_break()

doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("1.", "Introduction", "4"),
    ("1.1", "Purpose", "4"),
    ("1.2", "Scope", "4"),
    ("1.3", "Audience", "4"),
    ("2.", "Platform Overview", "5"),
    ("3.", "User Authentication Flow", "6"),
    ("4.", "Agent Creation Flow", "8"),
    ("5.", "Manual Agent Execution Flow", "10"),
    ("6.", "Scheduled Agent Execution Flow", "12"),
    ("7.", "Workflow Builder Flow", "14"),
    ("8.", "Agent Chat Flow", "16"),
    ("9.", "Execution Monitoring Flow", "18"),
    ("10.", "Reports & Analytics Flow", "20"),
    ("11.", "Flow Integration Map", "22"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}\t{title}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    if not "." in num or num.endswith("."):
        run.bold = True
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(0.6))


# ===================================================================
#  1. INTRODUCTION
# ===================================================================

doc.add_page_break()

doc.add_heading("1. Introduction", level=1)

doc.add_heading("1.1 Purpose", level=2)

add_body_paragraph(doc, (
    "This Application Flow document describes the step-by-step user journeys and system-level "
    "processes that occur within the GoalCert AutoMind platform. It traces the path from the moment "
    "a user opens the platform in a browser, through authentication, agent creation, workflow design, "
    "execution orchestration, real-time monitoring, and analytics consumption. Each flow is presented "
    "with a visual block diagram, a narrative description, and a detailed step table that maps actions "
    "to the specific frontend components, API endpoints, and backend services involved."
))

add_body_paragraph(doc, (
    "This document complements the High-Level Design (HLD), Low-Level Design (LLD), and Data Flow "
    "Diagram (DFD) by providing a user-centric view of the platform. While the DFD describes what "
    "data moves where, this document describes what happens from the user's perspective and how the "
    "system responds at each interaction point."
))

doc.add_heading("1.2 Scope", level=2)

add_body_paragraph(doc, (
    "This document covers eight primary application flows that represent the complete lifecycle of "
    "user interaction with AutoMind v1.0: authentication, agent creation (including AI generation), "
    "manual execution, scheduled execution, workflow building, agent chat, execution monitoring, and "
    "reports and analytics. Each flow is documented from the user's first action through the system's "
    "final response, including error paths and edge cases where they are architecturally significant."
))

doc.add_heading("1.3 Audience", level=2)

add_body_paragraph(doc, (
    "This document is intended for product managers, UI/UX designers, QA engineers, and developers "
    "who need to understand the end-to-end user experience and the system behaviour that supports it. "
    "It assumes familiarity with web application concepts (REST APIs, authentication tokens, single-page "
    "applications) but does not require deep knowledge of the platform's internal architecture."
))


# ===================================================================
#  2. PLATFORM OVERVIEW
# ===================================================================

doc.add_page_break()

doc.add_heading("2. Platform Overview", level=1)

add_body_paragraph(doc, (
    "AutoMind is an agentic AI workforce platform that enables users to create, configure, and deploy "
    "autonomous AI agents through a visual workflow builder. The platform consists of a React 19 "
    "single-page application (frontend), a FastAPI backend with Celery workers for async execution, "
    "PostgreSQL for persistent storage, and Redis for task queuing, scheduling, and real-time event "
    "streaming. Users interact with the platform through the following primary pages:"
))

# Page overview table
page_table = doc.add_table(rows=1, cols=3)
page_table.style = "Table Grid"
page_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(page_table, ["Page", "Route", "Purpose"])

page_data = [
    ("Login / Signup", "/login, /signup", "User authentication and registration"),
    ("Dashboard", "/", "Overview of agent statistics, activity feed, and quick actions"),
    ("Templates", "/templates", "Browse and create agents from pre-built templates"),
    ("Agent Detail", "/agents/:id", "View agent configuration, execution history, waterfall timeline, chat panel, and trigger execution"),
    ("Workflow Builder", "/agents/:id/builder", "Full-screen React Flow canvas for designing agent workflows with drag-and-drop nodes"),
    ("Execution Detail", "/executions/:executionId", "Detailed view of a single execution with per-node waterfall, logs, and output data"),
    ("Integrations", "/integrations", "Configure external service connections (Resend email, Slack webhook)"),
    ("Analytics", "/analytics", "Charts and heatmaps for agent performance, execution trends, and cost analysis"),
    ("Reports", "/reports", "Per-agent output summaries from recent executions"),
    ("Settings", "/settings", "User profile and API key configuration"),
]

for i, row_data in enumerate(page_data):
    row = page_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in page_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(3.5)

add_figure_caption(doc, "Table 1: AutoMind platform pages and their routes")

add_body_paragraph(doc, (
    "All protected routes are wrapped in a ProtectedRoute component that checks for a valid JWT token "
    "in the Zustand auth store. If no token is found, the user is redirected to the login page. The "
    "AppLayout component provides the sidebar navigation, top bar, and consistent page structure. The "
    "Workflow Builder is the only protected page that renders without the AppLayout, using the full "
    "viewport for the React Flow canvas."
))


# ===================================================================
#  3. USER AUTHENTICATION FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("3. User Authentication Flow", level=1)

add_body_paragraph(doc, (
    "The authentication flow is the gateway to the platform. Users can register a new account or log "
    "in with existing credentials. Both paths produce a JWT access token that is stored client-side "
    "and attached to every subsequent API request. The flow begins at the public login page and ends "
    "at the authenticated dashboard."
))

doc.add_heading("3.1 Flow Diagram", level=2)

create_flow_diagram(doc, [
    "User opens\n/login",
    "Enter\ncredentials",
    "POST\n/api/auth/login",
    "Verify\npassword",
    "Issue JWT\ntoken",
    "Store token\nin Zustand",
    "Redirect to\nDashboard /",
], "Figure 1: User Authentication Flow \u2014 Login path from credentials to dashboard")

doc.add_heading("3.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "Open Login Page",
     "User navigates to /login. If already authenticated (valid token in localStorage), "
     "the ProtectedRoute component redirects to the dashboard.",
     "LoginPage.tsx"),
    ("2", "Enter Credentials",
     "User enters email address and password into the login form. Client-side validation "
     "checks for non-empty fields before enabling the submit button.",
     "LoginPage.tsx"),
    ("3", "Submit Login Request",
     "Frontend sends POST /api/auth/login with JSON body {email, password}. The request "
     "does not include an Authorization header (public endpoint).",
     "authStore.ts"),
    ("4", "Verify Password",
     "Backend queries the users table for the email. If found, bcrypt.checkpw compares the "
     "submitted password against the stored password_hash. Invalid credentials return 401.",
     "auth.py router"),
    ("5", "Generate JWT",
     "On successful verification, the backend creates a JWT (HS256, 24h expiry) containing "
     "the user's UUID and email. Returns {user: {...}, token: 'jwt_string'}.",
     "security.py"),
    ("6", "Store Token",
     "Frontend stores the JWT token and user object in the Zustand auth store. The token is "
     "also persisted to localStorage for session recovery across page refreshes.",
     "authStore.ts"),
    ("7", "Redirect to Dashboard",
     "React Router navigates to / (dashboard). All subsequent API calls include the JWT in "
     "the Authorization: Bearer header via the configured axios/fetch interceptor.",
     "App.tsx routing"),
])

doc.add_paragraph()

add_body_paragraph(doc, (
    "The registration flow (POST /api/auth/register) follows an identical pattern with an additional "
    "name field. Duplicate email addresses return a 409 Conflict error. On successful registration, "
    "the user is immediately issued a JWT and redirected to the dashboard without needing to log in "
    "separately."
))


# ===================================================================
#  4. AGENT CREATION FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("4. Agent Creation Flow", level=1)

add_body_paragraph(doc, (
    "Agent creation is the first step in building an AI workflow. AutoMind supports three creation "
    "paths: manual creation (specifying name, type, and description), template-based creation (cloning "
    "a pre-built workflow from the template gallery), and AI-powered generation (describing what the "
    "agent should do in natural language and having GPT generate the complete workflow). All three "
    "paths converge at the agent detail page where the user can view and modify the agent."
))

doc.add_heading("4.1 Flow Diagram \u2014 AI Generation Path", level=2)

create_flow_diagram(doc, [
    "User opens\nDashboard /",
    "Click\nCreate Agent",
    "Describe agent\nin natural\nlanguage",
    "POST /api/\nagents/generate",
    "GPT generates\nworkflow JSON",
    "Create Agent\n+ Workflow",
    "Redirect to\n/agents/:id",
    "View agent\ndetail page",
], "Figure 2: Agent Creation Flow \u2014 AI-powered generation from natural language description")

doc.add_heading("4.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "Open Dashboard",
     "User lands on the dashboard which displays agent statistics (total agents, active agents, "
     "tasks completed, estimated savings) and a recent activity feed.",
     "DashboardPage.tsx"),
    ("2", "Initiate Creation",
     "User clicks the 'Create Agent' button. A modal or form appears with options: manual creation, "
     "template selection, or AI generation. For AI generation, a text area accepts a free-form description.",
     "DashboardPage.tsx"),
    ("3", "Describe Agent",
     "User writes a natural language description of what the agent should do. Example: 'Research "
     "trending AI topics, summarise them, and email a weekly digest to the team.'",
     "CreateAgentModal"),
    ("4", "Generate via API",
     "Frontend sends POST /api/agents/generate with {description: '...'}. The backend calls "
     "generate_agent_from_description() which uses GPT to produce agent metadata (name, type, "
     "description) and a complete workflow definition (nodes + edges in React Flow format).",
     "agents.py router"),
    ("5", "AI Workflow Generation",
     "The agent_generator service constructs a prompt asking GPT to output JSON with agent_name, "
     "agent_type, agent_description, schedule (optional cron), and a workflow object containing "
     "nodes (with type, label, config) and edges (with source, target, optional labels).",
     "agent_generator.py"),
    ("6", "Persist Agent + Workflow",
     "Backend creates an Agent record (name, type, description, optional schedule_cron) and a "
     "Workflow record with the generated definition JSONB. Both are committed in a single transaction. "
     "If the agent name already exists for this user, a 409 Conflict is returned.",
     "agents.py router"),
    ("7", "Navigate to Agent Detail",
     "Frontend receives the AgentResponse and navigates to /agents/:id. The agent detail page "
     "loads agent metadata, execution history (initially empty), and provides access to the "
     "workflow builder, chat, and execution controls.",
     "AgentDetailPage.tsx"),
    ("8", "Review and Deploy",
     "User can open the workflow builder to review or modify the generated workflow, then click "
     "Deploy to activate the agent. Deploy sets workflow.status = 'active', agent.status = 'active', "
     "and registers any cron schedule with RedBeat.",
     "WorkflowBuilderPage.tsx"),
])

doc.add_paragraph()

add_body_paragraph(doc, (
    "For template-based creation, the user navigates to /templates, browses the template gallery, "
    "and selects a template. The POST /api/agents request includes a template_id field. The backend "
    "reads the template's workflow_definition from the agent_templates table and clones it into the "
    "new agent's workflow. The user can then customise the workflow in the builder before deploying."
))


# ===================================================================
#  5. MANUAL AGENT EXECUTION FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("5. Manual Agent Execution Flow", level=1)

add_body_paragraph(doc, (
    "Manual execution allows users to trigger an agent's workflow on demand. This is the primary way "
    "to test and run agents during development. The flow begins with a button click on the agent "
    "detail page, dispatches the workflow to a Celery worker for asynchronous execution, and streams "
    "real-time progress back to the browser via Server-Sent Events (SSE)."
))

doc.add_heading("5.1 Flow Diagram", level=2)

create_flow_diagram(doc, [
    "Click\nRun Now",
    "POST /api/\nagents/:id/\nexecute",
    "Create\nExecution\n(pending)",
    "Dispatch\nCelery task",
    "BFS workflow\ntraversal",
    "Publish logs\nvia Redis",
    "SSE stream\nto browser",
    "Show results\non detail page",
], "Figure 3: Manual Agent Execution Flow \u2014 From trigger to real-time results")

doc.add_heading("5.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "Trigger Execution",
     "User clicks the 'Run Now' button on the agent detail page (/agents/:id). The button is "
     "disabled if the workflow has no nodes defined.",
     "AgentDetailPage.tsx"),
    ("2", "API Request",
     "Frontend sends POST /api/agents/:id/execute with Bearer JWT. The backend validates agent "
     "ownership, checks that the workflow exists and has nodes, and verifies the workflow status "
     "is 'active' or 'draft'.",
     "executions.py router"),
    ("3", "Create Execution Record",
     "Backend creates an Execution record in PostgreSQL with status='pending', triggered_by='manual', "
     "and empty variables. The execution is flushed to generate a UUID.",
     "executions.py router"),
    ("4", "Dispatch to Celery",
     "The API dispatches a Celery task (execute_workflow) with args=[execution_id, workflow_definition]. "
     "Returns 202 Accepted immediately with the ExecutionResponse.",
     "celery_app.py"),
    ("5", "Open SSE Stream",
     "Frontend opens an EventSource connection to GET /api/executions/:id/stream. This SSE endpoint "
     "subscribes to the Redis pub/sub channel execution:{id}:logs and forwards all published events.",
     "ExecutionStream"),
    ("6", "Workflow Execution",
     "Celery worker picks up the task, creates a fresh async DB engine (asyncpg engines cannot be "
     "shared across event loops). WorkflowExecutor marks execution as 'running', loads agent memory "
     "context (last 10 summaries), parses the workflow graph, and begins BFS traversal.",
     "executor.py"),
    ("7", "Node-by-Node Processing",
     "For each node in BFS order: the executor identifies the node type, resolves the appropriate "
     "NodeExecutor (trigger, ai_action, decision, integration, web_search, code_exec, escalation), "
     "interpolates {{variable}} placeholders, executes the node, merges output variables into the "
     "shared context, and logs the result to execution_node_logs.",
     "executor.py + nodes/"),
    ("8", "Real-time Streaming",
     "Each node's start and completion is published to the Redis pub/sub channel as a JSON event "
     "containing timestamp, message, node_id, and status (running, success, error). The SSE endpoint "
     "forwards these to the browser in real time.",
     "executor.py + Redis"),
    ("9", "Execution Completion",
     "When BFS completes (or a node fails), the executor marks the execution as 'success' or 'failed', "
     "records total duration and cumulative LLM cost, generates an AI execution summary via GPT-4o-mini, "
     "saves it to agent_memory, updates agent.last_execution_at, and publishes __STREAM_END__.",
     "executor.py"),
    ("10", "Display Results",
     "Frontend receives the __STREAM_END__ sentinel, closes the EventSource, and refetches the "
     "execution detail (GET /api/executions/:id) to display final status, per-node waterfall timeline, "
     "output data, LLM usage, and cost.",
     "AgentDetailPage.tsx"),
])


# ===================================================================
#  6. SCHEDULED AGENT EXECUTION FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("6. Scheduled Agent Execution Flow", level=1)

add_body_paragraph(doc, (
    "Scheduled execution enables agents to run automatically at configured intervals without user "
    "intervention. This is powered by RedBeat, a Celery Beat scheduler backed by Redis, which "
    "dispatches Celery tasks at cron-defined intervals. The flow begins when a user deploys an "
    "agent with a cron schedule and continues autonomously until the agent is paused or deleted."
))

doc.add_heading("6.1 Flow Diagram", level=2)

create_flow_diagram(doc, [
    "Deploy agent\nwith cron\nschedule",
    "Create\nRedBeat\nentry",
    "Cron timer\nfires",
    "Dispatch\nCelery task",
    "Load agent\n+ workflow",
    "Execute\nworkflow\n(BFS)",
    "Save results\n+ memory",
    "Wait for\nnext cron\ntrigger",
], "Figure 4: Scheduled Agent Execution Flow \u2014 Autonomous cron-based workflow execution")

doc.add_heading("6.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "Configure Schedule",
     "User sets a cron expression (e.g., '0 9 * * 1-5' for weekdays at 9 AM) and timezone "
     "in the agent configuration. The schedule is stored on the agent record as schedule_cron "
     "and schedule_timezone fields.",
     "AgentDetailPage.tsx"),
    ("2", "Deploy Workflow",
     "User clicks Deploy on the workflow builder. POST /api/agents/:id/workflow/deploy sets "
     "workflow.status='active', agent.status='active', records deployed_at timestamp. If the "
     "agent has a schedule_cron, the SchedulerService is invoked.",
     "workflows.py router"),
    ("3", "Register RedBeat Entry",
     "SchedulerService.schedule_agent() creates a RedBeatSchedulerEntry in Redis with the "
     "parsed crontab, timezone, and the Celery task name 'execute_workflow_scheduled' with "
     "args=[agent_id]. The entry key is redbeat:agent:{agent_id}.",
     "scheduler_service.py"),
    ("4", "Cron Timer Fires",
     "RedBeat (running as part of Celery Beat) monitors all schedule entries. When the cron "
     "expression matches the current time, it dispatches the execute_workflow_scheduled Celery "
     "task with the agent_id as the argument.",
     "RedBeat / Celery Beat"),
    ("5", "Load Agent and Workflow",
     "The Celery task loads the Agent and Workflow from PostgreSQL. If the agent is not in "
     "'active' status or the workflow has no nodes, the task exits silently without creating "
     "an execution record.",
     "tasks.py"),
    ("6", "Create Execution Record",
     "An Execution record is created with triggered_by='schedule' and status='pending'. The "
     "rest of the execution follows the same WorkflowExecutor path as manual execution.",
     "tasks.py"),
    ("7", "Execute Workflow",
     "WorkflowExecutor runs the BFS traversal, executing each node, publishing logs to Redis "
     "pub/sub (though no browser may be listening for scheduled runs), and recording node logs.",
     "executor.py"),
    ("8", "Save Results and Memory",
     "On completion, the executor saves the execution status, duration, cost, and variables. "
     "An AI-generated execution summary is stored in agent_memory for future context injection.",
     "executor.py + memory"),
    ("9", "Await Next Trigger",
     "The RedBeat entry remains active. The next cron match will fire another execution. "
     "The cycle continues until the user pauses the agent (POST /api/agents/:id/pause) or "
     "deletes it (DELETE /api/agents/:id), which removes the RedBeat entry via unschedule_agent().",
     "scheduler_service.py"),
])


# ===================================================================
#  7. WORKFLOW BUILDER FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("7. Workflow Builder Flow", level=1)

add_body_paragraph(doc, (
    "The workflow builder is the visual canvas where users design agent workflows. It is a full-screen "
    "React Flow 12.x editor that supports seven node types: trigger, ai_action, web_search, decision, "
    "integration, escalation, and code_exec. Users drag nodes from a palette, connect them with edges, "
    "configure each node through a side panel, and save the workflow definition as JSONB. The builder "
    "page does not use the standard AppLayout, giving the canvas the full browser viewport."
))

doc.add_heading("7.1 Flow Diagram", level=2)

create_flow_diagram(doc, [
    "Open builder\n/agents/:id\n/builder",
    "Load\nworkflow\ndefinition",
    "Drag nodes\nonto canvas",
    "Connect\nnodes with\nedges",
    "Configure\nnode settings",
    "Save\nworkflow\n(PUT)",
    "Deploy\nworkflow\n(POST)",
], "Figure 5: Workflow Builder Flow \u2014 Visual workflow design and deployment")

doc.add_heading("7.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "Open Workflow Builder",
     "User clicks 'Edit Workflow' on the agent detail page. React Router navigates to "
     "/agents/:id/builder. The page renders without AppLayout, providing a full-screen canvas "
     "with a left-side node palette and a top toolbar.",
     "WorkflowBuilderPage.tsx"),
    ("2", "Load Workflow Definition",
     "Frontend sends GET /api/agents/:id/workflow to load the existing workflow definition. "
     "The definition contains nodes (array of {id, type, position, data}), edges (array of "
     "{id, source, target, label}), and viewport ({x, y, zoom}). React Flow renders the graph.",
     "WorkflowBuilderPage.tsx"),
    ("3", "Add Nodes",
     "User drags node types from the palette onto the canvas. Available types: Trigger (start node), "
     "AI Action (LLM call), Web Search (DuckDuckGo), Decision (conditional branch), Integration "
     "(email/Slack), Escalation (alert email), Code Exec (Python sandbox). Each node is assigned "
     "a unique ID and positioned at the drop coordinates.",
     "Node palette component"),
    ("4", "Connect Nodes",
     "User draws edges between nodes by dragging from a source handle to a target handle. "
     "Decision nodes have two output handles labeled 'true' and 'false' for conditional branching. "
     "Edges define the execution order during BFS traversal.",
     "React Flow canvas"),
    ("5", "Configure Nodes",
     "Clicking a node opens a configuration side panel. Each node type has specific settings: "
     "AI Action (prompt, model, temperature, max_tokens, system_prompt, output_variable), "
     "Web Search (query, max_results, output_variable), Decision (left_operand, operator, "
     "right_operand), Integration (service, recipients, subject, body/message), Code Exec "
     "(code, timeout, output_variable). Variables from upstream nodes are referenced via "
     "{{variable_name}} syntax.",
     "NodeConfigPanel"),
    ("6", "Save Workflow",
     "User clicks Save. Frontend serialises the React Flow state (nodes, edges, viewport) to "
     "JSON and sends PUT /api/agents/:id/workflow with {definition: {...}}. Backend updates "
     "the workflow.definition JSONB column.",
     "workflows.py router"),
    ("7", "Deploy Workflow",
     "User clicks Deploy. Frontend sends POST /api/agents/:id/workflow/deploy. Backend sets "
     "workflow.status='active', agent.status='active', records deployed_at, and registers any "
     "cron schedule with RedBeat. The workflow is now ready for execution.",
     "workflows.py router"),
])


# ===================================================================
#  8. AGENT CHAT FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("8. Agent Chat Flow", level=1)

add_body_paragraph(doc, (
    "The agent chat feature allows users to have conversational interactions with their agents. "
    "The chat is contextualised with the agent's execution memory \u2014 the last 15 execution summaries "
    "are injected into the system prompt, enabling the agent to answer questions about its past work, "
    "findings, and status. Responses are streamed token-by-token via Server-Sent Events using "
    "OpenAI's gpt-4o-mini model."
))

doc.add_heading("8.1 Flow Diagram", level=2)

create_flow_diagram(doc, [
    "Open chat\npanel on\nagent page",
    "Type and\nsend\nmessage",
    "POST /api/\nagents/:id\n/chat",
    "Load agent\nmemory\n(last 15)",
    "Build\nsystem\nprompt",
    "Stream via\nOpenAI API",
    "SSE tokens\nto browser",
    "Render\nresponse",
], "Figure 6: Agent Chat Flow \u2014 Contextual conversation with execution memory")

doc.add_heading("8.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "Open Chat Panel",
     "User clicks the Chat tab on the agent detail page (/agents/:id). The chat panel opens "
     "with a message input field and displays previous conversation messages (stored in local "
     "component state, not persisted to the backend).",
     "AgentDetailPage.tsx"),
    ("2", "Send Message",
     "User types a message and sends it. The frontend appends the message to the conversation "
     "history array [{role: 'user', content: '...'}] and prepares a ChatRequest payload with "
     "the message text and the last 10 conversation history entries.",
     "ChatPanel component"),
    ("3", "API Request",
     "Frontend sends POST /api/agents/:id/chat with Bearer JWT and body {message: '...', "
     "history: [{role, content}, ...]}. The backend verifies agent ownership.",
     "chat.py router"),
    ("4", "Load Memory Context",
     "Backend calls MemoryService.get_agent_context(db, agent_id, limit=15) to retrieve the "
     "last 15 execution summaries from the agent_memory table. These summaries are formatted "
     "as a context string with timestamps, summary text, and key outputs.",
     "memory_service.py"),
    ("5", "Construct System Prompt",
     "A system prompt is built containing the agent's name, type, description, status, schedule, "
     "and the formatted memory context. This gives the LLM full awareness of the agent's identity "
     "and past execution history.",
     "chat.py router"),
    ("6", "OpenAI Streaming Call",
     "Backend calls AsyncOpenAI.chat.completions.create() with model='gpt-4o-mini', the "
     "constructed messages array (system + last 10 history + new user message), stream=True, "
     "max_tokens=1024, and temperature=0.7. Tokens are yielded as they arrive.",
     "chat.py router"),
    ("7", "SSE Token Streaming",
     "Each token chunk from OpenAI is wrapped in a JSON SSE event {type: 'token', content: '...'} "
     "and sent to the client via sse-starlette EventSourceResponse. On completion, a {type: 'done'} "
     "event is sent. On error, a {type: 'error', content: '...'} event is sent.",
     "chat.py router"),
    ("8", "Render Response",
     "Frontend receives SSE events, accumulates token content into a growing response string, "
     "and renders it in the chat panel as an assistant message. Once the 'done' event is received, "
     "the complete response is added to the conversation history for future context.",
     "ChatPanel component"),
])

doc.add_paragraph()

add_body_paragraph(doc, (
    "If no OpenAI API key is configured (OPENAI_API_KEY environment variable is empty), the chat "
    "endpoint returns a non-streaming JSON response indicating that an API key is required. This "
    "ensures the platform does not crash in development environments without API keys configured."
))


# ===================================================================
#  9. EXECUTION MONITORING FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("9. Execution Monitoring Flow", level=1)

add_body_paragraph(doc, (
    "Execution monitoring allows users to inspect the progress and results of agent workflow runs. "
    "The monitoring experience spans two levels: the agent detail page (which shows execution history "
    "with a timeline view) and the execution detail page (which shows per-node waterfall charts, "
    "input/output data, LLM usage metrics, and error messages). Real-time monitoring during an active "
    "execution is handled via SSE streaming as described in the manual execution flow."
))

doc.add_heading("9.1 Flow Diagram", level=2)

create_flow_diagram(doc, [
    "Open\nDashboard /",
    "View activity\nfeed",
    "Click agent\nor execution",
    "Load\nexecution\nhistory",
    "Select\nexecution",
    "View node\nwaterfall\n+ logs",
    "Inspect node\noutput data",
], "Figure 7: Execution Monitoring Flow \u2014 From dashboard to per-node inspection")

doc.add_heading("9.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "View Dashboard",
     "User opens the dashboard (/). GET /api/dashboard/stats returns aggregate metrics: "
     "total_agents, active_agents, tasks_completed, estimated_savings (tasks * $25), and "
     "avg_response_time (ms). GET /api/dashboard/activity returns the 20 most recent executions "
     "with agent name, type, status, timing, and cost.",
     "DashboardPage.tsx"),
    ("2", "Browse Activity Feed",
     "The activity feed shows execution events with status badges (success/failed/running), "
     "trigger type (manual/schedule), duration, and cost. Each entry links to either the "
     "agent detail page or the execution detail page.",
     "DashboardPage.tsx"),
    ("3", "Navigate to Agent",
     "User clicks an agent name to navigate to /agents/:id. The agent detail page loads agent "
     "metadata (GET /api/agents/:id) and execution history (GET /api/agents/:id/executions "
     "with pagination).",
     "AgentDetailPage.tsx"),
    ("4", "View Execution History",
     "The execution history section shows a paginated list of past executions with status, "
     "trigger type, start/end times, duration, and cost. Each execution entry is clickable.",
     "AgentDetailPage.tsx"),
    ("5", "Open Execution Detail",
     "User clicks an execution to navigate to /executions/:executionId. GET /api/executions/:id "
     "returns the execution record with full node logs. The response includes per-node timing, "
     "status, input/output JSONB data, LLM usage, and error messages.",
     "ExecutionDetailPage.tsx"),
    ("6", "View Waterfall Timeline",
     "The execution detail page renders a waterfall chart showing each node's start time, "
     "duration, and status as colored bars. This visualises the execution order (BFS) and "
     "identifies bottlenecks (typically AI action nodes with LLM latency).",
     "ExecutionDetailPage.tsx"),
    ("7", "Inspect Node Data",
     "User clicks individual nodes in the waterfall to expand their details: input_data "
     "(the node configuration at execution time), output_data (the node's result), llm_usage "
     "(model, input_tokens, output_tokens, cost for AI nodes), and error_message (for failed nodes).",
     "ExecutionDetailPage.tsx"),
])


# ===================================================================
#  10. REPORTS & ANALYTICS FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("10. Reports & Analytics Flow", level=1)

add_body_paragraph(doc, (
    "The reports and analytics features provide aggregated views of agent performance and execution "
    "trends. The Reports page (/reports) shows per-agent output summaries from recent executions, "
    "enabling users to see what each agent has produced. The Analytics page (/analytics) provides "
    "charts, trend lines, and heatmaps for execution volume, success rates, cost analysis, and "
    "agent activity patterns over time."
))

doc.add_heading("10.1 Flow Diagram", level=2)

create_flow_diagram(doc, [
    "Open\n/reports or\n/analytics",
    "Fetch agent\nlist + stats",
    "Fetch\nexecution\nhistory",
    "Aggregate\nand compute\nmetrics",
    "Render\ncharts and\ntables",
], "Figure 8: Reports & Analytics Flow \u2014 Data aggregation and visualisation")

doc.add_heading("10.2 Detailed Steps", level=2)

create_steps_table(doc, [
    ("1", "Open Reports Page",
     "User navigates to /reports from the sidebar. The page fetches the user's agent list "
     "(GET /api/agents) and recent executions for each agent (GET /api/agents/:id/executions). "
     "Per-agent output summaries are extracted from the execution variables and node output data.",
     "ReportsPage.tsx"),
    ("2", "Open Analytics Page",
     "User navigates to /analytics from the sidebar. The page fetches dashboard stats "
     "(GET /api/dashboard/stats), activity feed (GET /api/dashboard/activity), and agent list "
     "(GET /api/agents) to build comprehensive analytics views.",
     "AnalyticsPage.tsx"),
    ("3", "Compute Aggregate Metrics",
     "Frontend computes derived metrics from the raw data: execution success rate per agent, "
     "average execution duration trends, total LLM cost per agent, execution frequency over "
     "time periods, and peak activity hours.",
     "AnalyticsPage.tsx"),
    ("4", "Render Charts",
     "The analytics page renders interactive charts: execution volume over time (bar/line chart), "
     "success/failure ratio (donut chart), cost breakdown by agent (bar chart), response time "
     "trends (line chart), and an activity heatmap showing execution density by hour and day.",
     "AnalyticsPage.tsx"),
    ("5", "Render Agent Reports",
     "The reports page renders per-agent summary cards showing the agent name, type, last "
     "execution timestamp, output summary text, and key metrics (total runs, success rate, "
     "average cost). Users can drill down to individual executions from each card.",
     "ReportsPage.tsx"),
])


# ===================================================================
#  11. FLOW INTEGRATION MAP
# ===================================================================

doc.add_page_break()

doc.add_heading("11. Flow Integration Map", level=1)

add_body_paragraph(doc, (
    "The eight application flows described in this document are not isolated paths \u2014 they interact "
    "and depend on each other. The following table maps the relationships between flows, showing "
    "how each flow feeds into or depends on other flows in the platform lifecycle."
))

# Integration map table
int_table = doc.add_table(rows=1, cols=3)
int_table.style = "Table Grid"
int_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(int_table, ["Flow", "Depends On", "Feeds Into"])

int_data = [
    ("Authentication", "None (entry point)",
     "All other flows (JWT required for every protected API call)"),
    ("Agent Creation", "Authentication (JWT required)",
     "Workflow Builder (new agent needs workflow design), Manual Execution (agent must exist)"),
    ("Workflow Builder", "Agent Creation (agent must exist)",
     "Manual Execution (workflow must have nodes), Scheduled Execution (deploy activates cron)"),
    ("Manual Execution", "Workflow Builder (workflow must be defined with nodes)",
     "Execution Monitoring (creates execution records), Agent Chat (execution memory accumulated), "
     "Reports & Analytics (execution data feeds aggregation)"),
    ("Scheduled Execution", "Workflow Builder (deploy with cron schedule)",
     "Execution Monitoring (creates execution records), Agent Chat (execution memory accumulated), "
     "Reports & Analytics (execution data feeds aggregation)"),
    ("Agent Chat", "Agent Creation (agent must exist), Manual/Scheduled Execution (memory context)",
     "None (terminal flow, informational only)"),
    ("Execution Monitoring", "Manual/Scheduled Execution (execution records must exist)",
     "None (terminal flow, read-only inspection)"),
    ("Reports & Analytics", "Manual/Scheduled Execution (execution data required)",
     "None (terminal flow, read-only aggregation)"),
]

for i, row_data in enumerate(int_data):
    row = int_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in int_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(2.5)
    row.cells[2].width = Inches(3.0)

add_figure_caption(doc, "Table 2: Flow Integration Map \u2014 Dependencies and downstream effects between application flows")

add_body_paragraph(doc, (
    "The integration map reveals a clear lifecycle pattern. Authentication is the universal prerequisite. "
    "Agent Creation and Workflow Builder form the setup phase. Manual and Scheduled Execution form the "
    "operational phase, producing execution records and memory that feed the observation phase (Monitoring, "
    "Chat, Reports, Analytics). This lifecycle can be summarised as: Authenticate \u2192 Create \u2192 Build \u2192 "
    "Execute \u2192 Observe."
))


# ===================================================================
#  SUMMARY TABLE - API ENDPOINTS PER FLOW
# ===================================================================

doc.add_page_break()

doc.add_heading("12. API Endpoints by Flow", level=1)

add_body_paragraph(doc, (
    "The following table provides a cross-reference of every API endpoint used across the eight "
    "application flows. This serves as a quick lookup for developers implementing or debugging "
    "specific user journeys."
))

api_table = doc.add_table(rows=1, cols=4)
api_table.style = "Table Grid"
api_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_purple_header_row(api_table, ["Flow", "Method", "Endpoint", "Purpose"])

api_data = [
    ("Authentication", "POST", "/api/auth/register", "Create new user account"),
    ("Authentication", "POST", "/api/auth/login", "Authenticate and receive JWT"),
    ("Authentication", "GET", "/api/auth/me", "Validate token and fetch profile"),
    ("Agent Creation", "POST", "/api/agents", "Create agent (manual or from template)"),
    ("Agent Creation", "POST", "/api/agents/generate", "AI-generate agent from description"),
    ("Agent Creation", "GET", "/api/templates", "List available agent templates"),
    ("Workflow Builder", "GET", "/api/agents/:id/workflow", "Load workflow definition"),
    ("Workflow Builder", "PUT", "/api/agents/:id/workflow", "Save workflow definition"),
    ("Workflow Builder", "POST", "/api/agents/:id/workflow/deploy", "Deploy and activate workflow"),
    ("Manual Execution", "POST", "/api/agents/:id/execute", "Trigger manual execution"),
    ("Manual Execution", "GET", "/api/executions/:id/stream", "SSE real-time log stream"),
    ("Scheduled Exec.", "POST", "/api/agents/:id/pause", "Pause agent and remove schedule"),
    ("Scheduled Exec.", "POST", "/api/agents/:id/resume", "Resume agent and restore schedule"),
    ("Agent Chat", "POST", "/api/agents/:id/chat", "Stream chat response via SSE"),
    ("Monitoring", "GET", "/api/agents/:id/executions", "List execution history"),
    ("Monitoring", "GET", "/api/executions/:id", "Get execution with node logs"),
    ("Monitoring", "GET", "/api/executions/:id/logs", "Get execution node logs"),
    ("Reports/Analytics", "GET", "/api/dashboard/stats", "Aggregate dashboard statistics"),
    ("Reports/Analytics", "GET", "/api/dashboard/activity", "Recent execution activity feed"),
    ("Integrations", "GET", "/api/integrations", "List configured integrations"),
    ("Integrations", "POST", "/api/integrations", "Connect/update integration"),
    ("Integrations", "DELETE", "/api/integrations/:id", "Disconnect integration"),
    ("Memory", "GET", "/api/agents/:id/memory", "List agent execution memories"),
    ("Memory", "DELETE", "/api/agents/:id/memory", "Clear agent memories"),
]

for i, row_data in enumerate(api_data):
    row = api_table.add_row()
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        if j == 0:
            run.bold = True
        elif j == 1:
            run.bold = True
            run.font.color.rgb = PURPLE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    if i % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F5F0FF")

for row in api_table.rows:
    row.cells[0].width = Inches(1.3)
    row.cells[1].width = Inches(0.7)
    row.cells[2].width = Inches(2.8)
    row.cells[3].width = Inches(2.2)

add_figure_caption(doc, "Table 3: Complete API endpoint cross-reference by application flow")


# ===================================================================
#  END OF DOCUMENT
# ===================================================================

doc.add_page_break()

doc.add_heading("End of Document", level=1)

add_body_paragraph(doc, (
    "This Application Flow document provides a complete mapping of user journeys through the GoalCert "
    "AutoMind platform. The eight flows cover the full lifecycle from authentication through agent "
    "creation, workflow design, execution (manual and scheduled), real-time monitoring, conversational "
    "chat with execution memory, and analytics consumption. Each flow is documented with a visual "
    "block diagram, a narrative description, and a detailed step table that maps actions to specific "
    "frontend components, API endpoints, and backend services."
))

add_body_paragraph(doc, (
    "This document should be read in conjunction with the High-Level Design (HLD), Low-Level Design "
    "(LLD), Data Flow Diagram (DFD), Business Requirements Document (BRD), and Functional Requirements "
    "Document (FRD) for a complete understanding of the AutoMind platform architecture and user experience."
))


# ===================================================================
#  SAVE
# ===================================================================

output_path = "/Users/narhen/automind/docs/GoalCert_AutoMind_Application_Flow_v1_0.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Document saved to: {output_path}")
